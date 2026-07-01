#!/usr/bin/env python3
# Copyright 2026 Adam Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Generate Imaging Intelligence Agent — Project Bible (DOCX + PDF).

Complete implementation reference for building the HCLS Imaging Intelligence
Agent on NVIDIA DGX Spark. Formatted in the same VCP-style theme as the
white paper and executive bullets.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Colors (identical to white paper) ──────────────────────
NAVY = RGBColor(0x1B, 0x23, 0x33)
TEAL = RGBColor(0x1A, 0xAF, 0xCC)
GREEN = RGBColor(0x76, 0xB9, 0x00)
GRAY_BODY = RGBColor(0x33, 0x33, 0x33)
GRAY_META = RGBColor(0x66, 0x66, 0x66)
GRAY_CODE = RGBColor(0x44, 0x44, 0x44)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_NAVY = "1B2333"
HEX_TEAL = "1AAFCC"
HEX_GREEN = "76B900"
HEX_LIGHT = "F8FAFB"
HEX_CODE_BG = "F5F5F5"
FONT = "Calibri"
CODE_FONT = "Consolas"

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
style.font.color.rgb = GRAY_BODY
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15


# ── Helper functions ───────────────────────────────────────

def R(p, text, bold=False, italic=False, size=None, color=None, font=None):
    r = p.add_run(text)
    r.font.name = font or FONT
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return r


def P(text="", bold=False, italic=False, size=10.5, color=GRAY_BODY,
      before=0, after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text:
        R(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def H1(text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        r.clear()
    R(p, text, bold=True, size=22, color=NAVY)
    return p


def H2(text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.clear()
    R(p, text, bold=True, size=16, color=NAVY)
    return p


def H3(text):
    return P(text, bold=True, size=13, color=TEAL, before=10, after=4)


def body(text, before=0, after=6):
    return P(text, size=10.5, color=GRAY_BODY, before=before, after=after)


def bullet(lead, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    R(p, lead, bold=True, size=10.5, color=GRAY_BODY)
    R(p, f"  {text}", size=10.5, color=GRAY_BODY)
    return p


def code_block(text, language=""):
    """Render a code block with monospace font and light gray background."""
    # Add a small label if language provided
    if language:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        R(p, language, bold=True, size=8, color=TEAL, font=FONT)

    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.2)
        R(p, line if line else " ", size=8, color=GRAY_CODE, font=CODE_FONT)


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="D0D0D0", sz="4"):
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>')
    tcPr.append(borders)


def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, HEX_NAVY)
        set_cell_borders(c, HEX_NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(p, h, bold=True, size=9, color=WHITE)
    for ri, row in enumerate(rows):
        bg = "FFFFFF" if ri % 2 == 0 else HEX_LIGHT
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            set_cell_shading(c, bg)
            set_cell_borders(c, "E0E0E0")
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if ci == 0:
                R(p, str(val), bold=True, size=9, color=TEAL)
            else:
                R(p, str(val), size=9, color=GRAY_BODY)
    return t


def spacer(pts=6):
    P("", after=pts)


# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════

# Green accent bar
accent = doc.add_table(rows=1, cols=1)
accent.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = accent.rows[0].cells[0]
set_cell_shading(cell, HEX_GREEN)
p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
R(p, " ", size=6, color=GREEN)

for _ in range(4):
    P("")

P("Project Bible", bold=True, size=13, color=TEAL, after=6)
P("Imaging Intelligence Agent",
  bold=True, size=32, color=NAVY, after=2)
P("Implementation Reference",
  bold=True, size=32, color=NAVY, after=6)
P("Complete architecture, code patterns, schemas, configurations, and "
  "implementation sequences for building the HCLS Imaging Intelligence "
  "Agent on NVIDIA DGX Spark.",
  italic=False, size=12, color=NAVY, after=6)
P("NVIDIA DGX Spark  |  MONAI Deploy  |  NVIDIA NIM  |  Open-Source Stack",
  italic=True, size=11, color=GRAY_META, after=12)
P("")
P("02/2026  |  Version 1.0  |  Apache 2.0 License", size=9, color=GRAY_META, after=2)
P("Author: Adam Jones", size=9, color=GRAY_META, after=2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
H1("Table of Contents")

toc_items = [
    "1.  Project Overview & Goals",
    "2.  DGX Spark Hardware Reference",
    "3.  Repository Layout",
    "4.  Docker Compose Services",
    "5.  PostgreSQL + pgvector Schema",
    "6.  Pydantic Data Models",
    "7.  Orthanc DICOM Server Configuration",
    "8.  MONAI Deploy MAP Pattern",
    "9.  CT Head \u2014 Hemorrhage Triage Workflow",
    "10. CT Chest \u2014 Lung Nodule Tracking Workflow",
    "11. CXR \u2014 Rapid Findings Workflow",
    "12. MRI Brain \u2014 MS Lesion Tracking Workflow",
    "13. DICOM SR Output (highdicom)",
    "14. FHIR DiagnosticReport Output",
    "15. LangGraph Agent Architecture",
    "16. NIM LLM Integration",
    "17. Embedding Service",
    "18. Nextflow Pipeline Orchestration",
    "19. Streamlit Portal",
    "20. Monitoring Stack",
    "21. Testing Strategy",
    "22. ARM64 Compatibility Guide",
    "23. Configuration Reference",
    "24. Implementation Sequence",
]
for item in toc_items:
    P(item, size=11, color=NAVY, before=2, after=2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW & GOALS
# ══════════════════════════════════════════════════════════════
H1("1. Project Overview & Goals")

H3("What This Agent Does")
body(
    "The Imaging Intelligence Agent processes CT, MRI, and X-ray studies using NVIDIA "
    "MONAI models on DGX hardware. It provides automated detection, segmentation, "
    "classification, longitudinal tracking, and clinical triage \u2014 from DICOM ingestion "
    "through agentic inference to structured clinical output."
)

H3("Four Reference Workflows")
add_table(
    ["Workflow", "Modality", "Target Latency", "Key Metric"],
    [
        ["Hemorrhage Triage", "CT Head", "< 90 seconds",
         "Sensitivity > 95% for hemorrhage > 5 mL"],
        ["Lung Nodule Tracking", "CT Chest", "< 5 minutes",
         "Detection > 90% for nodules \u2265 4 mm"],
        ["Rapid Findings", "CXR", "< 30 seconds",
         "Pneumothorax sensitivity > 95%"],
        ["MS Lesion Tracking", "MRI Brain", "< 5 minutes",
         "3D U-Net on FLAIR sequences"],
    ]
)

spacer()
H3("Pipeline Pattern")
body("Every workflow follows the same canonical pattern:")
bullet("1. Ingestion:", "DICOM arrives via DICOMweb STOW-RS or DIMSE C-STORE")
bullet("2. Event trigger:", "study.complete event fires from Orthanc")
bullet("3. Prior retrieval:", "Query PostgreSQL for prior studies, retrieve from Orthanc")
bullet("4. GPU inference:", "MONAI Deploy MAP executes on DGX Spark GPU")
bullet("5. Post-processing:", "Measurements, classifications, embeddings computed")
bullet("6. Persistence:", "Findings \u2192 PostgreSQL, embeddings \u2192 pgvector")
bullet("7. Output encoding:", "DICOM SR \u2192 PACS, FHIR DiagnosticReport \u2192 EHR")
bullet("8. Triage routing:", "Worklist prioritization, on-call alerts dispatched")
bullet("9. Provenance:", "Immutable audit trail persisted")

spacer()
H3("HCLS AI Factory Integration")
body(
    "This agent is one node in the broader HCLS AI Factory. Cross-modal triggers include:"
)
bullet("Imaging \u2192 Genomics (Parabricks):",
       "Lung-RADS 4B+ triggers somatic/germline tumor profiling. "
       "Parabricks: 30x WGS from ~30 hours (CPU) to ~10 minutes on DGX (8-GPU).")
bullet("Imaging \u2192 Drug Discovery (BioNeMo):",
       "Quantitative imaging endpoints (RECIST) for clinical trial candidate scoring.")
bullet("Imaging \u2192 Clinical Reasoning (NIM LLM):",
       "RAG-grounded clinical reports with guideline and literature evidence.")
bullet("Imaging \u2192 Biomarker Intelligence Agent:",
       "Genomic + imaging biomarker fusion on shared DGX Spark compute.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. DGX SPARK HARDWARE REFERENCE
# ══════════════════════════════════════════════════════════════
H1("2. DGX Spark Hardware Reference")

H3("Specifications")
add_table(
    ["Parameter", "Value"],
    [
        ["CPU", "NVIDIA Grace (ARM64 / aarch64)"],
        ["GPU", "NVIDIA Blackwell GB10, 1 GPU"],
        ["Memory", "128 GB unified LPDDR5x (CPU + GPU shared pool)"],
        ["Storage", "Up to 4 TB NVMe"],
        ["Storage Access", "GPUDirect Storage (zero-copy GPU access)"],
        ["Price", "$4,699"],
        ["OS", "Ubuntu-based (NVIDIA DGX OS)"],
        ["NVAIE Cost", "Zero at desktop-class"],
    ]
)

spacer()
H3("Critical: ARM64 Architecture")
body(
    "ALL containers must be ARM64-compatible. The Grace CPU is aarch64, not x86_64. "
    "This affects base Docker images (must use ARM64 variants), Python wheel availability, "
    "NIM containers (use -dgx-spark variant image tags), and any compiled C/C++ extensions."
)

H3("Unified Memory Model")
body(
    "The 128 GB is shared between CPU and GPU \u2014 there is no separate GPU VRAM. "
    "No explicit CPU\u2192GPU data transfers are needed for many operations. Memory pressure "
    "from CPU workloads reduces GPU-available memory. Monitor total system memory, "
    "not just \u201cGPU memory.\u201d"
)

H3("DGX Compute Progression")
add_table(
    ["Phase", "Hardware", "Price", "Scope"],
    [
        ["1 \u2014 Proof Build", "DGX Spark", "$4,699", "1\u20132 workflows"],
        ["2 \u2014 Departmental", "1\u20132x DGX B200", "$500K\u2013$1M",
         "All workflows, PACS integration"],
        ["3 \u2014 Multi-Site", "4\u20138x DGX B200 + InfiniBand", "$2M\u2013$4M",
         "FLARE federated learning"],
        ["4 \u2014 AI Factory", "DGX SuperPOD", "$7M\u2013$60M+",
         "Thousands of concurrent studies"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. REPOSITORY LAYOUT
# ══════════════════════════════════════════════════════════════
H1("3. Repository Layout")

body("The imaging agent follows the same conventions as the existing HCLS AI Factory "
     "(Nextflow DSL2, Docker Compose, Streamlit, Prometheus + Grafana).")

code_block("""hls-imaging-agent/
\u251c\u2500\u2500 main.nf                          # Nextflow DSL2 entry point
\u251c\u2500\u2500 nextflow.config                  # Profiles: docker, dgx_spark
\u251c\u2500\u2500 docker-compose.yml               # All services
\u251c\u2500\u2500 docker-compose.dev.yml           # Dev overrides (mock NIM, etc.)
\u251c\u2500\u2500 start-services.sh                # Service startup script
\u251c\u2500\u2500 demo.sh                          # Demo launcher
\u251c\u2500\u2500 .env.example                     # Environment variable template
\u2502
\u251c\u2500\u2500 modules/                         # Nextflow workflow modules
\u2502   \u251c\u2500\u2500 ct_head_hemorrhage.nf
\u2502   \u251c\u2500\u2500 ct_chest_lung_nodule.nf
\u2502   \u251c\u2500\u2500 cxr_rapid_findings.nf
\u2502   \u2514\u2500\u2500 mri_brain_ms_lesion.nf
\u2502
\u251c\u2500\u2500 maps/                            # MONAI Deploy Application Packages
\u2502   \u251c\u2500\u2500 ct_head_hemorrhage/  (Dockerfile, app.py, operators.py)
\u2502   \u251c\u2500\u2500 ct_chest_lung_nodule/
\u2502   \u251c\u2500\u2500 cxr_rapid_findings/
\u2502   \u2514\u2500\u2500 mri_brain_ms_lesion/
\u2502
\u251c\u2500\u2500 agent/                           # LangGraph clinical reasoning
\u2502   \u251c\u2500\u2500 graph.py, nodes.py, tools.py, state.py, prompts.py
\u2502
\u251c\u2500\u2500 services/                        # Microservices
\u2502   \u251c\u2500\u2500 dicom_listener/  (listener.py \u2014 Orthanc webhook)
\u2502   \u251c\u2500\u2500 fhir_publisher/  (publisher.py \u2014 FHIR output)
\u2502   \u251c\u2500\u2500 embedding_service/  (embedder.py \u2014 384-dim vectors)
\u2502   \u2514\u2500\u2500 portal/  (app.py \u2014 Streamlit dashboard)
\u2502
\u251c\u2500\u2500 src/                             # Shared library
\u2502   \u251c\u2500\u2500 models.py, db.py, dicom_utils.py, fhir_utils.py, config.py
\u2502
\u251c\u2500\u2500 db/init.sql                      # PostgreSQL + pgvector schema
\u251c\u2500\u2500 config/orthanc.json              # Orthanc DICOM server config
\u251c\u2500\u2500 config/prometheus.yml             # Monitoring
\u251c\u2500\u2500 tests/                           # pytest test suite
\u251c\u2500\u2500 scripts/download_models.sh        # MONAI Model Zoo downloads
\u2514\u2500\u2500 docs/diagrams/architecture.mmd   # Mermaid architecture diagram""", "")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. DOCKER COMPOSE SERVICES
# ══════════════════════════════════════════════════════════════
H1("4. Docker Compose Services")

H3("Port Allocation")
body("Existing HCLS AI Factory uses ports 8001\u20138510. Imaging agent uses 8520+:")
add_table(
    ["Service", "Port", "Protocol"],
    [
        ["Orthanc (DIMSE)", "4242", "DIMSE"],
        ["Orthanc (DICOMweb)", "8042", "HTTP REST"],
        ["PostgreSQL", "5432", "TCP"],
        ["NIM LLM", "8520", "HTTP (OpenAI-compatible)"],
        ["Embedding Service", "8521", "HTTP REST"],
        ["DICOM Listener", "8522", "HTTP webhook"],
        ["FHIR Publisher", "8523", "HTTP REST"],
        ["Agent API", "8524", "HTTP REST"],
        ["Streamlit Portal", "8525", "HTTP"],
        ["Prometheus", "9099", "HTTP"],
        ["Grafana", "3000", "HTTP"],
        ["DCGM Exporter", "9400", "HTTP"],
    ]
)

spacer()
H3("Service Architecture")
body(
    "All 11 services are defined in docker-compose.yml. Key patterns from the existing "
    "HCLS AI Factory: GPU resource reservations for NVIDIA containers, health checks with "
    "curl, named volumes for persistence, depends_on with service_healthy conditions, and "
    "environment variables sourced from .env."
)

bullet("Data Layer:", "Orthanc (DICOM server + DICOMweb) + PostgreSQL with pgvector extension")
bullet("Execution Layer:", "NIM LLM (ARM64 -dgx-spark variant), Embedding Service, DICOM Listener, "
       "FHIR Publisher, LangGraph Agent")
bullet("Presentation:", "Streamlit Portal with NVIDIA green theme")
bullet("Monitoring:", "DCGM Exporter + Prometheus + Grafana")

spacer()
H3("NIM Container on DGX Spark")
code_block("""# ARM64 NIM image for DGX Spark
image: nvcr.io/nvidia/nim/meta-llama3-8b-instruct:latest-dgx-spark
environment:
  - NGC_API_KEY=${NGC_API_KEY}
deploy:
  resources:
    reservations:
      devices:
        - driver: nvidia
          count: 1
          capabilities: [gpu]""", "yaml")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. POSTGRESQL + PGVECTOR SCHEMA
# ══════════════════════════════════════════════════════════════
H1("5. PostgreSQL + pgvector Schema")

body(
    "A single PostgreSQL instance with the pgvector extension handles structured queries, "
    "vector similarity search, and hybrid queries. No separate relational, vector, or "
    "analytics database required."
)

H3("Core Tables")
add_table(
    ["Table", "Purpose", "Key Columns"],
    [
        ["studies", "Imaging study metadata",
         "study_instance_uid, patient_id, modality, body_part, status"],
        ["series", "Series within a study",
         "series_instance_uid, study_id, modality, num_instances"],
        ["findings", "Clinical findings from workflows",
         "workflow, finding_type, severity, confidence, details (JSONB)"],
        ["measurements", "Quantitative measurements",
         "measurement_type, value, unit, prior_value, delta_percent"],
        ["embeddings", "384-dim vectors for similarity",
         "level (study/series/lesion), embedding vector(384), HNSW index"],
        ["provenance", "Audit trail per inference",
         "model_id, model_version, duration_ms, input_uids, inference_params"],
        ["worklist_entries", "Triage routing",
         "urgency, priority (P1\u2013P4), notification, routing, acknowledged"],
    ]
)

spacer()
H3("pgvector Index")
code_block("""-- HNSW index for fast approximate nearest neighbor search
CREATE INDEX idx_embeddings_hnsw ON embeddings
    USING hnsw (embedding vector_cosine_ops)
    WITH (m = 16, ef_construction = 64);""", "sql")

spacer()
H3("Hybrid Query Example")
code_block("""-- Growing nodules AND similar phenotype
WITH growing_nodules AS (
    SELECT f.study_id, f.id AS finding_id
    FROM findings f
    JOIN measurements m ON m.finding_id = f.id
    WHERE f.workflow = 'ct_chest_nodule'
      AND m.measurement_type = 'doubling_time'
      AND m.value < 400
)
SELECT s.patient_id, s.study_date,
       e.embedding <=> $1::vector AS phenotype_distance
FROM growing_nodules gn
JOIN studies s ON gn.study_id = s.id
JOIN embeddings e ON e.study_id = s.id AND e.level = 'study'
ORDER BY e.embedding <=> $1::vector
LIMIT 10;""", "sql")

spacer()
H3("Helper Views")
bullet("active_worklist:", "Unacknowledged entries ordered by priority (P1 first)")
bullet("study_summary:", "Study metadata with finding counts, critical/urgent counts, max confidence")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. PYDANTIC DATA MODELS
# ══════════════════════════════════════════════════════════════
H1("6. Pydantic Data Models")

body(
    "All domain objects are defined as Pydantic models in src/models.py, following the "
    "same patterns as the existing HCLS AI Factory (Field validators, enums for type safety, "
    "nested models). Loguru is used for logging throughout."
)

H3("Enums")
add_table(
    ["Enum", "Values"],
    [
        ["Modality", "CT, MR, CR, DX"],
        ["BodyPart", "HEAD, CHEST, BRAIN, ABDOMEN, MSK"],
        ["Severity", "critical, urgent, moderate, routine"],
        ["Priority", "P1, P2, P3, P4"],
        ["WorkflowType", "ct_head_hemorrhage, ct_chest_nodule, cxr_findings, mri_ms_lesion"],
        ["StudyStatus", "received, processing, completed, failed"],
        ["DiseaseActivity", "stable, active, highly_active"],
        ["LungRADS", "1, 2, 3, 4A, 4B, 4X"],
    ]
)

spacer()
H3("Core Models")
add_table(
    ["Model", "Purpose", "Key Fields"],
    [
        ["Study", "DICOM imaging study",
         "study_instance_uid, patient_id, modality, body_part, status"],
        ["Finding", "Clinical finding from a workflow",
         "workflow, finding_type, severity, confidence (0\u20131), details (dict)"],
        ["Measurement", "Quantitative measurement",
         "measurement_type, value, unit, prior_value, delta_percent"],
        ["EmbeddingRecord", "384-dim vector",
         "study_id, level, embedding (list[float])"],
        ["ProvenanceBundle", "Inference audit trail",
         "model_id, model_version, input_uids, duration_ms"],
        ["WorklistEntry", "Triage routing",
         "urgency, priority, notification, routing"],
    ]
)

spacer()
H3("Workflow-Specific Models")
add_table(
    ["Model", "Workflow", "Key Fields"],
    [
        ["HemorrhageResult", "CT Head",
         "detected, hemorrhage_type, volume_ml, midline_shift_mm, urgency"],
        ["NoduleResult", "CT Chest",
         "nodule_id, lung_rads, volume_mm3, doubling_time_days, malignancy_risk"],
        ["CXRFindingResult", "CXR",
         "finding_name, detected, confidence, gradcam_region"],
        ["MSLesionResult", "MRI Brain",
         "total_lesion_count, new_lesion_count, disease_activity, volume_change_percent"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. ORTHANC DICOM SERVER
# ══════════════════════════════════════════════════════════════
H1("7. Orthanc DICOM Server Configuration")

body(
    "Orthanc serves as the DICOM server and archive, providing DICOMweb endpoints "
    "(STOW-RS, WADO-RS, QIDO-RS) for standards-based integration with PACS and the "
    "imaging pipeline."
)

H3("Key Configuration (orthanc.json)")
bullet("AET:", "IMAGING_AGENT")
bullet("DIMSE Port:", "4242")
bullet("HTTP Port:", "8042")
bullet("DICOMweb:", "Enabled at /dicom-web/ root")
bullet("StableAge:", "10 seconds (triggers study.complete)")

spacer()
H3("study.complete Event (Lua Script)")
body(
    "When a study has been stable (no new instances) for StableAge seconds, Orthanc fires "
    "the OnStableStudy Lua callback. This POSTs a webhook to the DICOM listener service, "
    "which triggers the appropriate Nextflow pipeline."
)
code_block("""function OnStableStudy(studyId, tags, metadata)
    -- Port 8000 is the container-internal port; mapped to host 8522 in docker-compose
    local url = "http://dicom-listener:8000/webhook/study-complete"
    local body = '{"orthanc_id": "' .. studyId .. '"}'
    HttpPost(url, body, {["Content-Type"] = "application/json"})
end""", "lua")

spacer()
H3("DICOMweb Client Usage (Python)")
code_block("""from dicomweb_client import DICOMwebClient

client = DICOMwebClient(url="http://orthanc:8042/dicom-web")

# QIDO-RS: Query studies
studies = client.search_for_studies(search_filters={"PatientID": "IMG-2026-0142"})

# WADO-RS: Retrieve instances as pydicom datasets
instances = client.retrieve_series(
    study_instance_uid=study_uid, series_instance_uid=series_uid)

# STOW-RS: Store DICOM SR / SEG / GSPS
client.store_instances(datasets=[sr_dataset])""", "python")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. MONAI DEPLOY MAP PATTERN
# ══════════════════════════════════════════════════════════════
H1("8. MONAI Deploy MAP Pattern")

body(
    "Each imaging workflow is packaged as a MONAI Deploy Application Package (MAP) \u2014 "
    "a containerized, portable inference pipeline with standardized I/O conventions."
)

H3("Application Class Structure")
code_block("""from monai.deploy.core import Application, resource

@resource(cpu=4, gpu=1, memory="16Gi")
class WorkflowApp(Application):
    def compose(self):
        preprocess_op = PreprocessOperator()
        inference_op = InferenceOperator()
        postprocess_op = PostprocessOperator()
        self.add_flow(preprocess_op, inference_op)
        self.add_flow(inference_op, postprocess_op)""", "python")

spacer()
H3("Operator Pattern")
body(
    "Each operator defines setup() with input/output specs and compute() for processing. "
    "MONAI transforms (LoadImaged, Spacingd, ScaleIntensityRanged, etc.) handle "
    "preprocessing. Model weights are loaded from mounted volumes at /models/."
)

H3("I/O Conventions")
add_table(
    ["Path", "Purpose"],
    [
        ["/var/holoscan/input/", "DICOM input directory (mounted by orchestrator)"],
        ["/var/holoscan/output/", "Output directory (SR, SEG, measurements JSON)"],
        ["/models/", "Model weights directory (mounted volume)"],
    ]
)

spacer()
H3("MAP Dockerfile")
code_block("""FROM nvcr.io/nvidia/pytorch:24.01-py3
WORKDIR /app
RUN pip install --no-cache-dir \\
    monai-deploy-app-sdk>=0.6.0 monai>=1.3.0 \\
    pydicom>=2.4.0 highdicom>=0.22.0 numpy scipy
COPY app.py operators.py requirements.txt ./
ENV MONAI_DEPLOY_MODEL_PATH=/models
ENTRYPOINT ["python", "app.py"]""", "dockerfile")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. CT HEAD HEMORRHAGE
# ══════════════════════════════════════════════════════════════
H1("9. CT Head \u2014 Hemorrhage Triage Workflow")

bullet("Target:", "< 90 seconds end-to-end  |  Sensitivity: > 95% for hemorrhage > 5 mL")
bullet("Model:", "3D U-Net (MONAI) \u2014 hemorrhage-triage-v2.1, validated on RSNA ICH (752K slices)")

spacer()
H3("Pipeline Stages")
bullet("1. Preprocessing:", "Load NCCT head DICOM \u2192 reorient RAS \u2192 resample 1mm isotropic \u2192 "
       "window (W:80 L:40 for blood) \u2192 normalize [0,1]")
bullet("2. Segmentation:", "3D U-Net binary segmentation (hemorrhage vs. normal)")
bullet("3. Volume estimation:", "Count positive voxels \u00d7 voxel volume")
bullet("4. Midline shift:", "Compute brain centerline from falx cerebri, measure max lateral displacement")
bullet("5. Classification:", "Map volume + shift + thickness to urgency")
bullet("6. Output:", "Finding \u2192 Measurements \u2192 WorklistEntry \u2192 DICOM SR")

spacer()
H3("MONAI Transforms Pipeline")
code_block("""CT_HEAD_PREPROCESS = Compose([
    LoadImaged(keys=["image"], reader="PydicomReader"),
    EnsureChannelFirstd(keys=["image"]),
    Orientationd(keys=["image"], axcodes="RAS"),
    Spacingd(keys=["image"], pixdim=(1.0, 1.0, 1.0), mode="bilinear"),
    ScaleIntensityRanged(keys=["image"], a_min=0, a_max=80,
                         b_min=0.0, b_max=1.0, clip=True),
    CropForegroundd(keys=["image"], source_key="image", margin=10),
    EnsureTyped(keys=["image"], dtype="float32"),
])""", "python")

spacer()
H3("Model Architecture")
code_block("""model = UNet(
    spatial_dims=3, in_channels=1, out_channels=2,
    channels=(16, 32, 64, 128, 256), strides=(2, 2, 2, 2),
    num_res_units=2, norm="batch",
)""", "python")

spacer()
H3("Urgency Classification (Brain Trauma Foundation)")
add_table(
    ["Condition", "Severity", "Priority"],
    [
        ["Volume > 30 mL OR shift > 5 mm OR thickness > 10 mm",
         "Critical", "P1 \u2014 Stat"],
        ["Volume > 5 mL (below surgical thresholds)",
         "Urgent", "P2 \u2014 Urgent"],
        ["Volume \u2264 5 mL",
         "Routine", "P4 \u2014 Routine"],
    ]
)

spacer()
H3("Midline Shift Measurement")
body(
    "Computed from the center of mass of the hemorrhage segmentation relative to the "
    "brain midline (center of the axial plane in RAS orientation). scipy.ndimage."
    "center_of_mass provides the hemorrhage centroid; displacement is multiplied by "
    "voxel spacing to obtain shift in mm."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 10. CT CHEST LUNG NODULE
# ══════════════════════════════════════════════════════════════
H1("10. CT Chest \u2014 Lung Nodule Tracking Workflow")

bullet("Target:", "< 5 minutes multi-stage  |  Detection: > 90% for nodules \u2265 4 mm")
bullet("Models:", "RetinaNet (detection, MONAI) + SegResNet (per-nodule segmentation, MONAI)")

spacer()
H3("Pipeline Stages (10 steps)")
bullet("1. Preprocessing:", "Load CT chest \u2192 RAS \u2192 1mm isotropic \u2192 lung window (W:1500 L:-600)")
bullet("2. Detection:", "RetinaNet detects candidate nodules with bounding boxes")
bullet("3. Segmentation:", "Per-nodule SegResNet segmentation within each bounding box")
bullet("4. Volumetrics:", "Voxel counting \u00d7 voxel spacing product")
bullet("5. Prior retrieval:", "Query PostgreSQL for prior CT chest \u2192 retrieve from Orthanc")
bullet("6. Registration:", "SyN diffeomorphic (ANTsPy) to align current and prior")
bullet("7. Volume Doubling Time:", "VDT = (\u0394t \u00d7 ln2) / ln(V2/V1)")
bullet("8. Lung-RADS:", "Rule-based assignment per ACR Lung-RADS v2022")
bullet("9. Risk scoring:", "Composite from VDT, morphology, location, patient factors")
bullet("10. Cross-modal trigger:", "Lung-RADS 4B+ \u2192 Parabricks genomics pipeline")

spacer()
H3("Volume Doubling Time")
code_block("""def calculate_vdt(current_volume, prior_volume, current_date, prior_date):
    delta_days = (current_date - prior_date).days
    vdt = (delta_days * math.log(2)) / math.log(current_volume / prior_volume)
    return round(vdt, 1)""", "python")

spacer()
H3("Lung-RADS Classification (ACR v2022)")
add_table(
    ["Nodule Type", "Size Range", "Lung-RADS", "Management"],
    [
        ["Solid", "< 4 mm", "1", "Annual LDCT"],
        ["Solid", "4\u20135 mm", "2", "Annual LDCT"],
        ["Solid", "6\u20137 mm", "3", "6-month follow-up"],
        ["Solid", "8\u201314 mm", "4A", "3-month LDCT; consider PET-CT"],
        ["Solid", "\u2265 15 mm", "4B", "Tissue sampling"],
        ["Ground-glass", "< 20 mm", "2", "Annual LDCT"],
        ["Ground-glass", "20\u201329 mm", "3", "6-month follow-up"],
        ["Ground-glass", "\u2265 30 mm", "4A", "3-month LDCT"],
        ["Part-solid", "< 6 mm", "2", "Annual LDCT"],
        ["Part-solid", "\u2265 8 mm", "4A", "3-month LDCT"],
    ]
)

spacer()
body(
    "Growth upgrade: VDT < 400 days upgrades category (2/3 \u2192 4A, 4A \u2192 4B). "
    "Lung-RADS 4B+ triggers Parabricks somatic variant analysis via Nextflow event."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 11. CXR RAPID FINDINGS
# ══════════════════════════════════════════════════════════════
H1("11. CXR \u2014 Rapid Findings Workflow")

bullet("Target:", "< 30 seconds end-to-end  |  Pneumothorax sensitivity: > 95%")
bullet("Model:", "DenseNet-121 (MONAI), 2D, 5 output heads. Validated on CheXpert + MIMIC-CXR (601K images).")

spacer()
H3("Multi-Label Classification")
add_table(
    ["Finding", "Confidence Threshold", "Clinical Significance"],
    [
        ["Pneumothorax", "0.50", "High-risk \u2014 lower threshold for safety"],
        ["Consolidation", "0.60", "Infection, atelectasis indicator"],
        ["Pleural Effusion", "0.55", "Fluid accumulation"],
        ["Cardiomegaly", "0.60", "Cardiac enlargement"],
        ["Fracture", "0.55", "Rib / clavicle fracture"],
    ]
)

spacer()
H3("GradCAM Heatmap Generation")
body(
    "GradCAM localizes each positive finding to an anatomic region. Implementation uses "
    "monai.visualize.GradCAM targeting the class_layers.relu layer in DenseNet-121. "
    "Heatmaps are stored as DICOM Secondary Capture images and GSPS graphic annotations "
    "for display in PACS."
)
code_block("""from monai.visualize import GradCAM

cam = GradCAM(nn_module=model, target_layers="class_layers.relu")
heatmap = cam(x=input_tensor, class_idx=target_class)  # (B, 1, H, W)""", "python")

spacer()
H3("Pipeline")
bullet("1. Preprocessing:", "Load CXR DICOM \u2192 resize 224\u00d7224 \u2192 normalize (ImageNet stats)")
bullet("2. Classification:", "DenseNet-121 multi-label with sigmoid activation")
bullet("3. GradCAM:", "Heatmap per positive finding")
bullet("4. Output:", "DICOM SR findings + Secondary Capture heatmaps + GSPS overlays")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 12. MRI BRAIN MS LESION
# ══════════════════════════════════════════════════════════════
H1("12. MRI Brain \u2014 MS Lesion Tracking Workflow")

bullet("Target:", "< 5 minutes multi-stage")
bullet("Model:", "3D U-Net (MONAI), ms-lesion-v2.3. Validated on ISBI MS Challenge + institutional (1,200 MRIs).")

spacer()
H3("Pipeline Stages")
bullet("1. Preprocessing:", "Load FLAIR DICOM \u2192 RAS \u2192 1mm isotropic \u2192 z-score normalize")
bullet("2. Segmentation:", "3D U-Net binary segmentation on FLAIR (lesion vs. normal)")
bullet("3. Connected components:", "scipy.ndimage.label to identify individual lesions")
bullet("4. Volumetrics:", "Per-lesion and total volume measurement")
bullet("5. Prior retrieval:", "Query PostgreSQL for prior MRI brain \u2192 retrieve from Orthanc")
bullet("6. Registration:", "Affine + SyN (ANTsPy SyNRA) to align current and prior FLAIR")
bullet("7. Lesion matching:", "Overlap analysis (Dice) to classify new / stable / enlarging")
bullet("8. Disease activity:", "stable / active / highly active based on lesion changes")

spacer()
H3("Disease Activity Classification")
add_table(
    ["Condition", "Classification"],
    [
        ["0 new + 0 enlarging + 0 enhancing", "Stable"],
        ["1\u20132 new/enlarging OR 1 enhancing", "Active"],
        ["\u2265 3 new/enlarging OR \u2265 2 enhancing", "Highly Active"],
    ]
)

spacer()
H3("Spatial Registration")
code_block("""import ants

result = ants.registration(
    fixed=current_flair, moving=prior_flair,
    type_of_transform="SyNRA",  # Rigid + Affine + SyN
)
warped_prior = result["warpedmovout"]
# Apply same transform to prior lesion mask for overlap analysis""", "python")

spacer()
H3("Lesion Matching")
body(
    "After registration, each current lesion is compared against warped prior lesions. "
    "Dice overlap \u2265 0.3 = matched. Volume increase > 50 mm\u00b3 = enlarging. "
    "No match = new lesion."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 13. DICOM SR OUTPUT
# ══════════════════════════════════════════════════════════════
H1("13. DICOM SR Output (highdicom)")

body(
    "Structured findings are encoded as DICOM Structured Reports (TID 1500 Measurement "
    "Report) using the highdicom library. SR objects are pushed to PACS via DICOMweb "
    "STOW-RS for native viewing in the radiologist\u2019s environment."
)

H3("Key Classes")
add_table(
    ["Class", "Purpose"],
    [
        ["Comprehensive3DSR", "Top-level SR document (TID 1500)"],
        ["MeasurementReport", "Container for imaging measurements"],
        ["NumericMeasurement", "Coded numeric value with units (UCUM)"],
        ["FindingSite", "Anatomic location (SNOMED CT coded)"],
        ["TrackingIdentifier", "Unique ID for longitudinal tracking"],
        ["ObserverContext", "Device observer (ImagingIntelligenceAgent)"],
    ]
)

spacer()
H3("Construction Pattern")
code_block("""import highdicom as hd
from highdicom.sr import Comprehensive3DSR, MeasurementReport, NumericMeasurement

# Create measurement
volume = NumericMeasurement(
    name=CodedConcept(value="118565006", meaning="Volume", scheme_designator="SCT"),
    value=12.3,
    unit=CodedConcept(value="mL", meaning="milliliter", scheme_designator="UCUM"),
)

# Build report + SR
report = MeasurementReport(imaging_measurements=[volume], ...)
sr = Comprehensive3DSR(evidence=[source_dataset], content=report, ...)

# Push to Orthanc
client.store_instances(datasets=[sr])""", "python")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 14. FHIR OUTPUT
# ══════════════════════════════════════════════════════════════
H1("14. FHIR DiagnosticReport Output")

body(
    "Clinical findings are also encoded as FHIR R4 resources for EHR integration. "
    "The fhir.resources library (Pydantic-based) constructs DiagnosticReport, "
    "Observation, and ImagingStudy references. Results are bundled into a FHIR "
    "transaction Bundle."
)

H3("Resource Mapping")
add_table(
    ["FHIR Resource", "Agent Output", "Coding System"],
    [
        ["DiagnosticReport", "Overall study report with conclusion", "LOINC 18748-4"],
        ["Observation", "Individual finding (one per detection)", "SNOMED CT"],
        ["Observation.component", "Quantitative measurement", "LOINC"],
        ["ImagingStudy", "Reference to DICOM study", "DICOM UID"],
        ["Bundle (transaction)", "Wraps all resources for submission", "\u2014"],
    ]
)

spacer()
H3("Construction Pattern")
code_block("""from fhir.resources.diagnosticreport import DiagnosticReport
from fhir.resources.observation import Observation
from fhir.resources.bundle import Bundle, BundleEntry, BundleEntryRequest

report = DiagnosticReport(
    status="final",
    code=CodeableConcept(coding=[Coding(system="http://loinc.org", code="18748-4")]),
    conclusion="Acute subdural hemorrhage, 12.3 mL...",
    result=[Reference(reference="Observation/0"), ...],
)

bundle = Bundle(type="transaction", entry=[
    BundleEntry(resource=report, request=BundleEntryRequest(method="POST", url="DiagnosticReport")),
    BundleEntry(resource=observation, request=BundleEntryRequest(method="POST", url="Observation")),
])""", "python")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 15. LANGGRAPH AGENT
# ══════════════════════════════════════════════════════════════
H1("15. LangGraph Agent Architecture")

body(
    "Multi-step clinical reasoning is orchestrated by a LangGraph StateGraph with "
    "four nodes: triage, longitudinal analysis, population search, and report generation. "
    "MCP tools provide database access. LangChain\u2019s ChatOpenAI client connects to the "
    "local NIM LLM via OpenAI-compatible API."
)

H3("Agent State (TypedDict)")
code_block("""class AgentState(TypedDict):
    messages: Annotated[list, add_messages]
    study_id: int
    patient_id: str
    modality: str
    findings: list[dict]
    measurements: list[dict]
    worklist_entries: list[dict]
    prior_studies: list[dict]
    similar_studies: list[dict]
    rag_context: str
    report_text: str
    triage_complete: bool""", "python")

spacer()
H3("Graph Topology")
add_table(
    ["Node", "Role", "Next"],
    [
        ["triage", "Classify findings by urgency, create worklist entries",
         "longitudinal (if critical/urgent) OR report"],
        ["longitudinal", "Retrieve prior studies, compute deltas",
         "population"],
        ["population", "Find similar studies via embedding search",
         "report"],
        ["report", "Generate evidence-grounded clinical summary via RAG",
         "END"],
    ]
)

spacer()
H3("MCP Tools")
add_table(
    ["Tool", "Purpose"],
    [
        ["query_findings", "Structured SQL query on findings table"],
        ["search_similar_studies", "pgvector nearest neighbor search"],
        ["get_prior_measurements", "Longitudinal measurement retrieval"],
    ]
)

spacer()
H3("Checkpointing")
body(
    "MemorySaver for development; PostgreSQL-backed checkpointer for production. "
    "Enables replay and debugging of agent reasoning chains."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 16. NIM LLM INTEGRATION
# ══════════════════════════════════════════════════════════════
H1("16. NIM LLM Integration")

body(
    "NVIDIA NIM serves the LLM on DGX Spark via an ARM64 container. The API is "
    "OpenAI-compatible, so LangChain\u2019s ChatOpenAI client connects directly."
)

H3("Deployment")
bullet("Image:", "nvcr.io/nvidia/nim/meta-llama3-8b-instruct:latest-dgx-spark (ARM64)")
bullet("Endpoint:", "http://nim-llm:8000/v1 (internal) or http://localhost:8520/v1 (external)")
bullet("Auth:", "NGC_API_KEY environment variable")
bullet("Health:", "GET /v1/health/ready")

spacer()
H3("RAG Pipeline")
bullet("1. Retrieve:", "Guidelines + prior measurements + similar cases from pgvector")
bullet("2. Construct:", "Structured prompt with findings context and retrieved evidence")
bullet("3. Generate:", "NIM LLM produces evidence-grounded clinical summary")
bullet("4. Ground:", "All recommendations cite specific measurements and confidence scores")

spacer()
H3("Client Pattern")
code_block("""from langchain_openai import ChatOpenAI

llm = ChatOpenAI(
    base_url="http://nim-llm:8000/v1",
    api_key="not-needed",
    model="meta-llama3-8b-instruct",
    temperature=0.1,
)""", "python")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 17. EMBEDDING SERVICE
# ══════════════════════════════════════════════════════════════
H1("17. Embedding Service")

body(
    "A FastAPI microservice generates 384-dim embeddings for imaging studies, series, "
    "and individual lesions. Embeddings enable cohort retrieval (\u201c10 most similar "
    "CT studies\u201d) and phenotype matching via pgvector."
)

H3("Architecture")
bullet("Model:", "BiomedCLIP (microsoft/BiomedCLIP-PubMedBERT_256-vit_base_patch16_224) "
       "or similar medical vision encoder")
bullet("Dimensions:", "384-dim vectors, L2-normalized")
bullet("Levels:", "Study-level, series-level, lesion-level embeddings")
bullet("Storage:", "pgvector with HNSW index (m=16, ef_construction=64)")
bullet("Endpoint:", "POST /embed with study_id, level, image_path")

spacer()
H3("Cohort Retrieval Query")
code_block("""-- Find 10 most similar CT chest studies
SELECT s.study_instance_uid, s.patient_id, s.study_date,
       e.embedding <=> $1::vector AS distance
FROM embeddings e
JOIN studies s ON e.study_id = s.id
WHERE e.level = 'study' AND s.modality = 'CT' AND s.body_part = 'CHEST'
ORDER BY e.embedding <=> $1::vector
LIMIT 10;""", "sql")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 18. NEXTFLOW PIPELINE ORCHESTRATION
# ══════════════════════════════════════════════════════════════
H1("18. Nextflow Pipeline Orchestration")

body(
    "Nextflow DSL2 orchestrates multi-stage imaging workflows. The same patterns used by "
    "the existing HCLS AI Factory (profiles, process labels, container assignments, "
    "completion handlers) are followed."
)

H3("main.nf Structure")
code_block("""nextflow.enable.dsl = 2

include { CT_HEAD_HEMORRHAGE } from './modules/ct_head_hemorrhage'
include { CT_CHEST_LUNG_NODULE } from './modules/ct_chest_lung_nodule'
include { CXR_RAPID_FINDINGS } from './modules/cxr_rapid_findings'
include { MRI_BRAIN_MS_LESION } from './modules/mri_brain_ms_lesion'

workflow {
    if (params.workflow == 'ct_head') {
        CT_HEAD_HEMORRHAGE(Channel.of(params.study_uid))
    } else if (params.workflow == 'ct_chest') {
        CT_CHEST_LUNG_NODULE(Channel.of(params.study_uid))
    } // ... cxr, mri_brain
}""", "groovy")

spacer()
H3("Workflow Module Pattern")
code_block("""workflow CT_HEAD_HEMORRHAGE {
    take: study_uid
    main:
        PREPROCESS(study_uid)
        INFERENCE(PREPROCESS.out.preprocessed)
        POSTPROCESS(INFERENCE.out.results)
    emit:
        output = POSTPROCESS.out.output
}""", "groovy")

spacer()
H3("nextflow.config Profiles")
add_table(
    ["Profile", "Use Case", "Key Settings"],
    [
        ["docker", "Standard Docker execution",
         "docker.enabled = true"],
        ["dgx_spark", "DGX Spark with GPU",
         "--gpus all, max_memory 128 GB, max_gpus 1"],
    ]
)

spacer()
H3("Event-Driven Trigger")
body(
    "Orthanc study.complete webhook \u2192 DICOM listener service \u2192 determines modality \u2192 "
    "invokes nextflow run main.nf -profile dgx_spark --workflow <type> --study_uid <uid>."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 19. STREAMLIT PORTAL
# ══════════════════════════════════════════════════════════════
H1("19. Streamlit Portal")

body(
    "The dashboard follows the existing HCLS AI Factory portal patterns: NVIDIA green "
    "(#76b900) theme, gradient header, sidebar navigation, service status indicators, "
    "and GPU metrics from DCGM exporter."
)

H3("Pages")
add_table(
    ["Page", "Content"],
    [
        ["Worklist", "Active findings ordered by priority (P1\u2013P4), urgency color coding"],
        ["Studies", "Processed studies with finding counts, status badges"],
        ["Agent Activity", "Provenance log: model versions, durations, pipeline status"],
        ["Monitoring", "GPU utilization, pipeline throughput, embedded Grafana link"],
    ]
)

spacer()
H3("Theme")
code_block("""st.set_page_config(
    page_title="Imaging Intelligence Agent",
    layout="wide", initial_sidebar_state="expanded")

# NVIDIA green accent, gradient header, card styling
# --nvidia-green: #76b900
# border-left: 5px solid #76b900""", "python")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 20. MONITORING STACK
# ══════════════════════════════════════════════════════════════
H1("20. Monitoring Stack")

body(
    "Prometheus + Grafana + DCGM Exporter, matching the existing HCLS AI Factory "
    "monitoring setup."
)

H3("Key DCGM Metrics")
add_table(
    ["Metric", "Description"],
    [
        ["DCGM_FI_DEV_GPU_UTIL", "GPU utilization %"],
        ["DCGM_FI_DEV_GPU_TEMP", "GPU temperature (\u00b0C)"],
        ["DCGM_FI_DEV_POWER_USAGE", "Power draw (W)"],
        ["DCGM_FI_DEV_FB_USED", "GPU memory used (MB)"],
        ["DCGM_FI_DEV_FB_FREE", "GPU memory free (MB)"],
    ]
)

spacer()
H3("Scrape Configuration")
bullet("DCGM Exporter:", "5-second scrape interval for GPU metrics")
bullet("Prometheus:", "15-second default, self-monitoring")
bullet("Application services:", "15-second interval on /metrics endpoints")

spacer()
H3("Alert Rules")
bullet("HighGPUMemory:", "GPU memory > 90% for 5 minutes \u2192 warning")
bullet("InferenceFailureRate:", "Failure rate > 10% over 5 minutes \u2192 critical")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 21. TESTING STRATEGY
# ══════════════════════════════════════════════════════════════
H1("21. Testing Strategy")

H3("Test Levels")
add_table(
    ["Level", "Scope", "Tools"],
    [
        ["Unit", "Transforms, measurement algorithms, classification logic",
         "pytest, numpy assertions"],
        ["Integration", "End-to-end workflow with synthetic DICOM",
         "pytest + testcontainers (PostgreSQL, Orthanc)"],
        ["Synthetic DICOM", "Generate test data with pydicom",
         "pydicom.dataset, pydicom.uid.generate_uid"],
    ]
)

spacer()
H3("Synthetic DICOM Generation")
body(
    "Test DICOM studies are generated using pydicom with correct UIDs, pixel data, "
    "and metadata. Fixtures in conftest.py provide database connections (via "
    "testcontainers) and synthetic studies."
)

H3("Example Assertions")
code_block("""def test_urgency_critical():
    assert classify_urgency(volume_ml=35.0, shift_mm=6.0, thickness_mm=12.0) == ("critical", "P1")

def test_urgency_urgent():
    assert classify_urgency(volume_ml=12.0, shift_mm=3.0, thickness_mm=8.0) == ("urgent", "P2")

def test_vdt_calculation():
    vdt = calculate_vdt(489, 268, date(2026, 2, 1), date(2025, 1, 28))
    assert 280 < vdt < 300  # ~287 days""", "python")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 22. ARM64 COMPATIBILITY
# ══════════════════════════════════════════════════════════════
H1("22. ARM64 Compatibility Guide")

body(
    "DGX Spark uses the NVIDIA Grace CPU (ARM64 / aarch64). All containers and "
    "Python packages must be ARM64-compatible."
)

H3("ARM64-Compatible Base Images")
add_table(
    ["Image", "ARM64", "Notes"],
    [
        ["nvcr.io/nvidia/pytorch:24.01-py3", "Yes", "Check NGC for -aarch64 tags"],
        ["pgvector/pgvector:pg16", "Yes", "Multi-arch"],
        ["orthancteam/orthanc:24.1.2", "Yes", "Multi-arch"],
        ["prom/prometheus:v2.48.0", "Yes", "Multi-arch"],
        ["grafana/grafana:10.2.2", "Yes", "Multi-arch"],
        ["python:3.11-slim", "Yes", "Multi-arch"],
    ]
)

spacer()
H3("Python Packages")
body(
    "All key packages have pre-built ARM64 wheels: torch, monai, numpy, scipy, "
    "scikit-image, pydicom, highdicom, psycopg2-binary, langchain, langgraph, "
    "fastapi, uvicorn, streamlit, transformers, pgvector."
)

H3("NIM Container Variants")
body(
    "Append -dgx-spark to standard NIM image tags for ARM64 builds:")
code_block("""# Standard (x86_64)
nvcr.io/nvidia/nim/meta-llama3-8b-instruct:latest

# DGX Spark (ARM64)
nvcr.io/nvidia/nim/meta-llama3-8b-instruct:latest-dgx-spark""", "")

spacer()
H3("Building Multi-arch Images")
code_block("""docker buildx create --use --name multiarch
docker buildx build --platform linux/arm64 -t imaging-agent/ct-head:latest --push .""", "bash")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 23. CONFIGURATION REFERENCE
# ══════════════════════════════════════════════════════════════
H1("23. Configuration Reference")

H3(".env.example")
code_block("""# PostgreSQL
POSTGRES_USER=imaging
POSTGRES_PASSWORD=imaging_secret
POSTGRES_DB=imaging_agent

# NVIDIA
NGC_API_KEY=your_ngc_api_key_here

# Monitoring
GRAFANA_USER=admin
GRAFANA_PASSWORD=changeme

# FHIR (optional)
FHIR_SERVER_URL=http://localhost:8080/fhir

# Service Host (auto-detected if not set)
# SERVICE_HOST=192.168.1.100""", "env")

spacer()
H3("Key Environment Variables")
add_table(
    ["Variable", "Service", "Default"],
    [
        ["NGC_API_KEY", "NIM LLM", "(required)"],
        ["POSTGRES_USER", "PostgreSQL", "imaging"],
        ["POSTGRES_PASSWORD", "PostgreSQL", "imaging_secret"],
        ["POSTGRES_DB", "PostgreSQL", "imaging_agent"],
        ["ORTHANC_URL", "All services", "http://orthanc:8042"],
        ["NIM_LLM_URL", "Agent", "http://nim-llm:8000/v1"],
        ["EMBEDDING_URL", "Agent", "http://embedding-service:8000"],
        ["FHIR_SERVER_URL", "FHIR Publisher", "http://localhost:8080/fhir"],
        ["GRAFANA_USER", "Grafana", "admin"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 24. IMPLEMENTATION SEQUENCE
# ══════════════════════════════════════════════════════════════
H1("24. Implementation Sequence")

body(
    "Build in this order \u2014 each step depends on the ones above it. "
    "Start with infrastructure, then the simplest workflow, then layer complexity."
)

add_table(
    ["Step", "Component", "Verification"],
    [
        ["1", "PostgreSQL + pgvector",
         "psql connect; \\dx shows vector extension"],
        ["2", "Orthanc DICOM server",
         "curl http://localhost:8042/system returns JSON"],
        ["3", "DICOM listener service",
         "Upload DICOM \u2192 webhook fires \u2192 study row in DB"],
        ["4", "CXR Rapid Findings MAP",
         "Upload CXR \u2192 findings in DB within 30 seconds"],
        ["5", "CT Head Hemorrhage MAP",
         "Upload CT head \u2192 findings + measurements in DB"],
        ["6", "Embedding service + vector search",
         "POST /embed \u2192 384-dim vector stored in pgvector"],
        ["7", "LangGraph agent",
         "Query findings, generate clinical summary"],
        ["8", "FHIR publisher",
         "POST /publish \u2192 FHIR Bundle created"],
        ["9", "Streamlit portal",
         "Worklist renders, study list populates"],
        ["10", "CT Chest Lung Nodule MAP",
         "Upload CT \u2192 nodule tracking with VDT calculation"],
        ["11", "MRI Brain MS Lesion MAP",
         "Upload MRI \u2192 lesion tracking with registration"],
        ["12", "NIM LLM + RAG pipeline",
         "RAG query returns evidence-grounded response"],
        ["13", "Nextflow orchestration",
         "nextflow run main.nf -profile dgx_spark --workflow ct_head"],
        ["14", "Monitoring stack",
         "Grafana dashboard shows GPU metrics and traces"],
        ["15", "End-to-end integration test",
         "Upload 4 studies \u2192 all findings + SR + FHIR + worklist"],
    ]
)

spacer()
H3("Quick Start Commands")
code_block("""# 1. Clone and configure
git clone <repo-url> hls-imaging-agent && cd hls-imaging-agent
cp .env.example .env  # Edit with your NGC_API_KEY

# 2. Start infrastructure
docker compose up -d orthanc postgres

# 3. Download models
bash models/download_models.sh

# 4. Start all services
docker compose up -d

# 5. Upload test DICOM
bash scripts/seed_test_data.sh

# 6. Run a workflow
nextflow run main.nf -profile dgx_spark --workflow cxr --study_uid <uid>

# 7. Open portal: http://localhost:8525""", "bash")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TECHNOLOGY STACK SUMMARY
# ══════════════════════════════════════════════════════════════
H1("Appendix A — Technology Stack Summary")

add_table(
    ["Component", "Role", "License"],
    [
        ["MONAI", "Medical imaging AI framework", "Apache 2.0"],
        ["MONAI Deploy", "Containerized inference packaging (MAPs)", "Apache 2.0"],
        ["NVIDIA FLARE", "Federated learning across institutions", "Apache 2.0"],
        ["Nextflow", "Pipeline DAG orchestration", "Apache 2.0"],
        ["LangChain / LangGraph", "Agent orchestration + MCP tools", "MIT"],
        ["PostgreSQL + pgvector", "Structured + vector query", "PostgreSQL License"],
        ["Orthanc", "DICOM server", "GPLv3"],
        ["dcm4chee", "DICOM archive (alternative)", "MPL 1.1 / GPL 2.0 / LGPL 2.1"],
        ["pydicom", "DICOM parsing and manipulation", "MIT"],
        ["highdicom", "DICOM SR / SEG construction", "MIT"],
        ["HAPI FHIR", "FHIR server for clinical integration", "Apache 2.0"],
        ["fhir.resources", "Pydantic FHIR models", "BSD"],
        ["Grafana + Prometheus", "Monitoring and observability", "AGPL / Apache 2.0"],
        ["NVIDIA NIM", "Inference microservices", "NVAIE ($4,500/GPU/yr)"],
        ["NVIDIA Parabricks", "GPU-accelerated genomics", "NVAIE ($4,500/GPU/yr)"],
        ["NVIDIA BioNeMo", "Drug discovery", "NVAIE ($4,500/GPU/yr)"],
    ]
)

# ── Footer ───────────────────────────────────────────────────
P("")
P("")
footer = doc.add_table(rows=1, cols=1)
footer.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = footer.rows[0].cells[0]
set_cell_shading(cell, HEX_NAVY)
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(10)
R(p, "HCLS AI Factory", bold=True, size=10, color=WHITE)
R(p, "  |  Open Source (Apache 2.0)  |  NVIDIA DGX Spark  |  02/2026  |  v1.0",
  size=9, color=WHITE)

# ── Save ─────────────────────────────────────────────────────
BASE = "core/engines/clinical-imaging/open_public"
path = f"{BASE}/HCLS_Imaging_AI_Agent_Project_Bible.docx"
doc.save(path)
print(f"Project Bible saved to {path}")
