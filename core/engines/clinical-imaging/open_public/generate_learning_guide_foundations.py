#!/usr/bin/env python3
# Copyright 2026 Adam Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Generate Imaging Intelligence Agent — Learning Guide Foundations (DOCX).

High-school-level learning guide explaining every concept behind the
HCLS Imaging Intelligence Agent. VCP-style formatting matching the
Project Bible and Demo Guide.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Colors (identical to white paper) ──────────────────────
NAVY = RGBColor(0x1B, 0x23, 0x33)
TEAL = RGBColor(0x1A, 0xAF, 0xCC)
GREEN = RGBColor(0x76, 0xB9, 0x00)
GRAY_BODY = RGBColor(0x33, 0x33, 0x33)
GRAY_META = RGBColor(0x66, 0x66, 0x66)
GRAY_CODE = RGBColor(0x44, 0x44, 0x44)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_NAVY = "1B2333"
HEX_TEAL = "1AAFCC"
HEX_GREEN = "76B900"
HEX_LIGHT = "F8FAFB"
HEX_CODE_BG = "F5F5F5"
FONT = "Calibri"
CODE_FONT = "Consolas"

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
style.font.color.rgb = GRAY_BODY
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15


# ── Helper functions ───────────────────────────────────────

def R(p, text, bold=False, italic=False, size=None, color=None, font=None):
    r = p.add_run(text)
    r.font.name = font or FONT
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return r


def P(text="", bold=False, italic=False, size=10.5, color=GRAY_BODY,
      before=0, after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text:
        R(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def H1(text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        r.clear()
    R(p, text, bold=True, size=22, color=NAVY)
    return p


def H2(text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.clear()
    R(p, text, bold=True, size=16, color=NAVY)
    return p


def H3(text):
    return P(text, bold=True, size=13, color=TEAL, before=10, after=4)


def body(text, before=0, after=6):
    return P(text, size=10.5, color=GRAY_BODY, before=before, after=after)


def body_rich(parts, before=0, after=6):
    """Render a paragraph with mixed bold/italic/normal segments.

    parts is a list of tuples: (text, bold, italic)
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    for text, bold, italic in parts:
        R(p, text, bold=bold, italic=italic, size=10.5, color=GRAY_BODY)
    return p


def bullet(lead, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    R(p, lead, bold=True, size=10.5, color=GRAY_BODY)
    R(p, f"  {text}", size=10.5, color=GRAY_BODY)
    return p


def bullet_plain(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    R(p, "\u2022  ", bold=True, size=10.5, color=TEAL)
    R(p, text, size=10.5, color=GRAY_BODY)
    return p


def numbered(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    R(p, f"{num}. ", bold=True, size=10.5, color=TEAL)
    R(p, text, size=10.5, color=GRAY_BODY)
    return p


def callout(text):
    """Think of it this way... callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.35)
    R(p, "\U0001f4a1 Think of it this way: ", bold=True, size=10.5, color=TEAL)
    R(p, text, italic=True, size=10.5, color=GRAY_BODY)
    return p


def why_it_matters(text):
    """Why It Matters section."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    R(p, "Why it matters: ", bold=True, size=10.5, color=GREEN)
    R(p, text, size=10.5, color=GRAY_BODY)
    return p


def test_q(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    R(p, f"Q{num}. ", bold=True, size=10.5, color=TEAL)
    R(p, text, size=10.5, color=GRAY_BODY)
    return p


def code_block(text, language=""):
    if language:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        R(p, language, bold=True, size=8, color=TEAL, font=FONT)
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.2)
        R(p, line if line else " ", size=8, color=GRAY_CODE, font=CODE_FONT)


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="D0D0D0", sz="4"):
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>')
    tcPr.append(borders)


def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, HEX_NAVY)
        set_cell_borders(c, HEX_NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(p, h, bold=True, size=9, color=WHITE)
    for ri, row in enumerate(rows):
        bg = "FFFFFF" if ri % 2 == 0 else HEX_LIGHT
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            set_cell_shading(c, bg)
            set_cell_borders(c, "E0E0E0")
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if ci == 0:
                R(p, str(val), bold=True, size=9, color=TEAL)
            else:
                R(p, str(val), size=9, color=GRAY_BODY)
    return t


def spacer(pts=6):
    P("", after=pts)


# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════

accent = doc.add_table(rows=1, cols=1)
accent.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = accent.rows[0].cells[0]
set_cell_shading(cell, HEX_GREEN)
p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
R(p, " ", size=6, color=GREEN)

spacer(40)

P("Learning Guide", bold=True, size=36, color=TEAL,
  align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
P("Foundations", bold=True, size=28, color=TEAL,
  align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

spacer(12)

P("Imaging Intelligence Agent", bold=True, size=32, color=NAVY,
  align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
P("Every Concept Explained from First Principles",
  size=16, color=GRAY_META, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)

spacer(20)

P("HCLS AI Factory  |  NVIDIA DGX Platform",
  size=12, color=GRAY_META, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
P("Author: Adam Jones  |  February 2026",
  size=11, color=GRAY_META, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
P("Apache 2.0 License",
  size=10, color=GRAY_META, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════

H1("Table of Contents")
spacer(4)

toc_items = [
    "Chapter 1 \u2014 What Is Medical Imaging?",
    "Chapter 2 \u2014 How Computers See Images (AI/ML Basics)",
    "Chapter 3 \u2014 The DICOM Standard",
    "Chapter 4 \u2014 What Is a GPU and Why Does It Matter?",
    "Chapter 5 \u2014 Containers and Docker",
    "Chapter 6 \u2014 Databases",
    "Chapter 7 \u2014 The Four Clinical Workflows",
    "Chapter 8 \u2014 How the Agent Thinks (LangGraph and AI Reasoning)",
    "Chapter 9 \u2014 Embeddings and Similarity Search",
    "Chapter 10 \u2014 Talking to Hospital Systems (FHIR and Clinical Integration)",
    "Chapter 11 \u2014 RAG and Large Language Models",
    "Chapter 12 \u2014 Monitoring and Observability",
    "Chapter 13 \u2014 The HCLS AI Factory (The Bigger Picture)",
    "Chapter 14 \u2014 Trust, Safety, and Regulation",
    "Test Yourself \u2014 Comprehensive Review",
    "Glossary",
]
for i, item in enumerate(toc_items):
    numbered(i + 1, item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 1 — WHAT IS MEDICAL IMAGING?
# ══════════════════════════════════════════════════════════════

H1("Chapter 1 \u2014 What Is Medical Imaging?")

H2("X-Rays: Shadows of Your Body")
body("An X-ray works like shining a flashlight at your hand in a dark room. "
     "The light passes through your skin but gets blocked by your bones, "
     "so you see a shadow. In a real X-ray machine, the \"flashlight\" is a "
     "beam of X-ray radiation, and the \"shadow\" lands on a detector plate "
     "behind you. Dense things like bones appear white. Soft things like lungs "
     "appear dark. Air appears black.")
body("The most common X-ray is a chest X-ray (often shortened to CXR). "
     "Doctors order it to check the lungs, heart, and ribcage. A single "
     "chest X-ray produces one flat image \u2014 like a photograph.")

H2("CT Scans: Slicing a Loaf of Bread")
body("CT stands for Computed Tomography. A CT scanner is a large "
     "doughnut-shaped machine. You lie on a table that slides through "
     "the hole. Inside the doughnut, an X-ray tube spins around you, "
     "taking hundreds of X-ray images from every angle. A computer then "
     "stacks those images together to build a 3D picture.")
callout("If your body were a loaf of bread, each CT slice is like cutting "
        "one thin piece. You can look at any single slice, or you can stack "
        "all the slices to see the entire loaf in three dimensions.")
body("CT scans use X-ray radiation (like a regular X-ray) but take many "
     "more images. They are excellent for seeing bones, blood, and organs "
     "in detail. A single CT scan of the head can produce 200\u2013300 "
     "individual slices.")

H2("MRI: Magnets and Radio Waves")
body("MRI stands for Magnetic Resonance Imaging. Unlike X-rays and CT, "
     "MRI uses no radiation at all. Instead, it uses a very powerful "
     "magnet and radio waves.")
body("Your body is mostly water, and water contains hydrogen atoms. "
     "When you lie inside the MRI machine, the magnet lines up all the "
     "hydrogen atoms in your body. Then the machine sends radio wave "
     "pulses that knock those atoms out of alignment. When the atoms "
     "snap back into place, they release tiny signals. The machine "
     "detects those signals and uses them to build an image.")
body("MRI is especially good at showing soft tissue \u2014 the brain, "
     "spinal cord, muscles, and ligaments. It produces detailed images "
     "without any radiation exposure. The tradeoff is that MRI scans "
     "take longer (often 30\u201360 minutes) and the machines are more "
     "expensive.")
body_rich([
    ("One important MRI sequence is called ", False, False),
    ("FLAIR", True, False),
    (" (Fluid-Attenuated Inversion Recovery). FLAIR suppresses the signal "
     "from normal fluid in the brain so that abnormal spots \u2014 like "
     "damaged areas in multiple sclerosis \u2014 stand out brightly. "
     "The agent uses FLAIR images specifically for MS lesion tracking.", False, False),
])

H2("What Doctors Order and Why")
add_table(
    ["Exam", "Best For", "Radiation?", "Speed"],
    [
        ["X-Ray (CXR)", "Lungs, bones, heart size", "Yes (very low dose)", "Seconds"],
        ["CT Scan", "Brain bleeds, lung nodules, trauma", "Yes (moderate dose)", "1\u20135 minutes"],
        ["MRI", "Brain tissue, spinal cord, MS lesions", "No", "30\u201360 minutes"],
    ],
)
spacer(4)
body("A doctor picks the imaging type based on what they are looking for "
     "and how quickly they need the answer. A suspected brain bleed needs "
     "a CT because it is fast. A suspected MS flare-up needs an MRI "
     "because it shows soft tissue detail.")

H2("Studies and Series")
body_rich([
    ("When you get a scan, the hospital stores it as a ", False, False),
    ("study", True, False),
    (". A study is the complete collection of images from one scan session. "
     "For example, \"CT Head performed on January 15, 2026\" is one study.", False, False),
])
body_rich([
    ("Inside a study, there can be multiple ", False, False),
    ("series", True, False),
    (". Each series is a set of images acquired with the same settings. "
     "A CT head study might have a series taken without contrast dye and "
     "another series taken after injecting contrast dye. An MRI study "
     "might have a T1-weighted series, a T2-weighted series, and a FLAIR "
     "series \u2014 each one shows different tissue characteristics.", False, False),
])
callout("A study is like an album of photos from a vacation. Each series "
        "is a group of photos taken at the same location with the same "
        "camera settings.")

H2("What Radiologists Do")
body("A radiologist is a doctor who specializes in reading medical images. "
     "Their daily workflow looks like this:")
numbered(1, "A scan arrives from the scanner.")
numbered(2, "It appears on the radiologist\u2019s worklist \u2014 their to-do list.")
numbered(3, "The radiologist opens the images, scrolls through them, "
            "and looks for abnormalities.")
numbered(4, "They write a report describing what they found.")
numbered(5, "The report goes back to the ordering doctor, who uses it "
            "to make treatment decisions.")
body("A busy hospital radiologist may read hundreds of studies per day, "
     "each containing hundreds or thousands of images. That is millions "
     "of images per year. Fatigue and volume are real problems.")

why_it_matters(
    "The Imaging Intelligence Agent exists because radiologists are "
    "overwhelmed. It does not replace them. It acts like an assistant "
    "that pre-reads every scan, flags emergencies, measures changes "
    "over time, and writes draft findings \u2014 so the radiologist "
    "can work faster and more accurately."
)

H3("Test Yourself \u2014 Chapter 1")
test_q(1, "What is the difference between a CT scan and an MRI?")
test_q(2, "What does FLAIR stand for, and why is it useful for MS?")
test_q(3, "What is the difference between a study and a series?")
test_q(4, "Why can\u2019t AI replace radiologists entirely?")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 2 — HOW COMPUTERS SEE IMAGES
# ══════════════════════════════════════════════════════════════

H1("Chapter 2 \u2014 How Computers See Images")

H2("Pixels and Voxels: Images Are Numbers")
body("When you look at a photo on your phone, you see colors and shapes. "
     "But to a computer, that photo is just a grid of numbers. Each tiny "
     "square in the grid is called a pixel (picture element). A pixel "
     "stores a number representing brightness or color.")
body("A chest X-ray might be 2,048 \u00d7 2,048 pixels \u2014 that is "
     "about 4 million numbers arranged in a grid. A bright pixel (high "
     "number) means dense tissue. A dark pixel (low number) means air "
     "or soft tissue.")
body_rich([
    ("CT and MRI scans are three-dimensional, so instead of pixels they have ", False, False),
    ("voxels", True, False),
    (" (volume elements). A voxel is a tiny cube in 3D space, like one "
     "small block in a Minecraft world. A CT head scan might contain "
     "512 \u00d7 512 \u00d7 300 voxels \u2014 about 78 million numbers.", False, False),
])

H2("What a Neural Network Is")
body("A neural network is a type of computer program that learns patterns "
     "from examples. It is loosely inspired by how the brain works \u2014 "
     "layers of connected units that each do a small piece of math.")
callout("Imagine you are teaching a child to recognize dogs. You show "
        "them thousands of pictures labeled \"dog\" or \"not dog.\" At "
        "first they guess randomly. But every time they get it wrong, "
        "you give them feedback, and they adjust. After seeing thousands "
        "of examples, they get very good at it. That is exactly how a "
        "neural network learns.")
body("A neural network has layers:")
bullet("Input layer:", "Receives the image (all those pixel/voxel numbers).")
bullet("Hidden layers:", "Each layer extracts features \u2014 simple things "
       "like edges at first, then complex things like shapes and textures "
       "in deeper layers.")
bullet("Output layer:", "Produces the answer (a classification, a detection "
       "box, or a segmentation mask).")
body_rich([
    ("The more layers a network has, the \"deeper\" it is. That is why you "
     "hear the term ", False, False),
    ("deep learning", True, False),
    (" \u2014 it just means a neural network with many layers.", False, False),
])

H2("Training vs. Inference")
body_rich([
    ("Training", True, False),
    (" is like studying for a test. You feed the network thousands of "
     "labeled examples. The network adjusts its internal math to get "
     "better at predicting the right answer. Training is slow and "
     "computationally expensive \u2014 it can take hours or days.", False, False),
])
body_rich([
    ("Inference", True, False),
    (" is like taking the test. You give the trained network a new image "
     "it has never seen before, and it produces a prediction. Inference "
     "is fast \u2014 usually seconds per image.", False, False),
])
body("The Imaging Intelligence Agent only does inference. The models "
     "were already trained by researchers (often using MONAI, NVIDIA\u2019s "
     "medical imaging AI framework). The agent loads those pre-trained "
     "models and runs them on new patient scans.")

H2("Classification: Is There a Problem?")
body("Classification answers a yes/no (or multi-category) question about "
     "the entire image. The output is a confidence score \u2014 a number "
     "between 0 and 1 representing how sure the model is.")
body("The agent\u2019s CXR workflow runs a classifier called DenseNet-121. "
     "It looks at the whole chest X-ray and outputs confidence scores for "
     "five conditions: pneumothorax, consolidation, pleural effusion, "
     "cardiomegaly, and fracture.")

H2("Detection: Where Is the Problem?")
body("Detection goes further than classification. It not only says "
     "\"yes, there is a problem\" but also draws a bounding box \u2014 "
     "a rectangle \u2014 around where the problem is.")
callout("The difference between classification and detection is like "
        "someone saying \"there is a typo on this page\" (classification) "
        "versus circling the exact word (detection).")

H2("Segmentation: What Exact Shape Is the Problem?")
body("Segmentation is the most detailed level. Instead of drawing a box, "
     "it colors in the exact shape of the abnormality \u2014 pixel by "
     "pixel (or voxel by voxel). The result is called a segmentation "
     "mask.")
body("Models used for segmentation in this agent include 3D U-Net and "
     "SegResNet from the MONAI Model Zoo.")

H2("Confidence Scores")
body("Every prediction comes with a confidence score. If confidence is "
     "above a threshold (for example, 0.5), the finding is reported. "
     "If confidence is below the threshold, the finding is discarded "
     "or flagged as uncertain.")
body("Confidence scores are not the same as accuracy. A model can be "
     "95% confident and still be wrong. That is why a human radiologist "
     "always reviews the results.")

why_it_matters(
    "Classification, detection, and segmentation are the three fundamental "
    "AI techniques the agent uses on every scan. Understanding them helps "
    "you explain exactly what the agent is doing at each step."
)

H3("Test Yourself \u2014 Chapter 2")
test_q(1, "What is the difference between a pixel and a voxel?")
test_q(2, "Explain the difference between training and inference in your own words.")
test_q(3, "What does a segmentation mask look like?")
test_q(4, "Why are confidence scores not the same as accuracy?")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 3 — THE DICOM STANDARD
# ══════════════════════════════════════════════════════════════

H1("Chapter 3 \u2014 The DICOM Standard")

H2("What DICOM Is")
body("DICOM stands for Digital Imaging and Communications in Medicine. "
     "It is a universal standard for how medical images are stored, "
     "transmitted, and displayed. Every CT scanner, MRI machine, X-ray "
     "unit, and PACS system in the world uses DICOM.")
callout("JPEG is the format your phone uses for photos. DICOM is the "
        "format hospitals use for medical images. But DICOM is much "
        "more than just an image file \u2014 it includes the image "
        "data AND a large amount of metadata.")

H2("DICOM Tags: Metadata Attached to Every Image")
body("Every DICOM file has hundreds of tags \u2014 labeled pieces of "
     "information. Some important tags include:")
add_table(
    ["Tag", "Name", "Example"],
    [
        ["(0010,0010)", "Patient Name", "DOE^JOHN"],
        ["(0010,0020)", "Patient ID", "MRN12345"],
        ["(0008,0020)", "Study Date", "20260115"],
        ["(0008,0060)", "Modality", "CT, MR, CR, DX"],
        ["(0020,000D)", "Study Instance UID", "1.2.840.113619..."],
        ["(0020,000E)", "Series Instance UID", "1.2.840.113619..."],
        ["(0008,0018)", "SOP Instance UID", "1.2.840.113619..."],
    ],
)

H2("UIDs: Unique Identifiers")
bullet("Study Instance UID:", "Identifies the entire study (one scan session).")
bullet("Series Instance UID:", "Identifies a series within the study.")
bullet("SOP Instance UID:", "Identifies a single image file.")
callout("The Study UID is like the city, the Series UID is the street, "
        "and the SOP Instance UID is the house number.")

H2("DICOM Servers and Orthanc")
body("A DICOM server is a specialized database for storing and retrieving "
     "medical images. The agent uses Orthanc, an open-source DICOM server "
     "(GPLv3 license). Orthanc runs on port 8042 for its web interface "
     "and port 4242 for DICOM network communication.")
body("The Python library pydicom (MIT license) is what the agent uses "
     "to read and write DICOM files.")

H2("DICOMweb: The Modern Way")
body("DICOMweb lets you access DICOM images using normal web protocols "
     "(HTTP). The three main services are:")
bullet("STOW-RS (Store):", "Upload images to the server.")
bullet("WADO-RS (Retrieve):", "Download images from the server.")
bullet("QIDO-RS (Search):", "Search for studies on the server.")

H2("DICOM Structured Reports (DICOM SR)")
body("When the agent finishes analyzing a scan, it creates a DICOM SR "
     "\u2014 a structured document containing findings, measurements, "
     "and references. This SR appears alongside the original images in "
     "the PACS viewer.")
body("The agent creates DICOM SRs using the Python library highdicom.")

why_it_matters(
    "The agent must speak DICOM to work with any hospital. Without "
    "DICOM compliance, the agent\u2019s findings could not appear in "
    "the radiologist\u2019s viewer, could not be archived alongside "
    "original images, and could not be searched in the hospital\u2019s "
    "systems."
)

H3("Test Yourself \u2014 Chapter 3")
test_q(1, "What is the difference between DICOM and JPEG?")
test_q(2, "Name the three types of UIDs and what level each identifies.")
test_q(3, "What is the difference between DIMSE C-STORE and DICOMweb STOW-RS?")
test_q(4, "Why does the agent output findings as DICOM SR?")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 4 — WHAT IS A GPU AND WHY DOES IT MATTER?
# ══════════════════════════════════════════════════════════════

H1("Chapter 4 \u2014 What Is a GPU and Why Does It Matter?")

H2("CPU vs. GPU")
body("Every computer has a CPU (Central Processing Unit). The CPU is "
     "like one very smart worker \u2014 it can do almost any task, but "
     "it does them one at a time (or a few at a time with multiple "
     "cores). A modern CPU might have 8 to 64 cores.")
body("A GPU (Graphics Processing Unit) has thousands of small, simple "
     "cores that all work at the same time. Each individual core is "
     "less capable than a CPU core, but the sheer number of them "
     "working in parallel makes the GPU incredibly fast for certain tasks.")
callout("A CPU is like one expert chef who can make any dish. A GPU "
        "is like a kitchen with 1,000 line cooks \u2014 each one can "
        "only do simple tasks, but together they can prepare 1,000 "
        "plates simultaneously.")

H2("Why GPUs Are Great for AI")
body("Neural networks are built from millions of simple math operations "
     "\u2014 mostly multiplying numbers and adding them together. These "
     "operations are independent of each other, meaning they can all "
     "happen at the same time. That is exactly what a GPU is designed "
     "to do.")
body("Processing a CT scan through a segmentation model requires billions "
     "of multiplications. On a CPU, this might take minutes. On a GPU, "
     "it takes seconds.")

H2("The DGX Spark")
body("The DGX Spark is the specific computer this agent is designed to run on:")
add_table(
    ["Parameter", "Value"],
    [
        ["CPU", "NVIDIA Grace (ARM64 / aarch64)"],
        ["GPU", "NVIDIA Blackwell GB10, 1 GPU"],
        ["Memory", "128 GB unified LPDDR5x (CPU + GPU shared pool)"],
        ["Storage", "Up to 4 TB NVMe"],
        ["Storage Access", "GPUDirect Storage (zero-copy GPU access)"],
        ["Price", "$4,699"],
        ["OS", "Ubuntu-based (NVIDIA DGX OS)"],
        ["NVAIE Cost", "Zero at desktop-class"],
    ],
)
spacer(4)

bullet("ARM64 (aarch64):",
       "Most laptops use x86_64 processors (Intel or AMD). The DGX Spark "
       "uses an ARM processor \u2014 the same family of chips used in "
       "smartphones. All software must be compiled for ARM64.")
bullet("Unified memory:",
       "The 128 GB is shared between CPU and GPU \u2014 one pool, no "
       "copying back and forth. In most computers, CPU and GPU have "
       "separate memory pools.")
bullet("GPUDirect Storage:",
       "Data flows directly from NVMe drive into GPU memory, skipping "
       "the CPU entirely. Normally: Drive \u2192 CPU memory \u2192 GPU "
       "memory. With GPUDirect: Drive \u2192 GPU memory.")

H2("DGX Compute Progression")
add_table(
    ["Phase", "Hardware", "Price", "Scope"],
    [
        ["Proof Build", "DGX Spark (1 GPU)", "$4,699", "1\u20132 workflows"],
        ["Departmental", "1\u20132\u00d7 DGX B200 (8 GPUs each)", "$500K\u2013$1M", "All workflows, PACS integration"],
        ["Multi-Site", "4\u20138\u00d7 DGX B200 + InfiniBand", "$2M\u2013$4M", "Federated learning"],
        ["AI Factory", "DGX SuperPOD", "$7M\u2013$60M+", "Thousands of concurrent studies"],
    ],
)

why_it_matters(
    "The GPU is the engine that makes real-time medical image analysis "
    "possible. The DGX Spark\u2019s unified memory and GPUDirect Storage "
    "eliminate data transfer bottlenecks. Understanding this hardware "
    "explains why the agent can process a brain CT in under 90 seconds."
)

H3("Test Yourself \u2014 Chapter 4")
test_q(1, "What is the fundamental difference between a CPU and a GPU?")
test_q(2, "What does \"unified memory\" mean on the DGX Spark?")
test_q(3, "What is GPUDirect Storage and what bottleneck does it eliminate?")
test_q(4, "Why must all containers be ARM64-compatible on the DGX Spark?")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 5 — CONTAINERS AND DOCKER
# ══════════════════════════════════════════════════════════════

H1("Chapter 5 \u2014 Containers and Docker")

H2("What a Container Is")
body("A container is a self-contained package that includes everything "
     "a piece of software needs to run \u2014 the code, the libraries, "
     "the configuration files, and even a minimal operating system.")
callout("A container is like a lunchbox. Everything the program needs "
        "is packed inside. It does not matter what kitchen (computer) "
        "you open the lunchbox in \u2014 the meal is the same every time.")

H2("Docker: The Tool That Runs Containers")
bullet("Docker Image:", "A read-only template \u2014 like a recipe.")
bullet("Docker Container:", "A running instance of an image \u2014 like "
       "the actual meal cooked from the recipe.")
bullet("Dockerfile:", "A text file with step-by-step instructions for "
       "building an image.")
bullet("Docker Hub / Registry:", "A website where pre-built images are "
       "stored. NVIDIA hosts theirs at nvcr.io.")

H2("docker-compose: Running Many Containers Together")
body("The agent is not one container \u2014 it is eleven containers that "
     "work together. Docker Compose lets you define all containers in a "
     "single docker-compose.yml file and start them all with one command: "
     "docker compose up.")
callout("docker-compose is like an orchestra conductor. Each musician "
        "(container) plays their own instrument, but the conductor "
        "ensures they all start at the right time and play together.")

H2("Volumes and Health Checks")
body_rich([
    ("Volumes", True, False),
    (" are persistent storage that survives container restarts. Without "
     "volumes, all data inside a container is lost when it stops \u2014 "
     "like erasing a whiteboard.", False, False),
])
body_rich([
    ("Health checks", True, False),
    (" are automatic \"are you alive?\" tests that Docker runs every few "
     "seconds. If a container stops responding, Docker can automatically "
     "restart it.", False, False),
])

H2("The Agent\u2019s 11 Containers")
add_table(
    ["Container", "Port", "Purpose"],
    [
        ["imaging-orthanc", "4242, 8042", "DICOM server \u2014 stores medical images"],
        ["imaging-postgres", "5432", "Database \u2014 findings, measurements, embeddings"],
        ["imaging-nim-llm", "8520", "Large language model for clinical reports"],
        ["imaging-embedding", "8521", "Creates similarity vectors"],
        ["imaging-dicom-listener", "8522", "Listens for new studies in Orthanc"],
        ["imaging-fhir-publisher", "8523", "Sends results to EHR via FHIR"],
        ["imaging-agent", "8524", "LangGraph clinical reasoning engine"],
        ["imaging-portal", "8525", "Streamlit web dashboard"],
        ["imaging-dcgm", "9400", "GPU metrics exporter"],
        ["imaging-prometheus", "9099", "Metrics collector"],
        ["imaging-grafana", "3000", "Metrics visualization dashboard"],
    ],
)

why_it_matters(
    "Containers are how the agent is packaged and deployed. Every piece "
    "of the system runs in its own container. This makes the agent "
    "portable, reliable (containers restart automatically), and "
    "maintainable (update one piece without touching the others)."
)

H3("Test Yourself \u2014 Chapter 5")
test_q(1, "What is the difference between a Docker image and a container?")
test_q(2, "Why does the agent use 11 separate containers instead of one?")
test_q(3, "What happens to data inside a container when it stops?")
test_q(4, "Why is ARM64 compatibility important for the DGX Spark?")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 6 — DATABASES
# ══════════════════════════════════════════════════════════════

H1("Chapter 6 \u2014 Databases")

H2("What a Database Is")
body("A database is an organized collection of information that you can "
     "search, filter, and update efficiently. A spreadsheet is the "
     "simplest analogy \u2014 rows of data organized into columns. "
     "But a database handles millions of rows and complex queries.")

H2("Tables, Rows, and Columns")
body("A database stores data in tables. Each table is like a spreadsheet:")
bullet("Column (field):", "Defines what kind of information is stored.")
bullet("Row (record):", "One entry \u2014 one specific finding, study, or measurement.")
spacer(4)
body("The agent\u2019s database has seven main tables:")
add_table(
    ["Table", "What It Stores", "Example"],
    [
        ["studies", "One row per imaging study", "CT Head, 2026-01-15, Patient MRN12345"],
        ["series", "One row per image series", "Axial non-contrast, 300 slices"],
        ["findings", "One row per AI-detected finding", "Hemorrhage, 95% confidence, CRITICAL"],
        ["measurements", "One row per measurement", "Volume: 25.3 mL, Midline shift: 8.2 mm"],
        ["embeddings", "One row per similarity vector", "384-dim vector for study-level search"],
        ["provenance", "One row per processing event", "Model v2.1, processed at 14:32:05"],
        ["worklist_entries", "One row per worklist item", "CRITICAL priority, on-call radiologist"],
    ],
)

H2("SQL and PostgreSQL")
body("SQL (Structured Query Language) is the language for asking the "
     "database questions. It reads almost like English:")
code_block(
    "SELECT * FROM findings\n"
    "WHERE severity = 'CRITICAL'\n"
    "AND created_at >= '2026-02-02';",
    "SQL"
)
spacer(4)
body("The agent uses PostgreSQL (\"Postgres\") \u2014 free, open-source, "
     "trusted by banks and governments. It runs in the imaging-postgres "
     "container on port 5432.")

H2("Views and Indexes")
body_rich([
    ("A ", False, False),
    ("view", True, False),
    (" is a saved SQL query that looks like a table. The agent has two views: "
     "active_worklist (incomplete items sorted by urgency) and study_summary "
     "(complete picture of each study).", False, False),
])
body_rich([
    ("An ", False, False),
    ("index", True, False),
    (" is a data structure that makes searches fast \u2014 like the index "
     "at the back of a textbook. Without an index, finding all CRITICAL "
     "hemorrhages in 100,000 rows requires reading every single row.", False, False),
])

why_it_matters(
    "The database is the agent\u2019s memory. Every finding, measurement, "
    "embedding, and decision is stored here. It enables longitudinal "
    "tracking \u2014 comparing today\u2019s scan to last month\u2019s scan."
)

H3("Test Yourself \u2014 Chapter 6")
test_q(1, "Name the seven tables in the agent\u2019s database.")
test_q(2, "What is SQL and what does it do?")
test_q(3, "Why was PostgreSQL chosen over other databases?")
test_q(4, "What is a database index and why does it matter?")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 7 — THE FOUR CLINICAL WORKFLOWS
# ══════════════════════════════════════════════════════════════

H1("Chapter 7 \u2014 The Four Clinical Workflows")

body("Every workflow follows the same general pattern:")
numbered(1, "A new study arrives in Orthanc.")
numbered(2, "The DICOM listener detects it and triggers processing.")
numbered(3, "GPU inference runs AI models on the images.")
numbered(4, "Post-processing extracts measurements and classifications.")
numbered(5, "Results are stored in PostgreSQL.")
numbered(6, "Outputs are sent to PACS and EHR.")
numbered(7, "The worklist is updated with the appropriate priority.")

H2("Workflow 1: Chest X-Ray (CXR) Rapid Findings")
body("Target time: less than 30 seconds. Detects five conditions:")
add_table(
    ["Condition", "What It Means"],
    [
        ["Pneumothorax", "Collapsed lung \u2014 air trapped between lung and chest wall"],
        ["Consolidation", "Lung filled with fluid or pus (often pneumonia)"],
        ["Pleural Effusion", "Fluid buildup around the lungs"],
        ["Cardiomegaly", "Enlarged heart"],
        ["Fracture", "Broken rib or other bone"],
    ],
)
spacer(4)
body("The model is DenseNet-121 (a 121-layer classification network). "
     "After classification, the agent generates a GradCAM heatmap \u2014 "
     "a colored overlay showing which image regions influenced the "
     "model\u2019s decision. Red/yellow = high influence, blue = low.")

H2("Workflow 2: CT Head Hemorrhage Triage")
body("Target time: less than 90 seconds. Detects bleeding inside the skull.")
body("Step by step:")
numbered(1, "Classification: DenseNet-121 determines if hemorrhage is present.")
numbered(2, "Segmentation: 3D U-Net outlines the exact shape of the bleed.")
numbered(3, "Volume estimation: count marked voxels \u00d7 voxel size = volume in mL.")
numbered(4, "Midline shift measurement: how far the brain has been pushed to one side.")
numbered(5, "Urgency classification based on Brain Trauma Foundation thresholds:")
spacer(2)
add_table(
    ["Volume", "Midline Shift", "Urgency"],
    [
        ["> 30 mL", "Any", "CRITICAL \u2014 potential surgical emergency"],
        ["5\u201330 mL", "> 5 mm", "URGENT \u2014 needs immediate attention"],
        ["5\u201330 mL", "\u2264 5 mm", "URGENT \u2014 close monitoring"],
        ["< 5 mL", "\u2264 5 mm", "ROUTINE \u2014 standard review"],
    ],
)

H2("Workflow 3: CT Chest Lung Nodule Tracking")
body("Target time: less than 5 minutes. Detects small spots (nodules) "
     "in the lungs that could be early-stage cancer.")
numbered(1, "Detection: RetinaNet draws bounding boxes around nodules.")
numbered(2, "Segmentation: SegResNet outlines the exact 3D shape.")
numbered(3, "Volumetrics: volume (mm\u00b3) and longest diameter (mm).")
numbered(4, "Prior study retrieval: compare to previous CT.")
numbered(5, "Volume Doubling Time (VDT) calculation:")
spacer(2)
body("VDT = (\u0394t \u00d7 ln 2) / ln(V2 / V1)")
body("A short VDT (< 400 days) means the nodule is growing fast \u2014 "
     "suspicious for malignancy. VDT < 400 upgrades the risk category.")
spacer(2)
numbered(6, "Lung-RADS classification:")
add_table(
    ["Lung-RADS", "Meaning", "Solid Nodule Size"],
    [
        ["1", "Negative \u2014 no nodules", "N/A"],
        ["2", "Benign appearance", "< 6 mm"],
        ["3", "Probably benign", "6\u20138 mm"],
        ["4A", "Suspicious", "8\u201315 mm"],
        ["4B", "Very suspicious", "\u2265 15 mm"],
    ],
)
spacer(2)
numbered(7, "Genomics trigger: Lung-RADS 4B+ automatically triggers "
            "the Parabricks genomics pipeline for tumor profiling.")

H2("Workflow 4: MRI Brain MS Lesion Tracking")
body("Target time: less than 5 minutes. Multiple sclerosis (MS) is a "
     "disease where the immune system attacks nerve insulation (myelin). "
     "Each attack creates a lesion (damaged spot).")
numbered(1, "3D U-Net segmentation on FLAIR MRI.")
numbered(2, "Lesion counting and volume measurement.")
numbered(3, "Spatial registration to prior MRI (aligning the two scans).")
numbered(4, "Identify new, enlarging, and stable lesions.")
numbered(5, "Disease activity classification:")
spacer(2)
add_table(
    ["New / Enlarging Lesions", "Classification"],
    [
        ["0", "Stable"],
        ["1\u20132", "Active"],
        ["3 or more", "Highly Active"],
    ],
)

why_it_matters(
    "These four workflows demonstrate the agent\u2019s range \u2014 "
    "from 30-second emergency triage to 5-minute longitudinal tracking. "
    "Understanding each workflow lets you explain exactly what the agent "
    "does and why."
)

H3("Test Yourself \u2014 Chapter 7")
test_q(1, "What five conditions does the CXR workflow detect?")
test_q(2, "At what hemorrhage volume is a finding classified as CRITICAL?")
test_q(3, "What is VDT and why does a short VDT indicate concern?")
test_q(4, "How does the MRI workflow classify disease activity?")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 8 — HOW THE AGENT THINKS
# ══════════════════════════════════════════════════════════════

H1("Chapter 8 \u2014 How the Agent Thinks")

H2("What Is an AI Agent?")
body("A regular AI model works like a vending machine: you put in an "
     "image, you get out a prediction. An agent works more like a "
     "doctor: it receives information, decides what additional tests "
     "to run, interprets the results, and forms a conclusion.")
body("The Imaging Intelligence Agent does not just detect hemorrhages "
     "\u2014 it detects them, decides whether they are critical, looks "
     "up prior scans, searches for similar cases, and generates a "
     "clinical report. Each step depends on the previous step.")

H2("LangGraph: The Framework")
body("LangGraph is an open-source framework (MIT license) for building "
     "multi-step AI reasoning workflows. It uses a concept called a "
     "StateGraph \u2014 essentially a flowchart the agent follows.")
bullet("Nodes:", "Individual steps. Each node is a Python function.")
bullet("Edges:", "Connections between steps defining the order.")
bullet("Conditional edges:", "Decision points where the agent takes "
       "different paths depending on current state.")

H2("The Agent\u2019s Four Nodes")
numbered(1, "triage_node: Receives findings, determines severity "
            "(CRITICAL / URGENT / ROUTINE), decides downstream analyses.")
numbered(2, "longitudinal_node: Retrieves prior studies, compares "
            "findings, calculates changes over time.")
numbered(3, "population_node: Searches for similar cases using "
            "embedding similarity, retrieves outcomes.")
numbered(4, "report_node: Gathers all evidence, sends to LLM with "
            "guidelines, generates structured report.")

H2("Conditional Routing")
body("Not every study goes through every node:")
bullet("CRITICAL:", "Full analysis pipeline \u2014 all four nodes.")
bullet("ROUTINE:", "May skip longitudinal and population nodes, "
       "go straight to brief report.")

H2("State and Tools")
body("As the agent moves from node to node, it carries a state \u2014 "
     "a data structure that accumulates information (study_id, findings, "
     "prior_studies, similar_cases, severity, report, provenance).")
body("The agent has tools it can call: query_findings, "
     "search_similar_studies, calculate_vdt, retrieve_guidelines, "
     "generate_report. These tools use MCP (Model Context Protocol) "
     "\u2014 a standardized way for AI agents to declare and use tools.")

why_it_matters(
    "The agent is a reasoning system that orchestrates multiple models, "
    "databases, and services into a coherent clinical workflow. "
    "Understanding the graph structure explains how different cases "
    "trigger different levels of analysis."
)

H3("Test Yourself \u2014 Chapter 8")
test_q(1, "What is the difference between an AI model and an AI agent?")
test_q(2, "Name the four nodes in the agent\u2019s StateGraph.")
test_q(3, "Why does the agent skip some nodes for ROUTINE findings?")
test_q(4, "What is MCP and why is it useful?")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 9 — EMBEDDINGS AND SIMILARITY SEARCH
# ══════════════════════════════════════════════════════════════

H1("Chapter 9 \u2014 Embeddings and Similarity Search")

H2("What Is an Embedding?")
body("An embedding is a way to represent something complex (like a "
     "medical image) as a list of numbers \u2014 not pixel values, "
     "but numbers that capture the \"essence\" of the image.")
callout("Imagine GPS coordinates for images. Similar images have "
        "similar coordinates. A landscape and another landscape would "
        "be nearby. A landscape and a portrait would be far apart.")
body("The agent turns each study into a list of 384 numbers "
     "(a 384-dimensional vector). Studies with similar characteristics "
     "end up with similar vectors.")

H2("BiomedCLIP: The Embedding Model")
body("The specific model is BiomedCLIP "
     "(microsoft/BiomedCLIP-PubMedBERT_256-vit_base_patch16_224). "
     "It was trained on millions of medical images paired with text "
     "descriptions. It runs in the imaging-embedding container on "
     "port 8521.")

H2("Cosine Similarity")
body("Cosine similarity measures how similar two embeddings are by "
     "comparing the direction of two vectors. Close to 1.0 = very "
     "similar. Close to 0 = unrelated.")
callout("Two arrows pointing in almost the same direction are similar "
        "(cosine similarity close to 1.0). Two arrows at right angles "
        "are unrelated (cosine similarity 0).")

H2("pgvector and HNSW Index")
body("pgvector is a PostgreSQL extension that adds vector data types "
     "and similarity search. Without an index, finding the 10 most "
     "similar studies out of 100,000 requires 100,000 comparisons.")
body("The HNSW (Hierarchical Navigable Small World) index organizes "
     "vectors into a navigable graph structure so approximate nearest "
     "neighbors can be found much faster.")

H2("Three Levels of Embeddings")
bullet("Study-level:", "One vector per study. \"Find studies similar to this one.\"")
bullet("Series-level:", "One vector per series. \"Find series with similar characteristics.\"")
bullet("Lesion-level:", "One vector per finding. \"Find nodules that look like this one.\"")

H2("Hybrid Queries")
body("The real power: combining SQL filtering with vector similarity. "
     "Example: \"Find 10 CT chest studies most similar to this one "
     "that also had Lung-RADS 4A or higher.\"")

why_it_matters(
    "Embeddings enable \"patients like this\" queries. The agent can "
    "instantly retrieve similar historical cases with known outcomes, "
    "providing evidence for clinical decisions."
)

H3("Test Yourself \u2014 Chapter 9")
test_q(1, "In your own words, what is an embedding?")
test_q(2, "What does cosine similarity measure?")
test_q(3, "What is the HNSW index and why is it needed?")
test_q(4, "Give an example of a hybrid query.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 10 — TALKING TO HOSPITAL SYSTEMS
# ══════════════════════════════════════════════════════════════

H1("Chapter 10 \u2014 Talking to Hospital Systems")

H2("EHR and FHIR")
body("An EHR (Electronic Health Record) is the hospital\u2019s digital "
     "patient chart. FHIR (Fast Healthcare Interoperability Resources, "
     "pronounced \"fire\") is a modern standard for exchanging "
     "healthcare data using web technology (HTTP, JSON).")
callout("FHIR is like a shared language between hospital systems. "
        "Without FHIR, every system would need a custom translation "
        "for every other system.")

H2("FHIR Resources")
add_table(
    ["Resource", "What It Represents"],
    [
        ["Patient", "The person being imaged"],
        ["ImagingStudy", "Reference to the DICOM study"],
        ["DiagnosticReport", "The agent\u2019s findings and conclusions"],
        ["Observation", "Individual measurements (volume, diameter, shift)"],
    ],
)
spacer(4)
body("The agent creates a DiagnosticReport resource sent via the "
     "imaging-fhir-publisher container (port 8523).")

H2("Coding Systems: SNOMED CT and LOINC")
body("SNOMED CT codes clinical concepts (e.g., \"intracranial hemorrhage\"). "
     "LOINC codes observations (e.g., \"volume of hemorrhage\"). "
     "These codes ensure any EHR system interprets findings identically.")

H2("PACS, GSPS, and the Worklist")
body("PACS (Picture Archiving and Communication System) is where "
     "radiologists view images. The agent sends DICOM SR (findings) "
     "and GSPS (visual overlays/annotations) to PACS via DICOMweb "
     "STOW-RS.")
body("The worklist is the radiologist\u2019s prioritized to-do list: "
     "CRITICAL at the top, then URGENT, then ROUTINE. Stored in the "
     "worklist_entries table and displayed through the Streamlit "
     "portal (port 8525).")

H2("Clinician-in-the-Loop")
body("The agent is decision support, not autonomous diagnosis. "
     "Every output is a recommendation for a human expert to review. "
     "The radiologist confirms or overrides. The clinician is "
     "accountable for the final decision.")

why_it_matters(
    "An AI system that cannot communicate with existing hospital "
    "systems is useless in practice. FHIR, DICOM SR, GSPS, and "
    "worklist integration make the agent a practical clinical tool."
)

H3("Test Yourself \u2014 Chapter 10")
test_q(1, "What is the difference between an EHR and PACS?")
test_q(2, "What FHIR resource does the agent create to report findings?")
test_q(3, "Why are coding systems like SNOMED CT important?")
test_q(4, "What does \"clinician-in-the-loop\" mean?")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 11 — RAG AND LARGE LANGUAGE MODELS
# ══════════════════════════════════════════════════════════════

H1("Chapter 11 \u2014 RAG and Large Language Models")

H2("What Is a Large Language Model (LLM)?")
body("A Large Language Model is an AI that can read, understand, and "
     "generate text. The agent uses an LLM to write clinical reports. "
     "The specific model is Meta Llama 3 8B Instruct, served by "
     "NVIDIA NIM in the imaging-nim-llm container (port 8520).")

H2("The Problem: Hallucination")
body_rich([
    ("LLMs have a well-known weakness: they can make things up. This is called ", False, False),
    ("hallucination", True, False),
    (". In a hospital, hallucination is dangerous. A report that states "
     "the wrong measurement could lead to a wrong treatment decision.", False, False),
])

H2("What Is RAG?")
body("RAG stands for Retrieval-Augmented Generation. Instead of asking "
     "the LLM to generate a report from memory alone, you first "
     "retrieve all the relevant facts, then give those facts to the "
     "LLM as context.")
callout("Instead of asking someone to write a book report from memory "
        "(where they might misremember details), you give them the "
        "book, their notes, and a highlighter, and ask them to write "
        "with the book open in front of them.")

H2("The RAG Pipeline Step by Step")
numbered(1, "Retrieve findings from the database.")
numbered(2, "Retrieve prior measurements for comparison.")
numbered(3, "Retrieve similar cases and their outcomes via embedding search.")
numbered(4, "Retrieve relevant clinical guidelines (Lung-RADS, BTF, etc.).")
numbered(5, "Construct the prompt with all evidence and instructions.")
numbered(6, "Send to the LLM (NIM-served Llama 3).")
numbered(7, "Receive the grounded report. Every claim traceable to evidence.")

H2("Evidence Grounding and Cross-Modal Context")
body_rich([
    ("Evidence grounding", True, False),
    (" means every statement in the report is backed by data. The LLM "
     "is summarizing and contextualizing information that was already "
     "computed, measured, and retrieved.", False, False),
])
body("The RAG pipeline can pull evidence from multiple sources: "
     "imaging findings, longitudinal data, population data, "
     "clinical guidelines, and genomic data (when available).")

why_it_matters(
    "Reports with evidence are trustworthy. Hallucinated reports are "
    "dangerous. RAG prevents the LLM from inventing facts. Understanding "
    "RAG is critical for explaining why the agent\u2019s reports are "
    "reliable."
)

H3("Test Yourself \u2014 Chapter 11")
test_q(1, "What is hallucination and why is it dangerous in healthcare?")
test_q(2, "What does RAG stand for, and what problem does it solve?")
test_q(3, "List the types of evidence the RAG pipeline retrieves.")
test_q(4, "What is evidence grounding?")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 12 — MONITORING AND OBSERVABILITY
# ══════════════════════════════════════════════════════════════

H1("Chapter 12 \u2014 Monitoring and Observability")

H2("Metrics, Prometheus, and Grafana")
body("Monitoring means continuously watching a system to ensure it "
     "works correctly. The agent tracks metrics \u2014 numbers about "
     "system health (inference latency, throughput, error rates, "
     "GPU utilization, temperature).")
bullet("Prometheus:", "Collects metrics every 15\u201330 seconds by "
       "\"scraping\" each container\u2019s metrics endpoint. "
       "Runs on port 9099.")
bullet("Grafana:", "Visualizes metrics as charts and dashboards. "
       "Runs on port 3000.")
callout("Prometheus is like a hall monitor who checks every classroom "
        "every 15 seconds and writes down what they find. Grafana "
        "is the bulletin board that displays those notes as charts.")

H2("DCGM Exporter")
body("DCGM (Data Center GPU Manager) exports GPU-specific metrics "
     "in Prometheus format from the imaging-dcgm container (port 9400): "
     "GPU utilization, memory usage, temperature, power consumption, "
     "and error counts.")

H2("Alerting")
body("Grafana sends alerts when metrics cross thresholds. Examples: "
     "GPU temperature exceeded 85\u00b0C, error rate exceeded 5%, "
     "no studies processed in 30 minutes during business hours.")

H2("Provenance and Reproducibility")
body("Provenance is a complete record of how a result was produced: "
     "model ID/version, inference parameters, processing duration, "
     "input data lineage (DICOM UIDs), timestamps, operator approvals, "
     "and predetermined change control plans.")
body("Reproducibility means the ability to re-run the exact same "
     "analysis and get the exact same result. Achieved by storing "
     "immutable originals, recording every parameter, and versioning "
     "every model.")

why_it_matters(
    "In healthcare, you must prove the system is reliable and auditable. "
    "Monitoring provides real-time visibility. Provenance provides a "
    "complete audit trail. Reproducibility provides verification."
)

H3("Test Yourself \u2014 Chapter 12")
test_q(1, "What is the difference between Prometheus and Grafana?")
test_q(2, "What does DCGM monitor and why is it important?")
test_q(3, "What information does a provenance record include?")
test_q(4, "Why is reproducibility essential in medical AI?")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 13 — THE HCLS AI FACTORY
# ══════════════════════════════════════════════════════════════

H1("Chapter 13 \u2014 The HCLS AI Factory")

H2("The Bigger Picture")
body("The Imaging Intelligence Agent is one piece of the HCLS AI "
     "Factory \u2014 a collection of specialized AI agents that work "
     "together on shared NVIDIA DGX hardware.")
callout("Like a hospital with departments: imaging handles scans, "
        "genetics handles DNA, pharmacy handles drug matching. "
        "Each is specialized, but they share records and collaborate.")

H2("Cross-Modal Integration")
bullet("Imaging \u2192 Genomics (Parabricks):",
       "Lung-RADS 4B+ triggers tumor profiling. Parabricks processes "
       "30x WGS in ~10 minutes on DGX (vs. ~30 hours on CPU). "
       "Up to 50% lower compute cost.")
bullet("Imaging \u2192 Drug Discovery (BioNeMo):",
       "200+ adopters (Eli Lilly, Astellas, Insilico, Recursion). "
       "Drug candidates scored by imaging phenotype. RECIST endpoints "
       "for clinical trials.")
bullet("Imaging \u2192 Clinical Reasoning (NIM LLM):",
       "RAG-grounded clinical reports.")
bullet("Imaging \u2192 Longitudinal Care:",
       "Tracking patients across multiple time points.")
bullet("Imaging \u2192 Outcomes:",
       "Cohort retrieval linking imaging patterns to outcomes.")

H2("The Biomarker Agent")
body("Another AI Factory agent \u2014 the Biomarker Intelligence Agent "
     "\u2014 combines genomic biomarkers with imaging biomarkers for "
     "combined phenotype profiling. Both agents run on the same DGX "
     "Spark and share data.")

H2("Scaling Path")
add_table(
    ["Phase", "Hardware", "Price", "NVAIE Cost/Year"],
    [
        ["Proof Build", "DGX Spark", "$4,699", "$0"],
        ["Departmental", "1\u20132\u00d7 DGX B200", "$500K\u2013$1M", "$36K\u2013$72K"],
        ["Multi-Site", "4\u20138\u00d7 DGX B200", "$2M\u2013$4M", "$144K\u2013$288K"],
        ["AI Factory", "DGX SuperPOD", "$7M\u2013$60M+", "$576K\u2013$1.15M"],
    ],
)

H2("NVIDIA FLARE: Federated Learning")
body("At Phase 3+, NVIDIA FLARE enables multiple hospitals to "
     "collaboratively improve AI models without sharing patient data. "
     "Each hospital trains locally and shares only model parameters. "
     "A central server aggregates updates. Patient data never leaves "
     "institutional control. FLARE is free (Apache 2.0).")

why_it_matters(
    "The imaging agent is designed from day one to be part of "
    "something bigger. Understanding the AI Factory context explains "
    "why the architecture uses standardized formats and why the "
    "$4,699 DGX Spark is a starting point, not the destination."
)

H3("Test Yourself \u2014 Chapter 13")
test_q(1, "What happens when the agent detects Lung-RADS 4B+?")
test_q(2, "What is federated learning and how does FLARE protect privacy?")
test_q(3, "Name three cross-modal integrations.")
test_q(4, "Why is the DGX Spark a \"proof build\" not a production system?")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 14 — TRUST, SAFETY, AND REGULATION
# ══════════════════════════════════════════════════════════════

H1("Chapter 14 \u2014 Trust, Safety, and Regulation")

H2("FDA AI/ML SaMD Framework")
body("The FDA regulates medical AI as Software as a Medical Device "
     "(SaMD). Key concepts:")
bullet("Intended use:", "Decision support for radiologists \u2014 not "
       "autonomous diagnosis.")
bullet("Risk categorization:", "Higher-risk applications face stricter "
       "requirements.")
bullet("Predetermined Change Control Plans (PCCP):", "Pre-approved "
       "rules for when/how AI can be updated without a new submission.")

H2("Decision Support vs. Autonomous Diagnosis")
body("Decision support: AI assists the clinician, who makes the final "
     "decision and is accountable. Autonomous diagnosis: AI makes the "
     "diagnosis without human review. The Imaging Intelligence Agent "
     "is decision support only. Every output includes language like "
     "\"recommend review.\"")

H2("Audit Trails and Immutability")
body("Every output is traceable to the exact model version, input "
     "data (DICOM UIDs), configuration, timestamps, and any human "
     "approvals. This audit trail is stored in the provenance table "
     "and is immutable \u2014 once written, it cannot be changed.")
body("Original DICOM images are never modified. Derived artifacts "
     "are stored alongside originals but never overwrite them.")
callout("Like a crime scene: you take photos (original DICOM), "
        "you write analysis notes (derived artifacts), but you "
        "never alter the scene itself.")

H2("Patient Data Security")
bullet("Data stays local:", "Patient data remains within institutional "
       "control.")
bullet("Least-privilege access:", "Each container only has access to "
       "the data it needs.")
bullet("Tenant isolation:", "Different departments cannot see each "
       "other\u2019s data.")

H2("Controlled Rollouts")
body("When a new model version is deployed:")
numbered(1, "New version deployed alongside existing version.")
numbered(2, "Both process the same studies in parallel.")
numbered(3, "Results compared to verify no regression.")
numbered(4, "Only after validation is the new version promoted.")
numbered(5, "Previous version outputs preserved \u2014 never deleted.")

why_it_matters(
    "Without trust, governance, and regulatory alignment, the most "
    "brilliant AI system will sit unused. The answer to \"can we "
    "trust it?\" is: \"here is exactly how every result is produced, "
    "audited, and safeguarded.\""
)

H3("Test Yourself \u2014 Chapter 14")
test_q(1, "What is the difference between decision support and autonomous diagnosis?")
test_q(2, "What is a Predetermined Change Control Plan?")
test_q(3, "Why are original DICOM images stored as immutable files?")
test_q(4, "How does tenant isolation protect patient data?")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TEST YOURSELF — COMPREHENSIVE REVIEW
# ══════════════════════════════════════════════════════════════

H1("Test Yourself \u2014 Comprehensive Review")

body("These 30 questions cover all 14 chapters. Try answering in your "
     "own words before checking the answer key.")

H2("Questions")
questions = [
    "What imaging modality uses magnets and radio waves instead of radiation?",
    "What is the difference between a pixel and a voxel?",
    "Name the three fundamental AI operations used in the agent.",
    "What does DICOM stand for?",
    "What are the three types of UIDs in DICOM?",
    "What is DICOMweb and how does it differ from DIMSE?",
    "How many cores does a GPU typically have compared to a CPU?",
    "What does \"unified memory\" mean on the DGX Spark?",
    "What is GPUDirect Storage?",
    "What is a Docker container and how does it differ from a Docker image?",
    "How many containers does the agent run, and why are they separated?",
    "What is a Docker volume and why is it needed?",
    "Name the seven tables in the agent\u2019s database.",
    "What is SQL used for?",
    "What is a database index and what real-world object is it analogous to?",
    "What five conditions does the CXR workflow detect?",
    "At what hemorrhage volume is a finding classified as CRITICAL?",
    "What is VDT and what does a short VDT indicate?",
    "How does the MRI workflow determine disease activity?",
    "What is a StateGraph and what are its main components?",
    "Name the four nodes in the agent\u2019s reasoning graph.",
    "What is an embedding?",
    "What does cosine similarity measure?",
    "What is a hybrid query?",
    "What is FHIR and what problem does it solve?",
    "What is the difference between an EHR and PACS?",
    "What does RAG stand for and what problem does it solve?",
    "What is the difference between Prometheus and Grafana?",
    "What is federated learning and how does NVIDIA FLARE implement it?",
    "What is the difference between decision support and autonomous diagnosis?",
]

for i, q in enumerate(questions, 1):
    test_q(i, q)

doc.add_page_break()

H2("Answer Key")

answers = [
    ("MRI", "Magnetic Resonance Imaging uses magnets and radio waves. No radiation."),
    ("Pixel vs. voxel", "A pixel is a 2D picture element. A voxel is a 3D volume element."),
    ("Three operations", "Classification (is there a problem?), detection (where?), segmentation (exact shape)."),
    ("DICOM", "Digital Imaging and Communications in Medicine."),
    ("Three UIDs", "Study Instance UID (scan session), Series Instance UID (image set), SOP Instance UID (single file)."),
    ("DICOMweb", "Uses web protocols (HTTP/REST). DIMSE uses traditional networking. Both supported."),
    ("GPU cores", "Thousands of simple cores vs. 8\u201364 powerful CPU cores."),
    ("Unified memory", "128 GB shared between CPU and GPU \u2014 one pool, no copying."),
    ("GPUDirect Storage", "Data flows directly from NVMe to GPU memory, skipping the CPU."),
    ("Container vs. image", "Image is a read-only template (recipe). Container is a running instance (cooked meal)."),
    ("11 containers", "Separated so each is a specialist; failure isolation and independent updates."),
    ("Docker volume", "Persistent storage that survives container restarts."),
    ("Seven tables", "studies, series, findings, measurements, embeddings, provenance, worklist_entries."),
    ("SQL", "Language for querying and modifying data in a relational database."),
    ("Database index", "Speeds up searches \u2014 like the index at the back of a textbook."),
    ("Five CXR conditions", "Pneumothorax, consolidation, pleural effusion, cardiomegaly, fracture."),
    ("CRITICAL threshold", "> 30 mL (any midline shift)."),
    ("VDT", "Volume Doubling Time. Short VDT (< 400 days) = rapid growth = suspicious."),
    ("Disease activity", "0 new/enlarging = Stable, 1\u20132 = Active, 3+ = Highly Active."),
    ("StateGraph", "A flowchart with nodes (steps), edges (connections), and conditional routing."),
    ("Four nodes", "triage_node, longitudinal_node, population_node, report_node."),
    ("Embedding", "A list of numbers (vector) capturing the \"essence\" of an image."),
    ("Cosine similarity", "Measures how similar two vectors are by angle. 1.0 = very similar."),
    ("Hybrid query", "Combines SQL filtering with vector similarity search."),
    ("FHIR", "Standard for exchanging healthcare data via web technology. Solves interoperability."),
    ("EHR vs. PACS", "EHR stores all clinical data. PACS stores and displays medical images."),
    ("RAG", "Retrieval-Augmented Generation. Prevents LLM hallucination by grounding in evidence."),
    ("Prometheus vs. Grafana", "Prometheus collects/stores metrics. Grafana visualizes them."),
    ("Federated learning", "Multiple hospitals train models without sharing data. FLARE coordinates."),
    ("Decision support", "AI assists clinician (accountable). Autonomous: AI decides without human."),
]

for i, (lead, text) in enumerate(answers, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    R(p, f"{i}. ", bold=True, size=10, color=TEAL)
    R(p, f"{lead}: ", bold=True, size=10, color=GRAY_BODY)
    R(p, text, size=10, color=GRAY_BODY)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# GLOSSARY
# ══════════════════════════════════════════════════════════════

H1("Glossary")

body("Alphabetical list of every technical term used in this guide.")
spacer(4)

glossary = [
    ("Agent", "An AI program that makes decisions about what to do next based on what it finds."),
    ("ARM64 / aarch64", "Processor architecture used by DGX Spark. Different from x86_64 in most PCs."),
    ("BiomedCLIP", "Medical image embedding model that creates 384-dimensional vectors."),
    ("BioNeMo", "NVIDIA platform for drug discovery and molecular modeling."),
    ("Bounding Box", "A rectangle drawn around a detected object in an image."),
    ("Classification", "AI task answering \"is there a problem?\" with a confidence score."),
    ("Clinician-in-the-Loop", "Design principle where AI assists but never replaces human judgment."),
    ("Confidence Score", "Number between 0 and 1 representing model certainty."),
    ("Container", "Self-contained software package with everything needed to run."),
    ("Cosine Similarity", "Mathematical measure of vector similarity. 1.0 = identical direction."),
    ("CT", "Computed Tomography. X-rays from many angles creating a 3D picture."),
    ("DCGM", "Data Center GPU Manager. NVIDIA tool for monitoring GPU health."),
    ("Deep Learning", "Machine learning with neural networks having many layers."),
    ("DenseNet-121", "121-layer classification neural network for CXR and hemorrhage."),
    ("Detection", "AI task answering \"where is the problem?\" with a bounding box."),
    ("DGX Spark", "NVIDIA $4,699 desktop GPU computer with Grace CPU and GB10 GPU."),
    ("DICOM", "Digital Imaging and Communications in Medicine. Universal medical image standard."),
    ("DICOM SR", "Structured Report in DICOM format containing findings and measurements."),
    ("DICOMweb", "Web-based protocol for storing, retrieving, and searching DICOM images."),
    ("DIMSE", "Traditional DICOM networking protocol (C-STORE, C-FIND, etc.)."),
    ("Docker", "Tool for building and running containers."),
    ("docker-compose", "Tool for running multi-container applications from one config file."),
    ("Embedding", "List of numbers representing the \"essence\" of a complex object."),
    ("EHR", "Electronic Health Record. Hospital's comprehensive digital patient chart."),
    ("Evidence Grounding", "Ensuring every AI-generated claim is supported by retrieved data."),
    ("FDA", "Food and Drug Administration. U.S. medical device regulator."),
    ("Federated Learning", "Multi-institution model training without sharing raw data."),
    ("FHIR", "Fast Healthcare Interoperability Resources. Healthcare data exchange standard."),
    ("FLAIR", "MRI sequence highlighting brain lesions by suppressing normal fluid."),
    ("FLARE", "NVIDIA Federated Learning Application Runtime Environment (Apache 2.0)."),
    ("GPU", "Graphics Processing Unit. Thousands of cores for parallel computation."),
    ("GPUDirect Storage", "NVMe-to-GPU data path bypassing the CPU."),
    ("GradCAM", "Technique highlighting image regions that influenced an AI decision."),
    ("Grafana", "Open-source visualization tool for metrics dashboards."),
    ("GSPS", "Grayscale Softcopy Presentation State. DICOM annotation overlays."),
    ("Hallucination", "When an LLM generates factually incorrect or fabricated text."),
    ("Health Check", "Automatic Docker test verifying a container is responsive."),
    ("highdicom", "Python library for creating DICOM SR and other DICOM objects."),
    ("HNSW", "Hierarchical Navigable Small World. Fast vector similarity search index."),
    ("Hybrid Query", "Query combining SQL filtering with vector similarity search."),
    ("Immutable", "Cannot be changed after creation. Original DICOM files are immutable."),
    ("Index", "Data structure speeding up database searches."),
    ("Inference", "Running a trained model on new data. The \"taking the test\" phase."),
    ("LangGraph", "Open-source framework for multi-step AI reasoning workflows (MIT)."),
    ("LLM", "Large Language Model. AI that reads, understands, and generates text."),
    ("LOINC", "Coding system for lab tests and observations."),
    ("Lung-RADS", "Standardized scoring system for lung nodule risk assessment."),
    ("MAP", "MONAI Deploy Application Package. Containerized inference pipeline."),
    ("MCP", "Model Context Protocol. Standardized way for AI agents to use tools."),
    ("Midline Shift", "Brain center-line displacement (mm) from swelling or hemorrhage."),
    ("MONAI", "Medical Open Network for AI. NVIDIA medical imaging framework (Apache 2.0)."),
    ("MRI", "Magnetic Resonance Imaging. Magnets and radio waves, no radiation."),
    ("MS", "Multiple Sclerosis. Immune system attacks nerve insulation (myelin)."),
    ("Neural Network", "Program that learns patterns from examples in connected layers."),
    ("NIM", "NVIDIA Inference Microservice. Production model serving software."),
    ("NVMe", "Non-Volatile Memory Express. Very fast solid-state storage."),
    ("NVAIE", "NVIDIA AI Enterprise. Software license at $4,500/GPU/year."),
    ("Orthanc", "Open-source DICOM server (GPLv3)."),
    ("PACS", "Picture Archiving and Communication System. Image viewing for radiologists."),
    ("Parabricks", "NVIDIA GPU-accelerated genomics analysis platform."),
    ("PCCP", "Predetermined Change Control Plan. Pre-approved AI update rules."),
    ("pgvector", "PostgreSQL extension adding vector data types and similarity search."),
    ("Pixel", "Picture element. One point in a 2D image."),
    ("PostgreSQL", "Open-source relational database. Port 5432."),
    ("Prometheus", "Open-source monitoring system collecting metrics at intervals."),
    ("Provenance", "Complete record of how a result was produced."),
    ("pydicom", "Python library for reading and writing DICOM files (MIT)."),
    ("RAG", "Retrieval-Augmented Generation. Evidence-first report generation."),
    ("Radiologist", "Doctor specializing in reading medical images."),
    ("Reproducibility", "Ability to re-run analysis with same inputs and get same result."),
    ("RetinaNet", "Neural network for object detection (finding and boxing nodules)."),
    ("SaMD", "Software as a Medical Device. FDA classification for medical AI."),
    ("Segmentation", "AI task labeling every pixel/voxel as abnormal or normal."),
    ("SegResNet", "Neural network for medical image segmentation."),
    ("SNOMED CT", "Coding system for clinical concepts."),
    ("SQL", "Structured Query Language for databases."),
    ("StateGraph", "LangGraph flowchart with nodes, edges, and conditional routing."),
    ("Study", "Complete image collection from one scan session."),
    ("Series", "Set of images within a study, acquired with same settings."),
    ("Training", "Teaching a model with labeled examples. \"Studying for the test.\""),
    ("U-Net (3D)", "Neural network for 3D medical image segmentation."),
    ("UID", "Unique Identifier for DICOM studies, series, and images."),
    ("Unified Memory", "CPU and GPU sharing the same physical memory pool."),
    ("VDT", "Volume Doubling Time. Time for a nodule to double in volume."),
    ("View", "Saved SQL query that looks like a table."),
    ("Volume", "Docker: persistent storage. Medical: 3D measurement (mL or mm\u00b3)."),
    ("Voxel", "Volume element. One point in a 3D image."),
    ("Worklist", "Prioritized queue of studies awaiting radiologist review."),
    ("X-Ray", "Imaging using radiation to create a 2D shadow image."),
]

add_table(
    ["Term", "Definition"],
    [[term, defn] for term, defn in glossary],
)

spacer(12)

# ── Footer ─────────────────────────────────────────────────
P("\u2500" * 60, size=8, color=GRAY_META,
  align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=4)
P("HCLS AI Factory \u2014 Imaging Intelligence Agent",
  size=9, color=GRAY_META, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
P("Learning Guide (Foundations)  |  Apache 2.0 License  |  "
  "Author: Adam Jones  |  February 2026",
  size=8, color=GRAY_META, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)


# ══════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════

out = "HCLS_Imaging_AI_Agent_Learning_Guide_Foundations.docx"
doc.save(out)
print(f"\u2705  Saved {out}")
