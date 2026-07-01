#!/usr/bin/env python3
# Copyright 2026 Adam Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Generate NVIDIA DGX Spark Imaging Intelligence Executive Bullets.

Open/public version for HCLS AI Factory on GitHub (Apache 2.0).
Framed entirely around NVIDIA DGX Spark and open-source tooling.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Colors ────────────────────────────────────────────────────
DARK = RGBColor(0x1B, 0x1B, 0x1B)
TEAL = RGBColor(0x1A, 0xAF, 0xCC)
GREEN = RGBColor(0x76, 0xB9, 0x00)
GRAY_BODY = RGBColor(0x33, 0x33, 0x33)
GRAY_META = RGBColor(0x66, 0x66, 0x66)
GRAY_LIGHT = RGBColor(0x99, 0x99, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_DARK = "1B1B1B"
HEX_TEAL = "1AAFCC"
HEX_GREEN = "76B900"
HEX_TABLE_HDR = "1B2333"
HEX_TABLE_ALT = "F8FAFB"
FONT = "Calibri"

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
style.font.color.rgb = GRAY_BODY
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def R(p, text, bold=False, italic=False, size=None, color=None):
    r = p.add_run(text)
    r.font.name = FONT
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return r


def P(text="", bold=False, italic=False, size=10.5, color=GRAY_BODY,
      before=0, after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text:
        R(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def add_section_header(text):
    P("", before=24, after=4)
    p = P("", before=0, after=12)
    R(p, text, bold=True, size=22, color=DARK)


def add_sub_header(text):
    p = P("", before=12, after=6)
    R(p, text, bold=True, size=13, color=TEAL)


def add_bullet(lead, body_text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.2
    R(p, lead, bold=True, size=10.5, color=GRAY_BODY)
    R(p, body_text, bold=False, size=10.5, color=GRAY_BODY)


def add_styled_table(headers, data, col_widths=None):
    n_cols = len(headers)
    n_rows = len(data) + 1
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        R(p, h, bold=True, size=9, color=WHITE)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, HEX_TABLE_HDR)

    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            is_bold = c_idx == 0
            R(p, str(val), bold=is_bold, size=9, color=GRAY_BODY)
            if r_idx % 2 == 1:
                set_cell_shading(cell, HEX_TABLE_ALT)

    P("", before=4, after=8)
    return table


def add_footer_line():
    P("", before=24, after=0)
    p = P("", before=0, after=0, align=WD_ALIGN_PARAGRAPH.LEFT)
    R(p, "HCLS AI Factory", bold=True, size=8, color=DARK)
    R(p, "  |  ", size=8, color=GRAY_LIGHT)
    R(p, "Open Source (Apache 2.0)", size=8, color=GRAY_META)
    R(p, "  |  ", size=8, color=GRAY_LIGHT)
    R(p, "NVIDIA DGX Spark", size=8, color=GRAY_META)
    R(p, "  |  ", size=8, color=GRAY_LIGHT)
    R(p, "02/2026", size=8, color=GRAY_META)
    R(p, "  |  ", size=8, color=GRAY_LIGHT)
    R(p, "Version 1.0", size=8, color=GRAY_META)


# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════

P("", before=120, after=0)

P("Executive talking points", bold=False, size=13, color=TEAL,
  before=0, after=12, align=WD_ALIGN_PARAGRAPH.LEFT)

p = P("", before=0, after=8, align=WD_ALIGN_PARAGRAPH.LEFT)
R(p, "NVIDIA DGX Spark:\nImaging Intelligence\nfor Clinical AI", bold=True, size=36, color=DARK)

P("Fifteen technically grounded points for the open HCLS AI Factory",
  italic=False, size=14, color=GRAY_META, before=16, after=8,
  align=WD_ALIGN_PARAGRAPH.LEFT)

P("", before=120, after=0)

p = P("", before=0, after=0, align=WD_ALIGN_PARAGRAPH.LEFT)
R(p, "02/2026", bold=True, size=9, color=GRAY_META)
R(p, "  |  ", size=9, color=GRAY_LIGHT)
R(p, "Version 1.0", size=9, color=GRAY_META)
R(p, "  |  ", size=9, color=GRAY_LIGHT)
R(p, "Apache 2.0 License", size=9, color=GRAY_META)
R(p, "  |  ", size=9, color=GRAY_LIGHT)
R(p, "Author: Adam Jones", size=9, color=GRAY_META)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CONTENT PAGES
# ══════════════════════════════════════════════════════════════

P("The following fifteen points describe how the Imaging Intelligence "
  "Agent runs on NVIDIA DGX Spark as part of the open HCLS AI Factory. "
  "Each point is technically grounded and designed for executive and "
  "technical conversations about clinical AI on accelerated computing.",
  size=10.5, color=GRAY_BODY, before=8, after=16)

# ══════════════════════════════════════════════════════════════
# SECTION 1: ARCHITECTURE
# ══════════════════════════════════════════════════════════════
add_section_header("Architecture \u2014 NVIDIA Software Stack on DGX Spark")

add_bullet(
    "DGX Spark delivers a complete AI workstation for clinical imaging. ",
    "The NVIDIA DGX Spark (GB10 Grace Blackwell Superchip, 128 GB "
    "unified memory, $4,699) provides enough compute to run multi-stage "
    "imaging inference pipelines \u2014 detection, segmentation, classification, "
    "and clinical reasoning \u2014 on a single desktop device. CT, MRI, and "
    "X-ray studies are ingested, processed, and analyzed without leaving "
    "the local system."
)

add_bullet(
    "MONAI Deploy packages inference as portable, containerized applications. ",
    "MONAI Deploy Application Packages (MAPs) encapsulate each imaging "
    "pipeline stage \u2014 preprocessing, model inference, postprocessing \u2014 "
    "in a versioned, containerized unit. MAPs are portable across DGX "
    "Spark, DGX B200, DGX Cloud, and any CUDA-capable system. "
    "Orchestration via Nextflow or Airflow manages pipeline DAGs and "
    "event-driven triggers (e.g., study.complete \u2192 inference launch)."
)

add_bullet(
    "MONAI Model Zoo provides validated starting points for every modality. ",
    "Pre-trained architectures \u2014 3D U-Net (segmentation), RetinaNet "
    "(detection), SegResNet (segmentation), DenseNet-121 (classification), "
    "SwinUNETR (transformer-based) \u2014 are available for fine-tuning on "
    "institutional data. Each model is packaged with training scripts, "
    "configuration, and evaluation benchmarks."
)

add_bullet(
    "NVIDIA NIM serves models as standardized inference microservices. ",
    "NIM wraps each model in a production-grade microservice with "
    "versioning, health checks, auto-scaling, and standardized APIs. "
    "Imaging inference models and LLMs for clinical reasoning both "
    "deploy through the same NIM interface, simplifying operations "
    "and enabling consistent monitoring."
)

add_bullet(
    "PostgreSQL + pgvector provides unified structured and semantic query. ",
    "An open-source query layer replaces the need for separate relational "
    "and vector databases. Structured queries (\u201call Lung-RADS 4A+\u201d), "
    "vector similarity search (\u201c10 most similar CT studies\u201d), and "
    "hybrid queries (\u201cgrowing nodules AND similar phenotype\u201d) execute "
    "against the same canonical data. pgvector supports HNSW and "
    "IVFFlat indexing for fast retrieval."
)

add_bullet(
    "RAG pipelines ground every finding in clinical evidence. ",
    "Retrieval-augmented generation powered by NVIDIA NIM LLM serving "
    "anchors outputs in evidence \u2014 ACR guidelines, prior measurements, "
    "similar patient outcomes. Longitudinal delta analysis, cross-modal "
    "enrichment (imaging + genomics + biomarkers), and evidence-grounded "
    "reporting all flow through the RAG pipeline without external "
    "reasoning frameworks."
)

add_bullet(
    "LangChain / LangGraph orchestrates multi-step agent workflows. ",
    "Agent personas \u2014 triage agent, longitudinal tracker, population "
    "analyst \u2014 are defined with distinct roles and tool access via MCP "
    "(Model Context Protocol). Multi-step clinical reasoning, tool "
    "integration, checkpointing, and observability are handled by "
    "open-source agent frameworks running natively on DGX Spark."
)

add_bullet(
    "Canonical imaging state makes every artifact continuously computable. ",
    "CT/MRI/X-ray studies ingested via DICOMweb STOW-RS or DIMSE C-STORE "
    "are preserved as immutable evidence alongside derived artifacts \u2014 "
    "segmentation masks, volumetric measurements, GradCAM heatmaps, "
    "semantic embeddings, and provenance bundles. Every artifact is a "
    "first-class object, not a transient intermediate."
)

# ══════════════════════════════════════════════════════════════
# SECTION 2: HARDWARE ACCELERATION
# ══════════════════════════════════════════════════════════════
add_section_header("Hardware Acceleration \u2014 DGX Spark to SuperPOD")

add_bullet(
    "DGX Spark proves the architecture at desktop scale. ",
    "The GB10 Grace Blackwell Superchip with 128 GB unified memory "
    "runs 1\u20132 complete imaging workflows (CT head hemorrhage triage, "
    "CXR rapid findings) as a proof build. Zero NVIDIA AI Enterprise "
    "software cost at desktop-class. The canonical data model and "
    "pipeline architecture validated here carry forward unchanged to "
    "department and enterprise scale."
)

add_bullet(
    "DGX B200 scales to departmental multi-user, multi-modality workloads. ",
    "1\u20132x DGX B200 clusters ($500K\u2013$1M) support shared namespaces, "
    "PACS integration via DICOMweb, and clinical validation with "
    "radiologists. NVIDIA AI Enterprise licensing ($4,500/GPU/year) "
    "enables production NIM serving and model management."
)

add_bullet(
    "DGX SuperPOD delivers enterprise-scale AI factory operations. ",
    "4\u20138x DGX B200 + InfiniBand ($2M\u2013$4M) for multi-site deployments "
    "with continuous reprocessing, population-scale cohort retrieval, "
    "and NVIDIA FLARE federated learning. DGX SuperPOD ($7M\u2013$60M+) "
    "for thousands of concurrent studies across imaging, genomics, "
    "and drug discovery."
)

# ══════════════════════════════════════════════════════════════
# SECTION 3: CLINICAL WORKFLOWS + INTEGRATION
# ══════════════════════════════════════════════════════════════
add_section_header("Clinical Workflows \u2014 Speed, Governance, and Integration")

add_bullet(
    "Agentic diagnostic workflows deliver speed and consistency. ",
    "MONAI Deploy MAPs automate multi-stage pipelines \u2014 detection, "
    "segmentation, measurement, classification, triage. CT head "
    "hemorrhage triage in <90 seconds; CT chest lung nodule tracking "
    "in <5 minutes; CXR rapid findings in <30 seconds. Each stage "
    "produces canonical artifacts for downstream reasoning."
)

add_bullet(
    "Governance, reproducibility, and auditability are built in. ",
    "Provenance bundles record model ID, version, parameters, "
    "timestamps, and DICOM UIDs. Deterministic re-runs with coexisting "
    "version outputs. FDA AI/ML SaMD predetermined change control, "
    "controlled rollouts, and immutable audit trails. Fine-grained "
    "access control and tenant isolation at the data layer."
)

add_bullet(
    "The Imaging Agent is one node in a unified HCLS AI Factory. ",
    "Imaging signals link directly with genomics (Parabricks: 30x WGS "
    "from ~30 hours CPU to ~10 minutes GPU), clinical reasoning "
    "(NIM-served LLMs), and drug discovery (BioNeMo: 200+ adopters). "
    "A Lung-RADS 4B+ finding triggers Parabricks tumor profiling. "
    "Cohort retrieval via pgvector embeddings enables outcomes research. "
    "All agents share the same DGX Spark compute during proof build."
)

add_bullet(
    "NVIDIA FLARE enables federated learning without centralizing patient data. ",
    "Multi-site model improvement via privacy-preserving federated "
    "learning. Institutions contribute to model training without "
    "sharing raw imaging data. Free under Apache 2.0 \u2014 no software "
    "licensing cost for federated learning capabilities."
)

# ══════════════════════════════════════════════════════════════
# TECHNOLOGY STACK TABLE
# ══════════════════════════════════════════════════════════════
add_section_header("Open-Source Technology Stack")

P("The table below lists the core components of the Imaging Intelligence "
  "Agent, their roles, and licensing. All foundational components are "
  "open-source; NVIDIA AI Enterprise licensing applies only at "
  "production scale.",
  size=10.5, color=GRAY_BODY, before=0, after=12)

add_styled_table(
    ["Component", "Role", "License"],
    [
        ["MONAI",
         "Medical imaging AI framework (training + inference)",
         "Apache 2.0"],
        ["MONAI Deploy",
         "Containerized inference packaging (MAPs)",
         "Apache 2.0"],
        ["NVIDIA NIM",
         "Standardized inference microservices",
         "NVAIE ($4,500/GPU/yr)"],
        ["NVIDIA FLARE",
         "Federated learning across institutions",
         "Apache 2.0"],
        ["Nextflow",
         "Pipeline DAG orchestration",
         "Apache 2.0"],
        ["LangChain / LangGraph",
         "Agent orchestration + MCP tool integration",
         "MIT"],
        ["PostgreSQL + pgvector",
         "Structured + vector query (SQL + similarity)",
         "PostgreSQL License"],
        ["Orthanc",
         "DICOM server",
         "GPLv3"],
        ["dcm4chee",
         "DICOM archive",
         "MPL 1.1 / GPL 2.0 / LGPL 2.1"],
        ["pydicom",
         "DICOM parsing and manipulation",
         "MIT"],
        ["HAPI FHIR",
         "FHIR server for clinical integration",
         "Apache 2.0"],
        ["NVIDIA Parabricks",
         "GPU-accelerated genomics (30x WGS in ~10 min)",
         "NVAIE ($4,500/GPU/yr)"],
        ["NVIDIA BioNeMo",
         "Drug discovery (MolMIM + DiffDock)",
         "NVAIE ($4,500/GPU/yr)"],
    ]
)

# ══════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════
add_footer_line()

# ── Save ───────────────────────────────────────────────────
out = "core/engines/clinical-imaging/open_public/NVIDIA_DGX_Spark_Imaging_Executive_Bullets.docx"
doc.save(out)
print(f"Executive bullets saved to {out}")
