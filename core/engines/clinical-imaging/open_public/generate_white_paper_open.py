#!/usr/bin/env python3
# Copyright 2026 Adam Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Generate HCLS Imaging Intelligence Agent White Paper — Open Architecture on NVIDIA DGX Spark.

Open/public version for HCLS AI Factory on GitHub (Apache 2.0).
Framed entirely around NVIDIA DGX Spark and open-source tooling.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Colors ─────────────────────────────────────────────────
NAVY = RGBColor(0x1B, 0x23, 0x33)
TEAL = RGBColor(0x1A, 0xAF, 0xCC)
GREEN = RGBColor(0x76, 0xB9, 0x00)
GRAY_BODY = RGBColor(0x33, 0x33, 0x33)
GRAY_META = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_NAVY = "1B2333"
HEX_TEAL = "1AAFCC"
HEX_GREEN = "76B900"
HEX_LIGHT = "F8FAFB"
FONT = "Calibri"

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
style.font.color.rgb = GRAY_BODY
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15


def R(p, text, bold=False, italic=False, size=None, color=None):
    r = p.add_run(text)
    r.font.name = FONT
    if bold: r.bold = True
    if italic: r.italic = True
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return r


def P(text="", bold=False, italic=False, size=10.5, color=GRAY_BODY,
      before=0, after=6, align=None):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text: R(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def H1(text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        r.clear()
    R(p, text, bold=True, size=22, color=NAVY)
    return p


def H2(text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.clear()
    R(p, text, bold=True, size=16, color=NAVY)
    return p


def H3(text):
    return P(text, bold=True, size=13, color=TEAL, before=10, after=4)


def body(text, before=0, after=6):
    return P(text, size=10.5, color=GRAY_BODY, before=before, after=after)


def bullet(lead, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    R(p, lead, bold=True, size=10.5, color=GRAY_BODY)
    R(p, f"  {text}", size=10.5, color=GRAY_BODY)
    return p


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="D0D0D0", sz="4"):
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>')
    tcPr.append(borders)


def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, HEX_NAVY)
        set_cell_borders(c, HEX_NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(p, h, bold=True, size=9, color=WHITE)
    for ri, row in enumerate(rows):
        bg = "FFFFFF" if ri % 2 == 0 else HEX_LIGHT
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            set_cell_shading(c, bg)
            set_cell_borders(c, "E0E0E0")
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if ci == 0:
                R(p, str(val), bold=True, size=9, color=TEAL)
            else:
                R(p, str(val), size=9, color=GRAY_BODY)
    return t


# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
for _ in range(6):
    P("")

P("White paper", bold=True, size=13, color=TEAL, after=6)
P("Imaging Intelligence Agent in the HCLS AI Factory",
  bold=True, size=30, color=NAVY, after=4)
P("Open Architecture on NVIDIA DGX Spark \u2014 CT, MRI, and X-Ray as Canonical, "
  "Continuously Computable Evidence",
  bold=False, size=14, color=NAVY, after=4)
P("NVIDIA DGX Spark  |  MONAI Deploy  |  NVIDIA NIM  |  Open-Source Stack",
  italic=True, size=11, color=GRAY_META, after=8)
P("")
P("02/2026  |  Version 1.0  |  Apache 2.0 License", size=9, color=GRAY_META, after=2)
P("Author: Adam Jones", size=9, color=GRAY_META, after=2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════
H1("1. Executive Summary")

body(
    "The HCLS AI Factory follows a reusable pattern: identify a canonical artifact, build a "
    "persistent data model around it, and let agentic workflows operate continuously on that model. "
    "In genomics, the canonical artifact is the VCF. In imaging, it is canonical imaging data\u2014"
    "the original DICOM study plus every derived artifact: segmentation masks, volumetric "
    "measurements, semantic embeddings, structured findings, and provenance bundles."
)

body(
    "This white paper describes an Imaging Intelligence Agent that processes CT, MRI, and X-ray "
    "studies using NVIDIA MONAI models on DGX hardware, with an open-source software stack "
    "providing persistence, orchestration, query, and reasoning\u2014from DICOM ingestion through "
    "agentic inference to structured clinical output. The architecture is designed to run "
    "end-to-end on a single NVIDIA DGX Spark ($4,699) for proof builds and scale to DGX "
    "SuperPOD for enterprise deployments."
)

body(
    "Imaging studies arrive once, persist as canonical evidence, and are continuously "
    "re-computable as models evolve\u2014scaling from a proof build on DGX Spark to thousands of "
    "concurrent studies on DGX SuperPOD. All components are open-source or NVIDIA-licensed, "
    "released under Apache 2.0 as part of the HCLS AI Factory."
)

# ══════════════════════════════════════════════════════════════
# 2. CLINICAL PROBLEM & OPPORTUNITY
# ══════════════════════════════════════════════════════════════
H1("2. The AI Data Challenge in Medical Imaging")

body(
    "Medical imaging generates more data per patient encounter than any other clinical domain. "
    "A single CT chest produces 300\u2013500 DICOM slices; a multi-sequence MRI brain study can "
    "exceed 1,000 images. Across a mid-sized health system, tens of thousands of studies arrive "
    "daily\u2014yet the infrastructure that stores, processes, and reasons over this data was never "
    "designed for agentic AI."
)

H3("The Limits of Legacy Imaging Infrastructure")

body(
    "Today\u2019s imaging AI pipelines assemble disconnected components: a PACS or VNA for storage, "
    "an orchestration layer (Airflow, Prefect, or custom scripts) to stage data, GPU servers "
    "running containerized models, a relational database for findings, a vector database for "
    "embeddings, and an ad-hoc reasoning layer. Each must be independently provisioned, secured, "
    "monitored, and connected through brittle ETL."
)

body(
    "This assembled approach introduces three structural problems:"
)

bullet("Data movement friction.",
       "Every inference job requires copying DICOM data from archive to staging to GPU to results "
       "store\u2014a copy-and-stage pattern that consumes bandwidth, introduces latency, and creates "
       "governance gaps where data exists in uncontrolled locations.")
bullet("Pipeline fragility.",
       "Glue code between components becomes the most maintenance-intensive part of the system. "
       "Model updates, schema changes, or infrastructure migrations cascade across the stack.")
bullet("Reasoning fragmentation.",
       "Structured findings live in one database, embeddings in another, longitudinal priors in a "
       "third. Cross-modal queries (e.g., \u2018find all patients with growing lung nodules and elevated "
       "inflammatory markers\u2019) require joining across systems that were never designed to interoperate.")

body(
    "The market opportunity is significant\u2014AI in medical imaging is projected to reach $2\u20134 billion "
    "by 2028 (Signify Research, MarketsandMarkets)\u2014but capturing it requires a coherent "
    "data architecture, not more glue code."
)

# ══════════════════════════════════════════════════════════════
# 3. DGX SPARK SOFTWARE ARCHITECTURE
# ══════════════════════════════════════════════════════════════
H1("3. DGX Spark Software Architecture")

body(
    "The Imaging Intelligence Agent runs on a layered software architecture designed to operate "
    "entirely on NVIDIA DGX Spark for proof builds and scale to multi-node DGX clusters for "
    "production. The architecture uses open-source components organized into three functional layers: "
    "data, execution, and reasoning."
)

# Data Layer
H2("Data Layer: Local NVMe Storage with GPUDirect")

body(
    "DGX Spark provides 128 GB of unified memory and local NVMe storage with GPUDirect Storage "
    "support, enabling zero-copy data access from flash to GPU memory. For the Imaging AI Agent, "
    "this layer provides:"
)

bullet("Immutable DICOM archive:",
       "CT, MRI, and X-ray studies ingested via DICOMweb STOW-RS or DIMSE C-STORE are preserved "
       "as immutable clinical evidence on local NVMe. Orthanc or dcm4chee serves as the DICOM "
       "server and archive.")
bullet("Derived artifact persistence:",
       "Segmentation masks, volumetric measurements, GradCAM heatmaps, semantic embeddings, and "
       "provenance bundles are persisted alongside the source study with full lineage.")
bullet("GPUDirect Storage access:",
       "DICOM data flows from NVMe directly to GPU memory without CPU bounce buffers, minimizing "
       "latency for inference workloads.")

body(
    "At scale (Phase 2+), shared storage replaces local NVMe to enable multi-user access and "
    "longitudinal study retrieval across a department or institution."
)

# Execution Layer
H2("Execution Layer: MONAI Deploy + Nextflow Orchestration")

body(
    "The execution layer manages containerized inference pipelines and agent workflows:"
)

bullet("MONAI Deploy Application Packages (MAPs):",
       "Each imaging pipeline stage\u2014preprocessing, model inference, postprocessing\u2014is packaged "
       "in a versioned, containerized MAP. MAPs are portable across DGX Spark, DGX B200, DGX "
       "Cloud, and any CUDA-capable system.")
bullet("Nextflow / Airflow orchestration:",
       "Pipeline DAGs define multi-stage workflows (detection \u2192 segmentation \u2192 classification "
       "\u2192 triage). Event-driven triggers launch pipelines when new studies arrive (study.complete "
       "events). Nextflow DSL2 provides execution profiles for Docker, Singularity, and Slurm.")
bullet("NVIDIA NIM inference microservices:",
       "NIM wraps each model in a production-grade microservice with versioning, health checks, "
       "auto-scaling, and standardized APIs. Both imaging models and LLMs for clinical reasoning "
       "deploy through the same NIM interface.")
bullet("LangChain / LangGraph agent framework:",
       "Multi-step clinical reasoning workflows are orchestrated by open-source agent frameworks. "
       "Agent personas\u2014triage agent, longitudinal tracker, population analyst\u2014are defined with "
       "distinct roles and tool access via MCP (Model Context Protocol). Checkpointing and "
       "observability ensure production resilience.")

# Reasoning Layer
H2("Reasoning Layer: PostgreSQL + pgvector and RAG Pipeline")

body(
    "The reasoning layer provides structured query, semantic search, and evidence-grounded "
    "clinical reasoning:"
)

bullet("PostgreSQL + pgvector:",
       "A single open-source database handles structured finding queries (\u201call Lung-RADS 4A+\u201d), "
       "vector similarity search (\u201c10 most similar CT studies\u201d), and hybrid queries (\u201cgrowing "
       "nodules AND similar phenotype AND APOE4 carrier\u201d). pgvector supports HNSW and IVFFlat "
       "indexing for fast retrieval. No separate relational, vector, or analytics database required.")
bullet("RAG pipeline (NVIDIA NIM LLM serving):",
       "Retrieval-augmented generation anchors clinical reasoning in evidence\u2014ACR guidelines, "
       "prior measurements, similar patient outcomes. Longitudinal delta analysis and cross-modal "
       "enrichment (imaging + genomics + biomarkers) flow through the RAG pipeline. NVIDIA NIM "
       "serves the LLM inference layer.")

# Three-layer table
H2("Three-Layer Architecture Summary")

add_table(
    ["Layer", "Primary Responsibilities", "Components"],
    [
        ["Data Layer",
         "Canonical persistence, DICOM archive, "
         "derived artifact storage, GPUDirect access",
         "Local NVMe (DGX Spark); Orthanc / dcm4chee (DICOM server); "
         "GPUDirect Storage for zero-copy GPU access"],
        ["Execution Layer",
         "Event-driven orchestration, containerized "
         "inference, agent workflows, model serving",
         "MONAI Deploy MAPs; Nextflow / Airflow (orchestration); "
         "NVIDIA NIM (model serving); LangChain / LangGraph (agents); "
         "NVIDIA FLARE (federated learning)"],
        ["Reasoning Layer",
         "Structured + vector queries, RAG reasoning, "
         "longitudinal comparison, cohort retrieval",
         "PostgreSQL + pgvector (unified query); "
         "RAG pipeline with NVIDIA NIM LLM serving"],
    ]
)

# ══════════════════════════════════════════════════════════════
# 4. CANONICAL IMAGING DATA MODEL
# ══════════════════════════════════════════════════════════════
H1("4. Canonical Imaging Data Model")

body(
    "The Imaging Intelligence Agent operates on a canonical imaging data model\u2014a structured, "
    "persistent, continuously computable representation of everything the system knows about an "
    "imaging study. This model is the imaging equivalent of the VCF in genomics: a single, "
    "persistent artifact that captures the full computable state of the evidence and grows "
    "richer with each inference pass."
)

bullet("Original DICOM (immutable):",
       "CT, MRI, and X-ray studies preserved as clinical evidence. Ingested via DICOMweb STOW-RS "
       "or DIMSE C-STORE, event-triggered pull from PACS/VNA.")
bullet("Derived artifacts:",
       "Segmentation masks (volumetric), measurements (diameters, volumes, midline shifts), "
       "heatmaps (GradCAM localization), and spatial registrations.")
bullet("Semantic embeddings:",
       "Study-level, series-level, and lesion-level vectors (384-dim) supporting cohort retrieval "
       "and case matching via pgvector similarity search.")
bullet("Structured findings:",
       "Queryable clinical state compatible with DICOM SR, including confidence scores, thresholds, "
       "and guideline mappings (Lung-RADS, RECIST, McDonald criteria).")
bullet("Provenance bundles:",
       "Model ID, version, inference parameters, durations, input data lineage (DICOM UIDs), "
       "timestamps, and operator approvals. Supports predetermined change control plans.")

# ══════════════════════════════════════════════════════════════
# 5. REFERENCE WORKFLOWS
# ══════════════════════════════════════════════════════════════
H1("5. Reference Agent Workflows")

body(
    "Four reference workflows demonstrate the Imaging Intelligence Agent operating across modalities. "
    "Each workflow runs as an event-driven pipeline using MONAI Deploy MAPs on DGX GPU compute, "
    "with all inputs read from and all outputs persisted to the canonical imaging archive."
)

# CT Head
H2("CT Head \u2014 Acute Hemorrhage Triage")

bullet("Target:", "<90 seconds end-to-end  |  Sensitivity: >95% for hemorrhage >5 mL")
bullet("Pipeline:",
       "3D U-Net segmentation-based hemorrhage detection \u2192 volumetric quantification \u2192 midline "
       "shift measurement \u2192 urgency classification (critical / urgent / routine)")
bullet("Output:",
       "DICOM SR structured report + DICOM SEG volumetric mask + FHIR critical alert. "
       "Worklist prioritization and on-call notification per policy thresholds.")
bullet("Trigger:", "study.complete event on non-contrast CT head acquisition.")

# CT Chest
H2("CT Chest \u2014 Lung Nodule Detection and Longitudinal Tracking")

bullet("Target:", "<5 minutes multi-stage  |  Detection: >90% for nodules \u22654 mm")
bullet("Pipeline:",
       "RetinaNet detection \u2192 SegResNet per-nodule segmentation \u2192 volumetrics \u2192 "
       "Lung-RADS assignment \u2192 longitudinal matching to prior CT \u2192 volume "
       "doubling time calculation \u2192 malignancy risk scoring")
bullet("Output:",
       "Per-nodule DICOM SR with Lung-RADS category, growth trajectory, and risk score. "
       "GSPS graphic annotation overlays. Triggers genomics pipeline (Parabricks) if Lung-RADS 4B+.")
bullet("Longitudinal:",
       "Automatic retrieval of prior CT chest for registration and volume comparison. "
       "pgvector hybrid query for phenotype-matched cohort retrieval.")

# CXR
H2("CXR \u2014 Rapid Findings for ED and Inpatient Rounds")

bullet("Target:", "<30 seconds end-to-end  |  Pneumothorax sensitivity: >95%")
bullet("Pipeline:",
       "DenseNet-121 multi-label classification (pneumothorax, consolidation, pleural effusion, "
       "cardiomegaly, fracture) \u2192 GradCAM heatmap localization \u2192 confidence scoring")
bullet("Output:",
       "DICOM SR findings + GradCAM heatmap images (Secondary Capture) pushed to PACS via "
       "DICOMweb STOW-RS. Immediate high-risk flagging for worklist prioritization.")

# MRI Brain
H2("MRI Brain \u2014 MS Lesion Segmentation and Disease Activity")

bullet("Target:", "<5 minutes multi-stage")
bullet("Pipeline:",
       "3D U-Net lesion segmentation on FLAIR \u2192 lesion count and volume \u2192 spatial registration "
       "to prior MRI \u2192 new/enlarging lesion identification \u2192 disease activity "
       "assessment (stable / active / highly active)")
bullet("Output:",
       "Longitudinal lesion report with per-lesion measurements, new lesion details, and disease "
       "activity classification. DICOM SEG volumetric masks. FHIR DiagnosticReport to EHR.")
bullet("RAG reasoning:",
       "Evidence-grounded reasoning over McDonald criteria and current treatment guidelines to "
       "contextualize disease activity relative to the patient\u2019s therapy.")

# ══════════════════════════════════════════════════════════════
# 6. NVIDIA ACCELERATED COMPUTING
# ══════════════════════════════════════════════════════════════
H1("6. NVIDIA Accelerated Computing and Model Stack")

body(
    "NVIDIA provides the accelerated compute and model infrastructure that powers every layer "
    "of the Imaging Intelligence Agent."
)

add_table(
    ["Component", "Role", "Imaging Agent Usage"],
    [
        ["DGX Spark",
         "Desktop AI supercomputer",
         "GB10 Grace Blackwell Superchip, 128 GB unified memory, $4,699. "
         "Proof build for 1\u20132 workflows. Zero NVAIE software cost at desktop-class"],
        ["DGX B200",
         "Departmental GPU cluster",
         "$500K\u2013$515K per 8-GPU system. Multi-user, multi-modality. "
         "PACS integration and clinical validation"],
        ["DGX SuperPOD",
         "Enterprise AI factory",
         "$7M\u2013$60M+. Thousands of concurrent studies. "
         "Continuous reprocessing and population-scale analytics"],
        ["MONAI Deploy",
         "Containerized inference (MAPs)",
         "Application Packages for each workflow. Portable, versioned, "
         "orchestrated by Nextflow"],
        ["MONAI Model Zoo",
         "Pre-trained architectures",
         "3D U-Net, RetinaNet, SegResNet, DenseNet-121, SwinUNETR. "
         "Fine-tuning starting points for institutional data"],
        ["NVIDIA NIM",
         "Inference microservices",
         "Standardized model serving with versioning, health checks, auto-scaling. "
         "NVAIE licensed ($4,500/GPU/year)"],
        ["NVIDIA FLARE",
         "Federated learning",
         "Multi-site model improvement without centralizing patient data. "
         "Apache 2.0 (free)"],
        ["NVIDIA Parabricks",
         "GPU-accelerated genomics",
         "30x WGS from ~30 hours (CPU) to ~10 minutes on DGX (8-GPU). "
         "Triggered by Lung-RADS 4B+ imaging findings"],
        ["NVIDIA BioNeMo",
         "Drug discovery",
         "MolMIM + DiffDock for candidate generation. 200+ adopters. "
         "Quantitative imaging endpoints (RECIST) for clinical trials"],
    ]
)

# ══════════════════════════════════════════════════════════════
# 7. END-TO-END DATA TRACE
# ══════════════════════════════════════════════════════════════
H1("7. End-to-End Data Trace: CT Chest Lung Nodule")

body("The following trace illustrates how components collaborate across a single workflow:")

bullet("1. Ingestion.",
       "CT chest DICOM arrives via DICOMweb STOW-RS \u2192 persisted in local NVMe archive as "
       "immutable canonical data.")
bullet("2. Event trigger.",
       "Nextflow detects study.complete \u2192 triggers lung nodule pipeline.")
bullet("3. Prior retrieval.",
       "Prior CT chest retrieved from the local archive for longitudinal comparison.")
bullet("4. Detection.",
       "RetinaNet MAP executes on DGX Spark GPU \u2192 reads DICOM via GPUDirect Storage.")
bullet("5. Segmentation.",
       "SegResNet MAP segments each detected nodule \u2192 volumetric masks persisted to archive.")
bullet("6. Longitudinal registration.",
       "SyN diffeomorphic registration aligns current and prior studies \u2192 per-nodule volume comparison.")
bullet("7. Classification.",
       "Lung-RADS assignment + malignancy risk scoring based on growth, morphology, and location.")
bullet("8. Embedding.",
       "Lesion-level 384-dim vectors generated via NVIDIA NIM \u2192 stored in pgvector for similarity retrieval.")
bullet("9. Reasoning.",
       "RAG pipeline retrieves ACR Lung-RADS guidelines and institutional protocols \u2192 grounds recommendation.")
bullet("10. Output encoding.",
       "DICOM SR + GSPS overlays pushed to PACS; FHIR DiagnosticReport to EHR; provenance bundle persisted.")
bullet("11. Cross-modal trigger.",
       "If Lung-RADS 4B+, Nextflow triggers Parabricks somatic variant analysis in the HCLS AI Factory.")
bullet("12. Audit.",
       "Complete trace (model version, parameters, inputs, outputs, timestamps) persisted as provenance bundle.")

# ══════════════════════════════════════════════════════════════
# 8. HCLS AI FACTORY INTEGRATION
# ══════════════════════════════════════════════════════════════
H1("8. Integration with the HCLS AI Factory")

body(
    "The Imaging Intelligence Agent is one node in a broader multimodal care fabric. Through "
    "event-driven triggers and cross-modal reasoning, imaging connects directly to other "
    "HCLS AI Factory agents\u2014all running on the same DGX Spark during proof build:"
)

bullet("Imaging \u2192 Genomics (Parabricks):",
       "A confirmed malignant lung nodule triggers tumor profiling\u2014somatic and germline variant "
       "calling\u2014linking imaging phenotype to molecular characterization. Parabricks accelerates "
       "30x WGS from ~30 hours (CPU) to ~10 minutes on DGX (8-GPU).")
bullet("Imaging \u2192 RAG / Clinical Chat:",
       "The RAG pipeline retrieves relevant guidelines and literature, grounding clinical chat "
       "responses in imaging evidence and cross-modal context.")
bullet("Imaging \u2192 Drug Discovery (BioNeMo):",
       "Quantitative imaging endpoints (tumor volume, response assessment) feed drug candidate "
       "scoring and treatment stratification pipelines. 200+ adopters including Eli Lilly, "
       "Astellas, Insilico, and Recursion.")
bullet("Imaging \u2192 Biomarker Intelligence Agent:",
       "Cross-agent data flows enable genomic + imaging biomarker fusion and combined phenotype "
       "profiling on the shared DGX Spark compute.")
bullet("Imaging \u2192 Longitudinal Care:",
       "Continuous monitoring across time points with automated detection of meaningful change. "
       "pgvector embeddings enable cohort retrieval for outcomes research.")

# ══════════════════════════════════════════════════════════════
# 9. CLINICAL INTEGRATION
# ══════════════════════════════════════════════════════════════
H1("9. Clinical Integration")

H3("Output to PACS")
body(
    "DICOM SR (TID 1500 Measurement Report) encodes structured findings with measurements, "
    "classifications, and confidence scores. DICOM SEG objects encode volumetric segmentation masks. "
    "GSPS overlays provide graphic annotations (contours, arrows, text). GradCAM heatmaps are "
    "stored as Secondary Capture images. All outputs are pushed to PACS via DICOMweb STOW-RS, "
    "making AI results available in the radiologist\u2019s native viewing environment."
)

H3("Output to EHR")
body(
    "FHIR DiagnosticReport (R4) resources encode narrative summaries with coded findings (SNOMED CT, "
    "LOINC). Critical finding alerts are dispatched via HL7v2 ORU or FHIR Communication resources "
    "to the clinical notification system."
)

H3("Triage and Worklist Routing")
body(
    "Each study is independently triaged by the agent and mapped to institutional notification "
    "policy thresholds. Urgency classifications drive worklist prioritization (P1 Stat through P4 "
    "Routine) and automated alerting to on-call specialists."
)

H3("Clinician-in-the-Loop")
body(
    "The Imaging Intelligence Agent operates as agentic automation for triage, consistency, and "
    "structured evidence\u2014not autonomous diagnosis. Clinicians remain the accountable "
    "decision-makers. All outputs are reviewable and configurable, aligned to the FDA AI/ML SaMD "
    "framework with predetermined change control plan support."
)

# ══════════════════════════════════════════════════════════════
# 10. TRUST, GOVERNANCE
# ══════════════════════════════════════════════════════════════
H1("10. Trust, Governance, and Clinical Readiness")

H3("Provenance by Default")
body(
    "Every output is traceable to the exact model version, input data, inference configuration, "
    "and timestamps. Provenance bundles are persisted as first-class canonical artifacts, "
    "creating an immutable audit trail."
)

H3("Reproducibility")
body(
    "Deterministic re-runs on canonical data are enabled by immutable DICOM persistence and "
    "fixed-seed inference mode. Previous version outputs are preserved alongside new results "
    "when models are updated under predetermined change control plans."
)

H3("Regulatory Alignment")
body(
    "The system supports the FDA AI/ML SaMD framework, including predetermined change control plans "
    "for controlled model rollouts. Provenance bundles and reproducible inference provide the "
    "documentation chain required for regulatory submissions."
)

H3("Security")
body(
    "Least-privilege access control, role-based permissions, and tenant isolation ensure patient "
    "data remains within institutional control. Agent personas operate with distinct security "
    "credentials scoped to their clinical role."
)

H3("Observability")
body(
    "Pipeline traces (durations, failures, throughput), model performance monitoring, and drift "
    "detection provide continuous visibility into agent health and performance. Grafana and "
    "Prometheus dashboards are included in the open-source deployment."
)

# ══════════════════════════════════════════════════════════════
# 11. DEPLOYMENT ROADMAP
# ══════════════════════════════════════════════════════════════
H1("11. Deployment Roadmap")

H3("Phase 1 \u2014 Proof Build")
bullet("Infrastructure:", "Single DGX Spark ($4,699, GB10 Grace Blackwell, 128 GB unified memory)")
bullet("Scope:", "1\u20132 workflows (CT head hemorrhage triage, CXR rapid findings)")
bullet("Validation:", "MONAI Deploy MAPs validated; canonical data model proven; event triggers tested")
bullet("Software cost:", "Zero NVAIE licensing (desktop-class)")

H3("Phase 2 \u2014 Departmental")
bullet("Infrastructure:", "1\u20132x DGX B200 cluster ($500K\u2013$1M)")
bullet("Scope:", "All four reference workflows; shared filesystem namespace; multi-user")
bullet("Integration:", "PACS integration via DICOMweb; clinical validation with radiologists")
bullet("Software cost:", "NVAIE $36K\u2013$72K/year (8\u201316 GPUs)")

H3("Phase 3 \u2014 Multi-Site Enterprise")
bullet("Infrastructure:", "4\u20138x DGX B200 + InfiniBand ($2M\u2013$4M)")
bullet("Scope:", "Cross-site namespace; continuous reprocessing; population-scale cohort retrieval")
bullet("Capabilities:", "NVIDIA FLARE federated learning; multi-site model improvement")
bullet("Software cost:", "NVAIE $144K\u2013$288K/year (32\u201364 GPUs)")

H3("Phase 4 \u2014 AI Factory at Scale")
bullet("Infrastructure:", "DGX SuperPOD ($7M\u2013$60M+)")
bullet("Scope:", "Unified multimodal agent fabric\u2014Imaging + Genomics + Biomarkers + Outcomes + Therapy")
bullet("Capability:", "Thousands of concurrent studies; complete HCLS AI Factory on one compute substrate")
bullet("Software cost:", "NVAIE $576K\u2013$1.15M/year (128\u2013256 GPUs)")

# ══════════════════════════════════════════════════════════════
# 12. OPEN-SOURCE TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════
H1("12. Open-Source Technology Stack")

body(
    "The Imaging Intelligence Agent is built on open-source foundations with NVIDIA acceleration. "
    "All core components are freely available; NVIDIA AI Enterprise licensing applies only at "
    "production scale beyond DGX Spark."
)

add_table(
    ["Component", "Role", "License"],
    [
        ["MONAI", "Medical imaging AI framework", "Apache 2.0"],
        ["MONAI Deploy", "Containerized inference packaging (MAPs)", "Apache 2.0"],
        ["NVIDIA FLARE", "Federated learning across institutions", "Apache 2.0"],
        ["Nextflow", "Pipeline DAG orchestration", "Apache 2.0"],
        ["LangChain / LangGraph", "Agent orchestration + MCP tools", "MIT"],
        ["PostgreSQL + pgvector", "Structured + vector query", "PostgreSQL License"],
        ["Orthanc", "DICOM server", "GPLv3"],
        ["dcm4chee", "DICOM archive", "MPL 1.1 / GPL 2.0 / LGPL 2.1"],
        ["pydicom", "DICOM parsing and manipulation", "MIT"],
        ["HAPI FHIR", "FHIR server for clinical integration", "Apache 2.0"],
        ["Grafana + Prometheus", "Monitoring and observability", "AGPL / Apache 2.0"],
        ["NVIDIA NIM", "Inference microservices", "NVAIE ($4,500/GPU/yr)"],
        ["NVIDIA Parabricks", "GPU-accelerated genomics", "NVAIE ($4,500/GPU/yr)"],
        ["NVIDIA BioNeMo", "Drug discovery", "NVAIE ($4,500/GPU/yr)"],
    ]
)

# ══════════════════════════════════════════════════════════════
# 13. GETTING STARTED
# ══════════════════════════════════════════════════════════════
H1("13. Getting Started")

body(
    "The HCLS AI Factory is an open-source project released under the Apache 2.0 license. "
    "The imaging pipeline, along with the genomics, RAG/clinical chat, and drug discovery "
    "pipelines, is available on GitHub."
)

H3("Prerequisites")
bullet("Hardware:", "NVIDIA DGX Spark (recommended) or any system with NVIDIA GPU and CUDA 12+")
bullet("Software:", "Docker, Nextflow, Python 3.10+")
bullet("Data:", "Sample DICOM studies included for demonstration")

H3("Quick Start")
body(
    "1. Clone the HCLS AI Factory repository\n"
    "2. Configure environment variables (.env)\n"
    "3. Launch services: docker-compose up\n"
    "4. Ingest sample DICOM studies via DICOMweb STOW-RS\n"
    "5. Monitor pipeline execution via Grafana dashboard\n"
    "6. Review generated reports (DICOM SR, FHIR, PDF)"
)

H3("Contributing")
body(
    "Contributions are welcome under the Apache 2.0 license. See CONTRIBUTING.md for guidelines "
    "on code style, testing requirements, and the pull request process."
)

# ── Footer ───────────────────────────────────────────────────
P("")
P("")
p = P("", after=0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
R(p, "HCLS AI Factory", bold=True, size=9, color=NAVY)
R(p, "  |  Open Source (Apache 2.0)  |  NVIDIA DGX Spark  |  02/2026  |  v1.0",
  size=9, color=GRAY_META)

# ── Save ─────────────────────────────────────────────────────
path = "core/engines/clinical-imaging/open_public/HCLS_Imaging_AI_Agent_White_Paper_DGX_Spark.docx"
doc.save(path)
print(f"White paper saved to {path}")
