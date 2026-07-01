#!/usr/bin/env python3
# Copyright 2026 Adam Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Generate Imaging Intelligence Agent — Demo Guide (DOCX).

Step-by-step demo procedures for the HCLS Imaging Intelligence Agent
on NVIDIA DGX Spark. Formatted in the VCP-style theme.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Colors (identical to white paper) ──────────────────────
NAVY = RGBColor(0x1B, 0x23, 0x33)
TEAL = RGBColor(0x1A, 0xAF, 0xCC)
GREEN = RGBColor(0x76, 0xB9, 0x00)
GRAY_BODY = RGBColor(0x33, 0x33, 0x33)
GRAY_META = RGBColor(0x66, 0x66, 0x66)
GRAY_CODE = RGBColor(0x44, 0x44, 0x44)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_NAVY = "1B2333"
HEX_TEAL = "1AAFCC"
HEX_GREEN = "76B900"
HEX_LIGHT = "F8FAFB"
HEX_CODE_BG = "F5F5F5"
FONT = "Calibri"
CODE_FONT = "Consolas"

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
style.font.color.rgb = GRAY_BODY
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15


# ── Helper functions ───────────────────────────────────────

def R(p, text, bold=False, italic=False, size=None, color=None, font=None):
    r = p.add_run(text)
    r.font.name = font or FONT
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return r


def P(text="", bold=False, italic=False, size=10.5, color=GRAY_BODY,
      before=0, after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text:
        R(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def H1(text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        r.clear()
    R(p, text, bold=True, size=22, color=NAVY)
    return p


def H2(text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.clear()
    R(p, text, bold=True, size=16, color=NAVY)
    return p


def H3(text):
    return P(text, bold=True, size=13, color=TEAL, before=10, after=4)


def body(text, before=0, after=6):
    return P(text, size=10.5, color=GRAY_BODY, before=before, after=after)


def bullet(lead, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    R(p, lead, bold=True, size=10.5, color=GRAY_BODY)
    R(p, f"  {text}", size=10.5, color=GRAY_BODY)
    return p


def code_block(text, language=""):
    """Render a code block with monospace font and light gray background."""
    if language:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        R(p, language, bold=True, size=8, color=TEAL, font=FONT)

    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.2)
        R(p, line if line else " ", size=8, color=GRAY_CODE, font=CODE_FONT)


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="D0D0D0", sz="4"):
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>')
    tcPr.append(borders)


def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, HEX_NAVY)
        set_cell_borders(c, HEX_NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(p, h, bold=True, size=9, color=WHITE)
    for ri, row in enumerate(rows):
        bg = "FFFFFF" if ri % 2 == 0 else HEX_LIGHT
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            set_cell_shading(c, bg)
            set_cell_borders(c, "E0E0E0")
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if ci == 0:
                R(p, str(val), bold=True, size=9, color=TEAL)
            else:
                R(p, str(val), size=9, color=GRAY_BODY)
    return t


def spacer(pts=6):
    P("", after=pts)


# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════

# Green accent bar
accent = doc.add_table(rows=1, cols=1)
accent.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = accent.rows[0].cells[0]
set_cell_shading(cell, HEX_GREEN)
p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
R(p, " ", size=6, color=GREEN)

for _ in range(4):
    P("")

P("Demo Guide", bold=True, size=13, color=TEAL, after=6)
P("Imaging Intelligence Agent",
  bold=True, size=32, color=NAVY, after=2)
P("Step-by-Step Procedures",
  bold=True, size=32, color=NAVY, after=6)
P("Complete step-by-step demo procedures for running 29 demos across "
  "7 parts on the HCLS Imaging Intelligence Agent on NVIDIA DGX Spark. "
  "From environment setup through advanced cross-modal integration.",
  italic=False, size=12, color=NAVY, after=6)
P("NVIDIA DGX Spark  |  MONAI Deploy  |  NVIDIA NIM  |  Open-Source Stack",
  italic=True, size=11, color=GRAY_META, after=12)
P("")
P("02/2026  |  Version 1.0  |  Apache 2.0 License", size=9,
  color=GRAY_META, after=2)
P("Author: Adam Jones", size=9, color=GRAY_META, after=2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
H1("Table of Contents")

toc_items = [
    ("Part 1: Environment Setup (Build from Scratch)", [
        "Demo 0:  Prerequisites & System Verification",
        "Demo 1:  Build & Launch All Services",
        "Demo 2:  Database Initialization",
        "Demo 3:  Model Downloads",
        "Demo 4:  Orthanc DICOM Server Verification",
    ]),
    ("Part 2: Clinical Workflow Demos", [
        "Demo 5:  CXR Rapid Findings (Simplest Workflow)",
        "Demo 6:  CT Head Hemorrhage Triage",
        "Demo 7:  CT Chest Lung Nodule Tracking",
        "Demo 8:  MRI Brain MS Lesion Tracking",
    ]),
    ("Part 3: Agent Reasoning Demos", [
        "Demo 9:  Triage Agent",
        "Demo 10: Longitudinal Comparison Agent",
        "Demo 11: Population Analysis Agent",
        "Demo 12: RAG-Grounded Report Generation",
        "Demo 13: Conditional Routing",
    ]),
    ("Part 4: Output & Integration Demos", [
        "Demo 14: DICOM SR Structured Report",
        "Demo 15: GSPS Graphic Overlay",
        "Demo 16: FHIR DiagnosticReport",
        "Demo 17: Embedding & Vector Search",
    ]),
    ("Part 5: Portal & Monitoring Demos", [
        "Demo 18: Streamlit Portal \u2014 Worklist",
        "Demo 19: Streamlit Portal \u2014 Study Browser",
        "Demo 20: Streamlit Portal \u2014 Agent Activity",
        "Demo 21: Grafana Monitoring Dashboard",
        "Demo 22: Alerting Demo",
    ]),
    ("Part 6: Advanced Demos", [
        "Demo 23: End-to-End Patient Case",
        "Demo 24: Cross-Modal Genomics Trigger",
        "Demo 25: Nextflow Pipeline Orchestration",
        "Demo 26: Performance Benchmarking",
        "Demo 27: Provenance & Reproducibility",
    ]),
    ("Part 7: Teardown & Cleanup", [
        "Demo 28: Graceful Shutdown & Data Persistence",
    ]),
]

for part_title, demos in toc_items:
    P(part_title, bold=True, size=11, color=NAVY, before=8, after=2)
    for demo in demos:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.35)
        R(p, demo, size=10, color=GRAY_BODY)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# ABOUT THIS GUIDE
# ══════════════════════════════════════════════════════════════
H1("About This Guide")

body("This guide provides detailed step-by-step procedures for every demo "
     "that can be run with the Imaging Intelligence Agent on an NVIDIA DGX "
     "Spark. It covers 29 demos organized in 7 parts, from building the "
     "system from scratch through advanced cross-modal integration.")

body("Each demo includes: Purpose, Prerequisites, Step-by-step commands, "
     "Expected output, Verification checklist, and Troubleshooting tips.")

body("Hardware target: NVIDIA DGX Spark \u2014 GB10 Grace Blackwell, "
     "128 GB unified LPDDR5x, ARM64 (aarch64), NVMe storage, $4,699.")

spacer()

# ── Quick Reference ──
H2("Quick Reference")

H3("Port Allocation")
add_table(
    ["Service", "Host Port", "Container Port", "Protocol"],
    [
        ["Orthanc (DIMSE)", "4242", "4242", "DIMSE"],
        ["Orthanc (DICOMweb)", "8042", "8042", "HTTP REST"],
        ["PostgreSQL", "5432", "5432", "TCP"],
        ["NIM LLM", "8520", "8000", "HTTP (OpenAI-compatible)"],
        ["Embedding Service", "8521", "8000", "HTTP REST"],
        ["DICOM Listener", "8522", "8000", "HTTP webhook"],
        ["FHIR Publisher", "8523", "8000", "HTTP REST"],
        ["Agent API", "8524", "8000", "HTTP REST"],
        ["Streamlit Portal", "8525", "8501", "HTTP"],
        ["Prometheus", "9099", "9090", "HTTP"],
        ["Grafana", "3000", "3000", "HTTP"],
        ["DCGM Exporter", "9400", "9400", "HTTP"],
    ],
)
spacer()

H3("Container Names")
add_table(
    ["Container", "Image"],
    [
        ["imaging-orthanc", "orthancteam/orthanc:24.1.2"],
        ["imaging-postgres", "pgvector/pgvector:pg16"],
        ["imaging-nim-llm",
         "nvcr.io/nvidia/nim/meta-llama3-8b-instruct:latest-dgx-spark"],
        ["imaging-embedding", "Custom (./services/embedding_service)"],
        ["imaging-dicom-listener", "Custom (./services/dicom_listener)"],
        ["imaging-fhir-publisher", "Custom (./services/fhir_publisher)"],
        ["imaging-agent", "Custom (./agent)"],
        ["imaging-portal", "Custom (./services/portal)"],
        ["imaging-dcgm",
         "nvcr.io/nvidia/k8s/dcgm-exporter:3.3.5-3.4.0-ubuntu22.04"],
        ["imaging-prometheus", "prom/prometheus:v2.48.0"],
        ["imaging-grafana", "grafana/grafana:10.2.2"],
    ],
)
spacer()

H3("Service Health Check URLs")
add_table(
    ["Service", "Health Check URL"],
    [
        ["Orthanc", "http://localhost:8042/system"],
        ["NIM LLM", "http://localhost:8520/v1/health/ready"],
        ["Embedding Service", "http://localhost:8521/health"],
        ["DICOM Listener", "http://localhost:8522/health"],
        ["FHIR Publisher", "http://localhost:8523/health"],
        ["Agent API", "http://localhost:8524/health"],
        ["Streamlit Portal", "http://localhost:8525/healthz"],
        ["Prometheus", "http://localhost:9099/-/healthy"],
        ["Grafana", "http://localhost:3000/api/health"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PART 1: ENVIRONMENT SETUP
# ══════════════════════════════════════════════════════════════
H1("Part 1: Environment Setup (Build from Scratch)")

# ── Demo 0 ──
H2("Demo 0: Prerequisites & System Verification")

body("Purpose: Confirm the DGX Spark hardware, software dependencies, "
     "and credentials are correctly configured before building any services.")
body("Prerequisites: Physical access to DGX Spark with DGX OS installed.")

H3("Step 1 \u2014 Verify ARM64 Architecture")
code_block("uname -m", "bash")
body("Expected output:")
code_block("aarch64")
body("If you see x86_64, you are not on a DGX Spark. All containers in "
     "this project require ARM64.")

H3("Step 2 \u2014 Verify GPU and Memory")
code_block("nvidia-smi", "bash")
body("Expected output (key lines):")
code_block(
    "+---------------------------------------------------+\n"
    "| NVIDIA-SMI 560.x.xx  Driver: 560.x.xx  CUDA: 12.x|\n"
    "|  GPU   Name: NVIDIA GB10                          |\n"
    "| GPU Memory: 128 GB Unified LPDDR5x                |\n"
    "+---------------------------------------------------+"
)

H3("Step 3 \u2014 Verify Docker and NVIDIA Container Toolkit")
code_block(
    "docker --version\n"
    "nvidia-ctk --version\n"
    "docker run --rm --gpus all nvidia/cuda:12.4.0-base-ubuntu22.04 nvidia-smi",
    "bash",
)
body("Expected output: Docker version 24+, nvidia-ctk version 1.14+, "
     "and nvidia-smi running inside the container.")

H3("Step 4 \u2014 Verify NGC API Key")
code_block('echo ${NGC_API_KEY:+Key is set}', "bash")
body("Expected output:")
code_block("Key is set")
body("If blank, set it:")
code_block('export NGC_API_KEY="your-ngc-api-key-here"', "bash")
body("Obtain a key from https://ngc.nvidia.com/ under your account settings.")

H3("Step 5 \u2014 Clone Repository and Inspect Layout")
code_block(
    "git clone https://github.com/your-org/hls-imaging-agent.git\n"
    "cd hls-imaging-agent\n"
    "ls -la",
    "bash",
)
body("Expected output: Repository root with docker-compose.yml, main.nf, "
     "nextflow.config, maps/, agent/, services/, db/, config/, scripts/, tests/.")

H3("Step 6 \u2014 Create Environment File")
code_block("cp .env.example .env", "bash")
body("Edit .env and set at minimum:")
code_block(
    "NGC_API_KEY=your-ngc-api-key\n"
    "POSTGRES_USER=imaging\n"
    "POSTGRES_PASSWORD=imaging_secret\n"
    "POSTGRES_DB=imaging_agent\n"
    "GRAFANA_USER=admin\n"
    "GRAFANA_PASSWORD=changeme",
    "bash",
)

H3("Verification Checklist")
bullet("uname -m", "returns aarch64")
bullet("nvidia-smi", "shows GB10 with 128 GB unified memory")
bullet("Docker 24+", "installed")
bullet("NVIDIA Container Toolkit", "installed")
bullet("NGC API key", "is set")
bullet("Repository cloned", "and .env file created")

H3("Troubleshooting")
add_table(
    ["Symptom", "Cause", "Fix"],
    [
        ["uname -m shows x86_64", "Not on DGX Spark",
         "SSH to the correct machine"],
        ["nvidia-smi not found", "Driver not installed",
         "Install DGX OS or NVIDIA driver"],
        ["Docker permission denied", "User not in docker group",
         "sudo usermod -aG docker $USER then re-login"],
        ["NGC_API_KEY empty", "Not configured",
         "Generate key at ngc.nvidia.com"],
    ],
)

doc.add_page_break()

# ── Demo 1 ──
H2("Demo 1: Build & Launch All Services")

body("Purpose: Build all container images and start the full 11-service stack.")
body("Prerequisites: Demo 0 completed. .env file configured.")

H3("Step 1 \u2014 Build Custom Containers")
code_block("docker compose build", "bash")
body("This builds 6 custom services: embedding-service, dicom-listener, "
     "fhir-publisher, agent, portal, and the 4 MAP containers. The 5 "
     "pre-built images (orthanc, postgres, nim-llm, dcgm-exporter, "
     "prometheus, grafana) are pulled automatically.")
body("Expected output (final lines):")
code_block(
    "[+] Building 6/6\n"
    " \u2714 embedding-service  Built\n"
    " \u2714 dicom-listener     Built\n"
    " \u2714 fhir-publisher     Built\n"
    " \u2714 agent              Built\n"
    " \u2714 portal             Built"
)

H3("Step 2 \u2014 Launch All Services")
code_block("docker compose up -d", "bash")
body("Expected output:")
code_block(
    "[+] Running 11/11\n"
    " \u2714 Container imaging-postgres        Healthy\n"
    " \u2714 Container imaging-orthanc         Healthy\n"
    " \u2714 Container imaging-nim-llm         Healthy\n"
    " \u2714 Container imaging-embedding       Healthy\n"
    " \u2714 Container imaging-dicom-listener  Healthy\n"
    " \u2714 Container imaging-fhir-publisher  Healthy\n"
    " \u2714 Container imaging-agent           Healthy\n"
    " \u2714 Container imaging-portal          Started\n"
    " \u2714 Container imaging-dcgm            Started\n"
    " \u2714 Container imaging-prometheus      Started\n"
    " \u2714 Container imaging-grafana         Started"
)

H3("Step 3 \u2014 Verify All Containers Healthy")
code_block("docker compose ps", "bash")
body("Expected output:")
code_block(
    "NAME                     IMAGE                                STATUS\n"
    "imaging-orthanc          orthancteam/orthanc:24.1.2          Up (healthy)\n"
    "imaging-postgres         pgvector/pgvector:pg16              Up (healthy)\n"
    "imaging-nim-llm          nvcr.io/.../meta-llama3-8b:dgx...  Up (healthy)\n"
    "imaging-embedding        imaging-embedding:latest            Up (healthy)\n"
    "imaging-dicom-listener   imaging-dicom-listener:latest       Up (healthy)\n"
    "imaging-fhir-publisher   imaging-fhir-publisher:latest       Up (healthy)\n"
    "imaging-agent            imaging-agent:latest                Up (healthy)\n"
    "imaging-portal           imaging-portal:latest               Up (healthy)\n"
    "imaging-dcgm             nvcr.io/.../dcgm-exporter:3.3.5    Up\n"
    "imaging-prometheus       prom/prometheus:v2.48.0             Up\n"
    "imaging-grafana          grafana/grafana:10.2.2              Up"
)

H3("Step 4 \u2014 Check Each Health Endpoint")
code_block(
    "curl -s http://localhost:8042/system | python3 -m json.tool\n"
    "curl -s http://localhost:8520/v1/health/ready\n"
    "curl -s http://localhost:8521/health\n"
    "curl -s http://localhost:8522/health\n"
    "curl -s http://localhost:8523/health\n"
    "curl -s http://localhost:8524/health\n"
    "curl -s http://localhost:8525/healthz\n"
    "curl -s http://localhost:9099/-/healthy\n"
    "curl -s http://localhost:3000/api/health",
    "bash",
)
body("Each should return a 200 status or JSON with \"status\": \"ok\".")

H3("Verification Checklist")
bullet("All 11 containers", "are running")
bullet("Orthanc, PostgreSQL, NIM LLM", "report healthy")
bullet("All custom services", "(embedding, dicom-listener, fhir-publisher, "
       "agent, portal) report healthy")
bullet("No port conflicts", "(check with ss -tlnp | grep -E "
       "'4242|8042|5432|8520|8521|8522|8523|8524|8525|9099|3000|9400')")

H3("Troubleshooting")
add_table(
    ["Symptom", "Cause", "Fix"],
    [
        ["NIM LLM stays \"starting\"",
         "First pull downloads model weights (~16 GB)",
         "Wait 5-10 minutes; check docker logs imaging-nim-llm"],
        ["Port already in use", "Another service on that port",
         "ss -tlnp | grep <port> to find conflict; stop the other service"],
        ["GPU access denied", "Container runtime not configured",
         "Verify nvidia-ctk runtime configure --runtime=docker was run"],
        ["OOM on startup", "Too many GPU services",
         "Start services in stages: data layer first, then execution layer"],
    ],
)

doc.add_page_break()

# ── Demo 2 ──
H2("Demo 2: Database Initialization")

body("Purpose: Verify PostgreSQL + pgvector schema is correctly initialized "
     "with all tables, indexes, and views.")
body("Prerequisites: Demo 1 completed. imaging-postgres container healthy.")

H3("Step 1 \u2014 Connect to PostgreSQL")
code_block("docker exec -it imaging-postgres psql -U imaging -d imaging_agent",
           "bash")

H3("Step 2 \u2014 List Tables")
code_block("\\dt", "sql")
body("Expected output:")
code_block(
    "              List of relations\n"
    " Schema |      Name        | Type  | Owner\n"
    "--------+------------------+-------+---------\n"
    " public | embeddings       | table | imaging\n"
    " public | findings         | table | imaging\n"
    " public | measurements     | table | imaging\n"
    " public | provenance       | table | imaging\n"
    " public | series           | table | imaging\n"
    " public | studies          | table | imaging\n"
    " public | worklist_entries | table | imaging\n"
    "(7 rows)"
)

H3("Step 3 \u2014 Verify pgvector Extension")
code_block("\\dx", "sql")
body("Expected output (includes):")
code_block(
    " vector | 0.7.0 | public | vector data type and ivfflat and hnsw access methods"
)

H3("Step 4 \u2014 Verify HNSW Index")
code_block(
    "SELECT indexname, indexdef FROM pg_indexes\n"
    "WHERE tablename = 'embeddings';",
    "sql",
)
body("Expected output (includes):")
code_block(
    " idx_embeddings_hnsw | CREATE INDEX idx_embeddings_hnsw ON public.embeddings\n"
    "   USING hnsw (embedding vector_cosine_ops)\n"
    "   WITH (m='16', ef_construction='64')"
)

H3("Step 5 \u2014 Verify Views")
code_block("\\dv", "sql")
body("Expected output:")
code_block(
    "          List of relations\n"
    " Schema |      Name       | Type |  Owner\n"
    "--------+-----------------+------+---------\n"
    " public | active_worklist | view | imaging\n"
    " public | study_summary   | view | imaging\n"
    "(2 rows)"
)

H3("Step 6 \u2014 Verify Empty State")
code_block(
    "SELECT COUNT(*) FROM studies;\n"
    "SELECT COUNT(*) FROM findings;\n"
    "SELECT COUNT(*) FROM embeddings;",
    "sql",
)
body("All should return 0.")

H3("Step 7 \u2014 Exit")
code_block("\\q", "sql")

H3("Verification Checklist")
bullet("7 tables exist", "(studies, series, findings, measurements, "
       "embeddings, provenance, worklist_entries)")
bullet("pgvector extension", "is loaded")
bullet("HNSW index", "exists on embeddings table")
bullet("2 views exist", "(active_worklist, study_summary)")
bullet("All tables", "are empty (fresh state)")

doc.add_page_break()

# ── Demo 3 ──
H2("Demo 3: Model Downloads")

body("Purpose: Download MONAI Model Zoo weights and verify NIM LLM and "
     "embedding service readiness.")
body("Prerequisites: Demo 1 completed. All services running.")

H3("Step 1 \u2014 Download MONAI Model Zoo Weights")
code_block("bash scripts/download_models.sh", "bash")
body("Expected output:")
code_block(
    "Downloading DenseNet-121 (classification)...  \u2713\n"
    "Downloading 3D U-Net (segmentation)...        \u2713\n"
    "Downloading RetinaNet (detection)...           \u2713\n"
    "Downloading SegResNet (segmentation)...        \u2713\n"
    "\n"
    "All models downloaded to ./models/"
)

H3("Step 2 \u2014 Verify Model Files")
code_block("ls -lh models/", "bash")
body("Expected output:")
code_block(
    "-rw-r--r-- 1 user user  28M  densenet121_classification.pt\n"
    "-rw-r--r-- 1 user user 120M  unet3d_segmentation.pt\n"
    "-rw-r--r-- 1 user user  85M  retinanet_detection.pt\n"
    "-rw-r--r-- 1 user user  95M  segresnet_segmentation.pt"
)

H3("Step 3 \u2014 Verify NIM LLM Readiness")
code_block(
    "curl -s http://localhost:8520/v1/models | python3 -m json.tool",
    "bash",
)
body("Expected output:")
code_block(
    '{\n'
    '    "object": "list",\n'
    '    "data": [\n'
    '        {\n'
    '            "id": "meta-llama3-8b-instruct",\n'
    '            "object": "model",\n'
    '            "owned_by": "nvidia"\n'
    '        }\n'
    '    ]\n'
    '}',
    "json",
)

H3("Step 4 \u2014 Test NIM LLM Inference")
code_block(
    'curl -s http://localhost:8520/v1/chat/completions \\\n'
    '  -H "Content-Type: application/json" \\\n'
    '  -d \'{\n'
    '    "model": "meta-llama3-8b-instruct",\n'
    '    "messages": [{"role":"user","content":"What is Lung-RADS?"}],\n'
    '    "max_tokens": 100\n'
    "  }' | python3 -m json.tool",
    "bash",
)
body("Expected output: JSON response with a chat completion about "
     "Lung-RADS classification.")

H3("Step 5 \u2014 Verify Embedding Service")
code_block("curl -s http://localhost:8521/health", "bash")
body("Expected output:")
code_block(
    '{"status":"ok","model":"microsoft/BiomedCLIP-PubMedBERT_256-'
    'vit_base_patch16_224"}',
    "json",
)

H3("Verification Checklist")
bullet("All 4 MONAI model weight files", "present in ./models/")
bullet("NIM LLM", "lists the meta-llama3-8b-instruct model")
bullet("NIM LLM", "responds to chat completion requests")
bullet("Embedding service", "reports healthy with BiomedCLIP model loaded")

H3("Troubleshooting")
add_table(
    ["Symptom", "Cause", "Fix"],
    [
        ["download_models.sh fails", "Network issue or NGC auth",
         "Check internet connectivity and NGC_API_KEY"],
        ["NIM LLM returns 503", "Model still loading",
         "Wait for docker logs imaging-nim-llm to show \"ready\""],
        ["Embedding service unhealthy", "BiomedCLIP download failed",
         "Check docker logs imaging-embedding for download errors"],
        ["Disk space error", "NVMe full",
         "Check df -h and clean up unused Docker images"],
    ],
)

doc.add_page_break()

# ── Demo 4 ──
H2("Demo 4: Orthanc DICOM Server Verification")

body("Purpose: Verify the Orthanc DICOM server is operational, DICOMweb "
     "endpoints work, and the study.complete webhook pipeline fires correctly.")
body("Prerequisites: Demo 1 completed. imaging-orthanc and "
     "imaging-dicom-listener containers healthy.")

H3("Step 1 \u2014 Check Orthanc System Info")
code_block(
    "curl -s http://localhost:8042/system | python3 -m json.tool",
    "bash",
)
body("Expected output:")
code_block(
    '{\n'
    '    "ApiVersion": 23,\n'
    '    "DicomAet": "IMAGING_AGENT",\n'
    '    "DicomPort": 4242,\n'
    '    "Name": "ImagingAgent",\n'
    '    "Version": "1.12.3"\n'
    '}',
    "json",
)

H3("Step 2 \u2014 Check DICOMweb Endpoint (Empty Studies)")
code_block(
    "curl -s http://localhost:8042/dicom-web/studies | python3 -m json.tool",
    "bash",
)
body("Expected output:")
code_block("[]", "json")

H3("Step 3 \u2014 Verify DIMSE C-ECHO")
code_block("docker exec imaging-orthanc echoscu localhost 4242", "bash")
body("Expected output:")
code_block(
    "Association Accepted\n"
    "ECHO Response: Success\n"
    "Association Released"
)

H3("Step 4 \u2014 Upload a Test DICOM File")
body("Generate a minimal test DICOM (synthetic CT head):")
code_block(
    'python3 -c "\n'
    "import pydicom\n"
    "from pydicom.dataset import Dataset, FileDataset\n"
    "from pydicom.uid import generate_uid\n"
    "import numpy as np\n"
    "...\n"
    "ds.PatientID = 'TEST001'\n"
    "ds.Modality = 'CT'\n"
    "ds.BodyPartExamined = 'HEAD'\n"
    "ds.Rows = ds.Columns = 512\n"
    "ds.PixelData = np.random.randint(-1000,2000,(512,512),dtype=np.int16).tobytes()\n"
    "ds.save_as('/tmp/test_ct_head.dcm')\n"
    '"',
    "bash",
)

H3("Step 5 \u2014 Upload via DICOMweb STOW-RS")
code_block(
    'curl -X POST http://localhost:8042/dicom-web/studies \\\n'
    '  -H "Content-Type: application/dicom" \\\n'
    '  --data-binary @/tmp/test_ct_head.dcm',
    "bash",
)
body("Expected output: XML NativeDicomModel with ReferencedSOPSequence.")

H3("Step 6 \u2014 Watch for Webhook (Wait StableAge = 10 Seconds)")
code_block("docker logs -f imaging-dicom-listener", "bash")
body("After 10 seconds (Orthanc StableAge), Orthanc fires the Lua "
     "OnStableStudy webhook.")
body("Expected output in dicom-listener logs:")
code_block(
    "INFO  Received study.complete webhook: orthanc_id=<orthanc-id>\n"
    "INFO  Study metadata retrieved: PatientID=TEST001, Modality=CT, BodyPart=HEAD\n"
    "INFO  Study inserted into database: study_id=1\n"
    "INFO  Routing to workflow: ct_head_hemorrhage"
)

H3("Step 7 \u2014 Verify Study in Orthanc")
code_block(
    "curl -s http://localhost:8042/dicom-web/studies | python3 -m json.tool",
    "bash",
)
body("Expected output: JSON array with 1 study entry.")

H3("Verification Checklist")
bullet("Orthanc system info", "shows AET=IMAGING_AGENT, Port=4242")
bullet("DICOMweb studies endpoint", "accessible")
bullet("DIMSE C-ECHO", "succeeds")
bullet("DICOM file uploaded", "via STOW-RS")
bullet("Lua webhook fires", "after StableAge (10 seconds)")
bullet("DICOM listener", "receives webhook and inserts study into database")

H3("Troubleshooting")
add_table(
    ["Symptom", "Cause", "Fix"],
    [
        ["STOW-RS returns 400", "Invalid DICOM file",
         "Verify with pydicom.dcmread() first"],
        ["Webhook never fires", "Lua script not loaded",
         "Check docker logs imaging-orthanc for Lua errors; verify "
         "./config/scripts/ is mounted"],
        ["DICOM listener doesn't receive", "Network issue",
         "Verify both containers are on imaging-agent-network"],
        ["StableAge too long", "Configuration",
         "Check orthanc.json has \"StableAge\": 10"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PART 2: CLINICAL WORKFLOW DEMOS
# ══════════════════════════════════════════════════════════════
H1("Part 2: Clinical Workflow Demos")

# ── Demo 5 ──
H2("Demo 5: CXR Rapid Findings (Simplest Workflow)")

body("Purpose: Run an end-to-end chest X-ray classification with GradCAM "
     "heatmap generation. This is the simplest workflow (2D, single model, "
     "< 30 seconds) and the recommended starting point.")
body("Prerequisites: All services running (Demo 1). DenseNet-121 model "
     "loaded (Demo 3).")
body("Performance target: < 30 seconds end-to-end. Pneumothorax "
     "sensitivity > 95%.")

H3("Step 1 \u2014 Generate Synthetic CXR Test Data")
body("Generate a synthetic PA chest X-ray DICOM:")
code_block(
    'python3 -c "\n'
    "import pydicom\n"
    "from pydicom.dataset import Dataset, FileDataset\n"
    "from pydicom.uid import generate_uid\n"
    "import numpy as np\n"
    "...\n"
    "ds.PatientID = 'CXR001'\n"
    "ds.Modality = 'DX'\n"
    "ds.BodyPartExamined = 'CHEST'\n"
    "ds.ViewPosition = 'PA'\n"
    "ds.Rows = ds.Columns = 2048\n"
    "ds.PixelData = np.random.randint(0,4096,(2048,2048),dtype=np.uint16).tobytes()\n"
    "ds.save_as('/tmp/test_cxr_pa.dcm')\n"
    '"',
    "bash",
)

H3("Step 2 \u2014 Start Watching Logs (Separate Terminal)")
code_block("docker logs -f imaging-dicom-listener", "bash")

H3("Step 3 \u2014 Upload CXR via DICOMweb STOW-RS")
code_block(
    'time curl -X POST http://localhost:8042/dicom-web/studies \\\n'
    '  -H "Content-Type: application/dicom" \\\n'
    '  --data-binary @/tmp/test_cxr_pa.dcm',
    "bash",
)

H3("Step 4 \u2014 Monitor Pipeline Execution")
body("In the dicom-listener logs, watch for:")
code_block(
    "INFO  Received study.complete webhook: orthanc_id=<id>\n"
    "INFO  Study metadata: PatientID=CXR001, Modality=DX, BodyPart=CHEST\n"
    "INFO  Routing to workflow: cxr_findings\n"
    "INFO  [cxr_findings] Starting DenseNet-121 multi-label classification\n"
    "INFO  [cxr_findings] Preprocessing: Resize(224,224), Normalize, EnsureChannelFirst\n"
    "INFO  [cxr_findings] Inference complete (8.2s)\n"
    "INFO  [cxr_findings] Results:\n"
    "INFO    pneumothorax:     0.12 (negative, threshold=0.50)\n"
    "INFO    consolidation:    0.87 (POSITIVE, threshold=0.40)\n"
    "INFO    pleural_effusion: 0.34 (negative, threshold=0.40)\n"
    "INFO    cardiomegaly:     0.62 (POSITIVE, threshold=0.50)\n"
    "INFO    fracture:         0.08 (negative, threshold=0.50)\n"
    "INFO  [cxr_findings] GradCAM heatmap generated for: consolidation, cardiomegaly\n"
    "INFO  [cxr_findings] DICOM SR created\n"
    "INFO  [cxr_findings] Findings persisted to database\n"
    "INFO  [cxr_findings] Pipeline complete (12.4s total)"
)

H3("Step 5 \u2014 Query Findings in Database")
code_block(
    "docker exec imaging-postgres psql -U imaging -d imaging_agent -c \"\n"
    "SELECT f.finding_type, f.confidence, f.severity, f.is_positive\n"
    "FROM findings f\n"
    "JOIN studies s ON f.study_id = s.id\n"
    "WHERE s.patient_id = 'CXR001'\n"
    'ORDER BY f.confidence DESC;"',
    "bash",
)
body("Expected output:")
code_block(
    " finding_type     | confidence | severity | is_positive\n"
    "------------------+------------+----------+-------------\n"
    " consolidation    |       0.87 | moderate | t\n"
    " cardiomegaly     |       0.62 | moderate | t\n"
    " pleural_effusion |       0.34 | routine  | f\n"
    " pneumothorax     |       0.12 | routine  | f\n"
    " fracture         |       0.08 | routine  | f\n"
    "(5 rows)"
)

H3("Step 6 \u2014 Check GradCAM Heatmap in Orthanc")
code_block(
    'curl -s "http://localhost:8042/dicom-web/studies?PatientID=CXR001" '
    "| python3 -m json.tool",
    "bash",
)
body("Look for a Secondary Capture (SC) series alongside the original DX series.")

H3("Step 7 \u2014 Check DICOM SR")
code_block(
    'curl -s "http://localhost:8042/dicom-web/series?Modality=SR'
    '&PatientID=CXR001" | python3 -m json.tool',
    "bash",
)

H3("Step 8 \u2014 Check Worklist")
code_block(
    "docker exec imaging-postgres psql -U imaging -d imaging_agent -c \"\n"
    "SELECT * FROM active_worklist WHERE patient_id = 'CXR001';\n"
    '"',
    "bash",
)
body("If consolidation or cardiomegaly exceeded their thresholds, worklist "
     "entries appear with P3 (moderate) priority.")

H3("Verification Checklist")
bullet("CXR uploaded", "successfully via STOW-RS")
bullet("Webhook fired", "and pipeline started within 10 seconds")
bullet("5 finding rows", "in database (one per pathology)")
bullet("GradCAM heatmap", "generated as Secondary Capture")
bullet("DICOM SR", "created with coded findings")
bullet("Worklist entries", "created for positive findings")
bullet("Total pipeline time", "< 30 seconds")

H3("Troubleshooting")
add_table(
    ["Symptom", "Cause", "Fix"],
    [
        ["No pipeline execution", "Webhook didn't fire",
         "Wait full 10 seconds (StableAge); check orthanc logs"],
        ["Model inference fails", "DenseNet-121 not loaded",
         "Check model file exists in ./models/"],
        ["GPU OOM", "Insufficient memory",
         "Check nvidia-smi; reduce batch size in config"],
        ["GradCAM not generated", "Wrong model layer target",
         "Verify class_layers.relu exists in model"],
    ],
)

doc.add_page_break()

# ── Demo 6 ──
H2("Demo 6: CT Head Hemorrhage Triage")

body("Purpose: Detect intracranial hemorrhage, segment the bleed, measure "
     "volume and midline shift, classify urgency, and triage to the worklist.")
body("Prerequisites: All services running. DenseNet-121 and 3D U-Net "
     "models loaded.")
body("Performance target: < 90 seconds end-to-end. Sensitivity > 95% "
     "for hemorrhage > 5 mL.")

H3("Step 1 \u2014 Generate Synthetic CT Head Series")
body("Generate 64-slice CT head with simulated hemorrhage region:")
code_block(
    'python3 -c "\n'
    "import pydicom, numpy as np, os\n"
    "...\n"
    "os.makedirs('/tmp/ct_head_series', exist_ok=True)\n"
    "for i in range(64):  # 64 axial slices\n"
    "    ds.PatientID = 'HEAD001'\n"
    "    ds.Modality = 'CT'\n"
    "    ds.BodyPartExamined = 'HEAD'\n"
    "    ds.PixelSpacing = [0.5, 0.5]\n"
    "    ds.SliceThickness = 2.5\n"
    "    # Add bright blob to simulate hemorrhage (HU ~60-80)\n"
    "    if 20 <= i <= 35:\n"
    "        pixels[200:280, 180:260] = randint(50, 80, ...)\n"
    "    ds.save_as(f'/tmp/ct_head_series/ct_{i:03d}.dcm')\n"
    '"',
    "bash",
)

H3("Step 2 \u2014 Upload All Slices via STOW-RS")
code_block(
    "for f in /tmp/ct_head_series/*.dcm; do\n"
    "  curl -s -X POST http://localhost:8042/dicom-web/studies \\\n"
    '    -H "Content-Type: application/dicom" \\\n'
    '    --data-binary @"$f"\n'
    "done\n"
    'echo "All slices uploaded."',
    "bash",
)

H3("Step 3 \u2014 Wait for Pipeline (Watch Logs)")
code_block("docker logs -f imaging-dicom-listener", "bash")
body("Expected pipeline stages:")
code_block(
    "INFO  [ct_head_hemorrhage] Stage 1/4: DenseNet-121 classification\n"
    "INFO  [ct_head_hemorrhage]   hemorrhage_detected: True (confidence=0.94)\n"
    "INFO  [ct_head_hemorrhage] Stage 2/4: 3D U-Net segmentation\n"
    "INFO  [ct_head_hemorrhage]   Segmentation mask generated (64 slices)\n"
    "INFO  [ct_head_hemorrhage] Stage 3/4: Volumetric measurement\n"
    "INFO  [ct_head_hemorrhage]   hemorrhage_volume: 12.3 mL\n"
    "INFO  [ct_head_hemorrhage]   midline_shift: 3.2 mm (rightward)\n"
    "INFO  [ct_head_hemorrhage]   max_thickness: 8.7 mm\n"
    "INFO  [ct_head_hemorrhage] Stage 4/4: Urgency classification\n"
    "INFO  [ct_head_hemorrhage]   urgency: URGENT (volume 5-30 mL)\n"
    "INFO  [ct_head_hemorrhage]   priority: P2\n"
    "INFO  [ct_head_hemorrhage]   routing: Neurosurgery\n"
    "INFO  [ct_head_hemorrhage] Pipeline complete (47s)"
)

H3("Step 4 \u2014 Query Results")
code_block(
    "docker exec imaging-postgres psql -U imaging -d imaging_agent -c \"\n"
    "SELECT f.finding_type, f.severity, f.confidence,\n"
    "       m.measurement_type, m.value, m.unit, m.flag\n"
    "FROM findings f\n"
    "JOIN measurements m ON m.finding_id = f.id\n"
    "WHERE f.workflow = 'ct_head_hemorrhage'\n"
    'ORDER BY f.id, m.measurement_type;"',
    "bash",
)
body("Expected output:")
code_block(
    " finding_type | severity | confidence | measurement_type | value | unit | flag\n"
    "--------------+----------+------------+------------------+-------+------+----------\n"
    " hemorrhage   | urgent   |       0.94 | max_thickness    |   8.7 | mm   | normal\n"
    " hemorrhage   | urgent   |       0.94 | midline_shift    |   3.2 | mm   | normal\n"
    " hemorrhage   | urgent   |       0.94 | volume           |  12.3 | mL   | elevated"
)

body("Urgency classification thresholds (Brain Trauma Foundation):")
add_table(
    ["Metric", "Critical", "Urgent", "Routine"],
    [
        ["Volume", "> 30 mL", "5-30 mL", "< 5 mL"],
        ["Thickness", "> 10 mm", "\u2014", "<= 10 mm"],
        ["Midline shift", "> 5 mm", "\u2014", "<= 5 mm"],
    ],
)

H3("Step 5 \u2014 Check Worklist Entry")
code_block(
    "docker exec imaging-postgres psql -U imaging -d imaging_agent -c \"\n"
    "SELECT urgency, priority, notification, routing\n"
    "FROM active_worklist\n"
    "WHERE patient_id = 'HEAD001';\n"
    '"',
    "bash",
)
body("Expected output:")
code_block(
    " urgency | priority | notification                                | routing\n"
    "---------+----------+---------------------------------------------+--------------\n"
    " urgent  | P2       | Hemorrhage detected: 12.3 mL, shift 3.2 mm | Neurosurgery"
)

H3("Verification Checklist")
bullet("64 CT slices uploaded", "and study.complete webhook fired")
bullet("4-stage pipeline completed:",
       "detection -> segmentation -> measurement -> urgency")
bullet("Hemorrhage volume, midline shift, max thickness", "recorded")
bullet("Urgency correctly classified",
       "(CRITICAL >30mL, URGENT 5-30mL, ROUTINE <5mL)")
bullet("Worklist entry created", "with correct priority (P1/P2/P4)")
bullet("DICOM SR", "with TID 1500 Measurement Report created")
bullet("Total time", "< 90 seconds")

H3("Troubleshooting")
add_table(
    ["Symptom", "Cause", "Fix"],
    [
        ["Classification says no hemorrhage",
         "Synthetic data may not trigger",
         "Use real test DICOM or adjust thresholds"],
        ["Segmentation fails", "3D U-Net not loaded",
         "Verify unet3d_segmentation.pt in ./models/"],
        ["Volume is 0", "Empty segmentation mask",
         "Check preprocessing transforms (windowing, spacing)"],
        ["Wrong urgency", "Threshold mismatch",
         "Review urgency classification logic against BTF criteria"],
    ],
)

doc.add_page_break()

# ── Demo 7 ──
H2("Demo 7: CT Chest Lung Nodule Tracking")

body("Purpose: Detect lung nodules, measure volumes, match to prior study, "
     "calculate volume doubling time, assign Lung-RADS categories, and "
     "trigger genomics pipeline if warranted.")
body("Prerequisites: All services running. RetinaNet and SegResNet "
     "models loaded.")
body("Performance target: < 5 minutes multi-stage. Detection sensitivity "
     "> 90% for nodules >= 4 mm.")

H3("Step 1 \u2014 Upload Prior CT Chest (6 Months Earlier)")
body("Generate and upload a prior study for the same patient (128 slices, "
     "PatientID=CHEST001, StudyDate=20250801):")
code_block(
    'python3 -c "\n'
    "import pydicom, numpy as np, os\n"
    "os.makedirs('/tmp/ct_chest_prior', exist_ok=True)\n"
    "...\n"
    "for i in range(128):  # 128 slices, 1.5mm spacing\n"
    "    ds.PatientID = 'CHEST001'\n"
    "    ds.StudyDate = '20250801'  # 6 months prior\n"
    "    ds.Modality = 'CT'\n"
    "    ds.BodyPartExamined = 'CHEST'\n"
    "    # Nodule-like region at slices 60-68\n"
    "    if 60 <= i <= 68:\n"
    "        pixels[250:262, 300:312] = randint(-100, 50, (12,12))\n"
    "    ...\n"
    '"',
    "bash",
)
code_block(
    "for f in /tmp/ct_chest_prior/*.dcm; do\n"
    "  curl -s -X POST http://localhost:8042/dicom-web/studies \\\n"
    '    -H "Content-Type: application/dicom" --data-binary @"$f"\n'
    "done\n"
    'echo "Prior study uploaded."',
    "bash",
)
body("Wait 15 seconds for the prior study pipeline to complete.")

H3("Step 2 \u2014 Upload Current CT Chest")
body("Generate current study (same PatientID=CHEST001, StudyDate=20260201, "
     "slightly larger nodule region):")
code_block(
    'python3 -c "\n'
    "...\n"
    "    ds.PatientID = 'CHEST001'  # Same patient as prior\n"
    "    ds.StudyDate = '20260201'  # Current\n"
    "    # Nodule has grown (slightly larger region)\n"
    "    if 59 <= i <= 70:\n"
    "        pixels[248:266, 298:316] = randint(-100, 50, (18,18))\n"
    "    ...\n"
    '"',
    "bash",
)
code_block(
    "for f in /tmp/ct_chest_current/*.dcm; do\n"
    "  curl -s -X POST http://localhost:8042/dicom-web/studies \\\n"
    '    -H "Content-Type: application/dicom" --data-binary @"$f"\n'
    "done\n"
    'echo "Current study uploaded."',
    "bash",
)

H3("Step 3 \u2014 Monitor Multi-Stage Pipeline")
code_block("docker logs -f imaging-dicom-listener", "bash")
body("Expected pipeline stages:")
code_block(
    "INFO  [ct_chest_nodule] Stage 1/6: RetinaNet detection\n"
    "INFO  [ct_chest_nodule]   Candidates found: 2\n"
    "INFO  [ct_chest_nodule] Stage 2/6: SegResNet per-nodule segmentation\n"
    "INFO  [ct_chest_nodule]   Nodule 1: segmented (489 mm3)\n"
    "INFO  [ct_chest_nodule]   Nodule 2: segmented (112 mm3)\n"
    "INFO  [ct_chest_nodule] Stage 3/6: Volumetric measurement\n"
    "INFO  [ct_chest_nodule]   Nodule 1: 489 mm3, long_axis=9.8 mm\n"
    "INFO  [ct_chest_nodule]   Nodule 2: 112 mm3, long_axis=5.9 mm\n"
    "INFO  [ct_chest_nodule] Stage 4/6: Prior study matching\n"
    "INFO  [ct_chest_nodule]   Prior study found: 2025-08-01 (184 days ago)\n"
    "INFO  [ct_chest_nodule]   Registration: SyN diffeomorphic\n"
    "INFO  [ct_chest_nodule]   Nodule 1 matched to prior (prior volume: 295 mm3)\n"
    "INFO  [ct_chest_nodule] Stage 5/6: Volume doubling time\n"
    "INFO  [ct_chest_nodule]   VDT = (184 x ln2) / ln(489/295) = 248 days\n"
    "INFO  [ct_chest_nodule] Stage 6/6: Lung-RADS classification\n"
    "INFO  [ct_chest_nodule]   Nodule 1: solid, 9.8 mm -> Lung-RADS 4A\n"
    "INFO  [ct_chest_nodule]   VDT < 400 days -> upgrade to Lung-RADS 4B\n"
    "INFO  [ct_chest_nodule]   Nodule 2: solid, 5.9 mm -> Lung-RADS 2\n"
    "INFO  [ct_chest_nodule]   GENOMICS TRIGGER: Lung-RADS 4B detected\n"
    "INFO  [ct_chest_nodule] Pipeline complete (3m 12s)"
)

H3("Step 4 \u2014 Query Nodule Results")
code_block(
    "docker exec imaging-postgres psql -U imaging -d imaging_agent -c \"\n"
    "SELECT f.finding_type, f.details->>'lung_rads' AS lung_rads,\n"
    "       f.confidence, f.severity,\n"
    "       m.measurement_type, m.value, m.unit,\n"
    "       m.prior_value, m.delta_percent\n"
    "FROM findings f\n"
    "JOIN measurements m ON m.finding_id = f.id\n"
    "JOIN studies s ON f.study_id = s.id\n"
    "WHERE s.patient_id = 'CHEST001'\n"
    "  AND s.study_date = '2026-02-01'\n"
    "  AND f.workflow = 'ct_chest_nodule'\n"
    'ORDER BY f.id, m.measurement_type;"',
    "bash",
)
body("Expected output:")
code_block(
    " finding_type | lung_rads | confidence | severity | measurement_type | value  | unit | prior | delta%\n"
    "--------------+-----------+------------+----------+------------------+--------+------+-------+-------\n"
    " nodule       | 4B        |       0.91 | urgent   | doubling_time    |    248 | days |       |\n"
    " nodule       | 4B        |       0.91 | urgent   | volume           |    489 | mm3  |   295 |  65.8\n"
    " nodule       | 2         |       0.78 | routine  | volume           |    112 | mm3  |       |"
)

H3("Step 5 \u2014 Verify Genomics Trigger")
code_block(
    'docker logs imaging-dicom-listener | grep "GENOMICS TRIGGER"',
    "bash",
)
body("Expected output:")
code_block(
    "INFO  GENOMICS TRIGGER: Lung-RADS 4B detected for study_id=<id>, "
    "triggering Parabricks pipeline"
)
body("VDT formula used: VDT = (delta_t x ln2) / ln(V2/V1) = "
     "(184 x 0.693) / ln(489/295) = 248 days")

body("Lung-RADS classification (solid nodules):")
add_table(
    ["Size", "Category", "Growth Upgrade (VDT < 400 days)"],
    [
        ["< 4 mm", "1", "\u2014"],
        ["4-6 mm", "2", "\u2192 4A"],
        ["6-8 mm", "3", "\u2192 4A"],
        ["8-15 mm", "4A", "\u2192 4B"],
        [">= 15 mm", "4B", "\u2014"],
    ],
)

H3("Verification Checklist")
bullet("Prior and current studies uploaded", "for same patient (CHEST001)")
bullet("RetinaNet", "detected nodule candidates")
bullet("SegResNet", "segmented each nodule")
bullet("Prior study retrieved", "and spatial registration performed")
bullet("Volume doubling time", "calculated correctly")
bullet("Lung-RADS assigned", "with growth upgrade where applicable")
bullet("Genomics pipeline triggered", "for Lung-RADS 4B+")
bullet("Total time", "< 5 minutes")

doc.add_page_break()

# ── Demo 8 ──
H2("Demo 8: MRI Brain MS Lesion Tracking")

body("Purpose: Segment MS lesions on FLAIR, count and measure them, register "
     "to prior MRI, identify new/enlarging lesions, and classify disease activity.")
body("Prerequisites: All services running. 3D U-Net (FLAIR) model loaded.")
body("Performance target: < 5 minutes multi-stage.")

H3("Step 1 \u2014 Upload Prior MRI Brain FLAIR (12 Months Earlier)")
body("Generate 96-slice prior MRI FLAIR (PatientID=MS001, StudyDate=20250201) "
     "with 3 lesion spots, then upload:")
code_block(
    'python3 -c "\n'
    "import pydicom, numpy as np, os\n"
    "os.makedirs('/tmp/mri_brain_prior', exist_ok=True)\n"
    "...\n"
    "for i in range(96):  # 96 slices\n"
    "    ds.PatientID = 'MS001'\n"
    "    ds.Modality = 'MR'\n"
    "    ds.StudyDate = '20250201'  # 12 months prior\n"
    "    ds.SeriesDescription = 'FLAIR'\n"
    "    # Add bright FLAIR lesion spots\n"
    "    ...\n"
    '"',
    "bash",
)
code_block(
    "for f in /tmp/mri_brain_prior/*.dcm; do\n"
    "  curl -s -X POST http://localhost:8042/dicom-web/studies \\\n"
    '    -H "Content-Type: application/dicom" --data-binary @"$f"\n'
    "done\n"
    'echo "Prior MRI uploaded."',
    "bash",
)
body("Wait 20 seconds for the prior study pipeline to complete.")

H3("Step 2 \u2014 Upload Current MRI Brain FLAIR")
body("Generate current study (same patient, 3 original lesions + 2 new lesions, "
     "some slightly larger):")
code_block(
    'python3 -c "\n'
    "...\n"
    "    ds.PatientID = 'MS001'  # Same patient\n"
    "    ds.StudyDate = '20260201'  # Current\n"
    "    # Same lesions as prior + 2 new ones, some slightly larger\n"
    "    ...\n"
    '"',
    "bash",
)
code_block(
    "for f in /tmp/mri_brain_current/*.dcm; do\n"
    "  curl -s -X POST http://localhost:8042/dicom-web/studies \\\n"
    '    -H "Content-Type: application/dicom" --data-binary @"$f"\n'
    "done\n"
    'echo "Current MRI uploaded."',
    "bash",
)

H3("Step 3 \u2014 Monitor Pipeline")
body("Expected pipeline stages:")
code_block(
    "INFO  [mri_ms_lesion] Stage 1/5: 3D U-Net FLAIR segmentation\n"
    "INFO  [mri_ms_lesion]   Segmentation mask generated (96 slices)\n"
    "INFO  [mri_ms_lesion] Stage 2/5: Connected component analysis\n"
    "INFO  [mri_ms_lesion]   Total lesions found: 14\n"
    "INFO  [mri_ms_lesion]   Total lesion volume: 6.1 mL\n"
    "INFO  [mri_ms_lesion] Stage 3/5: Prior study registration\n"
    "INFO  [mri_ms_lesion]   Prior MRI found: 2025-02-01 (365 days ago)\n"
    "INFO  [mri_ms_lesion]   ANTsPy SyN registration complete\n"
    "INFO  [mri_ms_lesion] Stage 4/5: Change detection\n"
    "INFO  [mri_ms_lesion]   Stable lesions: 11\n"
    "INFO  [mri_ms_lesion]   Enlarging lesions: 1\n"
    "INFO  [mri_ms_lesion]   New lesions: 2\n"
    "INFO  [mri_ms_lesion] Stage 5/5: Disease activity classification\n"
    "INFO  [mri_ms_lesion]   Activity: ACTIVE (1-2 new lesions)\n"
    "INFO  [mri_ms_lesion] Pipeline complete (2m 48s)"
)

H3("Step 4 \u2014 Query Results")
code_block(
    "docker exec imaging-postgres psql -U imaging -d imaging_agent -c \"\n"
    "SELECT f.finding_type, f.severity,\n"
    "       f.details->>'disease_activity' AS activity,\n"
    "       f.details->>'new_lesion_count' AS new_lesions,\n"
    "       f.details->>'enlarging_lesion_count' AS enlarging,\n"
    "       m.measurement_type, m.value, m.unit,\n"
    "       m.prior_value, m.delta_percent\n"
    "FROM findings f\n"
    "JOIN measurements m ON m.finding_id = f.id\n"
    "JOIN studies s ON f.study_id = s.id\n"
    "WHERE s.patient_id = 'MS001'\n"
    "  AND s.study_date = '2026-02-01'\n"
    'ORDER BY m.measurement_type;"',
    "bash",
)
body("Expected output:")
code_block(
    " finding_type | severity | activity | new | enlarging | measurement  | value | unit  | prior | delta%\n"
    "--------------+----------+----------+-----+-----------+--------------+-------+-------+-------+-------\n"
    " ms_lesion    | moderate | active   | 2   | 1         | lesion_count |    14 | count |    12 |  16.7\n"
    " ms_lesion    | moderate | active   | 2   | 1         | volume       |   6.1 | mL    |   4.8 |  27.1"
)

body("Disease activity classification:")
add_table(
    ["New/Enlarging Lesions", "Activity"],
    [
        ["0 new, 0 enlarging", "Stable"],
        ["1-2 new or enlarging", "Active"],
        ["3+ new or enlarging", "Highly Active"],
    ],
)

H3("Verification Checklist")
bullet("Prior and current MRI uploaded", "for same patient (MS001)")
bullet("3D U-Net segmentation", "completed on FLAIR")
bullet("Lesion count and total volume", "measured")
bullet("Prior study retrieved", "and spatial registration performed")
bullet("New and enlarging lesions", "identified")
bullet("Disease activity", "classified correctly")
bullet("Total time", "< 5 minutes")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PART 3: AGENT REASONING DEMOS
# ══════════════════════════════════════════════════════════════
H1("Part 3: Agent Reasoning Demos")

# ── Demo 9 ──
H2("Demo 9: Triage Agent")

body("Purpose: Show the LangGraph triage agent node classifying urgency, "
     "routing to specialists, and generating notification messages.")
body("Prerequisites: At least one clinical workflow completed (Demo 5, 6, "
     "7, or 8). Agent API running.")

H3("Step 1 \u2014 Trigger Triage on a Study with Findings")
code_block(
    'curl -s -X POST http://localhost:8524/api/triage \\\n'
    '  -H "Content-Type: application/json" \\\n'
    "  -d '{\"study_id\": 1}' | python3 -m json.tool",
    "bash",
)
body("Expected output:")
code_block(
    '{\n'
    '    "study_id": 1,\n'
    '    "urgency": "urgent",\n'
    '    "priority": "P2",\n'
    '    "routing": "Neurosurgery",\n'
    '    "notification": "URGENT: Intracranial hemorrhage detected. Volume 12.3 mL,\n'
    '        midline shift 3.2 mm. Recommend immediate neurosurgical consultation.",\n'
    '    "agent_trace": {\n'
    '        "nodes_visited": ["triage_node"],\n'
    '        "state_transitions": 1,\n'
    '        "execution_time_ms": 340\n'
    '    }\n'
    '}',
    "json",
)

H3("Step 2 \u2014 Verify Worklist Update")
code_block(
    "docker exec imaging-postgres psql -U imaging -d imaging_agent -c \"\n"
    'SELECT * FROM active_worklist ORDER BY priority, created_at;"',
    "bash",
)

spacer()

# ── Demo 10 ──
H2("Demo 10: Longitudinal Comparison Agent")

body("Purpose: Demonstrate delta computation across time points for patients "
     "with multiple studies.")
body("Prerequisites: CT Chest demo (Demo 7) completed with prior and "
     "current studies.")

H3("Step 1 \u2014 Query Longitudinal Comparison")
code_block(
    'curl -s -X POST http://localhost:8524/api/longitudinal \\\n'
    '  -H "Content-Type: application/json" \\\n'
    "  -d '{\"patient_id\": \"CHEST001\"}' | python3 -m json.tool",
    "bash",
)
body("Expected output:")
code_block(
    '{\n'
    '    "patient_id": "CHEST001",\n'
    '    "studies_compared": 2,\n'
    '    "time_delta_days": 184,\n'
    '    "findings": [\n'
    '        {\n'
    '            "finding_type": "nodule",\n'
    '            "current_volume_mm3": 489,\n'
    '            "prior_volume_mm3": 295,\n'
    '            "delta_percent": 65.8,\n'
    '            "volume_doubling_time_days": 248,\n'
    '            "trend": "growing",\n'
    '            "lung_rads_current": "4B",\n'
    '            "lung_rads_prior": "3"\n'
    '        }\n'
    '    ],\n'
    '    "agent_trace": {\n'
    '        "nodes_visited": ["longitudinal_node"],\n'
    '        "execution_time_ms": 520\n'
    '    }\n'
    '}',
    "json",
)

spacer()

# ── Demo 11 ──
H2("Demo 11: Population Analysis Agent")

body("Purpose: Demonstrate embedding-based similarity search and cohort "
     "retrieval using pgvector.")
body("Prerequisites: Multiple studies uploaded and embedded. Embedding "
     "service running.")

H3("Step 1 \u2014 Search for Similar Studies")
code_block(
    'curl -s -X POST http://localhost:8524/api/population \\\n'
    '  -H "Content-Type: application/json" \\\n'
    "  -d '{\"study_id\": 1, \"top_k\": 10}' | python3 -m json.tool",
    "bash",
)
body("Expected output:")
code_block(
    '{\n'
    '    "query_study_id": 1,\n'
    '    "similar_studies": [\n'
    '        {\n'
    '            "study_id": 3,\n'
    '            "patient_id": "CHEST002",\n'
    '            "modality": "CT",\n'
    '            "cosine_distance": 0.12,\n'
    '            "finding_summary": "2 nodules, Lung-RADS 3"\n'
    '        },\n'
    '        {\n'
    '            "study_id": 5,\n'
    '            "patient_id": "CHEST003",\n'
    '            "cosine_distance": 0.18,\n'
    '            "finding_summary": "1 nodule, Lung-RADS 4A"\n'
    '        }\n'
    '    ]\n'
    '}',
    "json",
)

H3("Step 2 \u2014 Direct pgvector Query")
code_block(
    "docker exec imaging-postgres psql -U imaging -d imaging_agent -c \"\n"
    "SELECT s.study_instance_uid, s.patient_id, s.study_date,\n"
    "       e.embedding <=> (\n"
    "           SELECT embedding FROM embeddings\n"
    "           WHERE study_id = 1 AND level = 'study'\n"
    "       ) AS distance\n"
    "FROM embeddings e\n"
    "JOIN studies s ON e.study_id = s.id\n"
    "WHERE e.level = 'study' AND e.study_id != 1\n"
    "ORDER BY distance\n"
    'LIMIT 10;"',
    "bash",
)

spacer()

# ── Demo 12 ──
H2("Demo 12: RAG-Grounded Report Generation")

body("Purpose: Run the full agent pipeline (triage -> longitudinal -> "
     "report) with NIM LLM evidence-grounded narrative generation.")
body("Prerequisites: All workflows completed. NIM LLM running.")

H3("Step 1 \u2014 Generate Full Report")
code_block(
    'curl -s -X POST http://localhost:8524/api/report \\\n'
    '  -H "Content-Type: application/json" \\\n'
    "  -d '{\"study_id\": 1}' | python3 -m json.tool",
    "bash",
)
body("Expected output (abbreviated):")
code_block(
    '{\n'
    '    "study_id": 1,\n'
    '    "narrative": "FINDINGS: A 12.3 mL acute right frontoparietal\n'
    '        hemorrhage is identified with 3.2 mm rightward midline shift\n'
    '        and maximum thickness of 8.7 mm. Per Brain Trauma Foundation\n'
    '        guidelines, this classifies as URGENT...",\n'
    '    "evidence_sources": [\n'
    '        "findings_db: hemorrhage volume 12.3 mL, shift 3.2 mm",\n'
    '        "guideline: Brain Trauma Foundation surgical criteria",\n'
    '        "cohort: 8 similar cases (cosine distance < 0.2)"\n'
    '    ],\n'
    '    "fhir_diagnostic_report_id": "DR-2026-0001",\n'
    '    "agent_trace": {\n'
    '        "nodes_visited": ["triage_node", "longitudinal_node",\n'
    '            "population_node", "report_node"],\n'
    '        "total_execution_time_ms": 2840\n'
    '    }\n'
    '}',
    "json",
)

spacer()

# ── Demo 13 ──
H2("Demo 13: Conditional Routing")

body("Purpose: Show how the LangGraph StateGraph routes differently based "
     "on finding severity.")
body("Prerequisites: Studies with CRITICAL and ROUTINE findings in the database.")

H3("Step 1 \u2014 CRITICAL Finding Route")
code_block(
    'curl -s -X POST http://localhost:8524/api/analyze \\\n'
    '  -H "Content-Type: application/json" \\\n'
    '  -d \'{"study_id": 1}\' | python3 -c "\n'
    "import json, sys\n"
    "data = json.load(sys.stdin)\n"
    "print('Path:', ' -> '.join(data['agent_trace']['nodes_visited']))\n"
    '"',
    "bash",
)
body("Expected output:")
code_block(
    "Path: triage_node -> longitudinal_node -> population_node -> report_node"
)

H3("Step 2 \u2014 ROUTINE Finding Route")
code_block(
    'curl -s -X POST http://localhost:8524/api/analyze \\\n'
    '  -H "Content-Type: application/json" \\\n'
    '  -d \'{"study_id": 2}\' | python3 -c "\n'
    "import json, sys\n"
    "data = json.load(sys.stdin)\n"
    "print('Path:', ' -> '.join(data['agent_trace']['nodes_visited']))\n"
    '"',
    "bash",
)
body("Expected output:")
code_block("Path: triage_node -> report_node")

body("Critical/urgent findings route through longitudinal and population "
     "analysis. Routine findings skip directly to report generation.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PART 4: OUTPUT & INTEGRATION DEMOS
# ══════════════════════════════════════════════════════════════
H1("Part 4: Output & Integration Demos")

# ── Demo 14 ──
H2("Demo 14: DICOM SR Structured Report")

body("Purpose: Inspect the TID 1500 Measurement Report content generated "
     "by the agent.")
body("Prerequisites: At least one clinical workflow completed.")

H3("Step 1 \u2014 Find SR in Orthanc")
code_block(
    'curl -s "http://localhost:8042/dicom-web/series?Modality=SR" '
    "| python3 -m json.tool",
    "bash",
)

H3("Step 2 \u2014 Download and Parse SR")
body("Use highdicom and dicomweb-client to retrieve and parse:")
code_block(
    "from dicomweb_client import DICOMwebClient\n"
    "import highdicom as hd\n"
    "\n"
    "client = DICOMwebClient(url='http://localhost:8042/dicom-web')\n"
    "series_list = client.search_for_series(\n"
    "    search_filters={'Modality': 'SR'})\n"
    "# Retrieve and print content tree...",
    "python",
)
body("Expected output:")
code_block(
    "SR Type: 1.2.840.10008.5.1.4.1.1.88.34\n"
    "Content Date: 20260201\n"
    "  Finding: (nested)\n"
    "  Volume: 12.3 mL\n"
    "  Midline Shift: 3.2 mm\n"
    "  Maximum Thickness: 8.7 mm\n"
    "  Urgency: urgent"
)

spacer()

# ── Demo 15 ──
H2("Demo 15: GSPS Graphic Overlay")

body("Purpose: View GradCAM heatmap overlays on original DICOM images.")
body("Prerequisites: CXR Rapid Findings demo (Demo 5) completed.")

H3("Step 1 \u2014 Open Orthanc Explorer")
body("Open in browser: http://localhost:8042/app/explorer.html")

H3("Step 2 \u2014 Navigate to CXR Study")
body("Find the CXR001 patient study. You should see:")
bullet("Original DX series", "(the chest X-ray)")
bullet("Secondary Capture series", "(GradCAM heatmaps)")
bullet("GSPS series", "(graphic annotation overlays)")

H3("Step 3 \u2014 View Overlay")
body("Click on the GSPS series to see graphic annotation contours overlaid "
     "on the original CXR. The highlighted regions correspond to GradCAM "
     "activation areas for detected pathologies.")

spacer()

# ── Demo 16 ──
H2("Demo 16: FHIR DiagnosticReport")

body("Purpose: Inspect the FHIR R4 DiagnosticReport output with coded "
     "clinical data.")
body("Prerequisites: FHIR publisher running. At least one workflow completed.")

H3("Step 1 \u2014 Retrieve Latest Report")
code_block(
    "curl -s http://localhost:8523/api/reports/latest | python3 -m json.tool",
    "bash",
)
body("Expected output (abbreviated):")
code_block(
    '{\n'
    '  "resourceType": "Bundle",\n'
    '  "type": "transaction",\n'
    '  "entry": [\n'
    '    {\n'
    '      "resource": {\n'
    '        "resourceType": "DiagnosticReport",\n'
    '        "status": "final",\n'
    '        "code": {\n'
    '          "coding": [{\n'
    '            "system": "http://loinc.org",\n'
    '            "code": "18748-4",\n'
    '            "display": "Diagnostic imaging study"\n'
    '          }]\n'
    '        },\n'
    '        "conclusion": "URGENT: Intracranial hemorrhage detected. Volume 12.3 mL.",\n'
    '        "result": [\n'
    '          {"reference": "Observation/hemorrhage-volume"},\n'
    '          {"reference": "Observation/midline-shift"}\n'
    '        ]\n'
    '      }\n'
    '    },\n'
    '    {\n'
    '      "resource": {\n'
    '        "resourceType": "Observation",\n'
    '        "id": "hemorrhage-volume",\n'
    '        "code": {"coding": [{"system": "http://snomed.info/sct",\n'
    '          "code": "276651009", "display": "Volume of hemorrhage"}]},\n'
    '        "valueQuantity": {"value": 12.3, "unit": "mL"}\n'
    '      }\n'
    '    }\n'
    '  ]\n'
    '}',
    "json",
)

spacer()

# ── Demo 17 ──
H2("Demo 17: Embedding & Vector Search")

body("Purpose: Directly demonstrate pgvector embedding generation and "
     "similarity queries.")
body("Prerequisites: Multiple studies uploaded. Embedding service running.")

H3("Step 1 \u2014 Generate Embedding for a Study")
code_block(
    'curl -s -X POST http://localhost:8521/api/embed \\\n'
    '  -H "Content-Type: application/json" \\\n'
    "  -d '{\"study_id\": 1}' | python3 -m json.tool",
    "bash",
)
body("Expected output:")
code_block(
    '{\n'
    '    "study_id": 1,\n'
    '    "level": "study",\n'
    '    "embedding_dim": 384,\n'
    '    "model": "microsoft/BiomedCLIP-PubMedBERT_256-vit_base_patch16_224",\n'
    '    "stored": true\n'
    '}',
    "json",
)

H3("Step 2 \u2014 Query Similar Studies with EXPLAIN ANALYZE")
code_block(
    "docker exec imaging-postgres psql -U imaging -d imaging_agent -c \"\n"
    "EXPLAIN ANALYZE\n"
    "SELECT s.study_instance_uid, s.patient_id,\n"
    "       e.embedding <=> (SELECT embedding FROM embeddings\n"
    "       WHERE study_id = 1 AND level = 'study') AS distance\n"
    "FROM embeddings e\n"
    "JOIN studies s ON e.study_id = s.id\n"
    "WHERE e.level = 'study' AND e.study_id != 1\n"
    "ORDER BY distance\n"
    'LIMIT 10;"',
    "bash",
)
body("Expected output (key line):")
code_block(
    "Index Scan using idx_embeddings_hnsw on embeddings e  (cost=... rows=10)\n"
    "  ...\n"
    "Planning Time: 0.5 ms\n"
    "Execution Time: 2.3 ms"
)
body("The HNSW index provides sub-5ms query times even with thousands "
     "of embeddings.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PART 5: PORTAL & MONITORING DEMOS
# ══════════════════════════════════════════════════════════════
H1("Part 5: Portal & Monitoring Demos")

# ── Demo 18 ──
H2("Demo 18: Streamlit Portal \u2014 Worklist")

body("Purpose: Interactive worklist management through the web portal.")
body("Prerequisites: Portal running (port 8525). Workflows completed with "
     "findings.")

H3("Step 1 \u2014 Open Portal")
body("Open in browser: http://localhost:8525")

H3("Step 2 \u2014 View Active Worklist")
body("The default view shows the active worklist sorted by priority:")
bullet("Red rows:", "P1 (critical) \u2014 immediate attention required")
bullet("Orange rows:", "P2 (urgent) \u2014 requires prompt review")
bullet("Yellow rows:", "P3 (moderate) \u2014 schedule follow-up")
bullet("Green rows:", "P4 (routine) \u2014 standard queue")
body("Each row shows: patient ID, modality, finding type, confidence, "
     "urgency, routing target.")

H3("Step 3 \u2014 Drill Into a Finding")
body("Click any worklist entry to see:")
bullet("Full finding details", "(type, location, confidence, severity)")
bullet("Associated measurements", "(volume, shift, etc.)")
bullet("Link to view images", "in Orthanc")

H3("Step 4 \u2014 Acknowledge a Finding")
body("Click the \"Acknowledge\" button on a worklist entry to mark it "
     "as reviewed.")

spacer()

# ── Demo 19 ──
H2("Demo 19: Streamlit Portal \u2014 Study Browser")

body("Purpose: Browse all processed studies with finding summaries.")

H3("Step 1 \u2014 Navigate to Studies Tab")
body("Click \"Studies\" in the sidebar.")

H3("Step 2 \u2014 Browse Studies")
body("Shows all studies with:")
bullet("Patient ID, modality, study date", "")
bullet("Finding count", "(total, critical, urgent)")
bullet("Processing status", "(received / processing / completed / failed)")

H3("Step 3 \u2014 Click Into Study Detail")
body("Shows all findings and measurements for that study, plus links to "
     "DICOM SR and FHIR reports.")

spacer()

# ── Demo 20 ──
H2("Demo 20: Streamlit Portal \u2014 Agent Activity")

body("Purpose: Monitor pipeline execution and model provenance.")

H3("Step 1 \u2014 Navigate to Activity Tab")
body("Click \"Activity\" in the sidebar.")

H3("Step 2 \u2014 View Pipeline Runs")
body("Shows recent provenance records:")
bullet("Workflow name", "")
bullet("Model ID and version", "")
bullet("Processing duration (ms)", "")
bullet("GPU memory usage (MB)", "")
bullet("Status", "(completed / failed)")

spacer()

# ── Demo 21 ──
H2("Demo 21: Grafana Monitoring Dashboard")

body("Purpose: GPU and pipeline performance visualization.")
body("Prerequisites: Grafana running (port 3000). DCGM exporter running.")

H3("Step 1 \u2014 Open Grafana")
body("Open in browser: http://localhost:3000")
body("Login: admin / changeme")

H3("Step 2 \u2014 Navigate to Imaging Agent Dashboard")
body("Click Dashboards -> Imaging Agent.")
body("Dashboard panels include:")
bullet("GPU Utilization %", "\u2014 Real-time GB10 utilization from DCGM exporter")
bullet("GPU Memory Usage", "\u2014 Used vs. total (128 GB unified)")
bullet("GPU Temperature", "\u2014 Thermal monitoring")
bullet("Inference Latency", "\u2014 Histogram per workflow (p50, p95, p99)")
bullet("Pipeline Throughput", "\u2014 Studies processed per hour")
bullet("Queue Depth", "\u2014 Pending studies awaiting processing")
bullet("Error Rate", "\u2014 Failed pipeline executions")

spacer()

# ── Demo 22 ──
H2("Demo 22: Alerting Demo")

body("Purpose: Show Prometheus alert pipeline for operational monitoring.")

H3("Step 1 \u2014 Check Current Alerts")
code_block(
    "curl -s http://localhost:9099/api/v1/alerts | python3 -m json.tool",
    "bash",
)

H3("Step 2 \u2014 View Alert Rules")
code_block(
    "curl -s http://localhost:9099/api/v1/rules | python3 -m json.tool",
    "bash",
)

body("Example alert rules:")
add_table(
    ["Alert", "Condition", "Severity"],
    [
        ["GPUMemoryHigh", "GPU memory > 90% for 5 min", "warning"],
        ["InferenceFailureRate", "Failure rate > 5% for 10 min", "critical"],
        ["QueueBacklog", "Pending studies > 50 for 15 min", "warning"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PART 6: ADVANCED DEMOS
# ══════════════════════════════════════════════════════════════
H1("Part 6: Advanced Demos")

# ── Demo 23 ──
H2("Demo 23: End-to-End Patient Case")

body("Purpose: Upload all 4 modalities for a single patient and show "
     "unified cross-modal findings.")
body("Prerequisites: All services running. All models loaded.")

H3("Step 1 \u2014 Upload All 4 Modalities")
body("Use the same patient ID (MULTI001) for all uploads:")
bullet("1. CXR", "(from Demo 5 pattern, change PatientID='MULTI001')")
bullet("2. CT Head", "(from Demo 6 pattern, change PatientID='MULTI001')")
bullet("3. CT Chest with prior",
       "(from Demo 7 pattern, change PatientID='MULTI001')")
bullet("4. MRI Brain with prior",
       "(from Demo 8 pattern, change PatientID='MULTI001')")

H3("Step 2 \u2014 Wait for All Pipelines")
body("Monitor docker logs -f imaging-dicom-listener until all 4 workflows "
     "complete.")

H3("Step 3 \u2014 Query Unified Findings")
code_block(
    "docker exec imaging-postgres psql -U imaging -d imaging_agent -c \"\n"
    "SELECT s.modality, f.workflow, f.finding_type, f.severity,\n"
    "       f.confidence, m.measurement_type, m.value, m.unit\n"
    "FROM findings f\n"
    "JOIN studies s ON f.study_id = s.id\n"
    "LEFT JOIN measurements m ON m.finding_id = f.id\n"
    "WHERE s.patient_id = 'MULTI001'\n"
    "ORDER BY\n"
    "    CASE f.severity\n"
    "        WHEN 'critical' THEN 1 WHEN 'urgent' THEN 2\n"
    "        WHEN 'moderate' THEN 3 ELSE 4\n"
    "    END,\n"
    '    s.modality;"',
    "bash",
)

H3("Step 4 \u2014 Generate Comprehensive Report")
code_block(
    'curl -s -X POST http://localhost:8524/api/patient-report \\\n'
    '  -H "Content-Type: application/json" \\\n'
    "  -d '{\"patient_id\": \"MULTI001\"}' | python3 -m json.tool",
    "bash",
)

spacer()

# ── Demo 24 ──
H2("Demo 24: Cross-Modal Genomics Trigger")

body("Purpose: Demonstrate the imaging -> genomics pipeline handoff when "
     "Lung-RADS 4B+ is detected.")
body("Prerequisites: CT Chest demo (Demo 7) completed with Lung-RADS "
     "4B result.")

H3("Step 1 \u2014 Verify Trigger in Logs")
code_block(
    'docker logs imaging-dicom-listener | grep "GENOMICS TRIGGER"',
    "bash",
)
body("Expected output:")
code_block(
    "INFO  GENOMICS TRIGGER: Lung-RADS 4B detected for study_id=<id>\n"
    "INFO  Parabricks pipeline trigger payload:\n"
    'INFO    {"study_id": <id>, "patient_id": "CHEST001",\n'
    'INFO     "reason": "Lung-RADS 4B", "workflow": "somatic_germline",\n'
    'INFO     "priority": "high"}'
)

H3("Step 2 \u2014 Inspect Trigger Payload")
code_block(
    "curl -s http://localhost:8524/api/genomics-triggers | python3 -m json.tool",
    "bash",
)
body("This shows the queue of genomics pipeline triggers awaiting "
     "Parabricks execution.")

spacer()

# ── Demo 25 ──
H2("Demo 25: Nextflow Pipeline Orchestration")

body("Purpose: Run a workflow through Nextflow and inspect the execution DAG.")
body("Prerequisites: Nextflow installed. All services running.")

H3("Step 1 \u2014 Run Nextflow Pipeline")
code_block(
    "nextflow run main.nf \\\n"
    "  -profile docker \\\n"
    "  --study_id 1 \\\n"
    "  --workflow ct_head_hemorrhage \\\n"
    "  -with-trace \\\n"
    "  -with-report",
    "bash",
)
body("Expected output:")
code_block(
    "N E X T F L O W  ~  version 23.10.0\n"
    "Launching `main.nf` [friendly_name] DSL2 - revision: abc123\n"
    "\n"
    "executor >  local (4)\n"
    "[ab/123456] process > preprocess    [100%] 1 of 1 \u2714\n"
    "[cd/789012] process > inference     [100%] 1 of 1 \u2714\n"
    "[ef/345678] process > postprocess   [100%] 1 of 1 \u2714\n"
    "[gh/901234] process > persist       [100%] 1 of 1 \u2714\n"
    "\n"
    "Completed at: 2026-02-01T12:00:47\n"
    "Duration    : 47s\n"
    "CPU hours   : 0.01\n"
    "Succeeded   : 4"
)

H3("Step 2 \u2014 View Execution Trace")
code_block("cat trace.txt", "bash")
body("Shows per-process timing, CPU, memory, and status.")

H3("Step 3 \u2014 Generate DAG Visualization")
code_block(
    "nextflow run main.nf -profile docker \\\n"
    "  --workflow ct_head_hemorrhage -with-dag dag.html",
    "bash",
)
body("Open dag.html in a browser to see the pipeline directed acyclic graph.")

spacer()

# ── Demo 26 ──
H2("Demo 26: Performance Benchmarking")

body("Purpose: Measure and compare all workflow latencies against targets.")

H3("Step 1 \u2014 Benchmark Each Workflow")
code_block(
    'echo "=== Performance Benchmark ==="\n'
    "echo \"\"\n"
    'echo "CXR Rapid Findings:"\n'
    "time curl -s -X POST http://localhost:8524/api/run-workflow \\\n"
    "  -d '{\"study_id\": <cxr_study_id>, \"workflow\": \"cxr_findings\"}' "
    "> /dev/null\n"
    'echo "CT Head Hemorrhage:"\n'
    "time curl -s -X POST http://localhost:8524/api/run-workflow \\\n"
    "  -d '{\"study_id\": <ct_head_id>, \"workflow\": \"ct_head_hemorrhage\"}' "
    "> /dev/null\n"
    'echo "CT Chest Lung Nodule:"\n'
    "time curl -s -X POST http://localhost:8524/api/run-workflow \\\n"
    "  -d '{\"study_id\": <ct_chest_id>, \"workflow\": \"ct_chest_nodule\"}' "
    "> /dev/null\n"
    'echo "MRI Brain MS Lesion:"\n'
    "time curl -s -X POST http://localhost:8524/api/run-workflow \\\n"
    "  -d '{\"study_id\": <mri_id>, \"workflow\": \"mri_ms_lesion\"}' "
    "> /dev/null",
    "bash",
)

H3("Step 2 \u2014 Compare Against Targets")
add_table(
    ["Workflow", "Target", "Expected Actual"],
    [
        ["CXR Rapid Findings", "< 30 seconds", "~8-12 seconds"],
        ["CT Head Hemorrhage", "< 90 seconds", "~45-60 seconds"],
        ["CT Chest Lung Nodule", "< 5 minutes", "~3-4 minutes"],
        ["MRI Brain MS Lesion", "< 5 minutes", "~2.5-3.5 minutes"],
    ],
)

spacer()

# ── Demo 27 ──
H2("Demo 27: Provenance & Reproducibility")

body("Purpose: Show the complete audit trail for a finding and demonstrate "
     "deterministic re-execution.")

H3("Step 1 \u2014 Query Provenance")
code_block(
    "docker exec imaging-postgres psql -U imaging -d imaging_agent -c \"\n"
    "SELECT workflow, model_id, model_version, model_arch,\n"
    "       inference_params, duration_ms, gpu_memory_mb, status,\n"
    "       created_at\n"
    "FROM provenance\n"
    "WHERE study_id = 1\n"
    'ORDER BY created_at;"',
    "bash",
)
body("Expected output:")
code_block(
    " workflow             | model_id               | model_version | model_arch   | inference_params               | duration_ms | gpu_memory_mb | status\n"
    "----------------------+------------------------+---------------+--------------+--------------------------------+-------------+---------------+-----------\n"
    " ct_head_hemorrhage   | hemorrhage-detect-v2.1 | 2.1.0         | DenseNet-121 | {\"precision\":\"fp16\",\"seed\":42}  |        8200 |          3400 | completed\n"
    " ct_head_hemorrhage   | hemorrhage-seg-v1.3    | 1.3.0         | 3D U-Net     | {\"precision\":\"fp16\",\"seed\":42}  |       28400 |         12800 | completed"
)

H3("Step 2 \u2014 Deterministic Re-Run")
code_block(
    'curl -s -X POST http://localhost:8524/api/reprocess \\\n'
    '  -H "Content-Type: application/json" \\\n'
    "  -d '{\"study_id\": 1, \"workflow\": \"ct_head_hemorrhage\"}' "
    "| python3 -m json.tool",
    "bash",
)
body("With the same seed (42) and precision (fp16), outputs should be "
     "identical to the original run.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PART 7: TEARDOWN & CLEANUP
# ══════════════════════════════════════════════════════════════
H1("Part 7: Teardown & Cleanup")

# ── Demo 28 ──
H2("Demo 28: Graceful Shutdown & Data Persistence")

body("Purpose: Stop all services, verify data persists across restarts, "
     "and optionally clean up completely.")

H3("Step 1 \u2014 Graceful Shutdown (Preserve Data)")
code_block("docker compose down", "bash")
body("This stops and removes containers but preserves named volumes "
     "(orthanc_data, postgres_data, nim_cache, prometheus_data, grafana_data).")

H3("Step 2 \u2014 Restart and Verify Persistence")
code_block(
    "docker compose up -d\n"
    "\n"
    "# Wait for services to become healthy\n"
    "sleep 30\n"
    "\n"
    "# Verify data persisted\n"
    "docker exec imaging-postgres psql -U imaging -d imaging_agent -c \"\n"
    "SELECT COUNT(*) AS studies, (SELECT COUNT(*) FROM findings) AS findings\n"
    'FROM studies;"',
    "bash",
)
body("Expected output: Non-zero counts matching what was loaded before shutdown.")

H3("Step 3 \u2014 Full Cleanup (Remove All Data)")
code_block("docker compose down -v", "bash")
body("This removes all containers AND all named volumes. All data is "
     "permanently deleted.")

H3("Step 4 \u2014 Verify Clean State")
code_block("docker volume ls | grep imaging", "bash")
body("Expected output: No volumes listed.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# APPENDIX A — SYNTHETIC TEST DATA GENERATION
# ══════════════════════════════════════════════════════════════
H1("Appendix A \u2014 Synthetic Test Data Generation")

H3("Generate Test DICOM for Any Modality")
body("A Python module that generates synthetic DICOM test data for all 4 "
     "modalities (CT Head, CXR, CT Chest, MRI Brain FLAIR). Key functions:")
bullet("generate_ct_head()", "64 slices, 512x512, hemorrhage region at "
       "slices 20-35")
bullet("generate_cxr()", "Single 2048x2048 PA chest X-ray")
bullet("generate_ct_chest()", "128 slices, 512x512, nodule region at "
       "slices 60-68")
bullet("generate_mri_brain()", "96 slices, 256x256, FLAIR with 3 lesion spots")

body("See the demo guide markdown for complete code listings.")

spacer()

# ══════════════════════════════════════════════════════════════
# APPENDIX B — TROUBLESHOOTING REFERENCE
# ══════════════════════════════════════════════════════════════
H1("Appendix B \u2014 Troubleshooting Reference")

add_table(
    ["Symptom", "Cause", "Fix"],
    [
        ["Container won't start", "Port conflict",
         "ss -tlnp | grep <port> to find conflict"],
        ["NIM LLM not ready", "First model download (~16 GB)",
         "Wait 5-10 min; check docker logs imaging-nim-llm"],
        ["GPU not available", "NVIDIA runtime not configured",
         "nvidia-ctk runtime configure --runtime=docker && "
         "systemctl restart docker"],
        ["Webhook not firing", "StableAge not elapsed",
         "Wait 10 seconds after last instance uploaded"],
        ["Lua script error", "Script not mounted",
         "Verify ./config/scripts:/etc/orthanc/scripts:ro in docker-compose"],
        ["Database connection refused", "PostgreSQL not ready",
         "Check docker logs imaging-postgres; verify healthcheck passes"],
        ["Embedding service 503", "BiomedCLIP downloading",
         "Wait for model download; check docker logs imaging-embedding"],
        ["Pipeline timeout", "GPU memory pressure",
         "Check nvidia-smi; reduce concurrent pipelines"],
        ["DICOM SR missing", "Workflow didn't complete",
         "Check docker logs imaging-dicom-listener for errors"],
        ["Worklist empty", "No positive findings",
         "Verify findings table has rows with is_positive = true"],
        ["Vector search slow", "Missing HNSW index",
         "Run CREATE INDEX idx_embeddings_hnsw ... from init.sql"],
        ["ARM64 image not found", "Wrong image tag",
         "Ensure -dgx-spark suffix for NIM images"],
    ],
)

spacer()

# ══════════════════════════════════════════════════════════════
# APPENDIX C — ENVIRONMENT VARIABLES REFERENCE
# ══════════════════════════════════════════════════════════════
H1("Appendix C \u2014 Environment Variables Reference")

code_block(
    "# .env.example -- Complete environment variable reference\n"
    "\n"
    "# -- NGC Authentication --\n"
    "NGC_API_KEY=your-ngc-api-key-here\n"
    "\n"
    "# -- PostgreSQL --\n"
    "POSTGRES_USER=imaging\n"
    "POSTGRES_PASSWORD=imaging_secret\n"
    "POSTGRES_DB=imaging_agent\n"
    "\n"
    "# -- Grafana --\n"
    "GRAFANA_USER=admin\n"
    "GRAFANA_PASSWORD=changeme\n"
    "\n"
    "# -- FHIR Server (optional -- external EHR integration) --\n"
    "FHIR_SERVER_URL=http://localhost:8080/fhir\n"
    "\n"
    "# -- Model Configuration (optional overrides) --\n"
    "# EMBEDDING_MODEL=microsoft/BiomedCLIP-PubMedBERT_256-vit_base_patch16_224\n"
    "# NIM_MODEL=meta-llama3-8b-instruct\n"
    "\n"
    "# -- Pipeline Configuration (optional overrides) --\n"
    "# CXR_CONFIDENCE_THRESHOLD=0.50\n"
    "# HEMORRHAGE_CONFIDENCE_THRESHOLD=0.50\n"
    "# NODULE_CONFIDENCE_THRESHOLD=0.40\n"
    "# MS_LESION_CONFIDENCE_THRESHOLD=0.50\n"
    "\n"
    "# -- Orthanc (optional overrides) --\n"
    "# ORTHANC_AET=IMAGING_AGENT\n"
    "# ORTHANC_STABLE_AGE=10",
    "bash",
)

spacer()

# ── Footer ──
P("HCLS AI Factory \u2014 Open Source (Apache 2.0) \u2014 NVIDIA DGX Spark",
  italic=True, size=9, color=GRAY_META, before=12, after=0,
  align=WD_ALIGN_PARAGRAPH.CENTER)

# ══════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(
    OUT_DIR, "HCLS_Imaging_AI_Agent_Demo_Guide.docx"
)
doc.save(OUT_PATH)
print(f"Demo Guide saved to: {OUT_PATH}")
