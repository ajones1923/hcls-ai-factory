#!/usr/bin/env python3
# Copyright 2026 Adam Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Generate Imaging Intelligence Agent — Learning Guide Advanced (DOCX).

Graduate-level learning guide with deep technical, mathematical, and
architectural detail. VCP-style formatting matching the Project Bible,
Demo Guide, and Foundations Guide.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Colors ────────────────────────────────────────────────
NAVY = RGBColor(0x1B, 0x23, 0x33)
TEAL = RGBColor(0x1A, 0xAF, 0xCC)
GREEN = RGBColor(0x76, 0xB9, 0x00)
GRAY_BODY = RGBColor(0x33, 0x33, 0x33)
GRAY_META = RGBColor(0x66, 0x66, 0x66)
GRAY_CODE = RGBColor(0x44, 0x44, 0x44)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_NAVY = "1B2333"
HEX_TEAL = "1AAFCC"
HEX_GREEN = "76B900"
HEX_LIGHT = "F8FAFB"
HEX_CODE_BG = "F5F5F5"
FONT = "Calibri"
CODE_FONT = "Consolas"

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
style.font.color.rgb = GRAY_BODY
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15


# ── Helper functions ──────────────────────────────────────

def R(p, text, bold=False, italic=False, size=None, color=None, font=None):
    r = p.add_run(text)
    r.font.name = font or FONT
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return r


def P(text="", bold=False, italic=False, size=10.5, color=GRAY_BODY,
      before=0, after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text:
        R(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def H1(text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        r.clear()
    R(p, text, bold=True, size=22, color=NAVY)
    return p


def H2(text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.clear()
    R(p, text, bold=True, size=16, color=NAVY)
    return p


def H3(text):
    return P(text, bold=True, size=13, color=TEAL, before=10, after=4)


def body(text, before=0, after=6):
    return P(text, size=10.5, color=GRAY_BODY, before=before, after=after)


def body_rich(parts, before=0, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    for text, bold, italic in parts:
        R(p, text, bold=bold, italic=italic, size=10.5, color=GRAY_BODY)
    return p


def bullet(lead, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    R(p, lead, bold=True, size=10.5, color=GRAY_BODY)
    R(p, f"  {text}", size=10.5, color=GRAY_BODY)
    return p


def bullet_plain(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    R(p, "\u2022  ", bold=True, size=10.5, color=TEAL)
    R(p, text, size=10.5, color=GRAY_BODY)
    return p


def numbered(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    R(p, f"{num}. ", bold=True, size=10.5, color=TEAL)
    R(p, text, size=10.5, color=GRAY_BODY)
    return p


def impl_note(text):
    """Implementation Note callout."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.35)
    R(p, "\u2699 Implementation Note: ", bold=True, size=10.5, color=GREEN)
    R(p, text, size=10.5, color=GRAY_BODY)
    return p


def deep_dive(text):
    """Deep Dive callout."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.35)
    R(p, "\u26a1 Deep Dive: ", bold=True, size=10.5, color=TEAL)
    R(p, text, italic=True, size=10.5, color=GRAY_BODY)
    return p


def discussion_q(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    R(p, f"Q{num}. ", bold=True, size=10.5, color=TEAL)
    R(p, text, size=10.5, color=GRAY_BODY)
    return p


def code_block(text, language=""):
    if language:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        R(p, language, bold=True, size=8, color=TEAL, font=FONT)
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.2)
        R(p, line if line else " ", size=8, color=GRAY_CODE, font=CODE_FONT)


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="D0D0D0", sz="4"):
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>')
    tcPr.append(borders)


def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, HEX_NAVY)
        set_cell_borders(c, HEX_NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(p, h, bold=True, size=9, color=WHITE)
    for ri, row in enumerate(rows):
        bg = "FFFFFF" if ri % 2 == 0 else HEX_LIGHT
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            set_cell_shading(c, bg)
            set_cell_borders(c, "E0E0E0")
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if ci == 0:
                R(p, str(val), bold=True, size=9, color=TEAL)
            else:
                R(p, str(val), size=9, color=GRAY_BODY)
    return t


def spacer(pts=6):
    P("", after=pts)


# ==============================================================
# COVER PAGE
# ==============================================================

accent = doc.add_table(rows=1, cols=1)
accent.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = accent.rows[0].cells[0]
set_cell_shading(cell, HEX_GREEN)
p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
R(p, " ", size=6, color=GREEN)

spacer(40)

P("Learning Guide", bold=True, size=36, color=TEAL,
  align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
P("Advanced", bold=True, size=28, color=TEAL,
  align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

spacer(12)

P("Imaging Intelligence Agent", bold=True, size=32, color=NAVY,
  align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
P("Graduate-Level Technical Deep Dive",
  size=16, color=GRAY_META, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)

spacer(20)

P("HCLS AI Factory  |  NVIDIA DGX Platform",
  size=12, color=GRAY_META, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
P("Author: Adam Jones  |  February 2026",
  size=11, color=GRAY_META, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
P("Apache 2.0 License",
  size=10, color=GRAY_META, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)

doc.add_page_break()

# ==============================================================
# TABLE OF CONTENTS
# ==============================================================

H1("Table of Contents")
spacer(4)

toc_items = [
    "Chapter 1 \u2014 Medical Imaging Physics and Acquisition",
    "Chapter 2 \u2014 Deep Learning Architectures for Medical Imaging",
    "Chapter 3 \u2014 DICOM Deep Dive \u2014 Data Model and Networking",
    "Chapter 4 \u2014 GPU Architecture and CUDA Computation",
    "Chapter 5 \u2014 Container Orchestration and Microservice Architecture",
    "Chapter 6 \u2014 PostgreSQL + pgvector \u2014 Schema Design and Query Optimization",
    "Chapter 7 \u2014 Clinical Workflow Implementation \u2014 Deep Dive",
    "Chapter 8 \u2014 LangGraph Agent Architecture \u2014 Design Patterns",
    "Chapter 9 \u2014 Embedding Models and Vector Retrieval",
    "Chapter 10 \u2014 Clinical Interoperability \u2014 FHIR, HL7, and Integration Engines",
    "Chapter 11 \u2014 RAG Architecture and LLM Serving",
    "Chapter 12 \u2014 Observability Stack \u2014 Prometheus, Grafana, and DCGM",
    "Chapter 13 \u2014 HCLS AI Factory \u2014 Multi-Agent Architecture",
    "Chapter 14 \u2014 Regulatory, Safety, and Deployment Engineering",
    "Discussion Questions \u2014 Comprehensive",
    "References",
]
for i, item in enumerate(toc_items):
    numbered(i + 1, item)

doc.add_page_break()

# ==============================================================
# CHAPTER 1
# ==============================================================

H1("Chapter 1 \u2014 Medical Imaging Physics and Acquisition")

H2("X-Ray Attenuation and the Beer-Lambert Law")
body("X-ray imaging exploits differential attenuation of photons through tissue. "
     "The Beer-Lambert law governs intensity reduction:")
code_block("I = I\u2080 \u00d7 e^(-\u03bcx)")
body("Where I\u2080 is the incident intensity, \u03bc is the linear attenuation "
     "coefficient (cm\u207b\u00b9), and x is the material thickness. Different tissues "
     "have different \u03bc values \u2014 bone attenuates heavily (high \u03bc), air barely "
     "attenuates (\u03bc \u2248 0), and soft tissue falls between.")

H2("Hounsfield Units (HU) and CT Windowing")
body("CT scanners reconstruct attenuation maps calibrated to Hounsfield units:")
code_block("HU = 1000 \u00d7 (\u03bc_tissue - \u03bc_water) / \u03bc_water")
body("Key reference values: water = 0 HU, air = -1000 HU, dense bone = +1000 HU, "
     "acute blood = +50 to +70 HU, white matter = +20 to +30 HU.")
body_rich([
    ("Windowing", True, False),
    (" maps a range of HU values to display brightness. The agent uses specific "
     "windows for each workflow:", False, False),
])
add_table(
    ["Window", "Width (W)", "Level (L)", "Use"],
    [
        ["Brain", "80", "40", "Hemorrhage detection \u2014 optimizes blood/brain contrast"],
        ["Lung", "1500", "-600", "Lung nodule detection \u2014 shows lung parenchyma"],
        ["Bone", "2000", "500", "Fracture assessment"],
        ["Soft Tissue", "400", "50", "General abdomen/pelvis"],
    ],
)
spacer(4)
body("In DICOM, raw pixel values are converted to HU using the rescale equation:")
code_block("hu = pixel_value * RescaleSlope + RescaleIntercept", "python")
body("The RescaleSlope and RescaleIntercept are stored in DICOM tags (0028,1053) "
     "and (0028,1052).")

H2("CT Reconstruction")
body("The CT scanner acquires projections \u2014 line integrals of attenuation along "
     "ray paths. The mathematical foundation is the Radon transform: the set of all "
     "line integrals through a 2D function. Reconstruction recovers the original "
     "function from its Radon transform.")
body_rich([
    ("Filtered Back-Projection (FBP): ", True, False),
    ("The classical reconstruction algorithm. Each projection is filtered with a ramp "
     "filter (to correct for blurring inherent in back-projection) and then smeared "
     "back across the image space. FBP is fast and analytically exact for "
     "parallel-beam geometry.", False, False),
])
body_rich([
    ("Iterative Reconstruction (IR): ", True, False),
    ("Modern scanners use iterative methods (ADMIRE, ASIR-V, IMR) that model the "
     "acquisition physics, noise statistics, and image prior information. IR produces "
     "lower-noise images at reduced radiation dose but is computationally heavier.", False, False),
])
impl_note("The agent does not perform reconstruction \u2014 it receives "
          "already-reconstructed images. But understanding reconstruction artifacts "
          "(beam hardening, partial volume, motion) helps interpret model behavior "
          "on imperfect inputs.")

H2("MRI Physics: T1/T2 Relaxation")
body("MRI signal originates from precessing hydrogen nuclei. After RF excitation:")
body_rich([
    ("T1 relaxation (longitudinal recovery): ", True, False),
    ("Protons return to thermal equilibrium along the main magnetic field (B\u2080). "
     "T1 is the time constant for this exponential recovery:", False, False),
])
code_block("Mz(t) = M\u2080(1 - e^(-t/T1))")
body_rich([
    ("T2 relaxation (transverse decay): ", True, False),
    ("Phase coherence is lost due to spin-spin interactions. T2 governs the "
     "exponential decay of the transverse magnetization:", False, False),
])
code_block("Mxy(t) = Mxy\u2080 \u00d7 e^(-t/T2)")
body("Different tissues have characteristic T1/T2 values. CSF: long T1, long T2 "
     "(bright on T2). White matter: shorter T1 and T2 than gray matter. MS lesions: "
     "prolonged T1 and T2 (bright on T2/FLAIR).")
body_rich([
    ("FLAIR Sequence Design: ", True, False),
    ("FLAIR uses an inversion recovery preparation pulse timed so that CSF signal "
     "is nulled (TI chosen such that CSF\u2019s longitudinal magnetization passes "
     "through zero at readout). This makes MS lesions \u2014 which have prolonged T2 "
     "but shorter T1 than CSF \u2014 appear bright against suppressed CSF. The agent "
     "specifically requires FLAIR sequences for MS lesion segmentation.", False, False),
])

H2("Voxel Anisotropy and Resampling")
body("CT and MRI voxels are often anisotropic \u2014 the in-plane pixel spacing "
     "(e.g., 0.5 \u00d7 0.5 mm) differs from the slice thickness (e.g., 5 mm). "
     "Neural networks generally perform better on isotropic inputs.")
impl_note("The agent\u2019s preprocessing pipelines resample all inputs to "
          "1.0 \u00d7 1.0 \u00d7 1.0 mm isotropic using MONAI\u2019s Spacingd transform "
          "with trilinear interpolation. This standardizes geometry across different "
          "scanners and protocols but introduces interpolation artifacts in the "
          "through-plane direction when slice thickness is large.")

H2("DICOM Pixel Data Encoding")
body("DICOM stores pixel data in the (7FE0,0010) PixelData element. "
     "Key attributes controlling interpretation:")
bullet("Transfer Syntax:", "Determines compression (1.2.840.10008.1.2 = implicit VR "
       "little-endian uncompressed; 1.2.840.10008.1.2.4.90 = JPEG 2000 lossless)")
bullet("Photometric Interpretation:", "MONOCHROME1 (white = minimum), "
       "MONOCHROME2 (white = maximum), RGB")
bullet("Bits Allocated / Bits Stored:", "Typically 16/12 or 16/16 for CT "
       "(12-bit HU range)")
bullet("Pixel Representation:", "0 = unsigned, 1 = signed (important for HU "
       "values below zero)")
code_block("import pydicom\nimport numpy as np\n\nds = pydicom.dcmread(\"slice.dcm\")\n"
           "pixels = ds.pixel_array  # Returns numpy array\n\n"
           "# Convert to Hounsfield units\n"
           "hu = pixels * ds.RescaleSlope + ds.RescaleIntercept", "python")

H3("Discussion Questions \u2014 Chapter 1")
discussion_q(1, "Why does the agent use a W:80 L:40 window for hemorrhage detection "
               "rather than a standard brain window?")
discussion_q(2, "What artifacts might isotropic resampling introduce when the source "
               "CT has 5 mm slice thickness?")
discussion_q(3, "How does the FLAIR inversion time determine which tissues are nulled?")

doc.add_page_break()

# ==============================================================
# CHAPTER 2
# ==============================================================

H1("Chapter 2 \u2014 Deep Learning Architectures for Medical Imaging")

H2("Convolutional Neural Networks (CNNs)")
body("The convolution operation slides a learnable kernel (filter) across the input. "
     "For a 2D input x and kernel w of size k\u00d7k:")
code_block("(x * w)[i,j] = \u03a3_m \u03a3_n x[i+m, j+n] \u00d7 w[m, n]")
body("Key hyperparameters: kernel size (typically 3\u00d73), stride (step size between "
     "kernel applications), padding (border handling), and number of output channels "
     "(feature maps). Each layer learns multiple kernels, producing a stack of "
     "feature maps.")

H2("DenseNet-121: Dense Connectivity")
body_rich([
    ("Reference: ", True, False),
    ("Huang et al., \u201cDensely Connected Convolutional Networks,\u201d CVPR 2017.", False, True),
])
body("DenseNet-121 connects each layer to every subsequent layer within a dense block. "
     "For L layers, there are L(L+1)/2 connections instead of L. Each layer receives "
     "feature maps from all preceding layers as input.")
H3("Architecture Specifics")
bullet("4 dense blocks", "with (6, 12, 24, 16) layers respectively \u2192 "
       "6+12+24+16 = 58 dense layers + transition layers + initial conv + "
       "classifier = 121 weight layers")
bullet("Growth rate k = 32:", "Each layer produces k new feature maps. With dense "
       "connectivity, the Lth layer receives k\u2080 + k(L-1) input channels "
       "(where k\u2080 is the initial channel count)")
bullet("Bottleneck layers:", "1\u00d71 convolution before each 3\u00d73 convolution "
       "to reduce computation (BN \u2192 ReLU \u2192 1\u00d71 Conv \u2192 BN \u2192 "
       "ReLU \u2192 3\u00d73 Conv)")
bullet("Transition layers:", "Between dense blocks \u2014 1\u00d71 convolution with "
       "compression factor \u03b8=0.5 (halves channels) followed by 2\u00d72 "
       "average pooling")
bullet("Final classification:", "Global average pooling \u2192 fully connected layer")
impl_note("The agent uses DenseNet-121 for CXR multi-label classification (5 output "
          "units with sigmoid activation) and as the initial classification gate for "
          "CT head hemorrhage detection.")

H2("U-Net: Encoder-Decoder with Skip Connections")
body_rich([
    ("Reference: ", True, False),
    ("Ronneberger et al., \u201cU-Net: Convolutional Networks for Biomedical Image "
     "Segmentation,\u201d MICCAI 2015.", False, True),
])
bullet("Encoder (contracting path):", "Repeated application of two 3\u00d73 "
       "convolutions (each followed by batch normalization and ReLU), then "
       "2\u00d72 max pooling for downsampling. Each downsampling step doubles "
       "the feature channels.")
bullet("Decoder (expanding path):", "2\u00d72 transposed convolution for upsampling, "
       "concatenation with the corresponding encoder feature maps (skip connections), "
       "then two 3\u00d73 convolutions.")
bullet("Skip connections:", "Copy feature maps from encoder to decoder at each "
       "resolution level. This preserves spatial detail that the encoder\u2019s "
       "pooling operations would otherwise lose.")
body("The 3D U-Net extension processes volumetric data (5D input: batch \u00d7 channels "
     "\u00d7 depth \u00d7 height \u00d7 width). Memory requirements scale cubically "
     "with spatial dimensions. The agent addresses this with isotropic resampling to "
     "1mm, sliding window inference, and reduced channel counts.")
code_block("UNet(spatial_dims=3, in_channels=1, out_channels=2,\n"
           "     channels=(16, 32, 64, 128, 256), strides=(2, 2, 2, 2),\n"
           "     num_res_units=2, norm=\"batch\")", "python")

H2("SegResNet: Residual Blocks in Segmentation")
body_rich([
    ("Reference: ", True, False),
    ("Myronenko, \u201c3D MRI Brain Tumor Segmentation Using Autoencoder "
     "Regularization,\u201d BrainLes@MICCAI 2018.", False, True),
])
body("SegResNet introduces residual blocks (from ResNet) into a U-Net-like "
     "architecture. Each block computes:")
code_block("output = F(x) + x")
body("Where F(x) is the block\u2019s learned transformation. The identity shortcut "
     "(+x) addresses the vanishing gradient problem \u2014 gradients flow directly "
     "through the skip connection during backpropagation, enabling training of "
     "deeper networks.")
code_block("SegResNet(spatial_dims=3, in_channels=1, out_channels=2,\n"
           "          init_filters=16, blocks_down=(1, 2, 2, 4),\n"
           "          blocks_up=(1, 1, 1), norm=\"batch\")", "python")

H2("RetinaNet: Feature Pyramid Network + Focal Loss")
body_rich([
    ("Reference: ", True, False),
    ("Lin et al., \u201cFocal Loss for Dense Object Detection,\u201d ICCV 2017.", False, True),
])
body("RetinaNet addresses the class imbalance problem in detection (overwhelmingly "
     "many background regions vs. few objects). Two key innovations:")
body_rich([
    ("Feature Pyramid Network (FPN): ", True, False),
    ("Builds a multi-scale feature pyramid from the backbone network (ResNet-50). "
     "The pyramid has levels P3 through P7, with each level operating at a different "
     "spatial resolution. Objects of different sizes are detected at different "
     "pyramid levels.", False, False),
])
body_rich([
    ("Focal Loss: ", True, False),
    ("Modifies the standard cross-entropy loss to down-weight easy examples:", False, False),
])
code_block("FL(p_t) = -\u03b1_t (1 - p_t)^\u03b3 \u00d7 log(p_t)")
body("Where p_t is the predicted probability for the correct class, \u03b3 is the "
     "focusing parameter (typically \u03b3=2), and \u03b1 is a class-balancing weight. "
     "When a sample is correctly classified with high confidence (p_t \u2192 1), the "
     "modulating factor (1-p_t)^\u03b3 \u2192 0, reducing the loss contribution.")
body_rich([
    ("Anchor design: ", True, False),
    ("At each FPN level, anchors of different scales and aspect ratios tile the "
     "feature map. The agent uses sizes (4, 6, 8, 12, 16) and aspect ratios "
     "(0.5, 1.0, 2.0) for nodule detection. Non-maximum suppression (NMS) with "
     "threshold 0.3 removes duplicate detections.", False, False),
])

H2("SwinUNETR: Transformer-Based Segmentation")
body_rich([
    ("Reference: ", True, False),
    ("Hatamizadeh et al., \u201cSwin UNETR: Swin Transformers for Semantic "
     "Segmentation of Brain Tumors in MRI Images,\u201d BrainLes@MICCAI 2021.", False, True),
])
bullet("Patch embedding:", "Input volume is partitioned into non-overlapping 3D "
       "patches (e.g., 2\u00d72\u00d72 voxels \u2192 tokens)")
bullet("Shifted window attention:", "Self-attention computed within local windows, "
       "with alternate layers shifting the window partition to enable cross-window connections")
bullet("Hierarchical features:", "4 stages with downsampling via patch merging \u2014 "
       "analogous to the CNN encoder\u2019s pooling")
bullet("Decoder:", "CNN-based with skip connections from each transformer stage")

H2("Loss Functions")
add_table(
    ["Loss", "Formula", "When to Use"],
    [
        ["Cross-Entropy", "-\u03a3 y_i log(p_i)", "Classification, per-pixel classification"],
        ["Dice Loss", "1 - 2|A\u2229B|/(|A|+|B|)", "Segmentation \u2014 handles class imbalance"],
        ["Focal Loss", "-\u03b1(1-p_t)^\u03b3 log(p_t)", "Detection \u2014 focuses on hard examples"],
        ["Combined (Dice+CE)", "\u03bb\u2081\u00d7Dice + \u03bb\u2082\u00d7CE",
         "Segmentation \u2014 benefits from both"],
    ],
)
spacer(4)

H2("Data Augmentation for Medical Imaging")
bullet("Affine:", "Random rotation (\u00b115\u00b0), scaling (0.85-1.15), "
       "translation (\u00b110 pixels)")
bullet("Elastic deformation:", "Random B-spline displacement field \u2014 "
       "simulates anatomical variability")
bullet("Intensity augmentation:", "Random brightness/contrast shifts, noise "
       "injection (Gaussian, Rician for MRI)")
bullet("Flipping:", "Left-right mirroring (valid for bilateral anatomy but must be "
       "avoided for laterality-dependent tasks)")

H2("Transfer Learning")
numbered(1, "ImageNet pre-training: General visual features. Wide availability. "
           "Significant domain gap from medical imaging.")
numbered(2, "Medical pre-training: Weights trained on large medical image datasets "
           "(e.g., CheXpert for CXR, BTCV for CT). MONAI Model Zoo provides "
           "medical pre-trained weights for most architectures.")

H2("MONAI Framework Patterns")
bullet("Dictionary-based transforms:", "Operations keyed by data role "
       "(keys=[\"image\", \"label\"]) enabling matched augmentation of "
       "image-label pairs")
bullet("Sliding window inference:", "For volumes larger than GPU memory \u2014 "
       "processes overlapping patches, then stitches outputs with weighted "
       "averaging in overlap regions")
bullet("Dataset classes:", "CacheDataset (pre-caches transforms in memory), "
       "PersistentDataset (caches to disk)")

H3("Discussion Questions \u2014 Chapter 2")
discussion_q(1, "Derive the total number of parameters in the classification head "
               "of DenseNet-121 with growth rate k=32 given 4 dense blocks of "
               "(6,12,24,16) layers.")
discussion_q(2, "Why does focal loss with \u03b3=2 outperform standard "
               "cross-entropy for lung nodule detection?")
discussion_q(3, "What is the memory complexity of 3D sliding window inference "
               "with window size W\u00b3 and overlap ratio r?")

doc.add_page_break()

# ==============================================================
# CHAPTER 3
# ==============================================================

H1("Chapter 3 \u2014 DICOM Deep Dive")

H2("Information Object Definitions (IODs)")
body("DICOM organizes data in a four-level hierarchy: Patient \u2192 Study \u2192 "
     "Series \u2192 Instance. Each level has specific Information Object Definitions "
     "(IODs) that define required and optional attributes.")
body_rich([
    ("Composite IODs ", True, False),
    ("(most common) include attributes from multiple Information Entities. A CT Image "
     "IOD includes Patient IE, Study IE, Series IE, Frame of Reference IE, Equipment "
     "IE, and Image IE attributes \u2014 all in a single DICOM file.", False, False),
])
body_rich([
    ("Normalized IODs ", True, False),
    ("represent a single entity (e.g., Modality Worklist) and are accessed via "
     "N-services (N-GET, N-SET, N-CREATE) rather than C-services.", False, False),
])

H2("Value Representations (VR)")
add_table(
    ["VR", "Name", "Example"],
    [
        ["PN", "Person Name", "DOE^JOHN^A^^DR"],
        ["DA", "Date", "20260115"],
        ["TM", "Time", "143205.123456"],
        ["UI", "Unique Identifier", "1.2.840.113619.2.55.1"],
        ["DS", "Decimal String", "\"25.3\" (stored as text)"],
        ["IS", "Integer String", "\"300\""],
        ["SQ", "Sequence", "Nested dataset(s)"],
        ["OW/OB", "Other Word/Byte", "Pixel data"],
    ],
)
spacer(4)
body("Transfer syntaxes determine byte ordering and whether VRs are encoded explicitly "
     "or implicitly. DICOM\u2019s default (1.2.840.10008.1.2) is Implicit VR Little "
     "Endian, but Explicit VR Little Endian (1.2.840.10008.1.2.1) is preferred because "
     "the VR is stored in the data stream, eliminating dictionary lookup requirements.")

H2("DIMSE Services")
add_table(
    ["Service", "Direction", "Purpose"],
    [
        ["C-ECHO", "SCU \u2192 SCP", "Verify connectivity (\"ping\")"],
        ["C-STORE", "SCU \u2192 SCP", "Send a DICOM instance"],
        ["C-FIND", "SCU \u2192 SCP", "Query for matching instances"],
        ["C-MOVE", "SCU \u2192 SCP \u2192 3rd party",
         "Request SCP to send instances to a third-party AE"],
        ["C-GET", "SCU \u2192 SCP",
         "Request SCP to send instances back to the SCU"],
    ],
)
spacer(4)
body_rich([
    ("Before any DIMSE operation, an ", False, False),
    ("Association negotiation ", True, False),
    ("occurs. The SCU proposes presentation contexts (SOP Class + Transfer Syntax "
     "pairs), and the SCP accepts or rejects each. Only accepted contexts can be "
     "used.", False, False),
])

H2("DICOMweb REST API")
body_rich([
    ("STOW-RS (Store): ", True, False),
    ("POST multipart/related request with DICOM instances in the request body. "
     "Content-Type per part: application/dicom.", False, False),
])
body_rich([
    ("WADO-RS (Retrieve): ", True, False),
    ("GET request with study/series/instance UIDs in the URL path. Supports metadata "
     "retrieval (JSON or XML), bulk data retrieval, and rendered frame retrieval.", False, False),
])
body_rich([
    ("QIDO-RS (Search): ", True, False),
    ("GET request with query parameters mapped from DICOM tag names. Supports limit, "
     "offset, includefield, and standard DICOM matching.", False, False),
])

H2("DICOM SR Template TID 1500")
body("The agent\u2019s findings are encoded as DICOM SR Measurement Reports "
     "following template TID 1500:")
code_block("Measurement Report (root)\n"
           "\u251c\u2500\u2500 Language of Content (EN)\n"
           "\u251c\u2500\u2500 Observation Context\n"
           "\u2502   \u251c\u2500\u2500 Observer Type (Device)\n"
           "\u2502   \u2514\u2500\u2500 Device Observer (Imaging Intelligence Agent v2.1)\n"
           "\u251c\u2500\u2500 Image Library (reference to source images)\n"
           "\u2514\u2500\u2500 Measurement Group(s)\n"
           "    \u251c\u2500\u2500 Tracking Identifier\n"
           "    \u251c\u2500\u2500 Finding (coded concept, e.g., SCT:50960005 = Hemorrhage)\n"
           "    \u251c\u2500\u2500 Measurement(s)\n"
           "    \u2502   \u251c\u2500\u2500 Concept Name (e.g., DCM:118565006 = Volume)\n"
           "    \u2502   \u251c\u2500\u2500 Measured Value (25.3)\n"
           "    \u2502   \u2514\u2500\u2500 Units (UCUM: mL)\n"
           "    \u2514\u2500\u2500 Qualitative Evaluation(s)\n"
           "        \u2514\u2500\u2500 Severity (CRITICAL)")
code_block("import highdicom as hd\n\nmeasurement = hd.sr.Measurement(\n"
           "    name=hd.sr.CodedConcept(\"118565006\", \"SCT\", \"Volume\"),\n"
           "    value=25.3,\n"
           "    unit=hd.sr.CodedConcept(\"mL\", \"UCUM\", \"milliliter\"),\n"
           ")", "python")

H2("Orthanc Architecture")
bullet("SQLite metadata index:", "Fast lookups by UID, patient ID, date")
bullet("On-disk storage:", "Each DICOM instance stored as a separate file, organized by hash")
bullet("Lua scripting hooks:", "Event-driven callbacks (OnStoredInstance, OnStableStudy) "
       "\u2014 the agent\u2019s trigger mechanism")
bullet("REST API:", "Full CRUD on patients, studies, series, instances, plus DICOMweb endpoints")
body("The agent\u2019s Lua callback fires when a study has been stable (no new "
     "instances received) for StableAge seconds (configured as 10):")
code_block("function OnStableStudy(studyId, tags, metadata)\n"
           "    -- Port 8000 is the container-internal port; mapped to host 8522\n"
           "    local url = \"http://dicom-listener:8000/webhook/study-complete\"\n"
           "    local body = '{\"orthanc_id\": \"' .. studyId .. '\"}'\n"
           "    HttpPost(url, body, {[\"Content-Type\"] = \"application/json\"})\n"
           "end", "lua")

H3("Discussion Questions \u2014 Chapter 3")
discussion_q(1, "Why does the agent use StableAge=10 seconds instead of processing "
               "each instance immediately?")
discussion_q(2, "Describe the Association negotiation required for Orthanc to accept "
               "a JPEG 2000 compressed CT series.")
discussion_q(3, "How would you extend the DICOM SR template to include lesion-level "
               "embedding references?")

doc.add_page_break()

# ==============================================================
# CHAPTER 4
# ==============================================================

H1("Chapter 4 \u2014 GPU Architecture and CUDA Computation")

H2("GPU Hardware Architecture")
body("A modern NVIDIA GPU consists of an array of Streaming Multiprocessors (SMs). "
     "Each SM contains:")
bullet("CUDA cores:", "FP32 units")
bullet("Tensor Cores:", "Matrix multiply-accumulate units optimized for deep learning")
bullet("Warp schedulers:", "Dispatch units that execute threads in groups of 32 "
       "called \"warps\"")
bullet("Shared memory and L1 cache:", "Configurable split")
bullet("Register file:", "Per-thread fast storage")
body("The thread hierarchy: a grid contains blocks, each block contains threads. "
     "The hardware maps blocks to SMs and threads to CUDA cores within an SM. "
     "Threads within a block can synchronize and share data via shared memory.")

H2("Memory Hierarchy")
numbered(1, "Registers (~1 cycle latency, per-thread)")
numbered(2, "Shared memory / L1 cache (~20-30 cycles, per-SM, configurable 48-128 KB)")
numbered(3, "L2 cache (~200 cycles, shared across all SMs)")
numbered(4, "Global memory (HBM2e on data center GPUs, LPDDR5x on DGX Spark, "
           "~400-600 cycles)")
deep_dive("On the DGX Spark, the 128 GB LPDDR5x is unified \u2014 the same physical "
          "memory is accessible by both the Grace CPU and the GB10 GPU. This "
          "eliminates explicit PCIe transfers but introduces coherence overhead "
          "when both processors access the same pages.")

H2("Grace Blackwell GB10")
bullet("Compute capability:", "10.0+")
bullet("5th-generation Tensor Cores:", "FP4, FP8, FP16, BF16, TF32 support")
bullet("Hardware-accelerated sparsity:", "2:4 structured sparsity for 2\u00d7 throughput")
bullet("Unified memory:", "with Grace CPU via NVLink-C2C (coherent interconnect, not PCIe)")
body("The NVLink-C2C interconnect between Grace and Blackwell provides ~900 GB/s "
     "bidirectional bandwidth \u2014 far exceeding PCIe Gen5 (128 GB/s). Combined "
     "with unified memory addressing, the GPU can directly access CPU memory "
     "allocations without explicit copy commands, though access patterns affect "
     "performance.")

H2("cuDNN Convolution Algorithms")
add_table(
    ["Algorithm", "Approach", "Best For"],
    [
        ["Implicit GEMM", "Converts conv to matrix multiply", "General purpose, reliable"],
        ["FFT", "Transform to frequency domain, multiply, transform back",
         "Large kernels, large batches"],
        ["Winograd", "Minimal filtering algorithm (fewer multiplications)",
         "3\u00d73 and 5\u00d75 kernels"],
    ],
)
spacer(4)
impl_note("cuDNN\u2019s autotuning (torch.backends.cudnn.benchmark = True) benchmarks "
          "all algorithms on the first forward pass and selects the fastest for each "
          "layer configuration. Recommended for the agent since input shapes are fixed "
          "per workflow.")

H2("TensorRT Optimization")
numbered(1, "Graph optimization: Operator fusion (Conv + BN + ReLU \u2192 single "
           "kernel), constant folding, dead code elimination")
numbered(2, "Precision calibration: FP32 \u2192 FP16 (2\u00d7 throughput, minimal "
           "accuracy loss) or INT8 (4\u00d7 throughput, requires calibration dataset)")
numbered(3, "Kernel autotuning: Selects optimal CUDA kernels per layer based on "
           "the specific GPU")
numbered(4, "Memory optimization: Layer-level scratch space allocation, concurrent "
           "kernel execution planning")
impl_note("For the agent\u2019s workflows, FP16 inference provides the best "
          "tradeoff \u2014 nearly 2\u00d7 speedup with negligible accuracy impact "
          "on segmentation tasks. INT8 is viable for classification (DenseNet-121) "
          "but requires careful validation for segmentation models.")

H2("GPUDirect Storage")
body("GPUDirect Storage (GDS) eliminates the CPU from the I/O path:")
body_rich([
    ("Without GDS: ", True, False),
    ("NVMe \u2192 OS page cache \u2192 CPU memory \u2192 PCIe/NVLink \u2192 GPU "
     "memory (3 copies)", False, False),
])
body_rich([
    ("With GDS: ", True, False),
    ("NVMe \u2192 GPU memory via DMA (1 copy, zero CPU bounce buffers)", False, False),
])
code_block("cuFileDriverOpen();\ncuFileHandleRegister(&fh, &descr);\n"
           "cuFileRead(fh, devPtr, size, file_offset, 0);", "c")

H2("Memory Budget Analysis")
code_block("Model memory \u2248 num_parameters \u00d7 bytes_per_param "
           "(4 for FP32, 2 for FP16)\nActivation memory \u2248 batch_size \u00d7 "
           "sum(output_size_per_layer \u00d7 bytes_per_element)")
body("The agent\u2019s 3D U-Net for hemorrhage segmentation (~2M parameters) at FP16 "
     "requires ~4 MB for weights. But activations for a 256\u00d7256\u00d7256 input "
     "volume can require 2-4 GB. On the DGX Spark\u2019s 128 GB unified memory, this "
     "is comfortable \u2014 but memory must be shared with the NIM LLM service (~16 GB "
     "for Llama 3 8B), the embedding model (~500 MB), and the OS.")

H2("ARM64 Considerations")
bullet("Python wheels:", "Most scientific packages now provide ARM64 wheels, but "
       "some niche packages may require compilation from source")
bullet("NEON SIMD:", "ARM\u2019s SIMD instruction set (128-bit vectors). "
       "Performance-sensitive C extensions must be compiled with NEON support")
bullet("SVE:", "Grace supports SVE2 with variable-width vectors. Not yet widely "
       "leveraged by Python ecosystem")
bullet("Docker base images:", "Must use ARM64 variants (e.g., --platform linux/arm64)")

H3("Discussion Questions \u2014 Chapter 4")
discussion_q(1, "Calculate the theoretical peak FP16 throughput for the GB10 given "
               "its SM count and tensor core capabilities.")
discussion_q(2, "What is the expected speedup from TensorRT FP16 optimization for "
               "a 3D U-Net inference compared to PyTorch eager mode FP32?")
discussion_q(3, "Under what circumstances would GPUDirect Storage not improve "
               "performance on the DGX Spark?")

doc.add_page_break()

# ==============================================================
# CHAPTER 5
# ==============================================================

H1("Chapter 5 \u2014 Container Orchestration and Microservice Architecture")

H2("OCI Container Specification")
bullet("Image manifest:", "JSON document listing layers, config, and platform")
bullet("Config:", "Entry point, environment variables, exposed ports, volumes")
bullet("Layers:", "Filesystem diffs stacked via union mount (overlay2 driver on Linux)")
bullet("Rootfs:", "The combined filesystem visible to the container")
body("Each layer is content-addressable (SHA-256 hash). Identical layers are shared "
     "across images, reducing disk usage.")

H2("Multi-Stage Builds")
body("The agent\u2019s Dockerfiles use multi-stage builds to minimize image size:")
code_block("# Stage 1: Build dependencies\n"
           "FROM python:3.11-slim AS builder\n"
           "RUN pip install --no-cache-dir --target=/deps -r requirements.txt\n\n"
           "# Stage 2: Runtime image\n"
           "FROM python:3.11-slim\n"
           "COPY --from=builder /deps /usr/local/lib/python3.11/site-packages\n"
           "COPY app.py .", "dockerfile")

H2("ARM64 Multi-Architecture Builds")
code_block("docker buildx create --use --name multiarch\n"
           "docker buildx build --platform linux/arm64 -t imaging-agent:latest .",
           "bash")
body("Alternatively, build natively on the DGX Spark itself (no cross-compilation "
     "needed). QEMU user-space emulation enables ARM64 builds on x86 hosts but is "
     "significantly slower.")

H2("The Agent\u2019s Service Topology")
code_block("orthanc \u2500\u2500\u2500 dicom-listener \u2500\u2500\u2500 agent "
           "\u2500\u2500\u2500 portal\n"
           "                                 \u2191\n"
           "postgres \u2500\u2500\u2500 embedding-service   \u2502\n"
           "    \u2191                            \u2502\n"
           "    \u2514\u2500\u2500 fhir-publisher           \u2502\n"
           "                                 \u2502\n"
           "nim-llm \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500"
           "\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518\n\n"
           "dcgm-exporter \u2500\u2500\u2500 prometheus \u2500\u2500\u2500 grafana")
body_rich([
    ("Startup ordering: ", True, False),
    ("docker-compose depends_on with condition: service_healthy ensures services "
     "start only after their dependencies pass health checks. The longest startup "
     "path is: postgres \u2192 nim-llm \u2192 agent \u2192 portal (NIM model loading "
     "dominates at ~60-120 seconds).", False, False),
])
body_rich([
    ("Failure domains: ", True, False),
    ("If postgres fails, all data services are affected. If nim-llm fails, report "
     "generation degrades but triage still works. If orthanc fails, no new studies "
     "can be ingested but existing data remains queryable. This separation limits "
     "blast radius.", False, False),
])

H2("Health Check Patterns")
add_table(
    ["Pattern", "Implementation", "Used By"],
    [
        ["HTTP GET", "curl -f http://localhost:PORT/health", "All custom services (8520-8525)"],
        ["TCP Socket", "pg_isready utility", "PostgreSQL"],
        ["DIMSE echo", "curl to Orthanc REST /system endpoint", "Orthanc"],
    ],
)
spacer(4)
body("Health check parameters: interval=30s, timeout=10s, retries=3 for most services. "
     "Shorter interval (10s) for PostgreSQL since it is a critical dependency.")

H2("Container Security")
bullet("Non-root execution:", "USER 1000:1000 in Dockerfile \u2014 prevents container "
       "escape privilege escalation")
bullet("Read-only filesystem:", "read_only: true in compose with explicit tmpfs for /tmp")
bullet("Capability dropping:", "cap_drop: [ALL] with selective cap_add")
bullet("No new privileges:", "security_opt: [no-new-privileges:true]")

H2("Production Path: Kubernetes")
body("Docker Compose is suitable for single-node deployment (DGX Spark proof build). "
     "For Phase 2+ (departmental / multi-site), migration to Kubernetes provides:")
bullet("Horizontal scaling:", "Multiple replicas of stateless services")
bullet("Service discovery:", "DNS-based service resolution, load balancing")
bullet("Rolling updates:", "Zero-downtime deployment with readiness probes")
bullet("Resource management:", "CPU/memory requests and limits, GPU scheduling")
bullet("Persistent storage:", "StorageClass provisioners for NVMe, network storage")
bullet("Secrets management:", "Kubernetes Secrets for NGC_API_KEY, database credentials")

H3("Discussion Questions \u2014 Chapter 5")
discussion_q(1, "Design a Kubernetes deployment for the agent that supports horizontal "
               "scaling of the inference pipeline while keeping the database as a StatefulSet.")
discussion_q(2, "What is the maximum theoretical container density on a DGX Spark "
               "given its 128 GB memory and 11 services?")
discussion_q(3, "How would you implement circuit-breaker patterns between the agent "
               "and NIM LLM service?")

doc.add_page_break()

# ==============================================================
# CHAPTER 6
# ==============================================================

H1("Chapter 6 \u2014 PostgreSQL + pgvector")

H2("Schema Design Principles")
body("The agent\u2019s schema follows Third Normal Form (3NF) with selective "
     "denormalization:")
bullet("studies and series tables:", "normalized (series FK \u2192 studies)")
bullet("findings table:", "denormalized details JSONB column for workflow-specific "
       "data (avoids schema changes per workflow)")
bullet("measurements table:", "normalized (FK \u2192 findings), enabling cross-workflow "
       "measurement queries")
bullet("embeddings table:", "FK \u2192 studies, FK \u2192 findings (nullable), with "
       "vector(384) column")

H2("Full DDL")
code_block("CREATE EXTENSION IF NOT EXISTS vector;\n\n"
           "CREATE TABLE studies (\n"
           "    id                  SERIAL PRIMARY KEY,\n"
           "    study_instance_uid  TEXT UNIQUE NOT NULL,\n"
           "    patient_id          TEXT NOT NULL,\n"
           "    patient_name        TEXT,\n"
           "    study_date          DATE NOT NULL,\n"
           "    study_description   TEXT,\n"
           "    modality            TEXT NOT NULL,\n"
           "    accession_number    TEXT,\n"
           "    referring_physician TEXT,\n"
           "    body_part           TEXT,\n"
           "    num_series          INT DEFAULT 0,\n"
           "    num_instances       INT DEFAULT 0,\n"
           "    orthanc_id          TEXT,\n"
           "    status              TEXT DEFAULT 'received',\n"
           "    created_at          TIMESTAMPTZ DEFAULT NOW(),\n"
           "    updated_at          TIMESTAMPTZ DEFAULT NOW()\n"
           ");\n\n"
           "CREATE TABLE series (\n"
           "    id                  SERIAL PRIMARY KEY,\n"
           "    series_instance_uid TEXT UNIQUE NOT NULL,\n"
           "    study_id            INT REFERENCES studies(id) ON DELETE CASCADE,\n"
           "    series_number       INT,\n"
           "    series_description  TEXT,\n"
           "    modality            TEXT NOT NULL,\n"
           "    num_instances       INT DEFAULT 0,\n"
           "    created_at          TIMESTAMPTZ DEFAULT NOW()\n"
           ");\n\n"
           "CREATE TABLE findings (\n"
           "    id              SERIAL PRIMARY KEY,\n"
           "    study_id        INT REFERENCES studies(id) ON DELETE CASCADE,\n"
           "    workflow         TEXT NOT NULL,\n"
           "    finding_type    TEXT NOT NULL,\n"
           "    finding_code    TEXT,\n"
           "    location        TEXT,\n"
           "    laterality      TEXT,\n"
           "    severity        TEXT,\n"
           "    confidence      FLOAT NOT NULL CHECK (confidence >= 0 AND confidence <= 1),\n"
           "    is_positive     BOOLEAN DEFAULT TRUE,\n"
           "    details         JSONB DEFAULT '{}',\n"
           "    created_at      TIMESTAMPTZ DEFAULT NOW()\n"
           ");\n\n"
           "CREATE TABLE measurements (\n"
           "    id              SERIAL PRIMARY KEY,\n"
           "    finding_id      INT REFERENCES findings(id) ON DELETE CASCADE,\n"
           "    measurement_type TEXT NOT NULL,\n"
           "    value           FLOAT NOT NULL,\n"
           "    unit            TEXT NOT NULL,\n"
           "    reference_range TEXT,\n"
           "    flag            TEXT,\n"
           "    prior_value     FLOAT,\n"
           "    prior_date      DATE,\n"
           "    delta_percent   FLOAT,\n"
           "    created_at      TIMESTAMPTZ DEFAULT NOW()\n"
           ");\n\n"
           "CREATE TABLE embeddings (\n"
           "    id              SERIAL PRIMARY KEY,\n"
           "    study_id        INT REFERENCES studies(id) ON DELETE CASCADE,\n"
           "    finding_id      INT REFERENCES findings(id) ON DELETE SET NULL,\n"
           "    level           TEXT NOT NULL,\n"
           "    model_name      TEXT NOT NULL,\n"
           "    embedding       vector(384) NOT NULL,\n"
           "    metadata        JSONB DEFAULT '{}',\n"
           "    created_at      TIMESTAMPTZ DEFAULT NOW()\n"
           ");\n\n"
           "CREATE TABLE provenance (\n"
           "    id              SERIAL PRIMARY KEY,\n"
           "    study_id        INT REFERENCES studies(id) ON DELETE CASCADE,\n"
           "    workflow         TEXT NOT NULL,\n"
           "    model_id        TEXT NOT NULL,\n"
           "    model_version   TEXT NOT NULL,\n"
           "    model_arch      TEXT,\n"
           "    inference_params JSONB DEFAULT '{}',\n"
           "    input_uids      TEXT[] DEFAULT '{}',\n"
           "    duration_ms     INT,\n"
           "    gpu_memory_mb   INT,\n"
           "    status          TEXT DEFAULT 'completed',\n"
           "    error_message   TEXT,\n"
           "    created_at      TIMESTAMPTZ DEFAULT NOW()\n"
           ");\n\n"
           "CREATE TABLE worklist_entries (\n"
           "    id              SERIAL PRIMARY KEY,\n"
           "    study_id        INT REFERENCES studies(id) ON DELETE CASCADE,\n"
           "    finding_id      INT REFERENCES findings(id) ON DELETE SET NULL,\n"
           "    urgency         TEXT NOT NULL,\n"
           "    priority        TEXT NOT NULL,\n"
           "    notification    TEXT,\n"
           "    routing         TEXT,\n"
           "    acknowledged    BOOLEAN DEFAULT FALSE,\n"
           "    acknowledged_by TEXT,\n"
           "    acknowledged_at TIMESTAMPTZ,\n"
           "    created_at      TIMESTAMPTZ DEFAULT NOW()\n"
           ");", "sql")

H2("pgvector Internals")
body("pgvector stores vectors as fixed-length arrays of float4 (32-bit). The "
     "vector(384) column occupies 384 \u00d7 4 + 8 = 1,544 bytes per row (plus "
     "PostgreSQL tuple header overhead).")
body("Distance operators:")
bullet("<->", "L2 distance (Euclidean)")
bullet("<=>", "Cosine distance (1 - cosine_similarity)")
bullet("<#>", "Negative inner product")
impl_note("The agent uses cosine distance (<=>) for all similarity queries, "
          "consistent with BiomedCLIP\u2019s training objective.")

H2("HNSW Index Tuning")
code_block("CREATE INDEX idx_embeddings_hnsw ON embeddings\n"
           "    USING hnsw (embedding vector_cosine_ops)\n"
           "    WITH (m = 16, ef_construction = 64);", "sql")
bullet("m = 16:", "Maximum number of bi-directional links per node per layer. "
       "Higher m \u2192 better recall, larger index, slower inserts. Range 4-64.")
bullet("ef_construction = 64:", "Size of the dynamic candidate list during index "
       "build. Higher \u2192 better index quality, slower build time. Must be \u2265 2m.")
bullet("ef_search (runtime):", "SET hnsw.ef_search = 40; Controls query candidate "
       "list size. Higher \u2192 better recall, slower queries.")
body("IVFFlat is the alternative: faster to build than HNSW but lower recall at "
     "the same query speed. HNSW is preferred for the agent\u2019s workload (moderate "
     "dataset size, high recall requirement).")

H2("Hybrid Query Pattern")
code_block("-- CTE filters by SQL predicates, then vector search on filtered set\n"
           "WITH candidates AS (\n"
           "    SELECT e.study_id, e.embedding\n"
           "    FROM embeddings e\n"
           "    JOIN studies s ON e.study_id = s.id\n"
           "    WHERE e.level = 'study'\n"
           "      AND s.modality = 'CT'\n"
           "      AND s.body_part = 'CHEST'\n"
           ")\n"
           "SELECT c.study_id, c.embedding <=> $1::vector AS distance\n"
           "FROM candidates c\n"
           "ORDER BY c.embedding <=> $1::vector\n"
           "LIMIT 10;", "sql")

H2("Query Plan Analysis")
code_block("EXPLAIN (ANALYZE, BUFFERS, FORMAT TEXT)\n"
           "SELECT * FROM embeddings\n"
           "ORDER BY embedding <=> '[0.1, 0.2, ...]'::vector\n"
           "LIMIT 10;", "sql")
body("Expected plan: Index Scan using idx_embeddings_hnsw with bounded heap scans. "
     "If the plan shows Seq Scan, verify the HNSW index exists and hnsw.ef_search "
     "is set.")

H3("Discussion Questions \u2014 Chapter 6")
discussion_q(1, "What are the trade-offs of storing workflow-specific data in a JSONB "
               "column versus dedicated columns per workflow?")
discussion_q(2, "Calculate the approximate disk space required for 1 million embeddings "
               "at 384 dimensions.")
discussion_q(3, "Design a schema migration strategy that supports adding a new workflow "
               "without altering existing tables.")

doc.add_page_break()

# ==============================================================
# CHAPTER 7
# ==============================================================

H1("Chapter 7 \u2014 Clinical Workflow Implementation")

H2("MONAI Deploy MAP Lifecycle")
body("A MAP (MONAI Deploy Application Package) follows this execution flow:")
numbered(1, "Input: DICOM files mounted at /var/holoscan/input/")
numbered(2, "compose(): Defines the operator DAG (directed acyclic graph)")
numbered(3, "Operators execute sequentially: Each receives InputContext, produces OutputContext")
numbered(4, "Output: Results written to /var/holoscan/output/")
code_block("@resource(cpu=4, gpu=1, memory=\"16Gi\")\n"
           "class WorkflowApp(Application):\n"
           "    def compose(self):\n"
           "        self.add_flow(PreprocessOp(), InferenceOp())\n"
           "        self.add_flow(InferenceOp(), PostprocessOp())", "python")

H2("CXR Rapid Findings \u2014 Implementation Detail")
body_rich([
    ("Preprocessing: ", True, False),
    ("Load CXR DICOM \u2192 resize to 224\u00d7224 (DenseNet-121 input) \u2192 "
     "normalize to [0,1] \u2192 ensure float32.", False, False),
])
body_rich([
    ("DenseNet-121 architecture detail: ", True, False),
    ("4 dense blocks \u2192 global average pooling \u2192 FC(num_features, 5) with "
     "sigmoid activation (not softmax \u2014 multi-label, not multi-class). Each "
     "output is independent \u2014 a patient can have both pneumothorax AND pleural "
     "effusion.", False, False),
])
body_rich([
    ("Threshold calibration: ", True, False),
    ("Thresholds per class are set at the operating point that maximizes the Youden "
     "index (sensitivity + specificity - 1) on the validation set. Pneumothorax has "
     "a lower threshold (0.50) because the clinical cost of a false negative is "
     "higher than a false positive.", False, False),
])
H3("GradCAM Implementation")
code_block("cam = GradCAM(nn_module=model, target_layers=\"class_layers.relu\")\n"
           "heatmap = cam(x=input_tensor, class_idx=target_class)", "python")
body("GradCAM computes: (1) forward pass to get activations at target layer, "
     "(2) backward pass to get gradients of target class score w.r.t. activations, "
     "(3) global average pool the gradients, (4) weighted sum of activations, "
     "(5) ReLU to keep only positive contributions.")

H2("CT Head Hemorrhage \u2014 Implementation Detail")
body_rich([
    ("Two-stage pipeline: ", True, False),
    ("Classification gate (DenseNet-121) \u2192 conditional segmentation (3D U-Net). "
     "The classification gate avoids running the expensive 3D segmentation on "
     "normal scans.", False, False),
])
H3("Volume Calculation")
code_block("volume_ml = np.sum(segmentation > 0) * np.prod(voxel_spacing) / 1000.0",
           "python")
body("Where voxel_spacing is in mm (from DICOM PixelSpacing and SpacingBetweenSlices), "
     "and division by 1000 converts mm\u00b3 to mL.")

H3("Midline Shift")
code_block("shift_mm = abs((center_of_mass[0] - axial_center) * voxel_spacing[0])",
           "python")

H3("Brain Trauma Foundation Decision Tree")
code_block("if volume_ml > 30 or shift_mm > 5 or thickness_mm > 10:\n"
           "    return (\"critical\", \"P1\")\n"
           "elif volume_ml > 5:\n"
           "    return (\"urgent\", \"P2\")\n"
           "else:\n"
           "    return (\"routine\", \"P4\")", "python")

H2("CT Chest Lung Nodule \u2014 Implementation Detail")
H3("VDT Derivation")
body("Assuming exponential growth: V(t) = V\u2080 \u00d7 2^(t/VDT). Solving for VDT:")
code_block("V2/V1 = 2^(\u0394t/VDT)\n"
           "ln(V2/V1) = (\u0394t/VDT) \u00d7 ln(2)\n"
           "VDT = (\u0394t \u00d7 ln(2)) / ln(V2/V1)")
body("VDT < 400 days is the commonly used threshold for suspicious growth rate "
     "(corresponds to approximately a 26% volume increase per year for a typical "
     "screening interval).")
body_rich([
    ("Lung-RADS v2022 decision matrix: ", True, False),
    ("The agent implements the full decision matrix including solid, ground-glass, "
     "and part-solid categories with different size thresholds for each. "
     "VDT < 400 upgrades the category.", False, False),
])
body_rich([
    ("Genomics trigger: ", True, False),
    ("Lung-RADS 4B or 4X \u2192 check_genomics_trigger() \u2192 POST to Nextflow "
     "API \u2192 Parabricks pipeline initiated.", False, False),
])

H2("MRI Brain MS Lesion \u2014 Implementation Detail")
body_rich([
    ("Preprocessing differences from CT: ", True, False),
    ("MRI requires Z-score normalization (NormalizeIntensityd) rather than HU "
     "windowing. N4 bias field correction (corrects spatial intensity inhomogeneity "
     "from RF coil sensitivity) is critical for accurate segmentation.", False, False),
])
H3("Registration with ANTsPy")
code_block("result = ants.registration(fixed=current, moving=prior,\n"
           "                           type_of_transform=\"SyNRA\")", "python")
body("SyNRA = Rigid + Affine + SyN (symmetric normalization diffeomorphic "
     "registration). The diffeomorphic component handles non-linear brain deformation. "
     "Registration quality directly affects lesion matching accuracy.")
body_rich([
    ("Lesion matching: ", True, False),
    ("After warping prior lesion masks to current space, overlap (Dice coefficient) "
     "> 0.3 between a current and prior lesion indicates a match. New lesion = "
     "current lesion with no matching prior. Enlarging lesion = matched but volume "
     "increase > 20%.", False, False),
])

H3("Discussion Questions \u2014 Chapter 7")
discussion_q(1, "Why is sigmoid activation used instead of softmax for CXR "
               "multi-label classification?")
discussion_q(2, "Derive the relationship between VDT and annual volume growth rate.")
discussion_q(3, "What registration failure modes could cause false-positive "
               "\u201cnew lesion\u201d detections in MS tracking?")

doc.add_page_break()

# ==============================================================
# CHAPTER 8
# ==============================================================

H1("Chapter 8 \u2014 LangGraph Agent Architecture")

H2("StateGraph Internals")
body("LangGraph\u2019s StateGraph compiles to a directed graph where:")
bullet("State channels:", "Each field in AgentState is a channel. Channels can have "
       "reducers that merge values from parallel branches.")
bullet("Nodes:", "Functions f(state) \u2192 partial_state that return updates to "
       "specific channels.")
bullet("Edges:", "Unconditional (always follow) or conditional (function evaluates "
       "state to choose next node).")
code_block("from langgraph.graph import StateGraph, END\n"
           "from typing import TypedDict, Annotated\n\n"
           "class AgentState(TypedDict):\n"
           "    study_id: int\n"
           "    findings: list[dict]\n"
           "    prior_studies: list[dict]\n"
           "    similar_cases: list[dict]\n"
           "    severity: str\n"
           "    report: str\n"
           "    provenance: dict\n\n"
           "graph = StateGraph(AgentState)\n"
           "graph.add_node(\"triage\", triage_node)\n"
           "graph.add_node(\"longitudinal\", longitudinal_node)\n"
           "graph.add_node(\"population\", population_node)\n"
           "graph.add_node(\"report\", report_node)\n\n"
           "graph.set_entry_point(\"triage\")\n"
           "graph.add_conditional_edges(\"triage\", route_by_severity,\n"
           "    {\"full_pipeline\": \"longitudinal\", \"brief_report\": \"report\"})\n"
           "graph.add_edge(\"longitudinal\", \"population\")\n"
           "graph.add_edge(\"population\", \"report\")\n"
           "graph.add_edge(\"report\", END)\n\n"
           "app = graph.compile()", "python")

H2("Conditional Routing Logic")
code_block("def route_by_severity(state: AgentState) -> str:\n"
           "    if state[\"severity\"] in (\"critical\", \"urgent\"):\n"
           "        return \"full_pipeline\"\n"
           "    return \"brief_report\"", "python")

H2("Tool Binding with MCP")
body("Tools are declared with Pydantic schemas and bound via the Model Context Protocol:")
code_block("from langchain_core.tools import tool\n"
           "from pydantic import BaseModel, Field\n\n"
           "class SimilarStudyInput(BaseModel):\n"
           "    embedding: list[float] = Field(description=\"384-dim query vector\")\n"
           "    modality: str = Field(description=\"CT, MR, CR, DX\")\n"
           "    limit: int = Field(default=10, description=\"Max results\")\n\n"
           "@tool(args_schema=SimilarStudyInput)\n"
           "def search_similar_studies(embedding, modality, limit=10):\n"
           "    \"\"\"Search for studies with similar imaging characteristics.\"\"\"\n"
           "    return db.hybrid_similarity_search(embedding, modality, limit)",
           "python")
body("MCP publishes a tool manifest (name, description, parameters JSON schema) that "
     "any compatible agent runtime can discover and invoke.")

H2("Checkpointing")
body("LangGraph supports state checkpointing for:")
bullet("Resumability:", "If the agent fails mid-graph, resume from the last checkpoint")
bullet("Human-in-the-loop:", "Pause at a node, present state to user, await approval "
       "before continuing")
bullet("Debugging:", "Replay graph execution from any checkpoint")
body("Backends: SQLite (proof build), PostgreSQL (production \u2014 reuses the "
     "agent\u2019s existing database).")

H2("Agent Personas")
body("Each agent node uses a tailored system prompt:")
code_block("TRIAGE_SYSTEM_PROMPT = \"\"\"You are a triage radiologist assistant.\n"
           "Your role: assess finding severity and determine routing.\n"
           "Use Brain Trauma Foundation guidelines for hemorrhage.\n"
           "Use Lung-RADS v2022 for lung nodules.\n"
           "Output: severity classification and recommended analysis path.\"\"\"",
           "python")
body("Different prompts for triage vs. longitudinal (compare to priors, calculate "
     "deltas) vs. population (interpret cohort data, outcome statistics) vs. report "
     "(synthesize all evidence into structured clinical narrative).")

H3("Discussion Questions \u2014 Chapter 8")
discussion_q(1, "Design a fan-out/fan-in pattern where longitudinal and population "
               "nodes execute in parallel and their results are merged before the "
               "report node.")
discussion_q(2, "What are the trade-offs of using PostgreSQL vs. Redis as a "
               "LangGraph checkpoint backend?")
discussion_q(3, "How would you implement rate limiting for LLM tool calls to avoid "
               "overwhelming the NIM service?")

doc.add_page_break()

# ==============================================================
# CHAPTER 9
# ==============================================================

H1("Chapter 9 \u2014 Embedding Models and Vector Retrieval")

H2("Contrastive Learning and CLIP")
body("CLIP (Contrastive Language-Image Pretraining) learns a joint embedding space "
     "for images and text. The training objective minimizes the InfoNCE loss:")
code_block("L = -log(exp(sim(I_i, T_i)/\u03c4) / \u03a3_j exp(sim(I_i, T_j)/\u03c4))")
body("Where sim() is cosine similarity between image embedding I and text embedding T, "
     "\u03c4 is a learnable temperature parameter, and the sum in the denominator runs "
     "over all image-text pairs in the batch. Matched pairs (I_i, T_i) are pulled "
     "together; mismatched pairs are pushed apart.")

H2("BiomedCLIP Architecture")
body_rich([
    ("Reference: ", True, False),
    ("Zhang et al., \u201cBiomedCLIP: A Multimodal Biomedical Foundation Model "
     "Pretrained from Fifteen Million Scientific Image-Text Pairs,\u201d 2023.", False, True),
])
bullet("Text encoder:", "PubMedBERT (BERT pretrained on PubMed abstracts)")
bullet("Image encoder:", "ViT-B/16 (Vision Transformer, base variant, 16\u00d716 patch size)")
bullet("Projection:", "Both encoders output vectors projected to 384-dimensional shared space")
body("The 384-dim output comes from the projection head design. The image encoder\u2019s "
     "CLS token (768-dim for ViT-B) is projected to 384-dim via a linear layer. This "
     "dimensionality balances representation capacity against storage and search efficiency.")

H2("Vision Transformer (ViT) Detail")
numbered(1, "Patch embedding: Image (224\u00d7224) split into 14\u00d714 = 196 patches "
           "of 16\u00d716 pixels. Each patch flattened and linearly projected to "
           "embedding dimension d=768.")
numbered(2, "Positional encoding: Learnable 1D position embeddings added to "
           "patch embeddings.")
numbered(3, "CLS token: A learnable token prepended to the sequence. Its output "
           "serves as the image representation.")
numbered(4, "Transformer blocks (12 for ViT-B): Each block applies multi-head "
           "self-attention (MHSA) and feed-forward network (FFN) with layer normalization.")
code_block("z'_l = MHSA(LN(z_{l-1})) + z_{l-1}\nz_l = FFN(LN(z'_l)) + z'_l")
H3("Scaled Dot-Product Attention")
code_block("Attention(Q, K, V) = softmax(QK^T / \u221ad_k) \u00d7 V")
body("Where Q, K, V are query, key, value matrices derived from linear projections "
     "of the input. d_k is the key dimension. The scaling by \u221ad_k prevents the "
     "dot products from growing large in magnitude, which would push the softmax "
     "into saturated regions.")

H2("Three-Level Embedding Strategy")
numbered(1, "Study-level: Mean-pool all series-level embeddings for the study. "
           "Used for broad case-matching.")
numbered(2, "Series-level: For each series, sample representative slices, encode each, "
           "mean-pool. Used for protocol-specific matching.")
numbered(3, "Lesion-level: Crop ROI around individual findings, resize to 224\u00d7224, "
           "encode. Used for finding-specific matching.")

H2("Vector Quantization")
body_rich([
    ("Product Quantization (PQ): ", True, False),
    ("Split 384-dim vector into m subvectors (e.g., m=48, 8 dims each). Quantize each "
     "subvector to its nearest centroid from a learned codebook of k centroids "
     "(typically k=256). Store only the centroid indices (48 bytes vs. 1,536 bytes "
     "for the original). Distance approximated by summing precomputed sub-distances.", False, False),
])
body_rich([
    ("HNSW + PQ: ", True, False),
    ("Build the HNSW graph on quantized vectors for faster construction and smaller "
     "memory footprint, with optional reranking using exact vectors for the top "
     "candidates.", False, False),
])
impl_note("At the agent\u2019s expected scale (tens of thousands to low millions of "
          "embeddings), exact HNSW with full vectors is sufficient. PQ becomes "
          "relevant at Phase 3+ (multi-site, millions of studies).")

H2("Embedding Drift")
body("If the patient population or imaging protocols change over time, the embedding "
     "distribution shifts. Monitoring strategies:")
bullet_plain("Track mean embedding per week/month \u2014 large shifts indicate "
             "distribution change")
bullet_plain("Monitor retrieval recall on a held-out validation set")
bullet_plain("Periodic reindexing when the distribution shift exceeds a threshold")

H3("Discussion Questions \u2014 Chapter 9")
discussion_q(1, "Derive the storage and memory savings of PQ with m=48, k=256 "
               "compared to exact vectors at 384 dimensions.")
discussion_q(2, "Why does BiomedCLIP use 384-dim embeddings instead of the "
               "encoder\u2019s native 768 dimensions?")
discussion_q(3, "How would you implement an embedding update pipeline that "
               "recomputes embeddings when a new model version is deployed "
               "without downtime?")

doc.add_page_break()

# ==============================================================
# CHAPTER 10
# ==============================================================

H1("Chapter 10 \u2014 Clinical Interoperability")

H2("FHIR R4 Resource Model")
body("The agent creates a FHIR Bundle containing DiagnosticReport and Observation resources:")
code_block("{\n"
           "  \"resourceType\": \"Bundle\",\n"
           "  \"type\": \"transaction\",\n"
           "  \"entry\": [\n"
           "    { \"resource\": { \"resourceType\": \"DiagnosticReport\", ... } },\n"
           "    { \"resource\": { \"resourceType\": \"Observation\", ... } }\n"
           "  ]\n"
           "}", "json")

H3("DiagnosticReport Structure")
bullet("status:", "\"final\" (post-review) or \"preliminary\" (pre-review)")
bullet("category:", "LOINC 18726-0 (Radiology studies)")
bullet("code:", "LOINC code for the specific exam type")
bullet("subject:", "Reference to Patient resource")
bullet("imagingStudy:", "Reference to ImagingStudy (maps to DICOM Study UID)")
bullet("result:", "Array of references to Observation resources")
bullet("conclusion:", "Free-text summary (LLM-generated)")
bullet("presentedForm:", "Base64-encoded PDF/HTML report attachment")

H3("Observation Structure")
bullet("code:", "LOINC code for the measurement type")
bullet("valueQuantity:", "{\"value\": 25.3, \"unit\": \"mL\", \"system\": "
       "\"http://unitsofmeasure.org\"}")
bullet("interpretation:", "Coded clinical significance (e.g., high, critical)")
bullet("component:", "Multi-value observations (e.g., nodule with separate diameter, "
       "volume, VDT components)")

H2("Terminology Binding")
body_rich([
    ("SNOMED CT ", True, False),
    ("hierarchy: Concepts have IS-A relationships (e.g., \u201cSubdural hemorrhage\u201d "
     "IS-A \u201cIntracranial hemorrhage\u201d IS-A \u201cHemorrhage\u201d). The agent "
     "codes findings at the most specific level available.", False, False),
])
body_rich([
    ("LOINC: ", True, False),
    ("Structured along axes: Component/Property/Time/System/Scale/Method. Example: "
     "\u201cVolume of hemorrhage in brain\u201d = Component(hemorrhage volume) + "
     "System(brain) + Scale(quantitative).", False, False),
])

H2("HL7v2 Integration (Legacy Pathway)")
body("Many hospitals still run HL7v2 (pipe-delimited messages). Common messages:")
bullet("ADT^A01 (Admit):", "New patient registration \u2192 creates Patient resource")
bullet("ORM^O01 (Order):", "Imaging order placed \u2192 creates ServiceRequest")
bullet("ORU^R01 (Result):", "Radiology report \u2014 the agent\u2019s findings "
       "could be wrapped here")
body("Integration engines (Mirth Connect, Rhapsody) translate between HL7v2 and FHIR. "
     "A FHIR fa\u00e7ade pattern exposes a FHIR API while converting to/from HL7v2 "
     "behind the scenes.")

H2("IHE AI Results Profile")
body("The Integrating the Healthcare Enterprise (IHE) organization publishes "
     "integration profiles. The IHE-AI Results supplement defines how AI findings "
     "should be communicated to PACS, EHR, and worklist systems \u2014 including "
     "provenance, confidence scores, and links to evidence images. The agent\u2019s "
     "architecture aligns with this profile.")

H2("SMART on FHIR Security")
bullet("OAuth2 authorization:", "with FHIR-specific scopes "
       "(e.g., patient/Observation.read)")
bullet("Launch context:", "which patient, which encounter")
bullet("Token-based access control:", "machine-to-machine OAuth2 flow for the "
       "agent\u2019s FHIR publisher")

H3("Discussion Questions \u2014 Chapter 10")
discussion_q(1, "Design the FHIR Observation resources needed to represent a lung "
               "nodule with diameter, volume, VDT, and Lung-RADS score as components.")
discussion_q(2, "How would you implement a FHIR fa\u00e7ade over the agent\u2019s "
               "PostgreSQL database to enable external EHR systems to query findings?")
discussion_q(3, "What IHE profiles are relevant for automated AI-to-PACS worklist "
               "prioritization?")

doc.add_page_break()

# ==============================================================
# CHAPTER 11
# ==============================================================

H1("Chapter 11 \u2014 RAG Architecture and LLM Serving")

H2("Transformer Architecture")
body("The transformer (Vaswani et al., 2017) processes sequences using self-attention:")
code_block("MultiHead(Q,K,V) = Concat(head_1, ..., head_h) \u00d7 W^O\n"
           "where head_i = Attention(QW_i^Q, KW_i^K, VW_i^V)")
body("Each transformer block: LayerNorm \u2192 MHSA \u2192 Residual \u2192 "
     "LayerNorm \u2192 FFN \u2192 Residual.")

H2("Llama 3 8B Architecture")
bullet("32 transformer layers", "")
bullet("32 attention heads, 8 KV heads", "(Grouped Query Attention \u2014 GQA)")
bullet("Hidden dimension:", "4096")
bullet("FFN dimension:", "14336 (SwiGLU activation)")
bullet("RoPE:", "Rotary Position Embedding for positional encoding")
bullet("Context window:", "8192 tokens")
bullet("Vocabulary:", "128K tokens (byte-level BPE)")
body_rich([
    ("GQA: ", True, False),
    ("Instead of separate KV projections per head, GQA groups 4 query heads per KV "
     "head (32 Q / 8 KV). This reduces KV cache memory by 4\u00d7 with minimal "
     "quality loss \u2014 critical for inference efficiency.", False, False),
])
body_rich([
    ("RoPE: ", True, False),
    ("Encodes position through rotation matrices applied to Q and K vectors. The "
     "rotation angle is proportional to position, enabling relative position encoding "
     "that generalizes to unseen sequence lengths.", False, False),
])

H2("NVIDIA NIM Serving")
bullet("Model loading:", "Loads weights into GPU memory, initializes KV cache")
bullet("Continuous batching:", "Dynamically adds/removes requests from the batch as "
       "they arrive/complete \u2014 maximizes GPU utilization")
bullet("KV cache management:", "Paged attention (inspired by vLLM) \u2014 KV cache "
       "allocated in pages, reducing fragmentation")
bullet("Speculative decoding:", "Small draft model generates candidate tokens; "
       "large model verifies in parallel \u2014 can accelerate generation 2-3\u00d7")
code_block("image: nvcr.io/nvidia/nim/meta-llama3-8b-instruct:latest-dgx-spark",
           "yaml")

H2("RAG Pipeline Design")
body("The agent\u2019s RAG pipeline follows a retriever \u2192 generator pattern:")
numbered(1, "Dense retrieval: Query embeddings search pgvector for similar studies")
numbered(2, "SQL retrieval: Direct database queries for structured findings, "
           "measurements, provenance")
numbered(3, "Evidence assembly: All retrieved evidence formatted into a structured prompt")
numbered(4, "Generation: NIM-served Llama 3 generates the report with evidence citations")

H3("Context Window Budget (8192 tokens)")
bullet("System prompt:", "~500 tokens")
bullet("Current findings:", "~300-500 tokens")
bullet("Prior measurements:", "~200-400 tokens")
bullet("Similar cases:", "~500-1000 tokens (top 5 cases, abbreviated)")
bullet("Guidelines:", "~500-800 tokens")
bullet("Generation headroom:", "~5000 tokens for output")

H2("Prompt Engineering")
code_block("You are a radiology AI assistant generating clinical reports.\n"
           "RULES:\n"
           "1. Only include facts supported by the evidence below.\n"
           "2. Never add information not present in the evidence.\n"
           "3. Cite specific measurements with units.\n"
           "4. Flag critical findings prominently.\n"
           "5. Compare to prior studies when available.\n"
           "6. Reference applicable clinical guidelines.\n\n"
           "EVIDENCE:\n"
           "{evidence_block}\n\n"
           "Generate a structured clinical report with sections:\n"
           "1. Clinical Indication\n"
           "2. Findings\n"
           "3. Measurements\n"
           "4. Comparison to Prior\n"
           "5. Impression\n"
           "6. Recommendation")

H2("Evaluation with RAGAS")
bullet("Faithfulness:", "What fraction of claims in the generated report are supported "
       "by the retrieved context?")
bullet("Answer relevancy:", "Does the report address the clinical question?")
bullet("Context recall:", "Did the retriever find all relevant evidence?")

H3("Discussion Questions \u2014 Chapter 11")
discussion_q(1, "Calculate the KV cache memory requirement for Llama 3 8B at full "
               "context length (8192 tokens) with GQA.")
discussion_q(2, "Design a hybrid retrieval strategy combining BM25 keyword search "
               "with dense embedding search using Reciprocal Rank Fusion.")
discussion_q(3, "How would you detect and prevent hallucinated measurements in "
               "the generated report?")

doc.add_page_break()

# ==============================================================
# CHAPTER 12
# ==============================================================

H1("Chapter 12 \u2014 Observability Stack")

H2("Prometheus Architecture")
body("Prometheus operates a pull model \u2014 it scrapes HTTP endpoints at "
     "configured intervals:")
code_block("# prometheus.yml\nscrape_configs:\n"
           "  - job_name: 'imaging-agent'\n"
           "    static_configs:\n"
           "      - targets: ['agent:8000']\n"
           "    scrape_interval: 15s\n"
           "  - job_name: 'dcgm'\n"
           "    static_configs:\n"
           "      - targets: ['dcgm-exporter:9400']\n"
           "    scrape_interval: 15s", "yaml")
body("Time series database (TSDB) stores samples as (timestamp, value) pairs, "
     "indexed by metric name and label set. Retention configured at 30 days.")

H2("Metric Types")
add_table(
    ["Type", "Description", "Example"],
    [
        ["Counter", "Monotonically increasing", "studies_processed_total"],
        ["Gauge", "Arbitrary value (up/down)", "gpu_utilization_percent"],
        ["Histogram", "Distribution in configurable buckets", "inference_duration_seconds"],
        ["Summary", "Distribution with configurable quantiles", "request_duration_seconds"],
    ],
)
spacer(4)
code_block("from prometheus_client import Counter, Histogram, Gauge\n\n"
           "STUDIES_TOTAL = Counter('studies_processed_total',\n"
           "    'Total studies processed', ['workflow'])\n"
           "INFERENCE_DURATION = Histogram('inference_duration_seconds',\n"
           "    'Inference time', ['workflow'],\n"
           "    buckets=[0.5, 1, 2, 5, 10, 30, 60, 120])\n"
           "GPU_UTIL = Gauge('gpu_utilization_percent',\n"
           "    'Current GPU utilization')", "python")

H2("PromQL Queries")
code_block("# Per-workflow inference latency (p95 over 5 minutes)\n"
           "histogram_quantile(0.95, rate(inference_duration_seconds_bucket[5m]))\n\n"
           "# Studies processed per hour\n"
           "rate(studies_processed_total[1h]) * 3600\n\n"
           "# GPU utilization (from DCGM)\n"
           "DCGM_FI_DEV_GPU_UTIL{gpu=\"0\"}", "promql")

H2("DCGM Telemetry Fields")
add_table(
    ["Field", "Description"],
    [
        ["DCGM_FI_DEV_GPU_UTIL", "GPU compute utilization (%)"],
        ["DCGM_FI_DEV_MEM_COPY_UTIL", "Memory controller utilization (%)"],
        ["DCGM_FI_DEV_GPU_TEMP", "GPU temperature (\u00b0C)"],
        ["DCGM_FI_DEV_POWER_USAGE", "Power consumption (W)"],
        ["DCGM_FI_DEV_FB_FREE", "Free framebuffer memory (MiB)"],
        ["DCGM_FI_DEV_FB_USED", "Used framebuffer memory (MiB)"],
        ["DCGM_FI_DEV_ECC_SBE_VOL", "Single-bit ECC errors (volatile)"],
        ["DCGM_FI_DEV_XID_ERRORS", "XID error count"],
    ],
)
spacer(4)

H2("OpenTelemetry Integration")
bullet("Spans:", "Timed operations (e.g., \"dicom_listener.process_study\", "
       "\"agent.triage_node\")")
bullet("Trace context propagation:", "W3C traceparent header passed between HTTP calls")
bullet("Exporters:", "Send traces to Jaeger, Zipkin, or Grafana Tempo")

H2("SLA Definition")
add_table(
    ["Workflow", "p50 Target", "p95 Target", "p99 Target"],
    [
        ["CXR Rapid Findings", "< 10s", "< 20s", "< 30s"],
        ["CT Head Hemorrhage", "< 45s", "< 75s", "< 90s"],
        ["CT Chest Lung Nodule", "< 2min", "< 4min", "< 5min"],
        ["MRI Brain MS Lesion", "< 2min", "< 4min", "< 5min"],
    ],
)
spacer(4)

H3("Discussion Questions \u2014 Chapter 12")
discussion_q(1, "Design a Grafana alerting rule that fires when the p95 inference "
               "latency exceeds the SLA target for any workflow.")
discussion_q(2, "What DCGM metrics would you monitor to predict GPU hardware "
               "failure before it occurs?")
discussion_q(3, "How would you implement distributed tracing across the agent\u2019s "
               "11 containers without modifying application code?")

doc.add_page_break()

# ==============================================================
# CHAPTER 13
# ==============================================================

H1("Chapter 13 \u2014 HCLS AI Factory")

H2("Multi-Agent Coordination")
body("The AI Factory agents communicate through:")
bullet("Shared PostgreSQL database:", "Common schema for cross-agent queries")
bullet("FHIR ServiceRequest:", "Trigger messages between agents (imaging \u2192 genomics)")
bullet("Event bus:", "Study completion events, finding alerts, genomics results "
       "(Kafka or Redis Streams at scale)")

H2("Parabricks Pipeline")
body("NVIDIA Parabricks GPU-accelerates genomics:")
bullet("BWA-MEM2:", "Read alignment to reference genome \u2014 GPU kernels for "
       "Smith-Waterman alignment")
bullet("HaplotypeCaller / DeepVariant:", "Variant calling (identify SNPs, indels, "
       "structural variants)")
bullet("Performance:", "30\u00d7 whole-genome sequencing in ~10 minutes on 8\u00d7A100 "
       "GPUs (vs. ~30 hours CPU)")
body("The imaging-to-genomics trigger:")
code_block("Lung nodule Lung-RADS 4B+ \u2192 FHIR ServiceRequest \u2192 "
           "Nextflow pipeline \u2192 Parabricks \u2192 VCF output \u2192 "
           "Variant annotation")

H2("Federated Learning: FedAvg Algorithm")
code_block("For each communication round t:\n"
           "    Server sends global model w_t to selected clients\n"
           "    Each client k:\n"
           "        w_k^{t+1} = w_t - \u03b7\u2207L_k(w_t)  "
           "(local training for E epochs)\n"
           "    Server aggregates:\n"
           "        w_{t+1} = \u03a3_k (n_k/n) \u00d7 w_k^{t+1}  "
           "(weighted average by local dataset size)")
body("NVIDIA FLARE implements this with:")
bullet("Scatter-and-Gather workflow:", "Server scatters model to clients, clients "
       "train locally, server gathers and aggregates")
bullet("Privacy modules:", "Differential privacy (add calibrated noise to gradients), "
       "secure aggregation (cryptographic protocols preventing server from seeing "
       "individual updates)")
bullet("Model filters:", "Client-side filters that clip gradients, quantize updates, "
       "or apply sparsification before transmission")

H2("RECIST 1.1 Criteria")
bullet("Complete Response (CR):", "Disappearance of all target lesions")
bullet("Partial Response (PR):", "\u226530% decrease in sum of longest diameters")
bullet("Progressive Disease (PD):", "\u226520% increase in sum of longest diameters "
       "AND \u22655mm absolute increase")
bullet("Stable Disease (SD):", "Neither PR nor PD")
impl_note("The imaging agent\u2019s automated measurements can serve as RECIST "
          "endpoints in clinical trials \u2014 replacing manual radiologist "
          "measurements with reproducible, automated quantification.")

H2("Nextflow DSL2 Pipeline Orchestration")
code_block("// main.nf\nnextflow.enable.dsl = 2\n\n"
           "include { CT_HEAD_HEMORRHAGE } from './modules/ct_head_hemorrhage'\n"
           "include { CT_CHEST_LUNG_NODULE } from './modules/ct_chest_lung_nodule'\n"
           "include { CXR_RAPID_FINDINGS } from './modules/cxr_rapid_findings'\n"
           "include { MRI_BRAIN_MS_LESION } from './modules/mri_brain_ms_lesion'\n\n"
           "workflow {\n"
           "    study_ch = Channel.fromPath(params.input_dir)\n\n"
           "    // Route by modality and body part\n"
           "    study_ch.branch {\n"
           "        ct_head: it.modality == 'CT' && it.body_part == 'HEAD'\n"
           "        ct_chest: it.modality == 'CT' && it.body_part == 'CHEST'\n"
           "        cxr: it.modality in ['CR', 'DX']\n"
           "        mri_brain: it.modality == 'MR' && it.body_part == 'BRAIN'\n"
           "    }.set { routed }\n\n"
           "    CT_HEAD_HEMORRHAGE(routed.ct_head)\n"
           "    CT_CHEST_LUNG_NODULE(routed.ct_chest)\n"
           "    CXR_RAPID_FINDINGS(routed.cxr)\n"
           "    MRI_BRAIN_MS_LESION(routed.mri_brain)\n"
           "}", "groovy")

H2("Scaling Analysis")
add_table(
    ["GPU Tier", "Studies/Hour (CXR)", "Studies/Hour (CT Head)", "Concurrent Pipelines"],
    [
        ["DGX Spark (1 GPU)", "~120", "~40", "1"],
        ["DGX B200 (8 GPU)", "~960", "~320", "8"],
        ["DGX SuperPOD (256 GPU)", "~30,000", "~10,000", "256"],
    ],
)
spacer(4)
body("Estimates assume single-study-per-GPU sequential processing. Throughput "
     "increases linearly with GPU count for embarrassingly parallel study processing.")

H3("Discussion Questions \u2014 Chapter 13")
discussion_q(1, "Design a differential privacy mechanism for the hemorrhage detection "
               "model using NVIDIA FLARE that provides \u03b5=1 privacy guarantee.")
discussion_q(2, "How would you implement RECIST 1.1 automated measurement with the "
               "lung nodule workflow?")
discussion_q(3, "Calculate the total compute cost (hardware + NVAIE license) for a "
               "Phase 3 multi-site deployment processing 500 studies/hour across "
               "4 hospitals.")

doc.add_page_break()

# ==============================================================
# CHAPTER 14
# ==============================================================

H1("Chapter 14 \u2014 Regulatory, Safety, and Deployment Engineering")

H2("FDA SaMD Risk Categorization")
body("The FDA categorizes SaMD based on two axes:")
add_table(
    ["Condition", "Treat/Diagnose", "Drive Management", "Inform"],
    [
        ["Critical", "IV (highest)", "III", "II"],
        ["Serious", "III", "II", "I (lowest)"],
        ["Non-serious", "II", "I", "I"],
    ],
)
spacer(4)
body("The imaging agent as triage decision support for serious conditions "
     "(hemorrhage, cancer) falls in Category II-III depending on the specific claim.")

H2("Predetermined Change Control Plan (PCCP)")
body("A PCCP consists of:")
body_rich([
    ("SaMD Pre-Specifications (SPS): ", True, False),
    ("Describe the types of changes anticipated:", False, False),
])
bullet_plain("\"The model may be retrained on expanded datasets including new "
             "scanner manufacturers\"")
bullet_plain("\"Confidence thresholds may be adjusted based on real-world "
             "performance data\"")
bullet_plain("\"New finding types may be added to the CXR workflow\"")
body_rich([
    ("Algorithm Change Protocol (ACP): ", True, False),
    ("Define the validation methodology:", False, False),
])
bullet_plain("\"Any retrained model must achieve sensitivity \u226595% and "
             "specificity \u226590% on the fixed validation set\"")
bullet_plain("\"Threshold changes must be validated on a held-out calibration "
             "set with \u22651000 positive cases\"")
bullet_plain("\"Changes exceeding the SPS scope require a new regulatory submission\"")

H2("Clinical Validation Methodology")
add_table(
    ["Metric", "Formula", "Interpretation"],
    [
        ["Sensitivity", "TP / (TP + FN)", "Fraction of true positives detected"],
        ["Specificity", "TN / (TN + FP)", "Fraction of true negatives correctly identified"],
        ["PPV", "TP / (TP + FP)", "Fraction of positive predictions that are correct"],
        ["NPV", "TN / (TN + FN)", "Fraction of negative predictions that are correct"],
        ["AUC-ROC", "Area under ROC curve", "Overall discrimination ability"],
        ["AUC-PR", "Area under precision-recall curve",
         "Performance on imbalanced datasets (preferred for rare findings)"],
    ],
)
spacer(4)
body("AUC-PR is preferred over AUC-ROC for medical imaging because class imbalance "
     "is extreme (e.g., <1% of CXRs have pneumothorax). AUC-ROC can appear favorable "
     "even with many false positives when the negative class is large.")

H2("Bias Assessment")
body("Demographic subgroup analysis: compute all metrics stratified by age group, "
     "sex, race/ethnicity, scanner manufacturer, and site. Report performance gaps "
     "with 95% confidence intervals.")
body("Fairness metrics:")
bullet("Equalized odds:", "P(\u0176=1|Y=1, A=a) = P(\u0176=1|Y=1, A=b) for all "
       "groups a, b (equal sensitivity)")
bullet("Demographic parity:", "P(\u0176=1|A=a) = P(\u0176=1|A=b) (equal positive "
       "prediction rate)")
bullet("Calibration:", "P(Y=1|\u0176=p) = p for all groups (confidence scores are "
       "accurate across groups)")

H2("IEC 62304 Software Lifecycle")
body("Medical device software development standard requiring:")
bullet_plain("Software development plan (architecture, design, testing strategy)")
bullet_plain("Software requirements specification (functional, performance, safety)")
bullet_plain("Software architecture design (modules, interfaces, data flow)")
bullet_plain("Unit testing, integration testing, system testing (with traceability "
             "to requirements)")
bullet_plain("Software maintenance plan (bug fixes, updates, monitoring)")
body("Classification: Class A (no injury), B (non-serious injury), C (death or "
     "serious injury). The agent\u2019s triage function is Class C.")

H2("Deployment Patterns")
body_rich([
    ("Shadow mode (parallel run): ", True, False),
    ("New model processes same studies as production model. Results compared offline. "
     "No impact on clinical workflow. Validates performance before promotion.", False, False),
])
body_rich([
    ("Canary deployment: ", True, False),
    ("New model serves a small percentage of studies (e.g., 5%). Monitor error rates "
     "and latency. Gradually increase if metrics are stable.", False, False),
])
body_rich([
    ("Blue-green deployment: ", True, False),
    ("Two identical environments. Switch traffic from blue (current) to green (new) "
     "atomically. Instant rollback by switching back.", False, False),
])

H2("Responsible AI")
H3("Explainability Methods")
bullet("GradCAM:", "Where is the model looking? (spatial explanation)")
bullet("SHAP:", "Which input features contribute most to the prediction? "
       "(feature attribution)")
bullet("Integrated Gradients:", "Axiomatic attribution method \u2014 satisfies "
       "sensitivity and implementation invariance")

H3("Uncertainty Quantification")
bullet("MC Dropout:", "Run inference N times with different dropout masks, compute "
       "prediction variance")
bullet("Deep Ensemble:", "Train N independent models, compute prediction variance "
       "across ensemble")
bullet("Epistemic vs. aleatoric:", "Epistemic uncertainty (model uncertainty) vs. "
       "aleatoric uncertainty (data noise)")

H3("Discussion Questions \u2014 Chapter 14")
discussion_q(1, "Draft a PCCP SPS and ACP for the CT head hemorrhage workflow that "
               "allows threshold adjustment and model retraining without a new 510(k).")
discussion_q(2, "How would you design a demographic bias audit for the CXR workflow "
               "given that training data underrepresents certain age groups?")
discussion_q(3, "Compare MC Dropout vs. Deep Ensemble for uncertainty quantification "
               "in terms of inference cost on the DGX Spark.")

doc.add_page_break()

# ==============================================================
# COMPREHENSIVE DISCUSSION QUESTIONS
# ==============================================================

H1("Discussion Questions \u2014 Comprehensive")
body("These 20 questions require synthesis across multiple chapters.", before=4, after=8)

discussion_q(1, "Design an end-to-end data flow from DICOM C-STORE receipt through "
               "LangGraph agent processing to FHIR DiagnosticReport output. Identify "
               "every container involved and every database write.")
discussion_q(2, "Calculate the total GPU memory budget for running the hemorrhage "
               "3D U-Net, the NIM LLM (Llama 3 8B), and the BiomedCLIP embedding "
               "model concurrently on the DGX Spark\u2019s 128 GB unified memory.")
discussion_q(3, "Design a federated learning protocol using NVIDIA FLARE that "
               "improves hemorrhage detection across 5 hospitals while preserving "
               "\u03b5=1 differential privacy. What FLARE components would you configure?")
discussion_q(4, "How would you modify the Lung-RADS classification algorithm to "
               "handle the case where prior study registration fails? What fallback "
               "strategy preserves clinical safety?")
discussion_q(5, "Propose a schema migration that adds a new \"cardiac_function\" "
               "workflow without modifying existing tables or breaking existing queries.")
discussion_q(6, "Design a Grafana dashboard for the hemorrhage workflow showing: "
               "p95 latency, throughput, CRITICAL finding rate, GPU utilization, "
               "and NIM token throughput.")
discussion_q(7, "Compare the annotation cost of training data for classification "
               "(DenseNet-121) vs. segmentation (3D U-Net) vs. detection (RetinaNet). "
               "How does this affect model iteration speed?")
discussion_q(8, "Design an A/B test protocol that compares radiologist performance "
               "with and without the agent\u2019s triage assistance while maintaining "
               "patient safety.")
discussion_q(9, "How would you implement a \"second opinion\" pattern where two "
               "independent models must agree on a CRITICAL classification before "
               "triggering an alert?")
discussion_q(10, "Calculate the cosine similarity search latency for 1 million "
                "embeddings using HNSW with m=16, ef_search=100. Compare to IVFFlat "
                "with lists=1000, probes=20.")
discussion_q(11, "Design the FHIR resources and terminology bindings needed to "
                "represent a longitudinal MS lesion tracking report with new, "
                "enlarging, and stable lesion counts.")
discussion_q(12, "How would you extend the RAG pipeline to include genomic variant "
                "data from Parabricks in the evidence context for a lung cancer "
                "staging report?")
discussion_q(13, "Propose a TensorRT optimization strategy for the agent\u2019s "
                "models. Which models benefit most from INT8 quantization? Which "
                "require FP16 minimum?")
discussion_q(14, "Design a container health check cascade that detects and "
                "automatically recovers from a NIM LLM out-of-memory error without "
                "losing in-flight studies.")
discussion_q(15, "How would the agent\u2019s architecture change if deployed on a "
                "DGX B200 (8 GPUs) instead of DGX Spark (1 GPU)? What parallelism "
                "strategies become available?")
discussion_q(16, "Design a provenance query that reconstructs the complete processing "
                "history for a specific patient across all modalities and workflows "
                "over 2 years.")
discussion_q(17, "How would you implement real-time embedding drift detection and "
                "automated reindexing without disrupting similarity search availability?")
discussion_q(18, "Compare the security implications of SMART on FHIR vs. API key "
                "authentication for the agent\u2019s FHIR publisher in a multi-tenant "
                "hospital deployment.")
discussion_q(19, "Design a Nextflow pipeline that implements the cross-modal trigger: "
                "imaging Lung-RADS 4B+ \u2192 genomics \u2192 combined report. Include "
                "error handling for Parabricks pipeline failure.")
discussion_q(20, "Estimate the total cost of ownership (hardware + NVAIE + operational) "
                "for a Phase 2 departmental deployment processing 200 studies/hour "
                "across all 4 workflows.")

doc.add_page_break()

# ==============================================================
# REFERENCES
# ==============================================================

H1("References")

H2("Deep Learning Architectures")
bullet_plain("Huang et al., \u201cDensely Connected Convolutional Networks,\u201d CVPR 2017")
bullet_plain("Ronneberger et al., \u201cU-Net: Convolutional Networks for Biomedical "
             "Image Segmentation,\u201d MICCAI 2015")
bullet_plain("\u00c7i\u00e7ek et al., \u201c3D U-Net: Learning Dense Volumetric "
             "Segmentation from Sparse Annotation,\u201d MICCAI 2016")
bullet_plain("Myronenko, \u201c3D MRI Brain Tumor Segmentation Using Autoencoder "
             "Regularization,\u201d BrainLes@MICCAI 2018")
bullet_plain("Lin et al., \u201cFocal Loss for Dense Object Detection,\u201d ICCV 2017")
bullet_plain("Lin et al., \u201cFeature Pyramid Networks for Object Detection,\u201d "
             "CVPR 2017")
bullet_plain("Hatamizadeh et al., \u201cSwin UNETR: Swin Transformers for Semantic "
             "Segmentation of Brain Tumors,\u201d BrainLes@MICCAI 2021")

H2("Foundation Models")
bullet_plain("Radford et al., \u201cLearning Transferable Visual Models From Natural "
             "Language Supervision,\u201d ICML 2021 (CLIP)")
bullet_plain("Zhang et al., \u201cBiomedCLIP: A Multimodal Biomedical Foundation "
             "Model,\u201d 2023")
bullet_plain("Dosovitskiy et al., \u201cAn Image is Worth 16x16 Words: Transformers "
             "for Image Recognition at Scale,\u201d ICLR 2021 (ViT)")
bullet_plain("Vaswani et al., \u201cAttention Is All You Need,\u201d NeurIPS 2017")
bullet_plain("Touvron et al., \u201cLlama 2: Open Foundation and Fine-Tuned Chat "
             "Models,\u201d 2023")
bullet_plain("Meta AI, \u201cLlama 3 Model Card,\u201d 2024")

H2("Medical Imaging Standards")
bullet_plain("DICOM Standard PS3.3: Information Object Definitions")
bullet_plain("DICOM Standard PS3.4: Service Class Specifications")
bullet_plain("DICOM Standard PS3.5: Data Structures and Encoding")
bullet_plain("DICOM Standard PS3.18: Web Services (DICOMweb)")
bullet_plain("HL7 FHIR R4 Specification (hl7.org/fhir/R4)")
bullet_plain("IHE Radiology Technical Framework")

H2("Clinical Guidelines")
bullet_plain("ACR Lung-RADS v2022 Assessment Categories")
bullet_plain("Brain Trauma Foundation, \u201cGuidelines for the Management of Severe "
             "Traumatic Brain Injury,\u201d 4th Edition")
bullet_plain("Thompson et al., \u201cDiagnosis of Multiple Sclerosis: 2017 Revisions "
             "of the McDonald Criteria,\u201d Lancet Neurology 2018")
bullet_plain("RECIST 1.1: Eisenhauer et al., \u201cNew Response Evaluation Criteria "
             "in Solid Tumours,\u201d European Journal of Cancer 2009")

H2("Regulatory")
bullet_plain("FDA, \u201cAI/ML-Based Software as a Medical Device (SaMD) Action "
             "Plan,\u201d 2021")
bullet_plain("FDA, \u201cMarketing Submission Recommendations for a Predetermined "
             "Change Control Plan for AI/ML-Enabled Device Software Functions,\u201d 2023")
bullet_plain("IEC 62304: Medical Device Software \u2014 Software Life Cycle Processes")
bullet_plain("ISO 14971: Medical Devices \u2014 Application of Risk Management")
bullet_plain("IEC 62366: Medical Devices \u2014 Usability Engineering")

H2("NVIDIA Platforms")
bullet_plain("NVIDIA DGX Spark Technical Specifications")
bullet_plain("NVIDIA MONAI Deploy SDK Documentation")
bullet_plain("NVIDIA NIM Documentation")
bullet_plain("NVIDIA Parabricks Documentation")
bullet_plain("NVIDIA FLARE Documentation (Apache 2.0)")
bullet_plain("NVIDIA BioNeMo Service Documentation")

spacer(20)
P("\u2500" * 60, size=8, color=GRAY_META, align=WD_ALIGN_PARAGRAPH.CENTER)
spacer(4)
body_rich([
    ("This document was created for the HCLS AI Factory \u2014 Imaging Intelligence "
     "Agent.", False, True),
])
body_rich([
    ("Apache 2.0 License | Author: Adam Jones | February 2026", False, True),
])

# ==============================================================
# SAVE
# ==============================================================

OUT = "core/engines/clinical-imaging/open_public/HCLS_Imaging_AI_Agent_Learning_Guide_Advanced.docx"
doc.save(OUT)
print(f"Saved: {OUT}")
