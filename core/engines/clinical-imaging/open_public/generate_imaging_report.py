#!/usr/bin/env python3
# Copyright 2026 Adam Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Generate HCLS Imaging Intelligence Report for Patient IMG-2026-0142."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Color palette ──────────────────────────────────────────
NAVY = RGBColor(0x1B, 0x23, 0x33)
TEAL = RGBColor(0x1A, 0xAF, 0xCC)
GREEN = RGBColor(0x76, 0xB9, 0x00)
GRAY_TEXT = RGBColor(0x33, 0x33, 0x33)
GRAY_SUB = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED_FLAG = RGBColor(0xDC, 0x26, 0x26)
AMBER_FLAG = RGBColor(0xF5, 0xA6, 0x23)
GREEN_FLAG = RGBColor(0x05, 0x96, 0x69)
GREEN_BRIGHT = RGBColor(0x76, 0xB9, 0x00)

HEX_TEAL = "1AAFCC"
HEX_GREEN = "76B900"
HEX_NAVY = "1B2333"
HEX_LIGHT_BG = "F8FAFB"
HEX_PIPELINE_BG = "F5F5F5"
HEX_YELLOW_BG = "FFF8E1"
HEX_RED_BG = "FEE2E2"
HEX_AMBER_BG = "FFF3CD"
HEX_GREEN_BG = "D1FAE5"

FONT = "Calibri"

doc = Document()

# ── Page setup ─────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
style.font.color.rgb = GRAY_TEXT
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)


# ── Helper functions ───────────────────────────────────────
def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="D0D0D0", sz="4"):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def add_run(paragraph, text, bold=False, italic=False, size=None, color=None, font_name=FONT):
    run = paragraph.add_run(text)
    run.font.name = font_name
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_paragraph(text="", bold=False, italic=False, size=10.5, color=GRAY_TEXT,
                  space_before=0, space_after=4, alignment=None):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def add_section_header(number, text):
    p = add_paragraph(space_before=24, space_after=8)
    add_run(p, f"{number}. ", bold=True, size=18, color=NAVY)
    add_run(p, text.upper(), bold=True, size=18, color=NAVY)
    return p


def add_sub_header(text):
    return add_paragraph(text, bold=True, size=13, color=TEAL, space_before=12, space_after=4)


def add_body(text, space_before=0, space_after=4):
    return add_paragraph(text, size=10.5, color=GRAY_TEXT, space_before=space_before, space_after=space_after)


def add_bullet(lead, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3)
    add_run(p, "\u25CF ", size=10, color=GREEN)
    add_run(p, lead, bold=True, size=10.5, color=GRAY_TEXT)
    add_run(p, f"  {body}", size=10.5, color=GRAY_TEXT)
    return p


def add_styled_table(headers, rows, col_widths=None, flag_col=None):
    """Create a table with teal headers, alternating rows, and optional flag coloring."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEX_TEAL)
        set_cell_borders(cell, HEX_TEAL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=9, color=WHITE)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        bg = "FFFFFF" if r_idx % 2 == 0 else HEX_LIGHT_BG
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_shading(cell, bg)
            set_cell_borders(cell, "E0E0E0")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            # First column bold teal
            if c_idx == 0:
                add_run(p, str(val), bold=True, size=9, color=TEAL)
            elif flag_col is not None and c_idx == flag_col:
                # Color-code flag cells
                flag_val = str(val).upper()
                if flag_val in ("CRITICAL", "HIGH", "URGENT", "POSITIVE", "ACTIVE", "MALIGNANT"):
                    set_cell_shading(cell, HEX_RED_BG)
                    add_run(p, str(val), bold=True, size=9, color=RED_FLAG)
                elif flag_val in ("ELEVATED", "WATCH", "MODERATE", "ENLARGED", "NEW", "SUSPICIOUS", "4A", "4B"):
                    set_cell_shading(cell, HEX_AMBER_BG)
                    add_run(p, str(val), bold=True, size=9, color=AMBER_FLAG)
                elif flag_val in ("NORMAL", "STABLE", "LOW", "ROUTINE", "NEGATIVE", "BENIGN", "1", "2", "3"):
                    set_cell_shading(cell, HEX_GREEN_BG)
                    add_run(p, str(val), bold=True, size=9, color=GREEN_FLAG)
                else:
                    add_run(p, str(val), size=9, color=GRAY_TEXT)
            else:
                add_run(p, str(val), size=9, color=GRAY_TEXT)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def add_info_table(data):
    """Two-column info box with alternating backgrounds."""
    table = doc.add_table(rows=len(data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(data):
        bg = "FFFFFF" if i % 2 == 0 else HEX_LIGHT_BG
        for j, cell in enumerate(table.rows[i].cells):
            set_cell_shading(cell, bg)
            set_cell_borders(cell, "E0E0E0")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if j == 0:
                add_run(p, label, bold=True, size=9.5, color=TEAL)
            else:
                add_run(p, value, size=9.5, color=GRAY_TEXT)
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.5)
    return table


# ══════════════════════════════════════════════════════════════
# REPORT CONTENT
# ══════════════════════════════════════════════════════════════

# ── Green accent bar ─────────────────────────────────────────
accent = doc.add_table(rows=1, cols=1)
accent.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = accent.rows[0].cells[0]
set_cell_shading(cell, HEX_GREEN)
p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
add_run(p, " ", size=28, color=GREEN)
accent.rows[0].height = Inches(0.6)

# ── Phase pipeline ───────────────────────────────────────────
add_paragraph("", space_after=4)
phases = ["ACQUISITION", "DETECTION", "ANALYSIS", "INTELLIGENCE"]
phase_desc = [
    "DICOM Ingestion\nfrom PACS/VNA",
    "AI Model Inference\nAcross Modalities",
    "Segmentation &\nClassification",
    "Longitudinal\nComparison & RAG"
]

pipeline = doc.add_table(rows=3, cols=len(phases))
pipeline.alignment = WD_TABLE_ALIGNMENT.CENTER
for i in range(len(phases)):
    # Header row
    cell_h = pipeline.rows[0].cells[i]
    set_cell_shading(cell_h, HEX_PIPELINE_BG)
    set_cell_borders(cell_h, "D0D0D0")
    p = cell_h.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"PHASE {i+1}", bold=True, size=8, color=GRAY_SUB)

    # Phase name row
    cell_p = pipeline.rows[1].cells[i]
    set_cell_shading(cell_p, HEX_GREEN)
    set_cell_borders(cell_p, HEX_GREEN)
    p = cell_p.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, phases[i], bold=True, size=13, color=WHITE)

    # Description row
    cell_d = pipeline.rows[2].cells[i]
    set_cell_shading(cell_d, "FFFFFF")
    set_cell_borders(cell_d, "D0D0D0")
    p = cell_d.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, phase_desc[i], size=7.5, color=GRAY_SUB)

# ── Title block ──────────────────────────────────────────────
add_paragraph("", space_after=4)
add_paragraph("IMAGING INTELLIGENCE", bold=True, size=36, color=NAVY,
              space_before=4, space_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph("REPORT", bold=True, size=36, color=NAVY,
              space_before=0, space_after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph("AI Factory Pipeline Report",
              italic=True, size=16, color=GRAY_SUB, space_after=8,
              alignment=WD_ALIGN_PARAGRAPH.CENTER)

# ── 3-cell summary header ────────────────────────────────────
summary_bar = doc.add_table(rows=1, cols=3)
summary_bar.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, value) in enumerate([
    ("Patient:", "IMG-2026-0142"),
    ("Studies:", "4 across 3 modalities"),
    ("Generated:", "February 1, 2026"),
]):
    cell = summary_bar.rows[0].cells[i]
    set_cell_borders(cell, "D0D0D0")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, label + " ", bold=True, size=10, color=NAVY)
    add_run(p, value, size=10, color=GRAY_TEXT)

add_paragraph("", space_after=4)

# ── Patient info box ─────────────────────────────────────────
add_info_table([
    ("Patient ID", "IMG-2026-0142"),
    ("Age / Sex", "58 years / Female"),
    ("Date of Birth", "March 14, 1967"),
    ("Referring Physician", "Dr. Sarah Chen, MD — Emergency Medicine"),
    ("Clinical Indication", "Acute-onset severe headache; known MS (natalizumab); lung nodule follow-up"),
    ("Studies Analyzed", "4 studies across 3 modalities (CT Head, CT Chest, CXR, MRI Brain FLAIR)"),
    ("Prior Studies Retrieved", "CT Chest (01/28/2025), MRI Brain FLAIR (07/15/2025)"),
    ("Report Generated", "February 1, 2026 at 14:23:07 UTC"),
    ("Pipeline Version", "imaging-agent-v1.2.0 (MONAI Deploy MAP 0.6.0)"),
    ("Classification", "For Clinical Decision Support — Clinician Review Required"),
])

# ══════════════════════════════════════════════════════════════
# SECTION 1 — Patient & Study Summary
# ══════════════════════════════════════════════════════════════
add_section_header(1, "Patient & Study Summary")
add_body(
    "This report presents AI-assisted imaging analysis for patient IMG-2026-0142, a 58-year-old female "
    "who presented to the emergency department with acute-onset severe headache. The patient has a known "
    "history of relapsing-remitting multiple sclerosis (RRMS), currently managed with natalizumab, and an "
    "8mm right upper lobe lung nodule identified on CT chest 12 months prior. A multi-modality workup was "
    "ordered including non-contrast CT head, CT chest with contrast, admission chest X-ray, and MRI brain "
    "with FLAIR sequences. All four studies were processed through the Imaging Intelligence Agent pipeline "
    "using MONAI Deploy application packages on NVIDIA DGX Spark GPU infrastructure. Prior studies were "
    "automatically retrieved for longitudinal comparison. Each workflow below reports detection results, "
    "segmentation volumes, classification scores, and clinical triage recommendations with full provenance."
)

# ══════════════════════════════════════════════════════════════
# SECTION 2 — Executive Findings
# ══════════════════════════════════════════════════════════════
add_section_header(2, "Executive Findings — Priority Signals")
add_body("The following findings were identified as the highest-priority clinical signals across all four imaging studies:", space_after=6)

add_bullet("URGENT — Acute Subdural Hemorrhage:",
           "Left frontoparietal acute subdural hematoma measuring 12.3 mL (thickness 8.7 mm, midline shift "
           "3.2 mm). Both measurements are below Brain Trauma Foundation surgical thresholds (> 10 mm thickness, "
           "> 5 mm shift), but neurosurgical awareness and close monitoring with repeat CT are recommended. "
           "Escalate to CRITICAL if neurological decline or GCS drop. Urgency: URGENT.")
add_bullet("URGENT — Lung Nodule Growth:",
           "Right upper lobe solid nodule increased from 268 mm\u00B3 to 489 mm\u00B3 (82.5% volume increase) "
           "over 12 months. Volume doubling time: 287 days. Reclassified from Lung-RADS 3 to Lung-RADS 4A. "
           "Per ACR Lung-RADS v2022: 3-month follow-up LDCT recommended; PET-CT may be considered given solid "
           "component \u2265 8 mm.")
add_bullet("URGENT — New MS Lesions:",
           "Two new periventricular FLAIR-hyperintense lesions identified since prior MRI (July 2025). "
           "Total lesion volume increased from 4.8 mL to 6.1 mL. Disease activity reclassified from "
           "Stable to Active. Treatment escalation discussion recommended.")
add_bullet("MODERATE — Pleural Effusion:",
           "Small left-sided pleural effusion identified on CXR with 0.91 confidence. Confirmed on CT chest. "
           "Clinical correlation with cardiac and hepatic status recommended.")
add_bullet("ROUTINE — Cardiomegaly:",
           "Mild cardiomegaly detected on CXR (cardiothoracic ratio 0.54). Stable compared to prior imaging. "
           "Echocardiography recommended if not recently performed.")
add_bullet("ROUTINE — MS Lesion Burden:",
           "Total 14 white matter lesions (up from 12). Stable lesions show no interval change in volume. "
           "Continued surveillance per standard MS monitoring protocol.")

# ══════════════════════════════════════════════════════════════
# SECTION 3 — CT Head — Hemorrhage Triage
# ══════════════════════════════════════════════════════════════
add_section_header(3, "CT Head \u2014 Hemorrhage Triage")
add_body(
    "Non-contrast CT head was acquired and processed through the hemorrhage triage pipeline. The workflow "
    "includes 3D U-Net hemorrhage detection, volumetric segmentation, midline shift measurement, and "
    "urgency classification. End-to-end processing completed in 47 seconds.",
    space_after=8
)

add_sub_header("Hemorrhage Detection & Classification")
add_styled_table(
    ["Finding", "Value", "Units", "Reference", "Flag"],
    [
        ["Hemorrhage Detected", "Yes", "—", "—", "Critical"],
        ["Hemorrhage Type", "Acute Subdural", "—", "—", "Critical"],
        ["Location", "Left frontoparietal", "—", "—", "Critical"],
        ["Hemorrhage Volume", "12.3", "mL", "BTF: thickness > 10 mm or shift > 5 mm for surgical indication", "Critical"],
        ["Maximum Thickness", "8.7", "mm", "< 10 mm", "Elevated"],
        ["Midline Shift", "3.2", "mm", "< 5 mm", "Elevated"],
        ["Midline Shift Direction", "Rightward", "—", "—", "—"],
        ["Ventricular Compression", "Mild left lateral", "—", "None", "Elevated"],
        ["Detection Confidence", "0.97", "—", "> 0.90 threshold", "Normal"],
        ["Sensitivity (validated)", "> 95%", "—", "for hemorrhage > 5 mL", "Normal"],
    ],
    flag_col=4
)

add_paragraph("", space_after=4)
add_sub_header("Urgency Classification")
add_styled_table(
    ["Parameter", "Result", "Clinical Action"],
    [
        ["Urgency Level", "URGENT", "Neurosurgical awareness; close monitoring with repeat CT"],
        ["Worklist Priority", "P2 — Urgent", "Prioritized on radiologist worklist"],
        ["Notification", "Alert sent", "On-call neurosurgeon notified per policy"],
        ["Recommended Follow-up", "Repeat CT Head 6h", "Monitor for expansion; escalate to CRITICAL if neurological decline, GCS drop, or ICP elevation"],
    ],
    flag_col=1
)

add_paragraph("", space_after=4)
add_sub_header("Regional Analysis")
add_styled_table(
    ["Brain Region", "Status", "Volume (mL)", "Confidence"],
    [
        ["Left frontoparietal subdural space", "Acute hemorrhage", "12.3", "0.97"],
        ["Right hemisphere", "No hemorrhage", "—", "0.99"],
        ["Posterior fossa", "No hemorrhage", "—", "0.98"],
        ["Subarachnoid spaces", "No hemorrhage", "—", "0.96"],
        ["Intraventricular", "No hemorrhage", "—", "0.97"],
        ["Intraparenchymal", "No hemorrhage", "—", "0.98"],
    ],
    flag_col=1
)

# ══════════════════════════════════════════════════════════════
# SECTION 4 — CT Chest — Lung Nodule Tracking
# ══════════════════════════════════════════════════════════════
add_section_header(4, "CT Chest \u2014 Lung Nodule Tracking")
add_body(
    "CT chest with contrast was acquired for lung nodule surveillance. The pipeline includes nodule detection "
    "(RetinaNet), per-nodule segmentation (SegResNet), volumetric measurement, longitudinal registration to "
    "the prior CT chest from January 28, 2025, Lung-RADS category assignment, and volume doubling time "
    "calculation. End-to-end processing completed in 3 minutes 12 seconds.",
    space_after=8
)

add_sub_header("Nodule Inventory — Current Study (February 1, 2026)")
add_styled_table(
    ["Nodule ID", "Location", "Type", "Long Axis (mm)", "Short Axis (mm)", "Volume (mm\u00B3)", "Lung-RADS"],
    [
        ["N1", "RUL — apical segment", "Solid", "10.2", "8.4", "489", "4A"],
        ["N2", "RML — lateral segment", "Ground-glass", "5.1", "4.3", "42", "2"],
        ["N3", "LLL — posterior basal", "Part-solid", "6.8", "5.2", "118", "3"],
    ],
    flag_col=6
)

add_paragraph("", space_after=4)
add_sub_header("Longitudinal Comparison — Nodule N1 (Primary Target)")
add_styled_table(
    ["Measurement", "Prior (01/28/2025)", "Current (02/01/2026)", "Change", "Flag"],
    [
        ["Long Axis", "7.8 mm", "10.2 mm", "+30.8%", "Elevated"],
        ["Short Axis", "6.1 mm", "8.4 mm", "+37.7%", "Elevated"],
        ["Volume", "268 mm\u00B3", "489 mm\u00B3", "+82.5%", "Urgent"],
        ["Volume Doubling Time", "—", "287 days", "< 400 days", "Elevated"],
        ["Density (HU mean)", "42 HU", "48 HU", "+6 HU", "Watch"],
        ["Morphology", "Smooth margins", "Slightly spiculated", "Changed", "Elevated"],
        ["Lung-RADS (prior)", "3", "4A", "Upgraded", "Urgent"],
    ],
    flag_col=4
)

add_paragraph("", space_after=4)
add_sub_header("Lung-RADS Classification Summary")
add_styled_table(
    ["Category", "Definition", "Management", "Probability of Malignancy"],
    [
        ["Lung-RADS 4A (N1)", "Suspicious finding", "3-month follow-up LDCT; consider PET-CT if solid component \u2265 8 mm", "5\u201315%"],
        ["Lung-RADS 2 (N2)", "Benign appearance", "Continue annual LDCT", "< 1%"],
        ["Lung-RADS 3 (N3)", "Probably benign", "6-month follow-up CT", "1\u20132%"],
    ],
    flag_col=0
)

add_paragraph("", space_after=4)
add_sub_header("Malignancy Risk Scoring — Nodule N1")
add_styled_table(
    ["Risk Factor", "Value", "Weight", "Contribution"],
    [
        ["Volume doubling time < 400 days", "287 days", "High", "Major contributor"],
        ["Spiculated margin (new)", "Present", "High", "Major contributor"],
        ["Upper lobe location", "RUL apical", "Moderate", "Supportive"],
        ["Patient age > 55", "58 years", "Moderate", "Supportive"],
        ["Solid nodule type", "Solid", "Moderate", "Supportive"],
        ["Growth > 50% in 12 months", "82.5%", "High", "Major contributor"],
        ["Composite Risk Score", "0.73", "—", "High risk — 3-month LDCT per Lung-RADS 4A; PET-CT warranted given size"],
    ],
    flag_col=2
)

add_paragraph("", space_after=4)
add_body(
    "Additional finding: Small left-sided pleural effusion identified, measuring approximately 150 mL. "
    "No mediastinal lymphadenopathy. Heart size at upper limits of normal. No pericardial effusion.",
    space_after=4
)

# ══════════════════════════════════════════════════════════════
# SECTION 5 — CXR — Rapid Findings
# ══════════════════════════════════════════════════════════════
add_section_header(5, "CXR \u2014 Rapid Findings")
add_body(
    "Admission PA chest X-ray was processed through the CXR rapid findings pipeline. The workflow includes "
    "multi-label classification across five pathology categories, GradCAM heatmap localization for each "
    "positive finding, and confidence scoring. End-to-end processing completed in 8 seconds.",
    space_after=8
)

add_sub_header("Multi-Label Classification Results")
add_styled_table(
    ["Finding", "Detected", "Confidence", "GradCAM Region", "Clinical Significance", "Flag"],
    [
        ["Pneumothorax", "No", "0.02", "—", "—", "Normal"],
        ["Consolidation", "No", "0.08", "—", "—", "Normal"],
        ["Pleural Effusion", "Yes", "0.91", "Left costophrenic angle", "Small, left-sided", "Elevated"],
        ["Cardiomegaly", "Yes", "0.78", "Cardiac silhouette", "Mild (CTR 0.54)", "Watch"],
        ["Fracture", "No", "0.03", "—", "—", "Normal"],
    ],
    flag_col=5
)

add_paragraph("", space_after=4)
add_sub_header("GradCAM Localization Details")
add_body(
    "GradCAM activation maps were generated for both positive findings. The pleural effusion localization "
    "correctly highlights the left costophrenic angle with a peak activation of 0.94, showing meniscus sign "
    "blunting. The cardiomegaly localization highlights the left heart border with peak activation of 0.81. "
    "Heatmap overlays are available as GSPS annotations (Series 1.2.840.10008.5.1.4.1.1.11.1) and can be "
    "displayed in PACS as standard DICOM overlays.",
    space_after=4
)

add_sub_header("CXR Quality Metrics")
add_styled_table(
    ["Metric", "Value", "Threshold", "Status"],
    [
        ["Image Quality Score", "0.92", "> 0.80", "Normal"],
        ["Rotation Assessment", "Minimal (< 2\u00B0)", "< 5\u00B0", "Normal"],
        ["Inspiration Level", "Adequate (9 posterior ribs)", "> 8 ribs", "Normal"],
        ["Pneumothorax Sensitivity (validated)", "> 95%", "> 95% target", "Normal"],
        ["Processing Time", "8 seconds", "< 30 seconds target", "Normal"],
    ],
    flag_col=3
)

# ══════════════════════════════════════════════════════════════
# SECTION 6 — MRI Brain — MS Lesion Tracking
# ══════════════════════════════════════════════════════════════
add_section_header(6, "MRI Brain \u2014 MS Lesion Tracking")
add_body(
    "MRI brain with FLAIR sequences was acquired for routine MS surveillance. The pipeline includes "
    "3D U-Net lesion segmentation on FLAIR, lesion counting and volumetric measurement, spatial registration "
    "to the prior MRI from July 15, 2025, identification of new and enlarging lesions, and disease activity "
    "classification. End-to-end processing completed in 2 minutes 48 seconds.",
    space_after=8
)

add_sub_header("Lesion Summary — Current Study")
add_styled_table(
    ["Parameter", "Value", "Prior (07/15/2025)", "Change", "Flag"],
    [
        ["Total Lesion Count", "14", "12", "+2 new lesions", "Urgent"],
        ["Total Lesion Volume", "6.1 mL", "4.8 mL", "+1.3 mL (+27.1%)", "Elevated"],
        ["New Lesions", "2", "—", "—", "Urgent"],
        ["Enlarging Lesions", "0", "—", "No enlargement of existing", "Normal"],
        ["Stable Lesions", "12", "12", "No interval change", "Normal"],
        ["Enhancing Lesions", "1 (of 2 new)", "0", "+1", "Elevated"],
        ["Disease Activity", "Active", "Stable", "Reclassified", "Urgent"],
    ],
    flag_col=4
)

add_paragraph("", space_after=4)
add_sub_header("New Lesion Details")
add_styled_table(
    ["Lesion ID", "Location", "Volume (mm\u00B3)", "Enhancement", "Morphology", "Flag"],
    [
        ["L13", "Left periventricular — frontal horn", "420", "Enhancing (Gd+)", "Ovoid, well-defined", "New"],
        ["L14", "Right periventricular — body of lateral ventricle", "310", "Non-enhancing", "Ovoid, well-defined", "New"],
    ],
    flag_col=5
)

add_paragraph("", space_after=4)
add_sub_header("Lesion Distribution by Region")
add_styled_table(
    ["Brain Region", "Lesion Count", "Volume (mL)", "New Lesions", "Status"],
    [
        ["Periventricular", "8", "3.8", "2", "Active"],
        ["Juxtacortical", "3", "1.2", "0", "Stable"],
        ["Infratentorial", "1", "0.4", "0", "Stable"],
        ["Deep white matter", "2", "0.7", "0", "Stable"],
    ],
    flag_col=4
)

add_paragraph("", space_after=4)
add_sub_header("Disease Activity Assessment")
add_styled_table(
    ["Criterion", "Finding", "Classification"],
    [
        ["New T2/FLAIR lesions since prior", "2 new lesions identified", "Active"],
        ["Enlarging T2/FLAIR lesions", "None", "Stable"],
        ["Gadolinium-enhancing lesions", "1 enhancing lesion (L13)", "Active"],
        ["Overall Disease Activity", "Active (2 new, 1 enhancing)", "Active — treatment review recommended"],
        ["Trajectory (6 months)", "Stable \u2192 Active", "Escalation discussion warranted"],
    ],
    flag_col=2
)

add_body(
    "Clinical note: Patient is currently on natalizumab. The emergence of 2 new lesions (1 enhancing) "
    "over a 6-month interval while on disease-modifying therapy suggests breakthrough disease activity. "
    "Discussion regarding treatment escalation or switch is recommended per current MS treatment guidelines.",
    space_before=4, space_after=4
)

# ══════════════════════════════════════════════════════════════
# SECTION 7 — Longitudinal Imaging Summary
# ══════════════════════════════════════════════════════════════
add_section_header(7, "Longitudinal Imaging Summary")
add_body(
    "The following timeline summarizes key imaging events and changes across all modalities for this patient:",
    space_after=8
)

add_styled_table(
    ["Date", "Study", "Key Finding", "Action Taken", "Status"],
    [
        ["01/28/2025", "CT Chest", "8mm RUL solid nodule (268 mm\u00B3)", "Lung-RADS 3 — 6-month follow-up", "Baseline"],
        ["07/15/2025", "MRI Brain FLAIR", "12 WM lesions, 4.8 mL total volume", "Stable disease — continue natalizumab", "Stable"],
        ["02/01/2026", "CT Head (non-contrast)", "12.3 mL acute subdural hemorrhage (sub-threshold for surgery)", "URGENT — neurosurgical awareness + repeat CT 6h", "Urgent"],
        ["02/01/2026", "CXR (PA)", "Small left pleural effusion, mild cardiomegaly", "Clinical correlation recommended", "Moderate"],
        ["02/01/2026", "CT Chest (contrast)", "RUL nodule grown to 489 mm\u00B3 (Lung-RADS 4A)", "3-month LDCT follow-up; consider PET-CT", "Urgent"],
        ["02/01/2026", "MRI Brain FLAIR", "14 WM lesions (+2 new), 6.1 mL total", "Active disease — treatment review", "Urgent"],
    ],
    flag_col=4
)

# ══════════════════════════════════════════════════════════════
# SECTION 8 — Clinical Triage & Routing
# ══════════════════════════════════════════════════════════════
add_section_header(8, "Clinical Triage & Routing")
add_body(
    "Each study was independently triaged based on AI findings and mapped to institutional notification "
    "policy thresholds. The following table summarizes triage actions:",
    space_after=8
)

add_styled_table(
    ["Study", "Urgency", "Worklist Priority", "Notification", "Routing"],
    [
        ["CT Head", "URGENT", "P2 — Urgent", "Alert: neurosurgeon on call", "ED attending + neurosurgery"],
        ["CT Chest", "URGENT", "P2 — Urgent", "Alert: ordering physician", "Pulmonology / interventional radiology"],
        ["CXR", "MODERATE", "P3 — Routine+", "Standard reporting queue", "ED attending"],
        ["MRI Brain", "URGENT", "P2 — Urgent", "Alert: ordering physician", "Neurology / MS clinic"],
    ],
    flag_col=1
)

add_paragraph("", space_after=4)
add_sub_header("Escalation Summary")
add_body(
    "Three studies triggered urgent notification: CT Head (acute subdural hemorrhage — below Brain Trauma "
    "Foundation surgical thresholds but requiring neurosurgical awareness and close monitoring), CT Chest "
    "(suspicious nodule growth requiring short-interval follow-up), and MRI Brain (breakthrough MS activity "
    "requiring neurology review). The CXR findings are additive context and do not independently require "
    "escalation. CT Head should be escalated to CRITICAL if the patient develops neurological decline, "
    "GCS drop, or signs of raised intracranial pressure."
)

# ══════════════════════════════════════════════════════════════
# SECTION 9 — Cross-Modal Correlations
# ══════════════════════════════════════════════════════════════
add_section_header(9, "Cross-Modal Correlations")
add_body(
    "The Imaging Intelligence Agent identifies clinically relevant correlations across modalities to support "
    "integrated clinical reasoning:",
    space_after=8
)

add_styled_table(
    ["Correlation", "Modalities", "Clinical Significance", "Recommendation"],
    [
        ["Subdural hemorrhage + natalizumab",
         "CT Head + clinical history",
         "Natalizumab is not directly associated with hemorrhage risk, but immunosuppression context is relevant for surgical planning",
         "Ensure hematology awareness for perioperative management"],
        ["Lung nodule growth + immunosuppression",
         "CT Chest + clinical history",
         "Long-term immunosuppressive therapy (natalizumab) may modulate malignancy risk; impaired immune surveillance is a recognized concern",
         "Factor into risk-benefit discussion for follow-up imaging and potential tissue sampling"],
        ["Pleural effusion (CXR) + CT confirmation",
         "CXR + CT Chest",
         "CXR finding of left pleural effusion confirmed on CT chest as small (approx. 150 mL); no loculation",
         "Clinical correlation; thoracentesis not indicated at this volume"],
        ["Active MS + breakthrough on therapy",
         "MRI Brain + treatment history",
         "2 new lesions (1 enhancing) on natalizumab suggests breakthrough activity; PML risk must also be considered with JCV status",
         "JCV antibody testing if not recent; consider treatment switch discussion"],
    ],
    flag_col=2
)

# ══════════════════════════════════════════════════════════════
# SECTION 10 — Structured Output Encoding
# ══════════════════════════════════════════════════════════════
add_section_header(10, "Structured Output Encoding")
add_body(
    "All findings have been encoded in standard clinical formats for integration with PACS, EHR, and "
    "clinical workflow systems:",
    space_after=8
)

add_styled_table(
    ["Output Type", "Format", "Destination", "Content"],
    [
        ["Structured Report", "DICOM SR (TID 1500)", "PACS", "Complete findings, measurements, and classifications per study"],
        ["Visual Annotations", "GSPS + Secondary Capture", "PACS", "Graphic annotation contours (GSPS), GradCAM heatmap images (Secondary Capture)"],
        ["Segmentation Objects", "DICOM SEG", "PACS / Research", "Volumetric segmentation masks for hemorrhage, nodules, MS lesions"],
        ["Diagnostic Report", "FHIR DiagnosticReport (R4)", "EHR", "Narrative summary with coded findings (SNOMED CT, LOINC)"],
        ["Critical Alert", "HL7v2 ORU / FHIR Communication", "EHR + Paging", "CRITICAL finding notification for subdural hemorrhage"],
    ],
    flag_col=0
)

add_paragraph("", space_after=4)
add_sub_header("DICOM UID References")
add_styled_table(
    ["Study", "Study Instance UID", "Series Count"],
    [
        ["CT Head", "1.2.840.113619.2.55.3.2831164062.142.1706698800.1", "2 (axial, coronal reformat)"],
        ["CT Chest", "1.2.840.113619.2.55.3.2831164062.142.1706698800.2", "3 (axial, coronal, sagittal)"],
        ["CXR", "1.2.840.113619.2.55.3.2831164062.142.1706698800.3", "1 (PA)"],
        ["MRI Brain", "1.2.840.113619.2.55.3.2831164062.142.1706698800.4", "4 (FLAIR, T1, T1+Gd, T2)"],
    ],
    flag_col=0
)

# ══════════════════════════════════════════════════════════════
# SECTION 11 — Provenance & Reproducibility
# ══════════════════════════════════════════════════════════════
add_section_header(11, "Provenance & Reproducibility")
add_body(
    "Every output in this report is traceable to the exact model version, input data, and inference "
    "configuration used. This provenance chain supports regulatory compliance (FDA AI/ML SaMD framework), "
    "predetermined change control plans, and deterministic re-runs on canonical data.",
    space_after=8
)

add_sub_header("Model Registry")
add_styled_table(
    ["Workflow", "Model", "Architecture", "Version", "Training Data", "Validation AUC"],
    [
        ["CT Head Hemorrhage", "hemorrhage-triage-v2.1", "3D U-Net (MONAI)", "2.1.0", "RSNA ICH Dataset (752K slices)", "0.972"],
        ["CT Lung Nodule Detection", "nodule-detect-v1.8", "RetinaNet (MONAI)", "1.8.0", "LIDC-IDRI (1,018 CTs) + LUNA16 subset", "0.944"],
        ["CT Lung Nodule Segmentation", "nodule-seg-v1.5", "SegResNet (MONAI)", "1.5.0", "LIDC-IDRI (2,635 nodules)", "0.918 (Dice)"],
        ["CXR Multi-Label", "cxr-classify-v3.0", "DenseNet-121 (MONAI)", "3.0.0", "CheXpert + MIMIC-CXR (601K images)", "0.961"],
        ["MRI MS Lesion Seg", "ms-lesion-v2.3", "3D U-Net (MONAI)", "2.3.0", "ISBI MS Challenge + institutional (1,200 MRIs)", "0.936 (Dice)"],
    ],
    flag_col=0
)

add_paragraph("", space_after=4)
add_sub_header("Inference Configuration")
add_styled_table(
    ["Parameter", "Value"],
    [
        ["GPU Hardware", "NVIDIA DGX Spark (GB10 Grace Blackwell Superchip, 128 GB unified memory)"],
        ["CUDA Version", "12.6"],
        ["MONAI Deploy MAP Version", "0.6.0"],
        ["NVIDIA NIM Runtime", "nim-inference-v1.4.0"],
        ["Inference Precision", "FP16 (TensorRT optimized)"],
        ["Deterministic Mode", "Enabled (fixed seed: 42)"],
        ["DICOM Ingestion", "DICOMweb STOW-RS from institutional PACS"],
        ["Registration Algorithm", "SyN diffeomorphic (ANTsPy 0.4.2)"],
    ],
    flag_col=0
)

# ══════════════════════════════════════════════════════════════
# SECTION 12 — Computational Pipeline Performance
# ══════════════════════════════════════════════════════════════
add_section_header(12, "Computational Pipeline Performance")
add_body(
    "The following table reports end-to-end processing times for each workflow, measured from DICOM "
    "ingestion to structured output encoding:",
    space_after=8
)

add_styled_table(
    ["Workflow", "Target", "Actual", "DICOM Ingestion", "Inference", "Post-Processing", "Status"],
    [
        ["CT Head Hemorrhage Triage", "< 90 sec", "47 sec", "6 sec", "28 sec", "13 sec", "Normal"],
        ["CT Chest Nodule Tracking", "< 5 min", "3 min 12 sec", "11 sec", "2 min 24 sec", "37 sec", "Normal"],
        ["CXR Rapid Findings", "< 30 sec", "8 sec", "2 sec", "4 sec", "2 sec", "Normal"],
        ["MRI Brain MS Tracking", "< 5 min", "2 min 48 sec", "14 sec", "1 min 52 sec", "42 sec", "Normal"],
    ],
    flag_col=6
)

add_paragraph("", space_after=4)
add_sub_header("Aggregate Pipeline Metrics")
add_styled_table(
    ["Metric", "Value"],
    [
        ["Total Studies Processed", "4"],
        ["Total Series Processed", "10"],
        ["Total Slices Analyzed", "1,847"],
        ["Aggregate Processing Time", "6 minutes 55 seconds"],
        ["GPU Utilization (peak)", "94%"],
        ["GPU Memory (peak)", "47.2 GB / 128 GB"],
        ["Prior Study Retrieval", "2 studies (CT Chest + MRI Brain)"],
        ["Spatial Registrations Computed", "2 (CT-to-CT, MRI-to-MRI)"],
        ["Structured Outputs Generated", "4 DICOM SR + 4 GSPS + 4 DICOM SEG + 1 FHIR DiagnosticReport"],
        ["Urgent Alerts Dispatched", "3 (subdural hemorrhage, lung nodule, MS activity)"],
    ],
    flag_col=0
)

# ══════════════════════════════════════════════════════════════
# Recommended Next Steps
# ══════════════════════════════════════════════════════════════
add_section_header(13, "Recommended Next Steps")

next_steps = [
    ("Immediate neurosurgical awareness", "for subdural hemorrhage with repeat CT in 6 hours; escalate to CRITICAL if neurological decline"),
    ("Schedule 3-month follow-up LDCT", "for Lung-RADS 4A nodule; consider PET-CT given solid component \u2265 8 mm and volume doubling time < 400 days"),
    ("Initiate neurology referral", "for MS treatment review given breakthrough disease activity (2 new lesions, 1 enhancing) on natalizumab"),
    ("Order JCV antibody testing", "if not recently performed to assess PML risk in the context of natalizumab therapy"),
    ("Clinical correlation for pleural effusion", "with cardiac and hepatic evaluation; echocardiography if not recently performed"),
]

for i, (bold_part, rest) in enumerate(next_steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    add_run(p, f"{i}. ", bold=True, size=10.5, color=NAVY)
    add_run(p, bold_part, bold=True, size=10.5, color=GRAY_TEXT)
    add_run(p, f" {rest}", size=10.5, color=GRAY_TEXT)

# ── NVIDIA footer bar ────────────────────────────────────────
add_paragraph("", space_after=8)
footer = doc.add_table(rows=1, cols=1)
footer.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = footer.rows[0].cells[0]
set_cell_shading(cell, HEX_NAVY)
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(2)
add_run(p, "HCLS AI Factory", bold=True, size=10, color=WHITE)
add_run(p, " | Imaging Intelligence Report", size=10, color=WHITE)

# Second line: Generated date
p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after = Pt(0)
add_run(p2, "Generated: February 1, 2026 at 14:23", size=9, color=WHITE)

# Third line: Platform
p3 = cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(0)
p3.paragraph_format.space_after = Pt(0)
add_run(p3, "Platform: NVIDIA DGX Spark | MONAI Deploy 0.6.0", size=9, color=WHITE)

# Fourth line: Powered by (green)
p4 = cell.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before = Pt(2)
p4.paragraph_format.space_after = Pt(10)
add_run(p4, "Powered by NVIDIA Accelerated Computing", bold=True, size=9, color=GREEN_BRIGHT)

# ── Disclaimer ───────────────────────────────────────────────
add_paragraph(
    "DEMONSTRATION ONLY \u2014 Synthetic data generated for workflow development and testing. "
    "Not medical advice. Do not use for clinical diagnosis or treatment decisions.",
    italic=True, size=8, color=GRAY_SUB, space_before=8, space_after=0,
    alignment=WD_ALIGN_PARAGRAPH.CENTER
)

# ── Save ─────────────────────────────────────────────────────
output_path = "core/engines/clinical-imaging/open_public/HCLS_Imaging_Intelligence_Report.docx"
doc.save(output_path)
print(f"Report saved to {output_path}")
