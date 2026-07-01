#!/usr/bin/env python3
"""
generate_tsc_docs.py — Render a Markdown document to a polished .docx (house style:
Calibri body, navy headings, real tables, page numbers) and then to .pdf via LibreOffice.

Usage:
    python3 generate_tsc_docs.py <input.md> <output_basename> [--no-pdf]

Produces <output_basename>.docx and (unless --no-pdf) <output_basename>.pdf.

Handles: title block, multi-level headings (# title / # Part -> H1 / ## -> H2 / ### -> H3 / #### -> H4),
GFM pipe tables, bullet/numbered lists (one nesting level), fenced code blocks, blockquotes,
horizontal rules, and inline **bold**, *italic*, `code`, and [text](url) links.

Author: Adam M. Jones / HCLS AI Factory — TSC Intelligence Engine (Engine 7).
"""
import sys
import os
import re
import subprocess

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x38, 0x64)
STEEL = RGBColor(0x2E, 0x5A, 0x88)
MONO = "Consolas"
BODY = "Calibri"

INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*.+?\*|_.+?_)")
LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]+\)")


def strip_inline(text):
    text = LINK_RE.sub(r"\1", text)
    return re.sub(r"[*`_]", "", text)


def add_runs(paragraph, text):
    """Add text to a paragraph, honoring **bold**, *italic*/_italic_, `code`, and links (text only)."""
    text = LINK_RE.sub(r"\1", text)
    for tok in INLINE_RE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**") and len(tok) >= 4:
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`") and len(tok) >= 2:
            r = paragraph.add_run(tok[1:-1]); r.font.name = MONO; r.font.size = Pt(9.5)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) >= 2:
            r = paragraph.add_run(tok[1:-1]); r.italic = True
        elif tok.startswith("_") and tok.endswith("_") and len(tok) >= 2:
            r = paragraph.add_run(tok[1:-1]); r.italic = True
        else:
            paragraph.add_run(tok)


def shade(cell_or_para_pr, color="F2F4F8"):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), color)
    cell_or_para_pr.append(sh)


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for kind, val in (("begin", None), ("instrText", "PAGE"), ("end", None)):
        el = OxmlElement(f"w:fldChar") if kind != "instrText" else OxmlElement("w:instrText")
        if kind == "instrText":
            el.set(qn("xml:space"), "preserve"); el.text = " PAGE "
        else:
            el.set(qn("w:fldCharType"), kind)
        run._r.append(el)


def setup_styles(doc):
    n = doc.styles["Normal"]
    n.font.name = BODY; n.font.size = Pt(11)
    n.paragraph_format.space_after = Pt(6); n.paragraph_format.line_spacing = 1.15
    for name, size, color in (("Heading 1", 18, NAVY), ("Heading 2", 14, NAVY),
                              ("Heading 3", 12, STEEL), ("Heading 4", 11, STEEL)):
        st = doc.styles[name]
        st.font.name = BODY; st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(12); st.paragraph_format.space_after = Pt(4)
    for s in doc.sections:
        s.page_width = Inches(8.5); s.page_height = Inches(11)
        s.left_margin = s.right_margin = Inches(1.0)
        s.top_margin = s.bottom_margin = Inches(1.0)
        add_page_number_footer(s)


def make_table(doc, rows):
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    ncol = max(len(r) for r in cells)
    t = doc.add_table(rows=len(cells), cols=ncol)
    try:
        t.style = "Light Grid Accent 1"
    except KeyError:
        t.style = "Table Grid"
    for i, row in enumerate(cells):
        for j in range(ncol):
            cell = t.cell(i, j)
            cell.paragraphs[0].text = ""
            add_runs(cell.paragraphs[0], row[j] if j < len(row) else "")
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                shade(cell._tc.get_or_add_tcPr(), "1F3864")
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()


def render(md_path, docx_path):
    doc = Document()
    setup_styles(doc)
    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    seen_title = False
    seen_first_h2 = False
    seen_subtitle = False
    i = 0
    n = len(lines)
    while i < n:
        raw = lines[i].rstrip("\n")
        line = raw.strip()

        # fenced code
        if line.startswith("```"):
            i += 1
            code = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i].rstrip("\n")); i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            shade(p._p.get_or_add_pPr())
            r = p.add_run("\n".join(code)); r.font.name = MONO; r.font.size = Pt(9)
            continue

        if not line:
            i += 1; continue

        # horizontal rule -> skip (Part page-breaks handle separation)
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", line):
            i += 1; continue

        # tables
        if line.startswith("|") and i + 1 < n and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]):
            block = [raw]
            i += 1  # skip separator
            i += 1
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i].rstrip("\n")); i += 1
            make_table(doc, [block[0]] + block[1:])
            continue

        # blockquote (collect consecutive)
        if line.startswith(">"):
            quote = []
            while i < n and lines[i].strip().startswith(">"):
                quote.append(re.sub(r"^\s*>\s?", "", lines[i].rstrip("\n"))); i += 1
            p = doc.add_paragraph(style="Intense Quote")
            add_runs(p, " ".join(q for q in quote if q.strip()))
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1)); text = strip_inline(m.group(2))
            if level == 1 and not seen_title:
                p = doc.add_paragraph(text, style="Title"); seen_title = True
                i += 1; continue
            if level == 1:  # Part divider
                h = doc.add_heading(text, level=1); h.paragraph_format.page_break_before = True
                i += 1; continue
            if level == 3 and seen_title and not seen_subtitle and not seen_first_h2:
                doc.add_paragraph(text, style="Subtitle"); seen_subtitle = True
                i += 1; continue
            wl = min(level, 4)
            h = doc.add_heading(text, level=wl)
            if wl == 2 and not seen_first_h2:
                h.paragraph_format.page_break_before = True; seen_first_h2 = True
            i += 1; continue

        # lists
        mb = re.match(r"^(\s*)([-*+])\s+(.*)$", raw)
        mn = re.match(r"^(\s*)(\d+)\.\s+(.*)$", raw)
        if mb or mn:
            indent = len((mb or mn).group(1))
            lvl2 = indent >= 2
            if mb:
                style = "List Bullet 2" if lvl2 else "List Bullet"
                content = mb.group(3)
            else:
                style = "List Number 2" if lvl2 else "List Number"
                content = mn.group(3)
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
            add_runs(p, content)
            i += 1; continue

        # normal paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, raw)
        i += 1

    doc.save(docx_path)
    return docx_path


def to_pdf(docx_path, outdir):
    profile = "file:///tmp/lo_tsc_profile"
    cmd = ["soffice", "--headless", f"-env:UserInstallation={profile}",
           "--convert-to", "pdf", "--outdir", outdir, docx_path]
    subprocess.run(cmd, check=True, timeout=300,
                   stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)


def main():
    args = [a for a in sys.argv[1:] if a != "--no-pdf"]
    no_pdf = "--no-pdf" in sys.argv
    if len(args) < 2:
        print("usage: generate_tsc_docs.py <input.md> <output_basename> [--no-pdf]"); sys.exit(1)
    md_path, base = args[0], args[1]
    docx_path = base + ".docx"
    render(md_path, docx_path)
    print(f"  docx: {docx_path}  ({os.path.getsize(docx_path)//1024} KB)")
    if not no_pdf:
        try:
            to_pdf(docx_path, os.path.dirname(os.path.abspath(docx_path)) or ".")
            pdf_path = base + ".pdf"
            print(f"  pdf:  {pdf_path}  ({os.path.getsize(pdf_path)//1024} KB)")
        except Exception as e:
            print(f"  pdf:  SKIPPED ({e})")


if __name__ == "__main__":
    main()
