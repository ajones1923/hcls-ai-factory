#!/usr/bin/env python3
# Copyright 2026 Adam Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Generate HCLS AI Factory — Demo Guide (DOCX + PDF).

Step-by-step walkthrough for demonstrating the HCLS AI Factory VCP/FTD demo.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

NAVY = RGBColor(0x1B, 0x23, 0x33)
TEAL = RGBColor(0x1A, 0xAF, 0xCC)
GREEN = RGBColor(0x76, 0xB9, 0x00)
GRAY_BODY = RGBColor(0x33, 0x33, 0x33)
GRAY_META = RGBColor(0x66, 0x66, 0x66)
GRAY_CODE = RGBColor(0x44, 0x44, 0x44)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_NAVY = "1B2333"
HEX_TEAL = "1AAFCC"
HEX_GREEN = "76B900"
HEX_LIGHT = "F8FAFB"
FONT = "Calibri"
CODE_FONT = "Consolas"

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
style.font.color.rgb = GRAY_BODY
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15


def R(p, text, bold=False, italic=False, size=None, color=None, font=None):
    r = p.add_run(text)
    r.font.name = font or FONT
    if bold: r.bold = True
    if italic: r.italic = True
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return r


def P(text="", bold=False, italic=False, size=10.5, color=GRAY_BODY,
      before=0, after=6, align=None):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text: R(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def H1(text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs: r.clear()
    R(p, text, bold=True, size=22, color=NAVY)
    return p


def H2(text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs: r.clear()
    R(p, text, bold=True, size=16, color=NAVY)
    return p


def H3(text):
    return P(text, bold=True, size=13, color=TEAL, before=10, after=4)


def body(text, before=0, after=6):
    return P(text, size=10.5, color=GRAY_BODY, before=before, after=after)


def bullet(lead, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    R(p, lead, bold=True, size=10.5, color=GRAY_BODY)
    R(p, f"  {text}", size=10.5, color=GRAY_BODY)
    return p


def code_block(text, language=""):
    if language:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        R(p, language, bold=True, size=8, color=TEAL, font=FONT)
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.2)
        R(p, line if line else " ", size=8, color=GRAY_CODE, font=CODE_FONT)


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="D0D0D0", sz="4"):
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>')
    tcPr.append(borders)


def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, HEX_NAVY)
        set_cell_borders(c, HEX_NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(p, h, bold=True, size=9, color=WHITE)
    for ri, row in enumerate(rows):
        bg = "FFFFFF" if ri % 2 == 0 else HEX_LIGHT
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            set_cell_shading(c, bg)
            set_cell_borders(c, "E0E0E0")
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if ci == 0:
                R(p, str(val), bold=True, size=9, color=TEAL)
            else:
                R(p, str(val), size=9, color=GRAY_BODY)
    return t


def spacer(pts=6):
    P("", after=pts)


# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════

accent = doc.add_table(rows=1, cols=1)
accent.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = accent.rows[0].cells[0]
set_cell_shading(cell, HEX_GREEN)
p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
R(p, " ", size=6, color=GREEN)

for _ in range(4):
    P("")

P("Demo Guide", bold=True, size=13, color=TEAL, after=6)
P("HCLS AI Factory",
  bold=True, size=32, color=NAVY, after=2)
P("VCP/FTD Demo Walkthrough",
  bold=True, size=24, color=NAVY, after=6)
P("Step-by-step guide for demonstrating the complete pipeline: "
  "from patient DNA to 100 ranked novel drug candidates.",
  italic=False, size=12, color=NAVY, after=6)
P("NVIDIA DGX Spark  |  Parabricks  |  BioNeMo  |  Milvus  |  Claude",
  italic=True, size=11, color=GRAY_META, after=12)
P("")
P("02/2026  |  Version 1.0  |  Apache 2.0 License", size=9, color=GRAY_META, after=2)
P("Author: Adam Jones", size=9, color=GRAY_META, after=2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# DEMO OVERVIEW
# ══════════════════════════════════════════════════════════════
H1("Demo Overview")

add_table(
    ["Parameter", "Value"],
    [
        ["Demo Duration", "15-20 minutes (live walkthrough)"],
        ["Pipeline Mode", "demo (pre-configured VCP/FTD)"],
        ["Hardware", "NVIDIA DGX Spark (GB10, 128 GB unified)"],
        ["Target Gene", "VCP \u2014 Frontotemporal Dementia"],
        ["End Result", "100 ranked novel drug candidates"],
    ],
)
spacer()

H2("What the Audience Will See")
bullet("1.", "Raw DNA data (FASTQ) entering the platform")
bullet("2.", "GPU-accelerated variant calling (Parabricks) completing in minutes")
bullet("3.", "11.7 million variant records annotated and indexed in a vector database")
bullet("4.", "Interactive Claude RAG chat identifying VCP as a drug target")
bullet("5.", "BioNeMo generating 100 novel VCP inhibitors")
bullet("6.", "Ranked candidates with docking scores and drug-likeness profiles")
bullet("7.", "PDF report generated automatically")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PRE-DEMO SETUP
# ══════════════════════════════════════════════════════════════
H1("Pre-Demo Setup")

H2("Step 1: Verify Hardware")
code_block(
    "nvidia-smi              # Verify GB10 GPU\n"
    "uname -m                # Expected: aarch64",
    "bash"
)
spacer()

H2("Step 2: Set Environment Variables")
code_block(
    "cp .env.example .env\n"
    "# Edit .env:\n"
    "# ANTHROPIC_API_KEY=sk-ant-...   (for Claude RAG)\n"
    "# NGC_API_KEY=nvapi-...          (for BioNeMo NIMs)",
    "bash"
)
spacer()

H2("Step 3: Start All Services")
code_block("./start-services.sh", "bash")
body("Services start in dependency order: infrastructure \u2192 Stage 1 \u2192 Stage 2 \u2192 Stage 3 \u2192 landing page.")

H2("Step 4: Verify All Services Healthy")
body("Open http://localhost:8080 \u2014 all 10 services should show green status.")
add_table(
    ["Service", "Port", "Expected"],
    [
        ["Parabricks Portal", "5000", "GREEN"],
        ["Milvus Vector DB", "19530", "GREEN"],
        ["RAG API", "5001", "GREEN"],
        ["Streamlit Chat", "8501", "GREEN"],
        ["MolMIM NIM", "8001", "GREEN"],
        ["DiffDock NIM", "8002", "GREEN"],
        ["Discovery UI", "8505", "GREEN"],
        ["Grafana", "3000", "GREEN"],
        ["Prometheus", "9099", "GREEN"],
        ["DCGM Exporter", "9400", "GREEN"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# DEMO SCRIPT
# ══════════════════════════════════════════════════════════════
H1("Demo Script")

# ── Opening ─────────────────────────────────────────────────
H2("Opening (1 minute)")
body("Show: Landing page at http://localhost:8080")
bullet("\u2022", '"This is the HCLS AI Factory \u2014 patient DNA to drug candidates in < 5 hours."')
bullet("\u2022", '"Everything runs on this single DGX Spark \u2014 a $4,699 desktop workstation."')
bullet("\u2022", '"Three stages: genomics, target identification, and drug discovery."')

# ── Stage 1 ─────────────────────────────────────────────────
H2("Stage 1: Genomics (3-4 minutes)")
body("Launch: python run_pipeline.py --mode demo")
body("Show: Genomics portal at http://localhost:5000")

H3("Talking Points")
bullet("\u2022", "~200 GB FASTQ input from Illumina sequencer (30\u00d7 WGS)")
bullet("\u2022", "Parabricks BWA-MEM2: 20-45 min on GPU (vs. 12-24 hours on CPU)")
bullet("\u2022", "DeepVariant: 10-35 min, >99% accuracy (CNN-based)")
bullet("\u2022", "Output: VCF with ~11.7 million variant records")
bullet("\u2022", "Show GPU utilization spiking on Grafana (http://localhost:3000)")

# ── Stage 2 ─────────────────────────────────────────────────
H2("Stage 2: RAG/Chat (5-6 minutes)")

H3("Annotation Pipeline")
bullet("\u2022", "ClinVar: 4.1M clinical variants \u2192 35,616 patient matches")
bullet("\u2022", "AlphaMissense: 71M AI predictions \u2192 6,831 scored variants")
bullet("\u2022", "VEP: functional consequences (HIGH/MODERATE/LOW/MODIFIER)")

H3("Vector Database")
body("Show: Attu UI at http://localhost:8000")
bullet("\u2022", "3.5M variant embeddings (BGE-small-en-v1.5, 384-dim)")
bullet("\u2022", "IVF_FLAT index, COSINE metric, 17 fields per record")
bullet("\u2022", '"This enables natural language queries over genomic data."')

H3("Interactive Chat")
body("Show: Streamlit chat at http://localhost:8501")
body('Type: "What are the most promising drug targets for neurodegenerative disease?"')
body("Claude identifies VCP with full evidence chain:")
bullet("\u2022", "rs188935092 \u2014 ClinVar Pathogenic, AlphaMissense 0.87")
bullet("\u2022", "HIGH impact missense variant in D2 ATPase domain")
bullet("\u2022", "Known target \u2014 CB-5083 (Phase I), druggability 0.92")
bullet("\u2022", "201 genes, 13 therapeutic areas, 171 druggable (85%)")

doc.add_page_break()

# ── Stage 3 ─────────────────────────────────────────────────
H2("Stage 3: Drug Discovery (4-5 minutes)")

H3("Structure Retrieval")
bullet("\u2022", "VCP \u2192 UniProt P55072 \u2192 RCSB PDB query")
bullet("\u2022", "4 structures: 8OOI, 9DIL, 7K56, 5FTK")
bullet("\u2022", "5FTK selected \u2014 2.3 \u00c5 X-ray with CB-5083 inhibitor bound")

H3("Molecule Generation")
body("Show: Discovery UI at http://localhost:8505")
bullet("\u2022", "MolMIM generates 100 novel analogs from CB-5083 seed")
bullet("\u2022", "98 pass RDKit chemical validity checks")

H3("Docking & Ranking")
bullet("\u2022", "DiffDock docks each candidate against VCP D2 domain")
bullet("\u2022", "34 candidates score below -8.0 kcal/mol (excellent)")
bullet("\u2022", "Composite: 30% generation + 40% docking + 30% QED")

H3("Key Demo Table")
add_table(
    ["Metric", "CB-5083 (Seed)", "Top Candidate", "Improvement"],
    [
        ["Dock Score", "-8.1 kcal/mol", "-11.4 kcal/mol", "+41% binding"],
        ["QED", "0.62", "0.81", "+31% drug-likeness"],
        ["MW", "487.2 Da", "423.5 Da", "-13% (better)"],
        ["Composite", "0.64", "0.89", "+39% overall"],
    ],
)
spacer()

body("PDF report generated automatically via ReportLab with full provenance.")

# ── Closing ─────────────────────────────────────────────────
H2("Closing (2 minutes)")
bullet("\u2022", '"< 5 hours, $4,699 desktop \u2192 raw DNA to 100 ranked drug candidates"')
bullet("\u2022", '"Collapses weeks/months to a single session"')
bullet("\u2022", '"Same Nextflow pipelines scale to DGX SuperPOD"')
bullet("\u2022", '"Open-source \u2014 Apache 2.0"')
spacer()

add_table(
    ["Phase", "Hardware", "Scale"],
    [
        ["Phase 1", "DGX Spark ($4,699)", "Proof build \u2014 what you just saw"],
        ["Phase 2", "DGX B200", "Department \u2014 multiple concurrent patients"],
        ["Phase 3", "DGX SuperPOD", "Enterprise \u2014 thousands, federated"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TROUBLESHOOTING & QUICK REFERENCE
# ══════════════════════════════════════════════════════════════
H1("Troubleshooting")

H2("Service Not Starting")
code_block(
    "docker compose ps                    # Check status\n"
    "docker compose logs <service-name>   # Check logs\n"
    "docker compose restart <service>     # Restart",
    "bash"
)

H2("BioNeMo NIM Not Ready")
body("NIMs require NGC API key and may take 2-5 minutes to initialize.")
code_block(
    "curl http://localhost:8001/v1/health/ready   # MolMIM\n"
    "curl http://localhost:8002/v1/health/ready   # DiffDock",
    "bash"
)

H2("GPU Out of Memory")
body("DeepVariant peaks at ~60 GB. Monitor with:")
code_block("watch -n 1 nvidia-smi", "bash")

spacer(6)

H1("Quick Reference")
add_table(
    ["Action", "Command / URL"],
    [
        ["Start services", "./start-services.sh"],
        ["Launch demo", "python run_pipeline.py --mode demo"],
        ["Landing page", "http://localhost:8080"],
        ["Genomics portal", "http://localhost:5000"],
        ["Chat interface", "http://localhost:8501"],
        ["Milvus UI", "http://localhost:8000"],
        ["Discovery UI", "http://localhost:8505"],
        ["Grafana", "http://localhost:3000"],
    ],
)

spacer(12)
p = P("", after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
R(p, "HCLS AI Factory \u2014 Apache 2.0  |  Author: Adam Jones  |  February 2026",
  italic=True, size=10, color=GRAY_META)

# ══════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════
import pathlib, subprocess, shutil

OUT = pathlib.Path(__file__).resolve().parent
docx_path = OUT / "HCLS_AI_Factory_Demo_Guide.docx"
pdf_path = OUT / "HCLS_AI_Factory_Demo_Guide.pdf"

doc.save(str(docx_path))
print(f"DOCX saved: {docx_path} ({docx_path.stat().st_size:,} bytes)")

if shutil.which("libreoffice"):
    subprocess.run([
        "libreoffice", "--headless", "--convert-to", "pdf",
        "--outdir", str(OUT), str(docx_path)
    ], capture_output=True, timeout=120)
    if pdf_path.exists():
        print(f"PDF  saved: {pdf_path} ({pdf_path.stat().st_size:,} bytes)")
elif shutil.which("docx2pdf"):
    subprocess.run(["docx2pdf", str(docx_path), str(pdf_path)],
                    capture_output=True, timeout=120)
    if pdf_path.exists():
        print(f"PDF  saved: {pdf_path} ({pdf_path.stat().st_size:,} bytes)")
else:
    print("No PDF converter found.")

print("Done.")
