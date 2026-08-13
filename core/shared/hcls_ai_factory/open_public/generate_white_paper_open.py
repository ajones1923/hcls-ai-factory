#!/usr/bin/env python3
# Copyright 2026 Adam Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Generate HCLS AI Factory White Paper — Open Architecture on NVIDIA DGX Spark.

From Patient DNA to Novel Drug Candidates in Under Five Hours.
Open/public version for GitHub (Apache 2.0).
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Colors ─────────────────────────────────────────────────
NAVY = RGBColor(0x1B, 0x23, 0x33)
TEAL = RGBColor(0x1A, 0xAF, 0xCC)
GREEN = RGBColor(0x76, 0xB9, 0x00)
GRAY_BODY = RGBColor(0x33, 0x33, 0x33)
GRAY_META = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_NAVY = "1B2333"
HEX_TEAL = "1AAFCC"
HEX_GREEN = "76B900"
HEX_LIGHT = "F8FAFB"
FONT = "Calibri"
CODE_FONT = "Consolas"

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
style.font.color.rgb = GRAY_BODY
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15


def R(p, text, bold=False, italic=False, size=None, color=None, font=None):
    r = p.add_run(text)
    r.font.name = font or FONT
    if bold: r.bold = True
    if italic: r.italic = True
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return r


def P(text="", bold=False, italic=False, size=10.5, color=GRAY_BODY,
      before=0, after=6, align=None):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text: R(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def H1(text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs: r.clear()
    R(p, text, bold=True, size=22, color=NAVY)
    return p


def H2(text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs: r.clear()
    R(p, text, bold=True, size=16, color=NAVY)
    return p


def H3(text):
    return P(text, bold=True, size=13, color=TEAL, before=10, after=4)


def body(text, before=0, after=6):
    return P(text, size=10.5, color=GRAY_BODY, before=before, after=after)


def bullet(lead, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    R(p, lead, bold=True, size=10.5, color=GRAY_BODY)
    R(p, f"  {text}", size=10.5, color=GRAY_BODY)
    return p


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="D0D0D0", sz="4"):
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>')
    tcPr.append(borders)


def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, HEX_NAVY)
        set_cell_borders(c, HEX_NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(p, h, bold=True, size=9, color=WHITE)
    for ri, row in enumerate(rows):
        bg = "FFFFFF" if ri % 2 == 0 else HEX_LIGHT
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            set_cell_shading(c, bg)
            set_cell_borders(c, "E0E0E0")
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if ci == 0:
                R(p, str(val), bold=True, size=9, color=TEAL)
            else:
                R(p, str(val), size=9, color=GRAY_BODY)
    return t


def spacer(pts=6):
    P("", after=pts)


# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════

accent = doc.add_table(rows=1, cols=1)
accent.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = accent.rows[0].cells[0]
set_cell_shading(cell, HEX_GREEN)
p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
R(p, " ", size=6, color=GREEN)

for _ in range(5):
    P("")

P("White Paper", bold=True, size=13, color=TEAL, after=6)
P("HCLS AI Factory",
  bold=True, size=30, color=NAVY, after=4)
P("Open Architecture on NVIDIA DGX Spark",
  bold=True, size=24, color=NAVY, after=4)
P("From Patient DNA to Novel Drug Candidates in Under Five Hours \u2014 "
  "GPU-Accelerated Genomics, RAG-Grounded Target Identification, and "
  "AI-Driven Drug Discovery",
  bold=False, size=12, color=NAVY, after=6)
P("NVIDIA DGX Spark  |  Parabricks  |  BioNeMo  |  Milvus  |  Claude",
  italic=True, size=11, color=GRAY_META, after=8)
P("")
P("02/2026  |  Version 1.0  |  Apache 2.0 License", size=9, color=GRAY_META, after=2)
P("Author: Adam Jones", size=9, color=GRAY_META, after=2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════
H1("1. Executive Summary")

body(
    "The HCLS AI Factory follows a reusable pattern: identify a canonical artifact, "
    "build a persistent data model around it, and let agentic workflows operate "
    "continuously on that model. In genomics, the canonical artifact is the VCF \u2014 "
    "a structured record of every genomic variant identified in a patient\u2019s DNA."
)

body(
    "This white paper describes an end-to-end platform that processes raw DNA "
    "sequencing data through GPU-accelerated variant calling, RAG-grounded clinical "
    "reasoning, and AI-driven drug discovery \u2014 all on a single NVIDIA DGX Spark "
    "desktop workstation."
)

body(
    "The platform transforms patient FASTQ files (~200 GB of raw sequencing data "
    "from a 30\u00d7 whole-genome study) into 100 ranked novel drug candidates in "
    "under 5 hours. Three stages execute sequentially: NVIDIA Parabricks performs "
    "GPU-accelerated alignment and variant calling (120-240 min), producing ~11.7 "
    "million variant records (~4.69 million of which pass filter). A RAG pipeline "
    "annotates variants with ClinVar, AlphaMissense, "
    "and VEP, indexes 3.5 million high-quality variants in a Milvus vector database, "
    "and uses Anthropic Claude to identify druggable gene targets. Finally, BioNeMo "
    "NIM services (MolMIM and DiffDock) generate novel molecules, predict binding "
    "affinities, and rank candidates by a composite drug-likeness score."
)

body(
    "The architecture is designed to run end-to-end on a $4,699 DGX Spark for proof "
    "builds and scale to DGX SuperPOD for enterprise deployments. All components are "
    "open-source or NVIDIA-licensed, released under Apache 2.0 as part of the HCLS "
    "AI Factory."
)

# ══════════════════════════════════════════════════════════════
# 2. THE PRECISION MEDICINE DATA CHALLENGE
# ══════════════════════════════════════════════════════════════
H1("2. The Precision Medicine Data Challenge")

body(
    "Precision medicine promises therapies tailored to an individual\u2019s genetic "
    "profile. A single 30\u00d7 whole-genome sequencing (WGS) run produces approximately "
    "200 GB of raw data and 11.7 million genomic variant records. The challenge is not "
    "generating this data \u2014 modern sequencers produce it reliably \u2014 but transforming "
    "it into actionable therapeutic hypotheses within a clinically relevant timeframe."
)

H3("The Limits of Traditional Bioinformatics")

body(
    "Today\u2019s genomic analysis pipelines assemble disconnected components: CPU-based "
    "alignment tools that take 12-24 hours, separate variant callers, annotation "
    "databases accessed through web APIs, and manual literature review for target "
    "identification."
)

body("This sequential, manual approach introduces three structural problems:")

bullet("Compute bottleneck.",
       "CPU-based BWA-MEM alignment of a 30\u00d7 WGS sample takes 12-24 hours on a "
       "32-core server. DeepVariant on CPU adds another 8-12 hours. The genomics "
       "stage alone consumes 1-2 days of wall time.")
bullet("Annotation fragmentation.",
       "Clinical variant databases (ClinVar), AI pathogenicity predictors "
       "(AlphaMissense), and functional annotation tools (VEP) exist as separate "
       "resources requiring bespoke ETL pipelines.")
bullet("Target-to-drug gap.",
       "Even after identifying a pathogenic variant in a druggable gene, the path "
       "to a lead compound requires separate molecular modeling tools and medicinal "
       "chemistry expertise \u2014 typically a months-long process.")

H3("The GPU-Accelerated Opportunity")

body(
    "NVIDIA DGX Spark collapses the compute bottleneck. Its GB10 GPU accelerates "
    "BWA-MEM2 alignment from hours to 20-45 minutes and DeepVariant variant calling "
    "from hours to 10-35 minutes \u2014 a 10-20\u00d7 speedup. More importantly, the same "
    "GPU that runs genomics can run vector similarity search (Milvus), molecular "
    "generation (MolMIM), and molecular docking (DiffDock). A single $4,699 desktop "
    "workstation handles the entire pipeline."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. ARCHITECTURE OVERVIEW
# ══════════════════════════════════════════════════════════════
H1("3. Architecture Overview")

H2("Three-Stage Pipeline")
add_table(
    ["Stage", "Technology", "Duration", "Input", "Output"],
    [
        ["1 \u2014 Genomics", "Parabricks 4.6", "120-240 min", "FASTQ (~200 GB)", "VCF (~11.7M variant records)"],
        ["2 \u2014 RAG/Chat", "Milvus + BGE + Claude", "Interactive", "VCF", "Target gene + evidence"],
        ["3 \u2014 Drug Discovery", "MolMIM + DiffDock + RDKit", "8-16 min", "Target gene", "100 ranked candidates"],
    ],
)
spacer()

H2("Technology Stack")
add_table(
    ["Layer", "Components"],
    [
        ["Compute", "NVIDIA DGX Spark (GB10 GPU, 128 GB unified, ARM64 cores)"],
        ["Genomics", "NVIDIA Parabricks 4.6.0-1, GRCh38 reference genome"],
        ["Annotation", "ClinVar (4.1M), AlphaMissense (71M), Ensembl VEP"],
        ["Vector DB", "Milvus 2.4, BGE-small-en-v1.5 (384-dim), IVF_FLAT"],
        ["LLM", "Anthropic Claude (claude-sonnet-4-20250514, temp=0.3)"],
        ["Drug Discovery", "BioNeMo MolMIM, BioNeMo DiffDock, RDKit"],
        ["Orchestration", "Nextflow DSL2, Docker Compose"],
        ["Monitoring", "Grafana, Prometheus, DCGM Exporter"],
    ],
)
spacer()

H2("Service Architecture")
body(
    "The platform runs 14 services across 14 ports, organized by stage:"
)
bullet("Orchestration:", "Landing page (8080) with 10-service health monitor")
bullet("Stage 1:", "Genomics portal (5000)")
bullet("Stage 2:", "Milvus (19530), Attu UI (8000), RAG API (5001), Chat (8501)")
bullet("Stage 3:", "MolMIM NIM (8001), DiffDock NIM (8002), Discovery UI (8505), Portal (8510)")
bullet("Monitoring:", "Grafana (3000), Prometheus (9099), Node Exporter (9100), DCGM (9400)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. STAGE 1 — GPU-ACCELERATED GENOMICS
# ══════════════════════════════════════════════════════════════
H1("4. Stage 1 \u2014 GPU-Accelerated Genomics")

H2("NVIDIA Parabricks 4.6")
body(
    "Parabricks provides GPU-accelerated implementations of standard bioinformatics "
    "tools. On DGX Spark\u2019s GB10 GPU, it delivers 10-20\u00d7 speedup over CPU "
    "implementations."
)

H3("BWA-MEM2 Alignment (fq2bam)")
body(
    "Aligns paired-end reads against the GRCh38 reference genome. GPU-accelerated "
    "implementation achieves 70-90% GPU utilization, producing a sorted BAM file "
    "with index in 20-45 minutes."
)

H3("Google DeepVariant")
body(
    "A CNN-based variant caller achieving >99% accuracy on the GIAB truth set. "
    "GPU-accelerated implementation achieves 80-95% GPU utilization, calling "
    "variants in 10-35 minutes. The deep learning approach outperforms traditional "
    "statistical callers (GATK HaplotypeCaller) on both SNPs and indels."
)

H2("Input: HG002 Reference Standard")
add_table(
    ["Parameter", "Value"],
    [
        ["Sample", "HG002 (NA24385, GIAB reference standard)"],
        ["Coverage", "30\u00d7 whole-genome sequencing (WGS)"],
        ["Read Length", "2\u00d7250 bp paired-end"],
        ["File Size", "~200 GB (FASTQ pair)"],
        ["Reference", "GRCh38 (3.1 GB, pre-indexed)"],
    ],
)
spacer()

H2("Output: Variant Call Format (VCF)")
add_table(
    ["Metric", "Count"],
    [
        ["Total Variants", "~11.7M"],
        ["High-Quality (QUAL>30)", "~3.5M"],
        ["SNPs", "~4.2M"],
        ["Indels", "~1.0M"],
        ["Coding Region", "~35,000"],
        ["Multi-allelic Sites", "~150,000"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. STAGE 2 — RAG-GROUNDED TARGET IDENTIFICATION
# ══════════════════════════════════════════════════════════════
H1("5. Stage 2 \u2014 RAG-Grounded Target Identification")

H2("Variant Annotation")
body(
    "Stage 2 begins by annotating the 3.5 million high-quality variants with "
    "three complementary databases:"
)

H3("ClinVar (NCBI)")
body(
    "4.1 million clinical variant records mapping genomic positions to clinical "
    "significance classifications (Pathogenic, Likely pathogenic, VUS, Likely "
    "benign, Benign). Approximately 35,616 patient variants match ClinVar entries."
)

H3("AlphaMissense (DeepMind)")
body(
    "71,697,560 AI-predicted pathogenicity scores for missense variants, derived "
    "from AlphaFold protein structure features. Thresholds: pathogenic >0.564, "
    "ambiguous 0.34-0.564, benign <0.34. Approximately 6,831 ClinVar-matched "
    "variants have AlphaMissense predictions."
)

H3("Ensembl VEP")
body(
    "Functional consequence annotation mapping variants to genes, transcripts, "
    "and impact levels (HIGH, MODERATE, LOW, MODIFIER). Identifies missense "
    "variants, stop gains, frameshift variants, and splice site disruptions."
)

H2("Vector Embedding and Indexing")
body(
    "Each annotated variant is transformed into a text summary and embedded using "
    "BGE-small-en-v1.5 (384 dimensions). The 3.5 million embeddings are indexed "
    "in Milvus 2.4 using IVF_FLAT (nlist=1024, COSINE metric) with 17 structured "
    "fields per record."
)

H2("RAG-Grounded Reasoning with Claude")
body(
    "User queries are expanded using 10 therapeutic area keyword maps, embedded, "
    "and used for approximate nearest-neighbor search in Milvus (top_k=20). "
    "Retrieved variant contexts are assembled into a RAG prompt and processed by "
    "Anthropic Claude (claude-sonnet-4-20250514, temperature=0.3). Claude generates "
    "structured target hypotheses: gene name, confidence level, evidence chain, "
    "therapeutic area, and recommended action."
)

H2("Knowledge Base: 201 Genes, 13 Therapeutic Areas")
add_table(
    ["Therapeutic Area", "Genes", "Examples"],
    [
        ["Neurology", "36", "VCP, APP, PSEN1, MAPT, SOD1"],
        ["Oncology", "27", "EGFR, BRAF, KRAS, TP53, BRCA1"],
        ["Metabolic", "22", "GCK, PPARG, SLC2A2, PCSK9"],
        ["Infectious Disease", "21", "ACE2, CCR5, IFITM3, TLR4"],
        ["Respiratory", "13", "CFTR, SERPINA1, MUC5B"],
        ["Rare Disease", "12", "VCP, HTT, SMN1, DMD"],
        ["Hematology", "12", "HBB, HBA1, F5, JAK2"],
        ["GI/Hepatology", "12", "HFE, ATP7B, NOD2"],
        ["Pharmacogenomics", "11", "CYP2D6, CYP2C19, CYP3A4"],
        ["Ophthalmology", "11", "RHO, RPE65, RS1, ABCA4"],
        ["Cardiovascular", "10", "LDLR, PCSK9, SCN5A"],
        ["Immunology", "9", "HLA-B, TNF, IL6, JAK1"],
        ["Dermatology", "9", "FLG, MC1R, TYR, KRT14"],
    ],
)
spacer()
body("Total: 201 genes, 171 druggable targets (85% druggability).")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. STAGE 3 — AI-DRIVEN DRUG DISCOVERY
# ══════════════════════════════════════════════════════════════
H1("6. Stage 3 \u2014 AI-Driven Drug Discovery")

H2("10-Stage Drug Discovery Pipeline")
add_table(
    ["Stage", "Process", "Description"],
    [
        ["1", "Initialize", "Load target hypothesis, validate inputs"],
        ["2", "Normalize Target", "Map gene \u2192 UniProt ID \u2192 PDB structures"],
        ["3", "Structure Discovery", "Query RCSB PDB for structures"],
        ["4", "Structure Prep", "Score by resolution, inhibitor, pockets"],
        ["5", "Molecule Generation", "MolMIM generates novel SMILES from seed"],
        ["6", "Chemistry QC", "RDKit validates chemical feasibility"],
        ["7", "Conformers", "RDKit 3D conformer embedding (ETKDG)"],
        ["8", "Docking", "DiffDock predicts binding poses and affinities"],
        ["9", "Ranking", "30% gen + 40% dock + 30% QED composite"],
        ["10", "Reporting", "PDF report via ReportLab"],
    ],
)
spacer()

H2("BioNeMo NIM Services")

H3("MolMIM (Port 8001)")
body(
    "A masked language model for molecular generation. Given a seed compound\u2019s "
    "SMILES string, it generates structurally novel analogs by masking and "
    "regenerating molecular tokens. Container: nvcr.io/nvidia/clara/bionemo-molmim:1.0"
)

H3("DiffDock (Port 8002)")
body(
    "A score-based generative diffusion model for molecular docking. It predicts "
    "the 3D binding pose and affinity of a ligand in a protein binding site without "
    "requiring pre-defined binding pockets. Container: nvcr.io/nvidia/clara/diffdock:1.0"
)

H2("Drug-Likeness Scoring")
body(
    "Each candidate is evaluated against Lipinski\u2019s Rule of Five (MW\u2264500, "
    "LogP\u22645, HBD\u22645, HBA\u226410), QED (>0.67 = drug-like), and TPSA "
    "(<140 \u00c5\u00b2 for oral bioavailability)."
)

H2("Composite Scoring Formula")
body(
    "The final ranking uses a weighted composite: 30% MolMIM generation confidence, "
    "40% DiffDock binding affinity (normalized: max(0, min(1, (10 + dock_score) / 20))), "
    "and 30% QED score. This balances novelty, binding prediction, and drug-likeness."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. VCP/FTD DEMONSTRATION
# ══════════════════════════════════════════════════════════════
H1("7. VCP/FTD Demonstration")

H2("Target: Valosin-Containing Protein (VCP/p97)")
body(
    "The platform ships with a pre-configured demonstration targeting VCP \u2014 "
    "a AAA+ ATPase involved in the ubiquitin-proteasome pathway. Pathogenic VCP "
    "mutations cause Frontotemporal Dementia (FTD), ALS, and IBMPFD."
)

add_table(
    ["Parameter", "Value"],
    [
        ["Gene", "VCP (UniProt P55072)"],
        ["Variant", "rs188935092 (chr9:35065263 G>A)"],
        ["ClinVar", "Pathogenic"],
        ["AlphaMissense", "0.87 (pathogenic, >0.564 threshold)"],
        ["Seed Compound", "CB-5083 (Phase I clinical VCP inhibitor)"],
        ["PDB Structures", "8OOI, 9DIL, 7K56, 5FTK"],
        ["Binding Site", "D2 ATPase domain (~450 \u00c5\u00b3, druggability 0.92)"],
    ],
)
spacer()

H2("Demo Results")
body(
    "The VCP/FTD demo produces 100 novel VCP inhibitor candidates. Typical "
    "results: 87 pass Lipinski\u2019s Rule of Five, 72 have QED >0.67, top 10 "
    "show docking scores from -8.2 to -11.4 kcal/mol, with composite scores "
    "ranging 0.68-0.89."
)

# ══════════════════════════════════════════════════════════════
# 8. CRYO-EM STRUCTURE EVIDENCE
# ══════════════════════════════════════════════════════════════
H1("8. Cryo-EM Structure Evidence")

H2("Automated Structure Retrieval and Scoring")
body(
    "The drug discovery pipeline automatically queries RCSB PDB for protein "
    "structures and scores them by resolution, inhibitor presence (+3 bonus), "
    "druggable pocket count (+0.5 each), and experimental method (Cryo-EM +0.5)."
)

add_table(
    ["PDB", "Resolution", "Method", "Key Feature"],
    [
        ["5FTK", "2.3 \u00c5", "X-ray", "CB-5083 inhibitor-bound (highest score)"],
        ["7K56", "2.5 \u00c5", "Cryo-EM", "VCP complex"],
        ["8OOI", "2.9 \u00c5", "Cryo-EM", "WT VCP hexamer"],
        ["9DIL", "3.2 \u00c5", "Cryo-EM", "Mutant VCP"],
    ],
)
spacer()
body(
    "The inhibitor-bound structure (5FTK) is preferred because it provides a "
    "pre-defined binding site and a reference ligand for molecular generation."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. ORCHESTRATION AND MONITORING
# ══════════════════════════════════════════════════════════════
H1("9. Orchestration and Monitoring")

H2("Nextflow DSL2 Orchestration")
body("The pipeline is orchestrated by Nextflow DSL2, supporting five modes:")
bullet("full:", "End-to-end (Stages 1\u21922\u21923)")
bullet("target:", "From existing VCF (Stages 2\u21923)")
bullet("drug:", "Known target to drug candidates (Stage 3 only)")
bullet("demo:", "Pre-configured VCP/FTD demonstration")
bullet("genomics_only:", "Variant calling only (Stage 1)")

body(
    "Six execution profiles (standard, docker, singularity, dgx_spark, slurm, test) "
    "adapt the pipeline to different infrastructure."
)

H2("Monitoring Stack")
body(
    "Grafana dashboards (port 3000) visualize GPU utilization, memory pressure, "
    "pipeline progress, and service health. Prometheus (port 9099) collects metrics "
    "from DCGM Exporter (port 9400) for GPU telemetry and Node Exporter (port 9100) "
    "for system resources. The landing page (port 8080) provides a 10-service "
    "health grid with real-time status indicators."
)

# ══════════════════════════════════════════════════════════════
# 10. CROSS-MODAL INTEGRATION
# ══════════════════════════════════════════════════════════════
H1("10. Cross-Modal Integration")

H2("Imaging Intelligence Agent Integration")
body(
    "The HCLS AI Factory ecosystem includes an Imaging Intelligence Agent for "
    "CT, MRI, and X-ray analysis. Cross-modal triggers connect imaging findings "
    "to genomic analysis:"
)
bullet("Imaging \u2192 Genomics:", "Lung-RADS 4B+ triggers FHIR ServiceRequest for tumor profiling")
bullet("Genomics \u2192 Drug Discovery:", "Pathogenic variants trigger targeted molecule generation")
bullet("Drug Discovery \u2192 Imaging:", "Candidates combined with imaging in clinical reports")

H2("NVIDIA FLARE for Federated Learning")
body(
    "Phase 3 deployments use NVIDIA FLARE for federated learning across institutions. "
    "Models train locally; only gradient updates are shared. Patient genomic data "
    "never leaves the originating institution."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 11. DEPLOYMENT ROADMAP
# ══════════════════════════════════════════════════════════════
H1("11. Deployment Roadmap")

H2("Three-Phase Scaling")
add_table(
    ["Phase", "Hardware", "Orchestration", "Scale"],
    [
        ["1 \u2014 Proof Build", "DGX Spark ($4,699)", "Docker Compose", "Single patient, sequential"],
        ["2 \u2014 Departmental", "1-2\u00d7 DGX B200", "Kubernetes", "Multiple concurrent patients"],
        ["3 \u2014 Enterprise", "DGX SuperPOD", "K8s + FLARE", "Thousands concurrent, federated"],
    ],
)
spacer()

H3("Phase 1: DGX Spark Proof Build")
body(
    "A single DGX Spark runs the complete pipeline: GB10 GPU handles Parabricks, "
    "Milvus, MolMIM, and DiffDock sequentially. Docker Compose manages all 14 "
    "services. The 128 GB unified memory accommodates all stages. Total cost: "
    "$4,699 hardware + API keys (Anthropic, NGC)."
)

H3("Phase 2: Departmental Scale")
body(
    "DGX B200 systems (8\u00d7 B200 GPUs, 1-2 TB HBM3e) enable parallel processing "
    "of multiple patients, GPU-dedicated Milvus instances, and multiple BioNeMo "
    "NIM replicas. Kubernetes replaces Docker Compose."
)

H3("Phase 3: Enterprise / Multi-Site")
body(
    "DGX SuperPOD with InfiniBand fabric, NVIDIA FLARE for federated learning, "
    "and institutional-scale variant databases. Thousands of patients processed "
    "concurrently with cross-institutional collaboration while maintaining "
    "data sovereignty."
)

# ══════════════════════════════════════════════════════════════
# 12. CONCLUSION
# ══════════════════════════════════════════════════════════════
H1("12. Conclusion")

body(
    "The HCLS AI Factory demonstrates that the full precision medicine pipeline \u2014 "
    "from raw DNA to novel drug candidates \u2014 can run on a single desktop workstation. "
    "GPU acceleration collapses genomics from days to hours. Vector databases and LLM "
    "reasoning transform annotation from manual curation to interactive exploration. "
    "Generative chemistry and molecular docking automate the target-to-lead transition "
    "that traditionally takes months."
)

body(
    "The three-stage architecture (Genomics \u2192 RAG/Chat \u2192 Drug Discovery) provides "
    "a reproducible, auditable, and scalable framework. The same Nextflow pipelines "
    "that run on a $4,699 DGX Spark scale to DGX SuperPOD for enterprise deployments. "
    "All components are open-source or NVIDIA-licensed under Apache 2.0."
)

body(
    "This is precision medicine as a continuous, computable workflow \u2014 not a "
    "disconnected collection of tools, but an integrated factory that transforms "
    "patient data into therapeutic hypotheses in a single session."
)

spacer(12)
p = P("", after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
R(p, "HCLS AI Factory \u2014 Apache 2.0  |  Author: Adam Jones  |  February 2026",
  italic=True, size=10, color=GRAY_META)

# ══════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════
import pathlib, subprocess, shutil

OUT = pathlib.Path(__file__).resolve().parent
docx_path = OUT / "HCLS_AI_Factory_White_Paper_DGX_Spark.docx"
pdf_path = OUT / "HCLS_AI_Factory_White_Paper_DGX_Spark.pdf"

doc.save(str(docx_path))
print(f"DOCX saved: {docx_path} ({docx_path.stat().st_size:,} bytes)")

if shutil.which("libreoffice"):
    subprocess.run([
        "libreoffice", "--headless", "--convert-to", "pdf",
        "--outdir", str(OUT), str(docx_path)
    ], capture_output=True, timeout=120)
    if pdf_path.exists():
        print(f"PDF  saved: {pdf_path} ({pdf_path.stat().st_size:,} bytes)")
    else:
        print("PDF conversion completed but file not found.")
elif shutil.which("docx2pdf"):
    subprocess.run(["docx2pdf", str(docx_path), str(pdf_path)],
                    capture_output=True, timeout=120)
    if pdf_path.exists():
        print(f"PDF  saved: {pdf_path} ({pdf_path.stat().st_size:,} bytes)")
else:
    print("No PDF converter found (install libreoffice or docx2pdf).")

print("Done.")
