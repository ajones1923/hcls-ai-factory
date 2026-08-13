#!/usr/bin/env python3
# Copyright 2026 Adam Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Generate HCLS AI Factory — Learning Guide: Foundations (DOCX + PDF).

High school level (Grades 9-12) introduction to the HCLS AI Factory.
Covers DNA, sequencing, GPU-accelerated variant calling, annotation,
RAG-grounded target identification, and generative drug discovery.
Formatted in VCP-style theme.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Colors (VCP palette) ─────────────────────────────────
NAVY = RGBColor(0x1B, 0x23, 0x33)
TEAL = RGBColor(0x1A, 0xAF, 0xCC)
GREEN = RGBColor(0x76, 0xB9, 0x00)
GRAY_BODY = RGBColor(0x33, 0x33, 0x33)
GRAY_META = RGBColor(0x66, 0x66, 0x66)
GRAY_CODE = RGBColor(0x44, 0x44, 0x44)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_NAVY = "1B2333"
HEX_TEAL = "1AAFCC"
HEX_GREEN = "76B900"
HEX_LIGHT = "F8FAFB"
HEX_CODE_BG = "F5F5F5"
FONT = "Calibri"
CODE_FONT = "Consolas"

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
style.font.color.rgb = GRAY_BODY
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15


# ── Helper functions ──────────────────────────────────────

def R(p, text, bold=False, italic=False, size=None, color=None, font=None):
    r = p.add_run(text)
    r.font.name = font or FONT
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return r


def P(text="", bold=False, italic=False, size=10.5, color=GRAY_BODY,
      before=0, after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text:
        R(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def H1(text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        r.clear()
    R(p, text, bold=True, size=22, color=NAVY)
    return p


def H2(text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.clear()
    R(p, text, bold=True, size=16, color=NAVY)
    return p


def H3(text):
    return P(text, bold=True, size=13, color=TEAL, before=10, after=4)


def body(text, before=0, after=6):
    return P(text, size=10.5, color=GRAY_BODY, before=before, after=after)


def bullet(lead, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    R(p, lead, bold=True, size=10.5, color=GRAY_BODY)
    R(p, f"  {text}", size=10.5, color=GRAY_BODY)
    return p


def code_block(text, language=""):
    if language:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        R(p, language, bold=True, size=8, color=TEAL, font=FONT)
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.2)
        R(p, line if line else " ", size=8, color=GRAY_CODE, font=CODE_FONT)


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="D0D0D0", sz="4"):
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>')
    tcPr.append(borders)


def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, HEX_NAVY)
        set_cell_borders(c, HEX_NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(p, h, bold=True, size=9, color=WHITE)
    for ri, row in enumerate(rows):
        bg = "FFFFFF" if ri % 2 == 0 else HEX_LIGHT
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            set_cell_shading(c, bg)
            set_cell_borders(c, "E0E0E0")
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if ci == 0:
                R(p, str(val), bold=True, size=9, color=TEAL)
            else:
                R(p, str(val), size=9, color=GRAY_BODY)
    return t


def spacer(pts=6):
    P("", after=pts)


# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════

accent = doc.add_table(rows=1, cols=1)
accent.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = accent.rows[0].cells[0]
set_cell_shading(cell, HEX_GREEN)
p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
R(p, " ", size=6, color=GREEN)

for _ in range(4):
    P("")

P("Learning Guide: Foundations", bold=True, size=13, color=TEAL, after=6)
P("HCLS AI Factory",
  bold=True, size=32, color=NAVY, after=2)
P("High School Level (Grades 9-12)",
  bold=True, size=18, color=NAVY, after=6)
P("How scientists use computers and AI to read DNA, find disease-causing "
  "changes, and design new medicines \u2014 all on a single desktop computer. "
  "No prior biology or computer science required.",
  italic=False, size=12, color=NAVY, after=6)
P("NVIDIA DGX Spark  |  Parabricks  |  BioNeMo  |  Milvus  |  Claude",
  italic=True, size=11, color=GRAY_META, after=12)
P("")
P("02/2026  |  Version 1.0  |  Apache 2.0 License", size=9, color=GRAY_META, after=2)
P("Author: Adam Jones", size=9, color=GRAY_META, after=2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
H1("Table of Contents")

toc_items = [
    "1.  What Is DNA?",
    "2.  Reading DNA \u2014 Sequencing",
    "3.  Stage 1 \u2014 Finding Variants with GPU Power",
    "4.  Stage 2 \u2014 Understanding What the Variants Mean",
    "5.  Stage 3 \u2014 Designing New Medicines",
    "6.  The VCP/FTD Demo",
    "7.  The Hardware \u2014 NVIDIA DGX Spark",
    "8.  Why This Matters",
    "Glossary",
    "Review Questions",
]
for item in toc_items:
    P(item, size=10.5, color=TEAL, after=3)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 1: WHAT IS DNA?
# ══════════════════════════════════════════════════════════════
H1("Chapter 1: What Is DNA?")

H2("Your Body\u2019s Instruction Manual")
body(
    "Every cell in your body contains DNA \u2014 a long molecule shaped like a "
    "twisted ladder (the famous \"double helix\"). DNA is like an instruction "
    "manual written in a four-letter alphabet: A (adenine), T (thymine), "
    "C (cytosine), and G (guanine)."
)
body(
    "These four letters combine to form \"words\" called genes. Humans have "
    "about 20,000 genes, and each gene contains instructions for building a "
    "specific protein \u2014 the molecular machines that do almost everything "
    "in your body."
)

H2("The Human Genome")
body(
    "Your complete set of DNA instructions is called your genome. It contains "
    "about 3.1 billion letter pairs, organized into 23 pairs of chromosomes "
    "(46 total). If you stretched out all the DNA in one cell, it would be "
    "about 6 feet long \u2014 but it\u2019s coiled so tightly that it fits inside "
    "a cell nucleus smaller than the period at the end of this sentence."
)

H2("What Are Variants?")
body(
    "No two people have exactly the same DNA (except identical twins). The "
    "differences between your DNA and a \"reference\" human genome are called "
    "variants. Most variants are harmless \u2014 they\u2019re what make you unique. "
    "But some variants can cause diseases by changing how a protein works."
)

H3("Key Vocabulary")
bullet("DNA \u2014", "the molecule that stores genetic instructions")
bullet("Gene \u2014", "a section of DNA that codes for one protein")
bullet("Genome \u2014", "your complete set of DNA (~3.1 billion letters)")
bullet("Chromosome \u2014", "a package of DNA (humans have 23 pairs)")
bullet("Variant \u2014", "a difference in your DNA compared to a reference")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 2: READING DNA — SEQUENCING
# ══════════════════════════════════════════════════════════════
H1("Chapter 2: Reading DNA \u2014 Sequencing")

H2("How Do Scientists Read DNA?")
body(
    "A machine called a DNA sequencer (made by a company called Illumina) "
    "reads your DNA by chopping it into millions of small pieces, reading "
    "each piece, and then using computers to put the puzzle back together."
)

H2("Whole-Genome Sequencing (WGS)")
body(
    "Whole-genome sequencing reads your entire genome \u2014 all 3.1 billion "
    "letters. The version used in this platform is called \"30\u00d7 coverage,\" "
    "which means every position in your genome is read about 30 times. "
    "Reading it multiple times helps catch errors."
)

H2("FASTQ Files")
body(
    "The sequencer produces huge files called FASTQ files. For a single "
    "person\u2019s whole genome at 30\u00d7 coverage, the FASTQ files are about "
    "200 gigabytes \u2014 that\u2019s roughly the same as 50 HD movies! The FASTQ "
    "file contains billions of short \"reads\" \u2014 each one about 250 letters "
    "long."
)

H3("Key Numbers")
add_table(
    ["Metric", "Value"],
    [
        ["Human genome size", "3.1 billion base pairs"],
        ["Data per WGS run", "~200 GB"],
        ["Read length", "250 letters (paired-end: two reads per fragment)"],
        ["Coverage", "30\u00d7 (every position read ~30 times)"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 3: STAGE 1 — FINDING VARIANTS WITH GPU POWER
# ══════════════════════════════════════════════════════════════
H1("Chapter 3: Stage 1 \u2014 Finding Variants with GPU Power")

H2("The Computer Challenge")
body(
    "Once you have 200 GB of raw sequencing data, you need to:"
)
bullet("1.", "Align each of the billions of short reads to the right position "
       "in the reference genome")
bullet("2.", "Call variants \u2014 figure out where your DNA differs from the "
       "reference")
body(
    "On a regular computer, this takes 1-2 days. That\u2019s too slow for "
    "clinical use."
)

H2("GPU Acceleration: NVIDIA Parabricks")
body(
    "A GPU (Graphics Processing Unit) is a special computer chip originally "
    "designed for video games. It turns out GPUs are also great at biology "
    "\u2014 they can process millions of DNA reads simultaneously."
)
body(
    "NVIDIA Parabricks is software that uses a GPU to do both alignment and "
    "variant calling. On the NVIDIA DGX Spark (a desktop computer that costs "
    "$4,699), Parabricks completes the entire process in about 1 hour instead "
    "of 1-2 days. That\u2019s a 10-20\u00d7 speedup!"
)

H2("BWA-MEM2: Alignment")
body(
    "BWA-MEM2 is the tool that aligns each short read to the reference "
    "genome. Think of it like finding where each puzzle piece goes in a "
    "3.1-billion-piece puzzle. On the GPU, this takes 20-45 minutes."
)
add_table(
    ["Metric", "Value"],
    [
        ["Duration", "20-45 minutes"],
        ["GPU Utilization", "70-90%"],
        ["Peak Memory", "~40 GB"],
        ["Output", "Sorted BAM + BAI index"],
    ],
)
spacer()

H2("DeepVariant: Finding Differences")
body(
    "Google DeepVariant uses a type of AI called a convolutional neural "
    "network (CNN) \u2014 the same kind of AI used to recognize faces in photos "
    "\u2014 to identify variants. It looks at the aligned reads and determines "
    "which differences are real variants versus sequencing errors. "
    "DeepVariant is >99% accurate and takes 10-35 minutes on the GPU."
)
add_table(
    ["Metric", "Value"],
    [
        ["Duration", "10-35 minutes"],
        ["GPU Utilization", "80-95%"],
        ["Peak Memory", "~60 GB"],
        ["Accuracy", ">99% (CNN-based)"],
    ],
)
spacer()

H2("The VCF File")
body(
    "The output is a VCF (Variant Call Format) file containing every variant "
    "found. For a typical human genome, this includes about 11.7 million "
    "variant records. That is a count of candidate rows in the file, not of "
    "confident calls \u2014 about 4.69 million carry the caller\u2019s PASS filter, "
    "and about 3.56 million clear the stricter quality threshold the platform uses."
)

H3("Key Concepts")
bullet("GPU \u2014", "a computer chip that\u2019s very fast at parallel processing")
bullet("Alignment \u2014", "matching short reads to the reference genome")
bullet("Variant calling \u2014", "identifying where your DNA differs from the "
       "reference")
bullet("VCF \u2014", "a file listing all the variants found")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 4: STAGE 2 — UNDERSTANDING WHAT THE VARIANTS MEAN
# ══════════════════════════════════════════════════════════════
H1("Chapter 4: Stage 2 \u2014 Understanding What the Variants Mean")

H2("Not All Variants Are Equal")
body(
    "Of the 11.7 million variant records found, most are harmless. The challenge is "
    "finding the few that actually cause disease. This is like finding a "
    "needle in a haystack \u2014 except the haystack has 11.7 million pieces "
    "of hay."
)

H2("Three Annotation Databases")

H3("ClinVar")
body(
    "A public database maintained by the National Institutes of Health (NIH). "
    "It contains 4.1 million variants that scientists have studied and "
    "classified as:"
)
bullet("Pathogenic \u2014", "known to cause disease")
bullet("Likely pathogenic \u2014", "probably causes disease")
bullet("VUS \u2014", "Variant of Uncertain Significance (we don\u2019t know yet)")
bullet("Likely benign \u2014", "probably harmless")
bullet("Benign \u2014", "known to be harmless")

H3("AlphaMissense")
body(
    "An AI tool from DeepMind (the same company that created AlphaFold, "
    "which won the Nobel Prize for predicting protein structures). "
    "AlphaMissense predicts how likely a variant is to cause disease, "
    "scoring each variant from 0 to 1. It covers 71 million missense "
    "variant predictions. A score above 0.564 means \"likely pathogenic.\""
)

H3("VEP (Variant Effect Predictor)")
body(
    "Tells you what the variant does to the protein. Does it change an "
    "amino acid? Does it break the protein? Does it have no effect? "
    "VEP classifies the impact as HIGH, MODERATE, LOW, or MODIFIER."
)

H2("The Annotation Funnel")
body(
    "Starting with 11.7 million variant records, the platform narrows down to "
    "the most important ones:"
)
add_table(
    ["Stage", "Variant Count", "Filter"],
    [
        ["Raw VCF", "~11.7M", "\u2014"],
        ["Quality filter", "~3.5M", "QUAL > 30"],
        ["ClinVar match", "~35,616", "Clinical significance annotated"],
        ["AlphaMissense match", "~6,831", "AI pathogenicity predicted"],
        ["High impact + pathogenic", "~2,400", "Actionable subset"],
        ["In druggable genes", "~847", "Targetable by medicines"],
    ],
)
spacer()

H2("Vector Database: Milvus")
body(
    "To search through 3.5 million annotated variants quickly, the platform "
    "uses a vector database called Milvus. Each variant is converted into a "
    "list of 384 numbers (called an \"embedding\") using a model called "
    "BGE-small-en-v1.5. These embeddings capture the meaning of each "
    "variant, so you can search for similar variants using natural language "
    "questions."
)
add_table(
    ["Parameter", "Value"],
    [
        ["Database", "Milvus"],
        ["Total embeddings", "~3.5M"],
        ["Embedding dimensions", "384"],
        ["Embedding model", "BGE-small-en-v1.5"],
        ["Distance metric", "COSINE"],
    ],
)
spacer()

H2("Claude: AI-Powered Reasoning")
body(
    "Anthropic Claude is a large language model (like ChatGPT) that reads "
    "the variant evidence and the knowledge base to identify the best drug "
    "targets. It\u2019s \"grounded\" in the actual data \u2014 it can only cite "
    "variants and evidence that actually exist in the database."
)
body(
    "The knowledge base contains 201 genes across 13 therapeutic areas "
    "(like neurology, oncology, and cardiovascular disease). Of these, "
    "171 genes (85%) are known to be \"druggable\" \u2014 meaning scientists "
    "know how to design medicines that target them."
)

H3("Key Vocabulary")
bullet("Annotation \u2014", "adding information about what each variant means")
bullet("Pathogenic \u2014", "disease-causing")
bullet("Embedding \u2014", "a mathematical representation of text")
bullet("Vector database \u2014", "a database that finds similar items by "
       "mathematical similarity")
bullet("RAG \u2014", "Retrieval-Augmented Generation \u2014 feeding real data "
       "to an AI to ground its answers")
bullet("Druggable \u2014", "a protein that can be targeted by a medicine")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 5: STAGE 3 — DESIGNING NEW MEDICINES
# ══════════════════════════════════════════════════════════════
H1("Chapter 5: Stage 3 \u2014 Designing New Medicines")

H2("From Gene to Drug")
body(
    "Once the AI identifies a disease-causing variant in a druggable gene, "
    "the next step is designing new medicines. This is normally a process "
    "that takes years and costs billions of dollars. The HCLS AI Factory "
    "does the first step \u2014 generating candidate molecules \u2014 in "
    "8-16 minutes."
)

H2("Step 1: Finding the Protein Structure")
body(
    "Before you can design a drug, you need to know what the target protein "
    "looks like in 3D. The platform queries the RCSB Protein Data Bank "
    "(PDB) \u2014 a public database of protein structures determined by X-ray "
    "crystallography and cryo-electron microscopy (Cryo-EM)."
)
body(
    "For the VCP protein (the demo target), four structures are available. "
    "The best one (called 5FTK) shows the protein with an existing drug "
    "(CB-5083) already bound to it. This tells us exactly where new drugs "
    "should attach."
)

H2("Step 2: Generating New Molecules (MolMIM)")
body(
    "BioNeMo MolMIM is an AI from NVIDIA that generates new molecule "
    "designs. You give it a \"seed\" molecule (like CB-5083) and it creates "
    "100 new molecules that are similar but different \u2014 like variations "
    "on a theme. It uses a technique called masked language modeling \u2014 "
    "the same approach that powers text AI, but applied to molecular "
    "structures instead of words."
)

H2("Step 3: Checking if They Work (DiffDock)")
body(
    "BioNeMo DiffDock is another NVIDIA AI that predicts whether each new "
    "molecule will actually bind to the target protein. It uses a diffusion "
    "model (similar to AI image generators like DALL-E) to predict the 3D "
    "binding pose and calculate a docking score \u2014 a number that indicates "
    "how strongly the molecule binds."
)
body(
    "A docking score below -8.0 kcal/mol is considered excellent. The best "
    "candidate in the VCP demo scores -11.4 kcal/mol \u2014 significantly "
    "better than the original CB-5083 drug (-8.1 kcal/mol)."
)
add_table(
    ["Score (kcal/mol)", "Interpretation"],
    [
        ["-12 to -8", "Excellent binding affinity"],
        ["-8 to -6", "Good binding affinity"],
        ["-6 to -4", "Moderate binding affinity"],
        ["> -4", "Weak binding affinity"],
    ],
)
spacer()

H2("Step 4: Drug-Likeness (RDKit)")
body(
    "Not every molecule that binds to a protein would make a good drug. "
    "RDKit is a chemistry toolkit that checks whether each molecule has "
    "properties that would make it a practical medicine:"
)
bullet("Lipinski\u2019s Rule of Five \u2014", "a set of rules about molecular "
       "weight, fat-solubility, and other properties that predict whether "
       "a drug can be taken as a pill")
bullet("QED \u2014", "Quantitative Estimate of Drug-likeness \u2014 a single "
       "number (0-1) that combines multiple drug-like properties. "
       "Above 0.67 is considered drug-like")
bullet("TPSA \u2014", "a measure of how well the molecule can cross cell "
       "membranes")
spacer()

H2("Step 5: Final Ranking")
body(
    "Each candidate gets a composite score based on:"
)
bullet("30%", "how confident the AI is in the molecule design (generation score)")
bullet("40%", "how well it binds to the protein (docking score)")
bullet("30%", "how drug-like it is (QED score)")
body(
    "The final output is 100 ranked drug candidates with a PDF report."
)

H3("Key Vocabulary")
bullet("Protein structure \u2014", "the 3D shape of a protein")
bullet("Cryo-EM \u2014", "a technique for determining protein structures "
       "using electron microscopes")
bullet("Docking \u2014", "predicting how a molecule fits into a protein\u2019s "
       "binding site")
bullet("Lipinski\u2019s Rule of Five \u2014", "rules for predicting if a "
       "molecule can be a pill")
bullet("QED \u2014", "a score measuring how drug-like a molecule is")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 6: THE VCP/FTD DEMO
# ══════════════════════════════════════════════════════════════
H1("Chapter 6: The VCP/FTD Demo")

H2("What Is VCP?")
body(
    "VCP (Valosin-Containing Protein, also called p97) is a protein that "
    "acts like a cellular recycling machine. It helps cells break down "
    "damaged or unwanted proteins through a process called the "
    "ubiquitin-proteasome system."
)
body(
    "When VCP has a disease-causing mutation, it can cause:"
)
bullet("Frontotemporal Dementia (FTD) \u2014", "a brain disease that affects "
       "personality, behavior, and language")
bullet("ALS \u2014", "Amyotrophic Lateral Sclerosis \u2014 a disease that "
       "destroys motor neurons")
bullet("IBMPFD \u2014", "a condition affecting muscles, bones, and the brain")

H2("The Demo Variant")
body(
    "The variant rs188935092 (at position chr9:35065263, where G changes "
    "to A) is classified as Pathogenic by ClinVar and scores 0.87 on "
    "AlphaMissense (well above the 0.564 threshold for pathogenic)."
)
add_table(
    ["Parameter", "Value"],
    [
        ["Gene", "VCP"],
        ["Protein", "p97 / Valosin-Containing Protein"],
        ["UniProt", "P55072"],
        ["Variant", "rs188935092 (chr9:35065263 G>A)"],
        ["ClinVar", "Pathogenic"],
        ["AlphaMissense", "0.87 (pathogenic, >0.564 threshold)"],
        ["Diseases", "FTD, ALS, IBMPFD"],
        ["Seed Compound", "CB-5083 (Phase I clinical VCP inhibitor)"],
    ],
)
spacer()

H2("Demo Results")
body(
    "The demo produces 100 novel VCP inhibitor candidates. The top "
    "candidate improves on the CB-5083 seed compound:"
)
add_table(
    ["Metric", "CB-5083 (Seed)", "Top Candidate", "Improvement"],
    [
        ["Composite score", "0.64", "0.89", "+39%"],
        ["Docking score", "-8.1 kcal/mol", "-11.4 kcal/mol", "41% stronger binding"],
        ["QED", "0.62", "0.81", "31% more drug-like"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 7: THE HARDWARE — NVIDIA DGX SPARK
# ══════════════════════════════════════════════════════════════
H1("Chapter 7: The Hardware \u2014 NVIDIA DGX Spark")

H2("A Supercomputer on Your Desk")
body(
    "The entire pipeline runs on the NVIDIA DGX Spark \u2014 a desktop "
    "computer that costs $4,699. Here are its specifications:"
)
add_table(
    ["Component", "Specification"],
    [
        ["GPU", "NVIDIA GB10"],
        ["Memory", "128 GB unified LPDDR5x (CPU + GPU shared)"],
        ["CPU", "NVIDIA Grace ARM64, ARM64 cores"],
        ["Storage", "NVMe (fast storage for ~200 GB genomic data)"],
        ["System RAM", "512 GB"],
        ["Price", "$4,699"],
    ],
)
spacer()

H2("Why \"Unified Memory\" Matters")
body(
    "In most computers, the CPU and GPU have separate memory, and data "
    "must be copied back and forth. In DGX Spark, the CPU and GPU share "
    "the same 128 GB of memory. This eliminates copying overhead and "
    "makes everything faster."
)

H2("Scaling Up")
body(
    "The same software that runs on a $4,699 DGX Spark can scale to "
    "larger systems:"
)
add_table(
    ["Phase", "Hardware", "Price", "Scale"],
    [
        ["1 \u2014 Proof Build", "DGX Spark", "$4,699", "Single patient, desktop"],
        ["2 \u2014 Departmental", "DGX B200", "$500K-$1M", "Multiple patients simultaneously"],
        ["3 \u2014 Enterprise", "DGX SuperPOD", "$7M-$60M+", "Thousands of patients"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 8: WHY THIS MATTERS
# ══════════════════════════════════════════════════════════════
H1("Chapter 8: Why This Matters")

H2("Traditional vs. HCLS AI Factory")
add_table(
    ["Step", "Traditional", "HCLS AI Factory"],
    [
        ["Sequence alignment", "12-24 hours (CPU)", "20-45 min (GPU)"],
        ["Variant calling", "8-12 hours (CPU)", "10-35 min (GPU)"],
        ["Annotation", "Days (manual)", "Minutes (automated)"],
        ["Target identification", "Weeks (literature review)", "Minutes (Claude RAG)"],
        ["Drug candidate generation", "Months (medicinal chemistry)", "8-16 min (BioNeMo AI)"],
        ["Total", "Weeks to months", "< 5 hours"],
    ],
)
spacer()

H2("The Bigger Picture")
body(
    "This platform is part of the HCLS AI Factory \u2014 a broader ecosystem "
    "that also includes:"
)
bullet("Imaging Intelligence Agent \u2014", "AI analysis of CT scans, MRI, "
       "and X-rays")
bullet("Cross-modal triggers \u2014", "for example, a suspicious lung nodule "
       "found on a CT scan can automatically trigger genomic analysis to "
       "look for cancer-related variants")
bullet("NVIDIA FLARE \u2014", "technology that lets multiple hospitals train "
       "AI models together without sharing patient data")

H2("Open Source")
body(
    "The entire platform is released under the Apache 2.0 license \u2014 "
    "meaning anyone can use, modify, and share it for free. This is "
    "important because it means:"
)
bullet("\u2022", "Any hospital can run it")
bullet("\u2022", "Researchers can verify and improve the methods")
bullet("\u2022", "No expensive software licenses required")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# GLOSSARY
# ══════════════════════════════════════════════════════════════
H1("Glossary")

add_table(
    ["Term", "Definition"],
    [
        ["Alignment", "Matching short DNA reads to the correct position in a reference genome"],
        ["AlphaMissense", "An AI tool that predicts whether a DNA variant is disease-causing"],
        ["ARM64", "A type of computer processor architecture (used in DGX Spark)"],
        ["ClinVar", "A public database of disease-related DNA variants"],
        ["Chromosome", "A package of DNA (humans have 23 pairs, 46 total)"],
        ["Composite Score", "A weighted combination of generation, docking, and QED scores"],
        ["Cryo-EM", "A technique for determining protein structures using frozen samples"],
        ["DeepVariant", "Google\u2019s AI for identifying DNA variants (>99% accuracy)"],
        ["DGX Spark", "NVIDIA\u2019s $4,699 desktop AI computer"],
        ["DiffDock", "NVIDIA\u2019s AI for predicting how molecules bind to proteins"],
        ["DNA", "Deoxyribonucleic acid \u2014 the molecule that stores genetic instructions"],
        ["Docking", "Predicting how a drug molecule fits into a protein"],
        ["Druggable", "A protein that can be targeted by a medicine"],
        ["Embedding", "A mathematical representation of text or data as a list of numbers"],
        ["FASTQ", "The file format for raw DNA sequencing data"],
        ["FTD", "Frontotemporal Dementia \u2014 a brain disease affecting personality"],
        ["Gene", "A section of DNA that codes for one protein"],
        ["Genome", "Your complete set of DNA (~3.1 billion letters)"],
        ["GPU", "Graphics Processing Unit \u2014 a chip for fast parallel computing"],
        ["Lipinski", "Rules predicting whether a molecule can be taken as a pill"],
        ["Milvus", "A vector database for fast similarity search"],
        ["MolMIM", "NVIDIA\u2019s AI for generating new molecule designs"],
        ["Parabricks", "NVIDIA\u2019s GPU-accelerated genomics software"],
        ["Pathogenic", "Disease-causing"],
        ["QED", "Quantitative Estimate of Drug-likeness (0-1 score)"],
        ["RAG", "Retrieval-Augmented Generation \u2014 grounding AI in real data"],
        ["TPSA", "Topological Polar Surface Area \u2014 predicts membrane permeability"],
        ["Variant", "A difference in your DNA compared to a reference genome"],
        ["VCF", "Variant Call Format \u2014 a file listing all detected DNA variants"],
        ["VCP", "Valosin-Containing Protein \u2014 the demo target gene"],
        ["VEP", "Variant Effect Predictor \u2014 classifies variant impacts"],
        ["WGS", "Whole-Genome Sequencing \u2014 reading the entire genome"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# REVIEW QUESTIONS
# ══════════════════════════════════════════════════════════════
H1("Review Questions")

body(
    "Test your understanding of the HCLS AI Factory by answering these "
    "questions. Refer back to the chapters for help."
)
spacer()

questions = [
    ("1.", "How many letters are in the human genome alphabet, and what are they?"),
    ("2.", "Why does whole-genome sequencing read each position ~30 times?"),
    ("3.", "What makes GPU processing faster than CPU processing for genomics?"),
    ("4.", "Name the three databases used to annotate variants and what each one does."),
    ("5.", "What does a QED score above 0.67 mean?"),
    ("6.", "How much does the DGX Spark hardware cost?"),
    ("7.", "What is the composite scoring formula for ranking drug candidates?"),
    ("8.", "Why is the VCP gene important in the demo?"),
    ("9.", "What does \"RAG\" stand for and why is it important for AI accuracy?"),
    ("10.", "How does the HCLS AI Factory connect to medical imaging?"),
]
for num, question in questions:
    bullet(num, question)

spacer(12)
p = P("", after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
R(p, "HCLS AI Factory Learning Guide: Foundations \u2014 Apache 2.0 | "
  "Author: Adam Jones | February 2026", italic=True, size=10, color=GRAY_META)

# ══════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════
import pathlib, subprocess, shutil

OUT = pathlib.Path(__file__).resolve().parent
docx_path = OUT / "HCLS_AI_Factory_Learning_Guide_Foundations.docx"
pdf_path = OUT / "HCLS_AI_Factory_Learning_Guide_Foundations.pdf"

doc.save(str(docx_path))
print(f"DOCX saved: {docx_path} ({docx_path.stat().st_size:,} bytes)")

# Attempt PDF conversion
if shutil.which("libreoffice"):
    subprocess.run([
        "libreoffice", "--headless", "--convert-to", "pdf",
        "--outdir", str(OUT), str(docx_path)
    ], capture_output=True, timeout=120)
    if pdf_path.exists():
        print(f"PDF  saved: {pdf_path} ({pdf_path.stat().st_size:,} bytes)")
    else:
        print("PDF conversion completed but file not found.")
elif shutil.which("docx2pdf"):
    subprocess.run(["docx2pdf", str(docx_path), str(pdf_path)],
                    capture_output=True, timeout=120)
    if pdf_path.exists():
        print(f"PDF  saved: {pdf_path} ({pdf_path.stat().st_size:,} bytes)")
else:
    print("No PDF converter found (install libreoffice or docx2pdf).")

print("Done.")
