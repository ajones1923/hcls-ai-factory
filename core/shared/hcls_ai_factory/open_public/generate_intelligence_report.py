#!/usr/bin/env python3
# Copyright 2026 Adam Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Generate HCLS AI Factory — Intelligence Report (DOCX + PDF).

Mock clinical intelligence report showing the complete pipeline output
for the VCP/FTD demo case: genomics, annotation, target identification,
and drug candidate ranking.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Colors ─────────────────────────────────────────────────
NAVY = RGBColor(0x1B, 0x23, 0x33)
TEAL = RGBColor(0x1A, 0xAF, 0xCC)
GREEN = RGBColor(0x76, 0xB9, 0x00)
GRAY_BODY = RGBColor(0x33, 0x33, 0x33)
GRAY_META = RGBColor(0x66, 0x66, 0x66)
GRAY_CODE = RGBColor(0x44, 0x44, 0x44)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED_FLAG = RGBColor(0xDC, 0x26, 0x26)
GREEN_FLAG = RGBColor(0x05, 0x96, 0x69)

HEX_NAVY = "1B2333"
HEX_TEAL = "1AAFCC"
HEX_GREEN = "76B900"
HEX_LIGHT = "F8FAFB"
HEX_RED_BG = "FEE2E2"
HEX_GREEN_BG = "D1FAE5"
FONT = "Calibri"
CODE_FONT = "Consolas"

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(1.5)
    s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(2.0)
    s.right_margin = Cm(2.0)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
style.font.color.rgb = GRAY_BODY
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15


def R(p, text, bold=False, italic=False, size=None, color=None, font=None):
    r = p.add_run(text)
    r.font.name = font or FONT
    if bold: r.bold = True
    if italic: r.italic = True
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return r


def P(text="", bold=False, italic=False, size=10.5, color=GRAY_BODY,
      before=0, after=4, align=None):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text: R(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def H1(text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs: r.clear()
    R(p, text, bold=True, size=20, color=NAVY)
    return p


def H2(text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs: r.clear()
    R(p, text, bold=True, size=14, color=NAVY)
    return p


def H3(text):
    return P(text, bold=True, size=12, color=TEAL, before=8, after=3)


def body(text, before=0, after=4):
    return P(text, size=10.5, color=GRAY_BODY, before=before, after=after)


def bullet(lead, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25)
    R(p, lead, bold=True, size=10, color=GRAY_BODY)
    R(p, f"  {text}", size=10, color=GRAY_BODY)
    return p


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="D0D0D0", sz="4"):
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>')
    tcPr.append(borders)


def add_table(headers, rows, compact=False):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    sz = 8.5 if compact else 9
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, HEX_NAVY)
        set_cell_borders(c, HEX_NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(p, h, bold=True, size=sz, color=WHITE)
    for ri, row in enumerate(rows):
        bg = "FFFFFF" if ri % 2 == 0 else HEX_LIGHT
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            set_cell_shading(c, bg)
            set_cell_borders(c, "E0E0E0")
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if ci == 0:
                R(p, str(val), bold=True, size=sz, color=TEAL)
            else:
                R(p, str(val), size=sz, color=GRAY_BODY)
    return t


def spacer(pts=4):
    P("", after=pts)


def status_badge(text, hex_bg, color):
    """Inline status badge."""
    p = P("", after=2)
    r = p.add_run(f"  {text}  ")
    r.font.name = FONT
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = color
    return p


# ══════════════════════════════════════════════════════════════
# COVER / HEADER
# ══════════════════════════════════════════════════════════════

# Green accent bar
accent = doc.add_table(rows=1, cols=1)
accent.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = accent.rows[0].cells[0]
set_cell_shading(cell, HEX_GREEN)
p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
R(p, " ", size=6, color=GREEN)

spacer(4)

P("Intelligence Report", bold=True, size=13, color=TEAL, after=4)
P("HCLS AI Factory \u2014 Genomics to Drug Discovery",
  bold=True, size=26, color=NAVY, after=4)
P("Pipeline Run: HCLS-VCP-2026-0087",
  bold=True, size=14, color=NAVY, after=6)

# Header table
add_table(
    ["Field", "Value"],
    [
        ["Patient ID", "GEN-2026-0087"],
        ["Run ID", "HCLS-VCP-2026-0087"],
        ["Pipeline Version", "HLS-Pipeline v1.0.0"],
        ["Pipeline Mode", "Full (Genomics \u2192 RAG/Chat \u2192 Drug Discovery)"],
        ["Hardware", "NVIDIA DGX Spark (GB10 GPU, 128 GB unified)"],
        ["Total Duration", "4 hours 12 minutes"],
        ["Report Date", "February 2026"],
        ["Status", "COMPLETE \u2014 100 novel drug candidates ranked"],
    ],
)

spacer(8)
p = P("", after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
R(p, "PIPELINE COMPLETE", bold=True, size=11, color=GREEN_FLAG)
R(p, "  |  ", size=11, color=GRAY_META)
R(p, "PRIMARY TARGET: VCP (Pathogenic)", bold=True, size=11, color=RED_FLAG)
R(p, "  |  ", size=11, color=GRAY_META)
R(p, "100 CANDIDATES RANKED", bold=True, size=11, color=TEAL)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. GENOMICS SUMMARY
# ══════════════════════════════════════════════════════════════
H1("1. Genomics Summary (Stage 1)")

H2("Input Data")
add_table(
    ["Parameter", "Value"],
    [
        ["Sample", "HG002 (NA24385, GIAB Ashkenazi male)"],
        ["Sequencing", "Illumina, 30\u00d7 WGS, 2\u00d7250 bp paired-end"],
        ["FASTQ Size", "198.7 GB (R1: 99.4 GB, R2: 99.3 GB)"],
        ["Reference", "GRCh38 (3.1 GB)"],
    ],
)
spacer()

H2("Parabricks Execution")
add_table(
    ["Step", "Tool", "Duration", "GPU Util", "Peak Memory"],
    [
        ["Alignment", "BWA-MEM2 (fq2bam)", "34 min", "82%", "38 GB"],
        ["Variant Calling", "Google DeepVariant", "22 min", "91%", "54 GB"],
        ["Total Stage 1", "", "56 min", "", ""],
    ],
)
spacer()

H2("VCF Output")
add_table(
    ["Metric", "Count"],
    [
        ["Total Variants Called", "11,724,891"],
        ["PASS Quality (QUAL>30)", "3,487,216"],
        ["SNPs", "4,198,433"],
        ["Indels", "1,012,548"],
        ["Multi-allelic Sites", "148,762"],
        ["Coding Region Variants", "35,616"],
        ["Ts/Tv Ratio", "2.07"],
    ],
)
spacer()

p = P("", after=4)
R(p, "Quality Assessment: ", bold=True, size=10, color=GRAY_BODY)
R(p, "PASS", bold=True, size=10, color=GREEN_FLAG)
R(p, " \u2014 Ts/Tv ratio within expected range (2.0-2.1), variant counts "
  "consistent with 30\u00d7 WGS of Ashkenazi ancestry sample.",
  size=10, color=GRAY_BODY)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. ANNOTATION & TARGET IDENTIFICATION
# ══════════════════════════════════════════════════════════════
H1("2. Annotation & Target Identification (Stage 2)")

H2("Annotation Funnel")
add_table(
    ["Stage", "Variants", "Filter Applied"],
    [
        ["Raw VCF", "11,724,891", "\u2014"],
        ["Quality Filter", "3,487,216", "QUAL > 30"],
        ["ClinVar Annotated", "35,616", "Clinical significance match"],
        ["AlphaMissense Scored", "6,831", "AI pathogenicity prediction"],
        ["HIGH Impact + Pathogenic", "2,412", "Actionable subset"],
        ["Druggable Gene Targets", "847", "Knowledge base match"],
    ],
)
spacer()

H2("Top 5 Target Hypotheses (Claude RAG Analysis)")
add_table(
    ["Rank", "Gene", "Variant", "ClinVar", "AM Score", "Area", "Druggability"],
    [
        ["1", "VCP", "rs188935092", "Pathogenic", "0.87", "Neurology", "0.92"],
        ["2", "EGFR", "rs121913229", "Pathogenic", "0.79", "Oncology", "0.95"],
        ["3", "BRCA1", "rs80357914", "Pathogenic", "0.72", "Oncology", "0.78"],
        ["4", "PCSK9", "rs11591147", "Pathogenic", "0.68", "Cardiovascular", "0.88"],
        ["5", "CFTR", "rs75527207", "Pathogenic", "0.81", "Respiratory", "0.71"],
    ],
    compact=True,
)
spacer()

H2("Primary Target: VCP")
add_table(
    ["Parameter", "Value"],
    [
        ["Gene", "VCP (Valosin-Containing Protein / p97)"],
        ["UniProt", "P55072"],
        ["Function", "AAA+ ATPase, ubiquitin-proteasome pathway"],
        ["Diseases", "Frontotemporal Dementia (FTD), ALS, IBMPFD"],
        ["Variant", "rs188935092 \u2014 missense, HIGH impact"],
        ["ClinVar", "Pathogenic (reviewed by expert panel)"],
        ["AlphaMissense", "0.87 (pathogenic, >0.564 threshold)"],
        ["Druggability", "0.92 (D2 ATPase domain, ~450 \u00c5\u00b3)"],
        ["Known Inhibitors", "CB-5083 (Phase I), NMS-873"],
        ["Confidence", "HIGH \u2014 multiple independent evidence sources"],
    ],
)
spacer()

H3("Evidence Chain")
bullet("1.", "Genomic: rs188935092 at chr9:35065263 (G>A), heterozygous, QUAL=892")
bullet("2.", "Clinical: ClinVar Pathogenic for FTD/ALS/IBMPFD (expert panel)")
bullet("3.", "AI Prediction: AlphaMissense 0.87 (>0.564 pathogenic threshold)")
bullet("4.", "Functional: VEP missense_variant, HIGH impact, D2 ATPase domain")
bullet("5.", "Druggability: Known target \u2014 CB-5083 reached Phase I clinical trial")
bullet("6.", "Structural: 4 PDB structures including inhibitor-bound 5FTK")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. DRUG DISCOVERY RESULTS
# ══════════════════════════════════════════════════════════════
H1("3. Drug Discovery Results (Stage 3)")

H2("Structure Evidence")
add_table(
    ["PDB ID", "Resolution", "Method", "Description", "Score"],
    [
        ["5FTK", "2.3 \u00c5", "X-ray", "VCP D2 + CB-5083 inhibitor", "13.2 (selected)"],
        ["7K56", "2.5 \u00c5", "Cryo-EM", "VCP complex", "10.8"],
        ["8OOI", "2.9 \u00c5", "Cryo-EM", "WT VCP hexamer", "8.9"],
        ["9DIL", "3.2 \u00c5", "Cryo-EM", "Mutant VCP", "7.4"],
    ],
)
spacer()
body("Selected: 5FTK \u2014 inhibitor-bound (CB-5083), X-ray at 2.3 \u00c5. "
     "Binding site: D2 ATPase domain, key residues ALA464, GLY479, ASP320, GLY215.")

H2("Molecule Generation (MolMIM)")
add_table(
    ["Parameter", "Value"],
    [
        ["Seed Compound", "CB-5083 (ATP-competitive VCP inhibitor)"],
        ["NIM Endpoint", "MolMIM (port 8001)"],
        ["Molecules Generated", "100"],
        ["Chemically Valid", "98 (2 rejected by RDKit)"],
        ["Generation Time", "2 min 14 sec"],
    ],
)
spacer()

H2("Drug-Likeness Profile")
add_table(
    ["Metric", "Pass", "Fail", "Pass Rate"],
    [
        ["Lipinski Rule of Five", "87", "11", "88.8%"],
        ["QED > 0.67 (drug-like)", "72", "26", "73.5%"],
        ["QED > 0.49 (moderate+)", "91", "7", "92.9%"],
        ["TPSA < 140 \u00c5\u00b2", "94", "4", "95.9%"],
    ],
)
spacer()

H2("Molecular Docking (DiffDock)")
add_table(
    ["Parameter", "Value"],
    [
        ["NIM Endpoint", "DiffDock (port 8002)"],
        ["Protein Target", "5FTK (VCP D2 domain)"],
        ["Candidates Docked", "98"],
        ["Docking Time", "8 min 42 sec"],
        ["Mean Dock Score", "-7.4 kcal/mol"],
        ["Best Dock Score", "-11.4 kcal/mol"],
        ["Excellent (< -8.0)", "34 candidates"],
        ["Good+ (< -6.0)", "78 candidates"],
    ],
)
spacer()

H2("Top 10 Ranked Candidates")
body("Composite scoring: 30% Generation + 40% Docking + 30% QED")
add_table(
    ["Rank", "Composite", "Gen", "Dock", "QED", "MW", "LogP", "Lipinski"],
    [
        ["1", "0.89", "0.92", "-11.4", "0.81", "423.5", "3.2", "PASS"],
        ["2", "0.86", "0.88", "-10.8", "0.79", "441.2", "3.7", "PASS"],
        ["3", "0.84", "0.85", "-10.2", "0.82", "398.7", "2.9", "PASS"],
        ["4", "0.82", "0.91", "-9.8", "0.74", "467.1", "4.1", "PASS"],
        ["5", "0.81", "0.83", "-9.5", "0.78", "412.3", "3.4", "PASS"],
        ["6", "0.79", "0.87", "-9.1", "0.71", "455.8", "3.8", "PASS"],
        ["7", "0.78", "0.80", "-8.9", "0.76", "389.2", "2.7", "PASS"],
        ["8", "0.76", "0.84", "-8.7", "0.69", "478.4", "4.3", "PASS"],
        ["9", "0.75", "0.79", "-8.5", "0.73", "401.6", "3.1", "PASS"],
        ["10", "0.74", "0.82", "-8.2", "0.68", "448.9", "3.9", "PASS"],
    ],
    compact=True,
)
spacer()

H2("CB-5083 Seed Comparison")
add_table(
    ["Metric", "CB-5083 (Seed)", "Top Candidate", "Improvement"],
    [
        ["Dock Score", "-8.1 kcal/mol", "-11.4 kcal/mol", "+41% binding"],
        ["QED", "0.62", "0.81", "+31% drug-likeness"],
        ["MW", "487.2 Da", "423.5 Da", "-13% (better absorption)"],
        ["Composite", "0.64", "0.89", "+39% overall"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. PIPELINE PERFORMANCE
# ══════════════════════════════════════════════════════════════
H1("4. Pipeline Performance")

H2("Stage Timing")
add_table(
    ["Stage", "Duration", "GPU Util", "Peak Memory"],
    [
        ["1 \u2014 Genomics (fq2bam)", "34 min", "82%", "38 GB"],
        ["1 \u2014 Genomics (DeepVariant)", "22 min", "91%", "54 GB"],
        ["2 \u2014 Annotation", "18 min", "15% (CPU)", "12 GB"],
        ["2 \u2014 Milvus Indexing", "24 min", "35%", "22 GB"],
        ["2 \u2014 RAG/Chat", "45 min", "5%", "8 GB"],
        ["3 \u2014 Structure Retrieval", "2 min", "0% (I/O)", "2 GB"],
        ["3 \u2014 MolMIM Generation", "2 min 14 sec", "78%", "18 GB"],
        ["3 \u2014 DiffDock Docking", "8 min 42 sec", "85%", "24 GB"],
        ["3 \u2014 Scoring + Reporting", "1 min 30 sec", "0% (CPU)", "4 GB"],
        ["Total", "~4 hr 12 min", "", ""],
    ],
)
spacer()

H2("All Services Healthy")
add_table(
    ["Service", "Port", "Status"],
    [
        ["Landing Page", "8080", "HEALTHY"],
        ["Genomics Portal", "5000", "HEALTHY"],
        ["Milvus", "19530", "HEALTHY"],
        ["RAG API", "5001", "HEALTHY"],
        ["Streamlit Chat", "8501", "HEALTHY"],
        ["MolMIM NIM", "8001", "HEALTHY"],
        ["DiffDock NIM", "8002", "HEALTHY"],
        ["Discovery UI", "8505", "HEALTHY"],
        ["Grafana", "3000", "HEALTHY"],
        ["Prometheus", "9099", "HEALTHY"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. CLINICAL INTERPRETATION
# ══════════════════════════════════════════════════════════════
H1("5. Clinical Interpretation")

H2("Summary")
body(
    "Patient GEN-2026-0087 carries a heterozygous pathogenic missense variant "
    "(rs188935092) in the VCP gene. This variant is associated with "
    "Frontotemporal Dementia (FTD), ALS, and Inclusion Body Myopathy with "
    "Paget Disease and Frontotemporal Dementia (IBMPFD). The variant is "
    "classified as Pathogenic by ClinVar expert panel review and scores 0.87 "
    "on the AlphaMissense pathogenicity scale."
)

H2("Drug Discovery Outcome")
body(
    "The AI-driven drug discovery pipeline identified 100 novel VCP inhibitor "
    "candidates with the top candidate showing a 39% improvement in composite "
    "score over the CB-5083 seed compound. All top 10 candidates pass "
    "Lipinski\u2019s Rule of Five and show favorable QED scores (>0.67), "
    "suggesting oral drug-likeness."
)

H2("Recommended Actions")
bullet("1.", "Genetic counseling for FTD/ALS risk assessment")
bullet("2.", "Experimental validation of top 5 candidates in VCP ATPase assays")
bullet("3.", "ADMET profiling for lead optimization")
bullet("4.", "Cross-modal follow-up with Imaging Intelligence Agent for neurological assessment")

# ══════════════════════════════════════════════════════════════
# 6. PROVENANCE
# ══════════════════════════════════════════════════════════════
H1("6. Provenance")

add_table(
    ["Item", "Value"],
    [
        ["Pipeline", "HLS-Pipeline v1.0.0 (Nextflow DSL2)"],
        ["Parabricks", "nvcr.io/nvidia/clara/clara-parabricks:4.6.0-1"],
        ["DeepVariant", "Google DeepVariant (via Parabricks, >99%)"],
        ["Reference", "GRCh38 (3.1 GB)"],
        ["ClinVar", "February 2026 release (4.1M variants)"],
        ["AlphaMissense", "v1.0 (71,697,560 predictions)"],
        ["VEP", "Ensembl VEP (GRCh38)"],
        ["Milvus", "v2.4 (IVF_FLAT, nlist=1024, COSINE)"],
        ["Embedding", "BGE-small-en-v1.5 (384-dim)"],
        ["LLM", "claude-sonnet-4-20250514 (temp=0.3)"],
        ["MolMIM", "nvcr.io/nvidia/clara/bionemo-molmim:1.0"],
        ["DiffDock", "nvcr.io/nvidia/clara/diffdock:1.0"],
        ["Hardware", "NVIDIA DGX Spark (GB10, 128 GB)"],
        ["Scoring", "30% gen + 40% dock + 30% QED"],
    ],
)

spacer(12)
p = P("", after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
R(p, "This is a demonstration intelligence report. All patient data is synthetic.",
  italic=True, size=9, color=GRAY_META)
spacer(4)
p = P("", after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
R(p, "HCLS AI Factory \u2014 Apache 2.0  |  Author: Adam Jones  |  February 2026",
  italic=True, size=9, color=GRAY_META)

# ══════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════
import pathlib, subprocess, shutil

OUT = pathlib.Path(__file__).resolve().parent
docx_path = OUT / "HCLS_AI_Factory_Intelligence_Report.docx"
pdf_path = OUT / "HCLS_AI_Factory_Intelligence_Report.pdf"

doc.save(str(docx_path))
print(f"DOCX saved: {docx_path} ({docx_path.stat().st_size:,} bytes)")

if shutil.which("libreoffice"):
    subprocess.run([
        "libreoffice", "--headless", "--convert-to", "pdf",
        "--outdir", str(OUT), str(docx_path)
    ], capture_output=True, timeout=120)
    if pdf_path.exists():
        print(f"PDF  saved: {pdf_path} ({pdf_path.stat().st_size:,} bytes)")
    else:
        print("PDF conversion completed but file not found.")
elif shutil.which("docx2pdf"):
    subprocess.run(["docx2pdf", str(docx_path), str(pdf_path)],
                    capture_output=True, timeout=120)
    if pdf_path.exists():
        print(f"PDF  saved: {pdf_path} ({pdf_path.stat().st_size:,} bytes)")
else:
    print("No PDF converter found (install libreoffice or docx2pdf).")

print("Done.")
