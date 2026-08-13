#!/usr/bin/env python3
# Copyright 2026 Adam Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Generate HCLS AI Factory — Executive Bullets (DOCX + PDF).

One-page reference for executives, stakeholders, and demo audiences.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

NAVY = RGBColor(0x1B, 0x23, 0x33)
TEAL = RGBColor(0x1A, 0xAF, 0xCC)
GREEN = RGBColor(0x76, 0xB9, 0x00)
GRAY_BODY = RGBColor(0x33, 0x33, 0x33)
GRAY_META = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_NAVY = "1B2333"
HEX_TEAL = "1AAFCC"
HEX_GREEN = "76B900"
HEX_LIGHT = "F8FAFB"
FONT = "Calibri"

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(1.5)
    s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(2.0)
    s.right_margin = Cm(2.0)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10)
style.font.color.rgb = GRAY_BODY
style.paragraph_format.space_after = Pt(3)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15


def R(p, text, bold=False, italic=False, size=None, color=None):
    r = p.add_run(text)
    r.font.name = FONT
    if bold: r.bold = True
    if italic: r.italic = True
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return r


def P(text="", bold=False, italic=False, size=10, color=GRAY_BODY,
      before=0, after=3, align=None):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text: R(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def H1(text):
    return P(text, bold=True, size=16, color=NAVY, before=12, after=4)


def H2(text):
    return P(text, bold=True, size=12, color=TEAL, before=8, after=3)


def body(text, before=0, after=3):
    return P(text, size=10, color=GRAY_BODY, before=before, after=after)


def bullet(lead, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.2)
    R(p, lead, bold=True, size=10, color=GRAY_BODY)
    R(p, f" {text}", size=10, color=GRAY_BODY)
    return p


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="D0D0D0", sz="4"):
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>')
    tcPr.append(borders)


def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, HEX_NAVY)
        set_cell_borders(c, HEX_NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(p, h, bold=True, size=8.5, color=WHITE)
    for ri, row in enumerate(rows):
        bg = "FFFFFF" if ri % 2 == 0 else HEX_LIGHT
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            set_cell_shading(c, bg)
            set_cell_borders(c, "E0E0E0")
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if ci == 0:
                R(p, str(val), bold=True, size=8.5, color=TEAL)
            else:
                R(p, str(val), size=8.5, color=GRAY_BODY)
    return t


def spacer(pts=3):
    P("", after=pts)


# ══════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════

accent = doc.add_table(rows=1, cols=1)
accent.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = accent.rows[0].cells[0]
set_cell_shading(cell, HEX_GREEN)
p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
R(p, " ", size=4, color=GREEN)

spacer(2)
P("Executive Bullets", bold=True, size=11, color=TEAL, after=2)
P("HCLS AI Factory", bold=True, size=24, color=NAVY, after=2)
P("From Patient DNA to Novel Drug Candidates in < 5 Hours",
  bold=True, size=14, color=NAVY, after=4)
P("NVIDIA DGX Spark  |  Parabricks  |  BioNeMo  |  Milvus  |  Claude",
  italic=True, size=9, color=GRAY_META, after=2)
P("02/2026  |  Version 1.0  |  Apache 2.0 License  |  Author: Adam Jones",
  size=8, color=GRAY_META, after=6)

# ── What It Is ─────────────────────────────────────────────
H1("What It Is")
body(
    "The HCLS AI Factory transforms patient DNA into ranked novel drug candidates "
    "in under 5 hours on a single NVIDIA DGX Spark ($4,699). Three GPU-accelerated "
    "stages run end-to-end with no manual intervention."
)

# ── The Problem ────────────────────────────────────────────
H1("The Problem")
bullet("\u2022", "CPU-based genomics: 12-36 hours for a single 30\u00d7 WGS sample")
bullet("\u2022", "Variant annotation fragmented across disconnected databases")
bullet("\u2022", "Variant-to-drug-lead gap: months of manual work")
bullet("\u2022", "No integrated platform connects genomics \u2192 reasoning \u2192 drug discovery")

# ── Three Stages ───────────────────────────────────────────
H1("The Solution \u2014 Three Stages")

H2("Stage 1: GPU-Accelerated Genomics (120-240 min)")
bullet("\u2022", "NVIDIA Parabricks 4.6 \u2014 10-20\u00d7 faster than CPU")
bullet("\u2022", "BWA-MEM2 alignment: 20-45 min | DeepVariant: 10-35 min, >99% accuracy")
bullet("\u2022", "Input: ~200 GB FASTQ | Output: VCF (~11.7M variant records)")

H2("Stage 2: RAG-Grounded Target Identification (Interactive)")
bullet("\u2022", "ClinVar (4.1M) + AlphaMissense (71M) + VEP annotation")
bullet("\u2022", "3.5M variant embeddings in Milvus | BGE-small-en-v1.5 (384-dim)")
bullet("\u2022", "Claude RAG reasoning | 201 genes, 13 therapeutic areas, 171 druggable (85%)")

H2("Stage 3: AI-Driven Drug Discovery (8-16 min)")
bullet("\u2022", "BioNeMo MolMIM (generation) + DiffDock (docking) + RDKit (scoring)")
bullet("\u2022", "Composite: 30% generation + 40% docking + 30% QED")
bullet("\u2022", "Output: 100 ranked novel drug candidates + PDF report")

doc.add_page_break()

# ── Key Numbers ────────────────────────────────────────────
H1("Key Numbers")
add_table(
    ["Metric", "Value"],
    [
        ["Total Pipeline Time", "< 5 hours"],
        ["Input Data", "~200 GB FASTQ (30\u00d7 WGS)"],
        ["Variants Called", "~11.7 million"],
        ["High-Quality Variants", "~3.5 million (QUAL>30)"],
        ["Knowledge Base", "201 genes, 13 therapeutic areas"],
        ["Druggable Targets", "171 (85%)"],
        ["Drug Candidates", "100 (ranked by composite score)"],
        ["Hardware Cost", "$4,699 (DGX Spark)"],
    ],
)
spacer()

# ── VCP Demo ───────────────────────────────────────────────
H1("VCP/FTD Demo Highlights")
bullet("\u2022", "Target: VCP gene \u2014 FTD, ALS, IBMPFD")
bullet("\u2022", "Variant: rs188935092 \u2014 ClinVar Pathogenic, AlphaMissense 0.87")
bullet("\u2022", "Seed: CB-5083 (Phase I VCP inhibitor)")
bullet("\u2022", "Top candidate: +39% composite improvement over seed")
bullet("\u2022", "Docking: -11.4 kcal/mol (vs. -8.1) | QED: 0.81 (vs. 0.62)")
bullet("\u2022", "All top 10 pass Lipinski's Rule of Five")

# ── Tech Stack ─────────────────────────────────────────────
H1("Technology Stack")
add_table(
    ["Layer", "Technology"],
    [
        ["Hardware", "NVIDIA DGX Spark (GB10, 128 GB unified, $4,699)"],
        ["Genomics", "Parabricks 4.6, DeepVariant (>99% accuracy)"],
        ["Annotation", "ClinVar, AlphaMissense, Ensembl VEP"],
        ["Vector DB", "Milvus 2.4, BGE-small-en-v1.5, IVF_FLAT"],
        ["LLM", "Anthropic Claude (RAG-grounded reasoning)"],
        ["Drug Discovery", "BioNeMo MolMIM, DiffDock, RDKit"],
        ["Orchestration", "Nextflow DSL2 (5 modes)"],
        ["Monitoring", "Grafana, Prometheus, DCGM Exporter"],
        ["License", "Apache 2.0 (fully open)"],
    ],
)
spacer()

# ── Deployment ─────────────────────────────────────────────
H1("Deployment Roadmap")
add_table(
    ["Phase", "Hardware", "Scale", "Cost"],
    [
        ["1 \u2014 Proof Build", "DGX Spark", "1 patient, Docker Compose", "$4,699"],
        ["2 \u2014 Departmental", "DGX B200", "Multiple concurrent, K8s", "$500K-$1M"],
        ["3 \u2014 Enterprise", "DGX SuperPOD", "Thousands, FLARE federated", "$7M-$60M+"],
    ],
)
spacer()

# ── Cross-Modal ────────────────────────────────────────────
H1("Cross-Modal Integration")
bullet("\u2022", "Imaging \u2192 Genomics: Lung-RADS 4B+ triggers tumor profiling")
bullet("\u2022", "Genomics \u2192 Drug Discovery: Pathogenic variants trigger molecule generation")
bullet("\u2022", "NVIDIA FLARE: Federated learning (data stays local)")

# ── Differentiation ────────────────────────────────────────
H1("Competitive Differentiation")
bullet("\u2022", "Only platform: genomics-to-drug-candidates on a single desktop GPU")
bullet("\u2022", "End-to-end: No manual handoffs between stages")
bullet("\u2022", "< 5 hours total (vs. weeks/months traditional)")
bullet("\u2022", "$4,699 proof build (vs. $100K+ CPU infrastructure)")
bullet("\u2022", "Open-source: Apache 2.0, reproducible, auditable")
bullet("\u2022", "Scalable: Same pipelines from DGX Spark to SuperPOD")

spacer(8)
p = P("", after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
R(p, "HCLS AI Factory \u2014 Apache 2.0  |  Author: Adam Jones  |  February 2026",
  italic=True, size=9, color=GRAY_META)

# ══════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════
import pathlib, subprocess, shutil

OUT = pathlib.Path(__file__).resolve().parent
docx_path = OUT / "HCLS_AI_Factory_Executive_Bullets.docx"
pdf_path = OUT / "HCLS_AI_Factory_Executive_Bullets.pdf"

doc.save(str(docx_path))
print(f"DOCX saved: {docx_path} ({docx_path.stat().st_size:,} bytes)")

if shutil.which("libreoffice"):
    subprocess.run([
        "libreoffice", "--headless", "--convert-to", "pdf",
        "--outdir", str(OUT), str(docx_path)
    ], capture_output=True, timeout=120)
    if pdf_path.exists():
        print(f"PDF  saved: {pdf_path} ({pdf_path.stat().st_size:,} bytes)")
elif shutil.which("docx2pdf"):
    subprocess.run(["docx2pdf", str(docx_path), str(pdf_path)],
                    capture_output=True, timeout=120)
    if pdf_path.exists():
        print(f"PDF  saved: {pdf_path} ({pdf_path.stat().st_size:,} bytes)")
else:
    print("No PDF converter found.")

print("Done.")
