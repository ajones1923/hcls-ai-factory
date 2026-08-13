#!/usr/bin/env python3
# Copyright 2026 Adam Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Generate HCLS AI Factory \u2014 Learning Guide: Advanced (DOCX + PDF).

Graduate / Professional level deep technical analysis of the HCLS AI
Factory architecture.  Covers computational genomics, variant annotation,
vector-database RAG, drug-discovery pipelines, Nextflow DSL2 orchestration,
clinical translation, scaling analysis, and advanced extensions.
Formatted in VCP-style theme.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# \u2500\u2500 Colors (VCP palette) \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500
NAVY = RGBColor(0x1B, 0x23, 0x33)
TEAL = RGBColor(0x1A, 0xAF, 0xCC)
GREEN = RGBColor(0x76, 0xB9, 0x00)
GRAY_BODY = RGBColor(0x33, 0x33, 0x33)
GRAY_META = RGBColor(0x66, 0x66, 0x66)
GRAY_CODE = RGBColor(0x44, 0x44, 0x44)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_NAVY = "1B2333"
HEX_TEAL = "1AAFCC"
HEX_GREEN = "76B900"
HEX_LIGHT = "F8FAFB"
HEX_CODE_BG = "F5F5F5"
FONT = "Calibri"
CODE_FONT = "Consolas"

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
style.font.color.rgb = GRAY_BODY
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15


# \u2500\u2500 Helper functions \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500

def R(p, text, bold=False, italic=False, size=None, color=None, font=None):
    r = p.add_run(text)
    r.font.name = font or FONT
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return r


def P(text="", bold=False, italic=False, size=10.5, color=GRAY_BODY,
      before=0, after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text:
        R(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def H1(text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        r.clear()
    R(p, text, bold=True, size=22, color=NAVY)
    return p


def H2(text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.clear()
    R(p, text, bold=True, size=16, color=NAVY)
    return p


def H3(text):
    return P(text, bold=True, size=13, color=TEAL, before=10, after=4)


def body(text, before=0, after=6):
    return P(text, size=10.5, color=GRAY_BODY, before=before, after=after)


def bullet(lead, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    R(p, lead, bold=True, size=10.5, color=GRAY_BODY)
    R(p, f"  {text}", size=10.5, color=GRAY_BODY)
    return p


def code_block(text, language=""):
    if language:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        R(p, language, bold=True, size=8, color=TEAL, font=FONT)
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.2)
        R(p, line if line else " ", size=8, color=GRAY_CODE, font=CODE_FONT)


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="D0D0D0", sz="4"):
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>')
    tcPr.append(borders)


def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, HEX_NAVY)
        set_cell_borders(c, HEX_NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(p, h, bold=True, size=9, color=WHITE)
    for ri, row in enumerate(rows):
        bg = "FFFFFF" if ri % 2 == 0 else HEX_LIGHT
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            set_cell_shading(c, bg)
            set_cell_borders(c, "E0E0E0")
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if ci == 0:
                R(p, str(val), bold=True, size=9, color=TEAL)
            else:
                R(p, str(val), size=9, color=GRAY_BODY)
    return t


def spacer(pts=6):
    P("", after=pts)


# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════

accent = doc.add_table(rows=1, cols=1)
accent.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = accent.rows[0].cells[0]
set_cell_shading(cell, HEX_GREEN)
p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
R(p, " ", size=6, color=GREEN)

for _ in range(4):
    P("")

P("Learning Guide: Advanced", bold=True, size=13, color=TEAL, after=6)
P("HCLS AI Factory",
  bold=True, size=32, color=NAVY, after=2)
P("Graduate / Professional Level",
  bold=True, size=18, color=NAVY, after=6)
P("Deep technical analysis of the HCLS AI Factory architecture, from "
  "BWA-MEM2 seed-and-extend algorithms through diffusion-based molecular "
  "docking, with emphasis on algorithmic design decisions, scaling "
  "bottlenecks, and clinical translation barriers.",
  italic=False, size=12, color=NAVY, after=6)
P("NVIDIA DGX Spark  |  Parabricks  |  BioNeMo  |  Milvus  |  Claude",
  italic=True, size=11, color=GRAY_META, after=12)
P("")
P("02/2026  |  Version 1.0  |  Apache 2.0 License", size=9, color=GRAY_META, after=2)
P("Author: Adam Jones", size=9, color=GRAY_META, after=2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
H1("Table of Contents")

toc_items = [
    "1.  Computational Genomics \u2014 From FASTQ to VCF",
    "2.  Variant Annotation \u2014 Multi-Database Integration",
    "3.  Vector Database Architecture \u2014 Milvus and RAG",
    "4.  Drug Discovery Pipeline \u2014 Deep Dive",
    "5.  Nextflow DSL2 Pipeline Architecture",
    "6.  Clinical Translation and Limitations",
    "7.  Scaling Analysis",
    "8.  Advanced Topics and Extensions",
    "Review Questions",
]
for item in toc_items:
    P(item, size=10.5, color=TEAL, after=3)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 1: COMPUTATIONAL GENOMICS — FROM FASTQ TO VCF
# ══════════════════════════════════════════════════════════════
H1("Chapter 1: Computational Genomics \u2014 From FASTQ to VCF")

H2("1.1 Sequencing Data Characteristics")
body(
    "The HCLS AI Factory processes Illumina short-read data: 2\u00d7250 bp "
    "paired-end reads from 30\u00d7 whole-genome sequencing of HG002 "
    "(NA24385), a GIAB Ashkenazi Jewish reference standard. The FASTQ "
    "files total approximately 200 GB and contain ~800 million read pairs."
)

H3("Why HG002?")
body(
    "The Genome in a Bottle (GIAB) Consortium provides extensively "
    "validated truth sets for HG002, enabling rigorous benchmarking. "
    "The high-confidence regions cover >95% of the GRCh38 reference, "
    "with variant calls validated by multiple orthogonal technologies "
    "(PacBio HiFi, Oxford Nanopore, Hi-C, optical mapping)."
)

H2("1.2 GPU-Accelerated Alignment: BWA-MEM2 on Parabricks")
body(
    "NVIDIA Parabricks 4.6.0-1 (container: nvcr.io/nvidia/clara/"
    "clara-parabricks:4.6.0-1) provides a GPU-accelerated "
    "implementation of BWA-MEM2."
)

H3("Algorithm Overview")
body(
    "BWA-MEM2 uses a seed-and-extend approach:"
)
bullet("1. Seeding:", "Extract fixed-length k-mers from the query read and "
       "look them up in the FM-index of the reference genome")
bullet("2. Chaining:", "Group collinear seeds into chains representing "
       "candidate alignment locations")
bullet("3. Extension:", "Perform Smith-Waterman local alignment around each "
       "chain to produce the final alignment")
bullet("4. Scoring:", "Select the best alignment and assign a MAPQ (mapping "
       "quality) score")

H3("GPU Acceleration Strategy")
body(
    "Parabricks parallelizes the computationally intensive Smith-Waterman "
    "extension step across GPU cores. The FM-index lookup (seeding) remains "
    "CPU-bound but constitutes a small fraction of total compute. The "
    "fq2bam command also integrates coordinate sorting and duplicate "
    "marking, eliminating separate samtools sort and picard MarkDuplicates "
    "steps."
)

H3("Performance on DGX Spark (GB10)")
add_table(
    ["Metric", "Value"],
    [
        ["Wall time", "20-45 minutes"],
        ["GPU utilization", "70-90%"],
        ["Peak memory", "~40 GB (of 128 GB unified)"],
        ["Output", "Sorted BAM + BAI index"],
        ["Mapping rate", ">99.5%"],
        ["Duplicate rate", "~8-12%"],
    ],
)
spacer()

H2("1.3 Deep Learning Variant Calling: DeepVariant")
body(
    "Google DeepVariant reframes variant calling as an image "
    "classification problem. For each candidate variant site, it "
    "constructs a pileup image \u2014 a visual representation of aligned "
    "reads at that position \u2014 and classifies it using a convolutional "
    "neural network (CNN)."
)

H3("Architecture Details")
bullet("Input:", "Pileup image (channels: read bases, base qualities, "
       "mapping qualities, strand, etc.)")
bullet("Network:", "Inception-v3 CNN architecture")
bullet("Output:", "Three-class softmax (homozygous reference, heterozygous "
       "variant, homozygous variant)")
bullet("Training:", "Supervised on GIAB truth sets, with data augmentation "
       "and hard example mining")

H3("Why DeepVariant Outperforms GATK HaplotypeCaller")
bullet("1.", "The CNN learns complex error patterns that statistical models "
       "cannot capture")
bullet("2.", "No explicit error model required \u2014 the network learns "
       "directly from data")
bullet("3.", "Better performance on indels and complex variants")
bullet("4.", "Transferable across sequencing platforms (Illumina, PacBio, ONT)")

H3("Performance")
add_table(
    ["Metric", "Value"],
    [
        ["Wall time", "10-35 minutes (GPU-accelerated via Parabricks)"],
        ["GPU utilization", "80-95%"],
        ["Peak memory", "~60 GB"],
        ["SNP F1", ">99.7% on HG002"],
        ["Indel F1", ">99.4% on HG002"],
        ["Total variants", "~11.7M (unfiltered)"],
        ["QUAL>30 variants", "~3.5M"],
    ],
)
spacer()

H2("1.4 VCF Quality Metrics")
add_table(
    ["Metric", "Expected Range", "Interpretation"],
    [
        ["Ti/Tv ratio", "2.0-2.1",
         "Transition/transversion ratio; deviation suggests systematic error"],
        ["Het/Hom ratio", "1.5-2.0",
         "Heterozygous/homozygous ratio; population-dependent"],
        ["SNP count", "~4.2M",
         "Consistent with Ashkenazi ancestry"],
        ["Indel count", "~1.0M",
         "Normal range for WGS"],
        ["Novel variant rate", "<5%",
         "Variants not in dbSNP; higher rates suggest error"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 2: VARIANT ANNOTATION — MULTI-DATABASE INTEGRATION
# ══════════════════════════════════════════════════════════════
H1("Chapter 2: Variant Annotation \u2014 Multi-Database Integration")

H2("2.1 ClinVar: Clinical Variant Classification")
body(
    "ClinVar (NCBI) is a freely accessible archive of relationships "
    "between human variants and phenotypes. The HCLS AI Factory integrates "
    "the February 2026 release containing 4.1 million variant-condition "
    "records."
)

H3("Classification System (ACMG/AMP)")
bullet("Pathogenic (P) \u2014", "Strong evidence of disease causation")
bullet("Likely Pathogenic (LP) \u2014", "Moderate evidence")
bullet("Variant of Uncertain Significance (VUS) \u2014",
       "Insufficient evidence")
bullet("Likely Benign (LB) \u2014",
       "Moderate evidence against pathogenicity")
bullet("Benign (B) \u2014", "Strong evidence against pathogenicity")

H3("Review Status Tiers")
body(
    "ClinVar classifies assertion confidence using star ratings (0-4 "
    "stars). The pipeline weights variants with \u22652 stars (multiple "
    "submitters with concordant interpretations) more heavily."
)

H3("Annotation Performance")
body(
    "Of ~3.5M QUAL>30 variants, approximately 35,616 (1.0%) match "
    "ClinVar entries. The low match rate reflects that most variants "
    "in a healthy individual are common polymorphisms not represented "
    "in a clinical database focused on rare disease."
)

H2("2.2 AlphaMissense: AI Pathogenicity Prediction")
body(
    "AlphaMissense (Cheng et al., Science 2023) predicts the "
    "pathogenicity of all possible human missense variants using "
    "features derived from AlphaFold protein structure predictions "
    "and evolutionary conservation."
)

H3("Model Architecture")
bullet("Input features:", "amino acid sequence context, evolutionary "
       "conservation (from MSA), and structural features from AlphaFold")
bullet("Output:", "pathogenicity score (0-1, continuous)")
bullet("Total predictions:", "71,697,560 unique missense variants")

H3("Calibrated Thresholds")
bullet("Pathogenic:", ">0.564 (90% precision on ClinVar pathogenic set)")
bullet("Ambiguous:", "0.34-0.564")
bullet("Benign:", "<0.34 (90% precision on ClinVar benign set)")

H3("Critical Limitation")
body(
    "AlphaMissense only predicts missense variant effects. Stop-gain, "
    "frameshift, splice site, and non-coding variants require other "
    "prediction tools. The pipeline uses VEP for functional consequence "
    "annotation to complement AlphaMissense."
)

H2("2.3 Ensembl VEP: Functional Consequence Prediction")
body(
    "The Variant Effect Predictor maps variants to genes, transcripts, "
    "and regulatory regions, annotating each with standardized Sequence "
    "Ontology (SO) terms."
)

H3("Impact Classification")
add_table(
    ["Impact Level", "Example Consequences", "Typical Action"],
    [
        ["HIGH", "stop_gained, frameshift_variant, splice_donor_variant",
         "Likely loss of function"],
        ["MODERATE", "missense_variant, inframe_deletion",
         "Protein function may change"],
        ["LOW", "synonymous_variant, splice_region_variant",
         "Unlikely to affect protein"],
        ["MODIFIER", "intron_variant, upstream_gene_variant",
         "Non-coding effects"],
    ],
)
spacer()

H2("2.4 The Annotation Pipeline Architecture")
body(
    "The three annotation databases are applied sequentially in "
    "annotator.py (23 KB):"
)

code_block(
    "VCF (11.7M variant records)\n"
    "  \u2192 parse_vcf(min_qual=30)     \u2192 3.5M variants\n"
    "  \u2192 annotate_clinvar()          \u2192 Clinical significance\n"
    "  \u2192 annotate_alphamissense()    \u2192 AI pathogenicity scores\n"
    "  \u2192 annotate_vep()              \u2192 Functional consequences\n"
    "  \u2192 generate_text_summary()     \u2192 Natural language descriptions\n"
    "  \u2192 embed_variants()            \u2192 384-dim BGE embeddings\n"
    "  \u2192 index_in_milvus()           \u2192 Searchable vector database",
    language="Annotation Pipeline Flow"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 3: VECTOR DATABASE ARCHITECTURE — MILVUS AND RAG
# ══════════════════════════════════════════════════════════════
H1("Chapter 3: Vector Database Architecture \u2014 Milvus and RAG")

H2("3.1 Milvus Schema Design")
body(
    "The genomic_evidence collection in Milvus 2.4 uses a 17-field "
    "schema designed to support both vector similarity search and "
    "scalar filtering:"
)

add_table(
    ["Field", "Type", "Rationale"],
    [
        ["id", "INT64 (PK, auto)", "Milvus-managed primary key"],
        ["embedding", "FLOAT_VECTOR(384)", "Semantic search vector"],
        ["chrom", "VARCHAR(10)", "Genomic coordinate filtering"],
        ["pos", "INT64", "Positional queries"],
        ["ref/alt", "VARCHAR(1000)", "Allele matching"],
        ["qual", "FLOAT", "Quality score filtering"],
        ["gene", "VARCHAR(100)", "Gene-level queries"],
        ["consequence", "VARCHAR(200)",
         "Functional filtering (e.g., missense only)"],
        ["impact", "VARCHAR(20)", "Impact level filtering"],
        ["genotype", "VARCHAR(10)", "Zygosity queries"],
        ["text_summary", "VARCHAR(2000)",
         "Human-readable context for RAG"],
        ["clinical_significance", "VARCHAR(200)",
         "ClinVar classification"],
        ["rsid", "VARCHAR(20)", "dbSNP lookup"],
        ["disease_associations", "VARCHAR(2000)",
         "Disease context for RAG"],
        ["am_pathogenicity", "FLOAT",
         "AlphaMissense score filtering"],
        ["am_class", "VARCHAR(20)",
         "Pathogenicity class filtering"],
    ],
)
spacer()

H2("3.2 Index Configuration and Performance")

H3("Index Type: IVF_FLAT (Inverted File with Flat Vectors)")
bullet("Why IVF_FLAT?", "At 3.5M vectors with 384 dimensions, IVF_FLAT "
       "provides the best recall-latency tradeoff. HNSW would use more "
       "memory; IVF_PQ would sacrifice recall.")
bullet("nlist=1024:", "Partitions vectors into 1024 clusters. Query "
       "searches ~16 clusters (nprobe=16), examining ~55K vectors per "
       "query.")
bullet("Metric:", "COSINE similarity (normalized dot product)")

H3("Search Performance")
add_table(
    ["Metric", "Value"],
    [
        ["Index build time", "~8 minutes (3.5M \u00d7 384-dim)"],
        ["Index memory", "~2 GB"],
        ["Search latency (nprobe=16)", "8-15 ms"],
        ["Recall@20", ">95%"],
    ],
)
spacer()

H2("3.3 RAG Architecture with Claude")
body(
    "The RAG pipeline in rag_engine.py (23 KB) implements a multi-stage "
    "retrieval strategy:"
)

H3("1. Query Expansion")
body(
    "User queries are enriched using 10 therapeutic area keyword maps. "
    "For example, a query about \"neurodegeneration\" is expanded with "
    "terms like \"frontotemporal dementia,\" \"ALS,\" \"motor neuron,\" "
    "\"tau protein.\""
)

H3("2. Hybrid Retrieval")
body(
    "The expanded query is embedded and used for vector search "
    "(top_k=20). Results are optionally filtered by scalar fields "
    "(e.g., impact=HIGH, am_class=pathogenic)."
)

H3("3. Context Assembly")
body(
    "Retrieved variants are formatted into structured context:"
)
code_block(
    "## Variant Evidence\n"
    "- chr9:35065263 G>A | Gene: VCP | Consequence: missense_variant\n"
    "  ClinVar: Pathogenic | AlphaMissense: 0.87 (pathogenic)\n"
    "  Disease: Frontotemporal Dementia, ALS, IBMPFD",
    language="Context Template"
)

H3("4. Claude Inference")
body(
    "The assembled context + knowledge base + user query are sent to "
    "claude-sonnet-4-20250514 (temperature=0.3, max_tokens=4096)."
)

H3("Why temperature=0.3?")
body(
    "Lower temperature produces more deterministic, factual responses. "
    "For clinical genomics, hallucination is dangerous \u2014 the model "
    "should report only what the evidence supports."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 4: DRUG DISCOVERY PIPELINE — DEEP DIVE
# ══════════════════════════════════════════════════════════════
H1("Chapter 4: Drug Discovery Pipeline \u2014 Deep Dive")

H2("4.1 The 10-Stage Architecture")
body(
    "The drug discovery pipeline in pipeline.py (18 KB) implements a "
    "sequential 10-stage workflow:"
)

add_table(
    ["Stage", "Module", "Key Algorithm"],
    [
        ["1. Initialize", "pipeline.py", "Pydantic model validation"],
        ["2. Normalize Target", "pipeline.py",
         "Gene \u2192 UniProt \u2192 PDB mapping"],
        ["3. Structure Discovery", "cryoem_evidence.py",
         "RCSB PDB REST API query"],
        ["4. Structure Preparation", "cryoem_evidence.py",
         "Multi-factor scoring"],
        ["5. Molecule Generation", "nim_clients.py",
         "MolMIM masked LM inference"],
        ["6. Chemistry QC", "molecule_generator.py",
         "RDKit valence/kekulization"],
        ["7. Conformer Generation", "molecule_generator.py",
         "RDKit ETKDG algorithm"],
        ["8. Molecular Docking", "nim_clients.py",
         "DiffDock diffusion inference"],
        ["9. Composite Ranking", "pipeline.py",
         "Weighted multi-objective"],
        ["10. Reporting", "pipeline.py",
         "ReportLab PDF generation"],
    ],
)
spacer()

H2("4.2 Cryo-EM Structure Scoring")
body(
    "The cryoem_evidence.py (6 KB) module implements a multi-factor "
    "structure scoring algorithm:"
)

code_block(
    "score += max(0, 5.0 - resolution)                 # Resolution: 0-5 scale\n"
    "if has_inhibitor_bound: score += 3.0                # Binding site defined\n"
    "score += num_druggable_pockets * 0.5                # Pocket count bonus\n"
    "if 'Cryo-EM' in method: score += 0.5                 # Method bonus",
    language="Python"
)

H3("Design Rationale")
bullet("Resolution:", "the primary factor (0-5 scale). The 5 \u00c5 "
       "cutoff excludes low-resolution structures unsuitable for docking.")
bullet("Inhibitor bonus (+3):", "Inhibitor-bound structures provide a "
       "pre-defined binding site and reference ligand geometry.")
bullet("Pocket count (+0.5 each):", "More druggable pockets increase "
       "therapeutic options.")
bullet("Cryo-EM bonus (+0.5):", "Reflects the growing prevalence and "
       "quality of Cryo-EM structures for drug targets.")

H2("4.3 MolMIM: Molecular Masked Inverse Modeling")
body(
    "MolMIM applies masked language modeling (the technique behind BERT "
    "in NLP) to molecular SMILES strings. Given a seed molecule, it:"
)
bullet("1.", "Tokenizes the SMILES into a vocabulary of molecular "
       "substructures")
bullet("2.", "Randomly masks 15-30% of tokens")
bullet("3.", "Predicts the masked tokens using a transformer architecture")
bullet("4.", "The predicted tokens create novel molecular structures")

H3("Critical Considerations")
bullet("SMILES output:", "MolMIM generates SMILES strings, not 3D "
       "structures. Chemical validity must be verified by RDKit.")
bullet("Stochastic generation:", "Different random seeds produce "
       "different molecules.")
bullet("Temperature control:", "Higher temperature = more diverse but "
       "potentially less valid molecules.")

H2("4.4 DiffDock: Diffusion-Based Molecular Docking")
body(
    "DiffDock (Corso et al., ICLR 2023) models molecular docking as a "
    "generative diffusion process over the product space of rotations, "
    "translations, and torsion angles."
)

H3("Key Innovation")
body(
    "Unlike grid-based docking methods (AutoDock Vina, Glide), DiffDock "
    "does not require a pre-defined search box around a binding site. "
    "It learns to predict binding poses directly from protein-ligand "
    "pairs, making it suitable for blind docking."
)

H3("Score Interpretation")
bullet("Confidence score (0-1):", "indicates the model's certainty about "
       "the predicted pose")
bullet("Binding affinity (kcal/mol):", "estimates the free energy of "
       "binding; more negative = stronger binding")

H3("Limitations")
bullet("Training bias:", "DiffDock was trained primarily on crystal "
       "structures; performance may degrade on Cryo-EM structures with "
       "lower resolution")
bullet("No kinetics:", "The model predicts pose and affinity but not "
       "binding kinetics (on/off rates)")
bullet("Rigid protein:", "Protein flexibility is not modeled \u2014 "
       "the protein is treated as rigid")

H2("4.5 Composite Scoring and Normalization")
body(
    "The composite scoring formula balances three objectives:"
)

code_block(
    "dock_normalized = max(0.0, min(1.0, (10.0 + dock_score) / 20.0))\n"
    "composite = 0.30 * gen_score + 0.40 * dock_normalized + 0.30 * qed_score",
    language="Python"
)

H3("Normalization Rationale")
bullet("Docking scores:", "range from ~-15 to ~0 kcal/mol. The formula "
       "(10 + dock) / 20 maps this to approximately 0-1, with -10 kcal/mol "
       "mapping to 0.0 and +10 mapping to 1.0.")
bullet("Generation scores:", "already 0-1 (MolMIM confidence).")
bullet("QED scores:", "inherently 0-1.")

H3("Weight Rationale")
bullet("Docking (40%):", "receives the highest weight because binding "
       "affinity is the most direct predictor of therapeutic activity")
bullet("Generation (30%):", "balances novelty of the molecular design")
bullet("QED (30%):", "balances practical drug-likeness")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 5: NEXTFLOW DSL2 PIPELINE ARCHITECTURE
# ══════════════════════════════════════════════════════════════
H1("Chapter 5: Nextflow DSL2 Pipeline Architecture")

H2("5.1 Module Design")
body(
    "The pipeline uses Nextflow DSL2's module system for composable "
    "workflow design:"
)

code_block(
    "hcls-orchestrator/\n"
    "\u251c\u2500\u2500 main.nf                 # Entry point, mode routing\n"
    "\u251c\u2500\u2500 nextflow.config         # Profiles, parameters\n"
    "\u251c\u2500\u2500 run_pipeline.py         # Python CLI launcher\n"
    "\u2514\u2500\u2500 modules/\n"
    "    \u251c\u2500\u2500 genomics.nf         # Stage 1 processes\n"
    "    \u251c\u2500\u2500 rag_chat.nf         # Stage 2 processes\n"
    "    \u251c\u2500\u2500 drug_discovery.nf   # Stage 3 processes\n"
    "    \u2514\u2500\u2500 reporting.nf        # Report generation",
    language="Directory Structure"
)

H2("5.2 Execution Modes and Data Flow")

add_table(
    ["Mode", "Data Flow", "Use Case"],
    [
        ["full", "FASTQ \u2192 VCF \u2192 Target \u2192 Candidates",
         "Complete pipeline"],
        ["target", "VCF \u2192 Target \u2192 Candidates",
         "Pre-existing VCF"],
        ["drug", "Target \u2192 Candidates",
         "Known gene target"],
        ["demo", "Pre-configured FASTQ \u2192 Candidates",
         "VCP/FTD demonstration"],
        ["genomics_only", "FASTQ \u2192 VCF",
         "Variant calling only"],
    ],
)
spacer()

H2("5.3 Profile Configuration")
body(
    "The nextflow.config defines six execution profiles optimized for "
    "different environments:"
)
bullet("dgx_spark:", "GPU resource requests, memory limits tuned for "
       "128 GB unified memory")
bullet("docker:", "Docker container execution with GPU passthrough")
bullet("singularity:", "Singularity containers for HPC environments "
       "without Docker")
bullet("slurm:", "SLURM scheduler integration for cluster execution")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 6: CLINICAL TRANSLATION AND LIMITATIONS
# ══════════════════════════════════════════════════════════════
H1("Chapter 6: Clinical Translation and Limitations")

H2("6.1 From Computational Hits to Drug Leads")
body(
    "The HCLS AI Factory generates computational drug candidates \u2014 "
    "not approved medications. The path from computational hit to "
    "clinical drug requires:"
)
bullet("1. In vitro validation:", "Test top candidates in biochemical "
       "assays (e.g., VCP ATPase activity inhibition)")
bullet("2. Cell-based assays:", "Confirm activity in relevant cell lines")
bullet("3. ADMET profiling:", "Absorption, Distribution, Metabolism, "
       "Excretion, and Toxicity studies")
bullet("4. Lead optimization:", "Iterative cycles of design, synthesis, "
       "and testing")
bullet("5. Preclinical studies:", "Animal models for efficacy and safety")
bullet("6. Clinical trials:", "Phase I (safety), Phase II (efficacy), "
       "Phase III (large-scale)")

H3("Estimated Timeline")
body(
    "10-15 years from computational hit to approved drug. The HCLS AI "
    "Factory accelerates the earliest stage \u2014 computational lead "
    "generation \u2014 from months to minutes."
)

H2("6.2 Limitations and Caveats")

H3("Genomics")
bullet("\u2022", "DeepVariant accuracy varies by variant type (SNPs > "
       "indels > structural variants)")
bullet("\u2022", "Short-read WGS has limited sensitivity for structural "
       "variants and repeat expansions")
bullet("\u2022", "Population-specific biases in GRCh38 may affect variant "
       "calling in non-European ancestries")

H3("RAG/Annotation")
bullet("\u2022", "ClinVar has known biases toward well-studied genes and "
       "European ancestry variants")
bullet("\u2022", "AlphaMissense is limited to missense variants; non-coding "
       "variants are not scored")
bullet("\u2022", "The 201-gene knowledge base covers common drug targets but "
       "not the full druggable genome")

H3("Drug Discovery")
bullet("\u2022", "MolMIM-generated molecules have not been synthesized or "
       "tested")
bullet("\u2022", "DiffDock docking scores are predictions, not experimental "
       "measurements")
bullet("\u2022", "Protein flexibility is not modeled; induced-fit effects "
       "are ignored")
bullet("\u2022", "The composite scoring weights (30/40/30) are heuristic, "
       "not optimized on clinical outcomes")

H2("6.3 Ethical Considerations")
bullet("Informed consent:", "Patient genomic data requires explicit "
       "consent for research use")
bullet("Data sovereignty:", "NVIDIA FLARE federated learning keeps data "
       "local; essential for HIPAA/GDPR compliance")
bullet("Return of results:", "Incidental findings (e.g., BRCA1 "
       "pathogenic variants) may require clinical reporting")
bullet("Equity:", "Pipeline performance should be validated across "
       "diverse ancestries to avoid exacerbating health disparities")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 7: SCALING ANALYSIS
# ══════════════════════════════════════════════════════════════
H1("Chapter 7: Scaling Analysis")

H2("7.1 DGX Spark Bottleneck Analysis")

add_table(
    ["Component", "Bottleneck", "Phase 1 Impact"],
    [
        ["Parabricks (fq2bam)", "GPU compute",
         "20-45 min, acceptable"],
        ["DeepVariant", "GPU memory (60 GB peak)",
         "Leaves 68 GB for other tasks"],
        ["Milvus indexing", "CPU + I/O",
         "24 min for 3.5M vectors"],
        ["MolMIM inference", "GPU compute",
         "2 min for 100 molecules"],
        ["DiffDock inference", "GPU compute + memory",
         "8 min for 98 candidates"],
        ["Sequential total", "GPU time-sharing",
         "~4 hours end-to-end"],
    ],
)
spacer()

H2("7.2 Phase 2: DGX B200 Scaling")
body(
    "With 8\u00d7 B200 GPUs and 1-2 TB HBM3e:"
)
bullet("Parallel Parabricks:", "4-8 simultaneous samples")
bullet("Dedicated Milvus GPU:", "GPU-accelerated vector search "
       "(sub-millisecond)")
bullet("NIM replicas:", "2-4 MolMIM + 2-4 DiffDock instances")
bullet("Estimated throughput:", "10-20 patients per day")

H2("7.3 Phase 3: DGX SuperPOD")
bullet("Hundreds of B200 GPUs", "with NVLink and InfiniBand")
bullet("Distributed Milvus cluster:", "Billions of variants across "
       "institutions")
bullet("NVIDIA FLARE:", "Federated model training without data sharing")
bullet("Estimated throughput:", "Thousands of patients per day")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 8: ADVANCED TOPICS AND EXTENSIONS
# ══════════════════════════════════════════════════════════════
H1("Chapter 8: Advanced Topics and Extensions")

H2("8.1 Alternative Embedding Strategies")
body(
    "BGE-small-en-v1.5 (384-dim) was chosen for its balance of quality "
    "and efficiency. Alternatives:"
)

add_table(
    ["Model", "Dimensions", "Size", "Trade-off"],
    [
        ["BGE-small-en-v1.5", "384", "33M params",
         "Current choice: fast, efficient"],
        ["BGE-base-en-v1.5", "768", "109M params",
         "Better recall, 2\u00d7 memory"],
        ["BGE-large-en-v1.5", "1024", "335M params",
         "Best recall, 3\u00d7 memory"],
        ["BiomedBERT", "768", "109M params",
         "Domain-specific, biomedical text"],
        ["PubMedBERT", "768", "109M params",
         "PubMed-trained, clinical text"],
    ],
)
spacer()

H2("8.2 Multi-Objective Optimization")
body(
    "The current composite scoring uses fixed weights (30/40/30). "
    "Advanced approaches:"
)
bullet("Pareto optimization:", "Identify the Pareto frontier of "
       "generation, docking, and QED")
bullet("Bayesian optimization:", "Learn optimal weights from "
       "experimental feedback")
bullet("Active learning:", "Prioritize candidates that reduce "
       "uncertainty in the scoring model")

H2("8.3 Long-Read Sequencing Integration")
body(
    "Oxford Nanopore and PacBio long-read technologies can detect "
    "structural variants (SVs) and repeat expansions that short-read "
    "WGS misses. Future extensions could:"
)
bullet("\u2022", "Add ONT/PacBio alignment with minimap2")
bullet("\u2022", "Detect SVs with Sniffles2 or PEPPER-Margin-DeepVariant")
bullet("\u2022", "Phase haplotypes for compound heterozygosity detection")

H2("8.4 Pharmacogenomics Integration")
body(
    "The knowledge base includes 11 pharmacogenomics genes (CYP2D6, "
    "CYP2C19, CYP3A4, DPYD, TPMT, etc.). Future extensions could:"
)
bullet("\u2022", "Star allele calling with PharmCAT")
bullet("\u2022", "Drug-drug interaction prediction")
bullet("\u2022", "Dosing recommendations based on metabolizer status")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# REVIEW QUESTIONS
# ══════════════════════════════════════════════════════════════
H1("Review Questions")

body(
    "Graduate-level questions for self-assessment and discussion. "
    "These require synthesis across multiple chapters."
)
spacer()

questions = [
    ("1.", "Explain why DeepVariant reframes variant calling as an image "
           "classification problem. What advantages does this provide over "
           "statistical models like GATK HaplotypeCaller?"),
    ("2.", "Describe the IVF_FLAT index configuration (nlist=1024, "
           "nprobe=16) and calculate the approximate number of vectors "
           "examined per query from a collection of 3.5M vectors."),
    ("3.", "Why does the RAG pipeline use temperature=0.3 for Claude? "
           "What are the trade-offs of lower vs. higher temperature in "
           "clinical genomics applications?"),
    ("4.", "Explain the docking score normalization formula max(0, min(1, "
           "(10 + dock_score) / 20)). What docking score maps to 0.5? "
           "Why is this mapping appropriate?"),
    ("5.", "Compare the AlphaMissense thresholds (pathogenic >0.564, "
           "benign <0.34) with ClinVar classifications. What does the "
           "\"ambiguous\" zone represent, and why is it clinically "
           "significant?"),
    ("6.", "Describe three limitations of DiffDock that could affect "
           "the reliability of the drug candidate rankings."),
    ("7.", "Explain why inhibitor-bound PDB structures (like 5FTK) "
           "receive a +3 bonus in the structure scoring algorithm. "
           "What information does an inhibitor-bound structure provide "
           "that an apo structure does not?"),
    ("8.", "Design a validation experiment to test the top 5 drug "
           "candidates from the VCP/FTD demo pipeline. What assays "
           "would you use, and what would constitute a positive result?"),
    ("9.", "Calculate the approximate memory budget for the Milvus "
           "vector index: 3.5M vectors \u00d7 384 dimensions \u00d7 4 bytes "
           "per float. How does this compare to the available memory "
           "on DGX Spark?"),
    ("10.", "The composite scoring weights (30% generation, 40% docking, "
            "30% QED) are heuristic. Propose an approach to optimize "
            "these weights using experimental feedback from in vitro "
            "screening."),
]
for num, question in questions:
    bullet(num, question)

spacer(12)
p = P("", after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
R(p, "HCLS AI Factory Learning Guide: Advanced \u2014 Apache 2.0 | "
  "Author: Adam Jones | February 2026", italic=True, size=10, color=GRAY_META)

# ══════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════
import pathlib, subprocess, shutil

OUT = pathlib.Path(__file__).resolve().parent
docx_path = OUT / "HCLS_AI_Factory_Learning_Guide_Advanced.docx"
pdf_path = OUT / "HCLS_AI_Factory_Learning_Guide_Advanced.pdf"

doc.save(str(docx_path))
print(f"DOCX saved: {docx_path} ({docx_path.stat().st_size:,} bytes)")

# Attempt PDF conversion
if shutil.which("libreoffice"):
    subprocess.run([
        "libreoffice", "--headless", "--convert-to", "pdf",
        "--outdir", str(OUT), str(docx_path)
    ], capture_output=True, timeout=120)
    if pdf_path.exists():
        print(f"PDF  saved: {pdf_path} ({pdf_path.stat().st_size:,} bytes)")
    else:
        print("PDF conversion completed but file not found.")
elif shutil.which("docx2pdf"):
    subprocess.run(["docx2pdf", str(docx_path), str(pdf_path)],
                    capture_output=True, timeout=120)
    if pdf_path.exists():
        print(f"PDF  saved: {pdf_path} ({pdf_path.stat().st_size:,} bytes)")
else:
    print("No PDF converter found (install libreoffice or docx2pdf).")

print("Done.")
