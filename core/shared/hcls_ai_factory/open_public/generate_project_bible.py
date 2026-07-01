#!/usr/bin/env python3
# Copyright 2026 Adam Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Generate HCLS AI Factory — Project Bible (DOCX + PDF).

Complete implementation reference for building the HCLS AI Factory on
NVIDIA DGX Spark. Covers all three stages: Secondary Genomics (Parabricks),
RAG/Chat (Milvus + Claude), and Drug Discovery (BioNeMo + RDKit).
Formatted in VCP-style theme.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Colors (VCP palette) ─────────────────────────────────
NAVY = RGBColor(0x1B, 0x23, 0x33)
TEAL = RGBColor(0x1A, 0xAF, 0xCC)
GREEN = RGBColor(0x76, 0xB9, 0x00)
GRAY_BODY = RGBColor(0x33, 0x33, 0x33)
GRAY_META = RGBColor(0x66, 0x66, 0x66)
GRAY_CODE = RGBColor(0x44, 0x44, 0x44)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_NAVY = "1B2333"
HEX_TEAL = "1AAFCC"
HEX_GREEN = "76B900"
HEX_LIGHT = "F8FAFB"
HEX_CODE_BG = "F5F5F5"
FONT = "Calibri"
CODE_FONT = "Consolas"

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
style.font.color.rgb = GRAY_BODY
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15


# ── Helper functions ──────────────────────────────────────

def R(p, text, bold=False, italic=False, size=None, color=None, font=None):
    r = p.add_run(text)
    r.font.name = font or FONT
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return r


def P(text="", bold=False, italic=False, size=10.5, color=GRAY_BODY,
      before=0, after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text:
        R(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def H1(text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        r.clear()
    R(p, text, bold=True, size=22, color=NAVY)
    return p


def H2(text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.clear()
    R(p, text, bold=True, size=16, color=NAVY)
    return p


def H3(text):
    return P(text, bold=True, size=13, color=TEAL, before=10, after=4)


def body(text, before=0, after=6):
    return P(text, size=10.5, color=GRAY_BODY, before=before, after=after)


def bullet(lead, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    R(p, lead, bold=True, size=10.5, color=GRAY_BODY)
    R(p, f"  {text}", size=10.5, color=GRAY_BODY)
    return p


def code_block(text, language=""):
    if language:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        R(p, language, bold=True, size=8, color=TEAL, font=FONT)
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.2)
        R(p, line if line else " ", size=8, color=GRAY_CODE, font=CODE_FONT)


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="D0D0D0", sz="4"):
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>')
    tcPr.append(borders)


def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, HEX_NAVY)
        set_cell_borders(c, HEX_NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(p, h, bold=True, size=9, color=WHITE)
    for ri, row in enumerate(rows):
        bg = "FFFFFF" if ri % 2 == 0 else HEX_LIGHT
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            set_cell_shading(c, bg)
            set_cell_borders(c, "E0E0E0")
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if ci == 0:
                R(p, str(val), bold=True, size=9, color=TEAL)
            else:
                R(p, str(val), size=9, color=GRAY_BODY)
    return t


def spacer(pts=6):
    P("", after=pts)


# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════

accent = doc.add_table(rows=1, cols=1)
accent.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = accent.rows[0].cells[0]
set_cell_shading(cell, HEX_GREEN)
p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
R(p, " ", size=6, color=GREEN)

for _ in range(4):
    P("")

P("Project Bible", bold=True, size=13, color=TEAL, after=6)
P("HCLS AI Factory",
  bold=True, size=32, color=NAVY, after=2)
P("Implementation Reference",
  bold=True, size=32, color=NAVY, after=6)
P("Complete architecture, pipeline stages, schemas, scoring formulas, "
  "and implementation sequences for building the HCLS AI Factory "
  "on NVIDIA DGX Spark \u2014 from patient DNA to novel drug candidates.",
  italic=False, size=12, color=NAVY, after=6)
P("NVIDIA DGX Spark  |  Parabricks  |  BioNeMo  |  Milvus  |  Claude",
  italic=True, size=11, color=GRAY_META, after=12)
P("")
P("02/2026  |  Version 1.0  |  Apache 2.0 License", size=9, color=GRAY_META, after=2)
P("Author: Adam Jones", size=9, color=GRAY_META, after=2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
H1("Table of Contents")

toc_items = [
    "1.  Project Overview & Goals",
    "2.  DGX Spark Hardware Reference",
    "3.  Repository Layout",
    "4.  Docker Compose Services",
    "5.  Stage 1: Genomics Pipeline",
    "6.  Stage 2: RAG/Chat Pipeline",
    "7.  Milvus Vector Database Schema",
    "8.  Variant Annotation Pipeline",
    "9.  Knowledge Base \u2014 201 Genes, 13 Therapeutic Areas",
    "10. Anthropic Claude LLM Integration",
    "11. Stage 3: Drug Discovery Pipeline",
    "12. BioNeMo NIM Services",
    "13. Drug-Likeness Scoring",
    "14. Cryo-EM Structure Evidence",
    "15. VCP/FTD Demo Walkthrough",
    "16. Pydantic Data Models",
    "17. Nextflow DSL2 Orchestration",
    "18. Landing Page & Service Health",
    "19. Monitoring Stack",
    "20. Cross-Modal Integration",
    "21. Configuration Reference",
    "22. Deployment Roadmap",
    "23. Testing Strategy",
    "24. Implementation Sequence",
]
for item in toc_items:
    P(item, size=10.5, color=TEAL, after=3)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW & GOALS
# ══════════════════════════════════════════════════════════════
H1("1. Project Overview & Goals")

H2("What This Platform Does")
body(
    "The HCLS AI Factory is an end-to-end precision medicine platform that "
    "takes a patient\u2019s raw DNA sequencing data (FASTQ) and produces ranked "
    "novel drug candidates \u2014 all on a single NVIDIA DGX Spark desktop "
    "workstation. Three GPU-accelerated stages execute sequentially: variant "
    "calling, RAG-grounded target identification, and generative drug discovery."
)

H2("Three-Stage Pipeline")
add_table(
    ["Stage", "Function", "Duration", "Key Output"],
    [
        ["1 \u2014 Genomics", "BWA-MEM2 alignment + DeepVariant calling", "120-240 min", "VCF (~11.7M variants)"],
        ["2 \u2014 RAG/Chat", "Annotation \u2192 Embedding \u2192 LLM reasoning", "Interactive", "Target gene + evidence"],
        ["3 \u2014 Drug Discovery", "MolMIM \u2192 DiffDock \u2192 RDKit scoring", "8-16 min", "100 ranked drug candidates"],
    ],
)
spacer()

H2("End-to-End Flow")
code_block(
    "Patient DNA \u2192 Illumina Sequencer \u2192 FASTQ (~200 GB)\n"
    "  \u2192 Parabricks fq2bam \u2192 BAM\n"
    "  \u2192 DeepVariant \u2192 VCF (11.7M variants)\n"
    "  \u2192 ClinVar + AlphaMissense + VEP annotation\n"
    "  \u2192 Milvus vector indexing (3.5M embeddings)\n"
    "  \u2192 Claude RAG reasoning \u2192 Target hypothesis\n"
    "  \u2192 RCSB PDB structure retrieval\n"
    "  \u2192 MolMIM molecule generation\n"
    "  \u2192 DiffDock molecular docking\n"
    "  \u2192 RDKit drug-likeness scoring\n"
    "  \u2192 100 ranked novel drug candidates + PDF report"
)
spacer()

H2("Design Principles")
bullet("GPU-first:", "Every compute-intensive step runs on the GB10 GPU")
bullet("Clinically grounded:", "ClinVar, AlphaMissense, and VEP provide evidence-based annotation")
bullet("Reproducible:", "Nextflow DSL2 orchestration with containerized processes")
bullet("Open:", "Apache 2.0 license, open-source tools, public reference databases")
bullet("Desktop-scale:", "Runs entirely on a $4,699 DGX Spark")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. DGX SPARK HARDWARE REFERENCE
# ══════════════════════════════════════════════════════════════
H1("2. DGX Spark Hardware Reference")

H2("Specifications")
add_table(
    ["Parameter", "Value"],
    [
        ["CPU", "NVIDIA Grace (ARM64 / aarch64), ARM64 cores"],
        ["GPU", "NVIDIA GB10, 1 GPU"],
        ["Memory", "128 GB unified LPDDR5x (CPU + GPU shared)"],
        ["System RAM", "512 GB"],
        ["Storage", "NVMe, high-throughput I/O"],
        ["Storage Access", "GPUDirect Storage (zero-copy GPU access)"],
        ["Price", "$4,699"],
        ["OS", "Ubuntu-based (NVIDIA DGX OS)"],
    ],
)
spacer()

H2("Critical: ARM64 Architecture")
body(
    "ALL containers must be ARM64-compatible. The Grace CPU is aarch64, not "
    "x86_64. This affects base Docker images, Python wheels, NVIDIA container "
    "images (NGC ARM64 variants), and any compiled C/C++ extensions."
)

H2("Unified Memory Model")
body(
    "The 128 GB LPDDR5x is shared between CPU and GPU \u2014 there is no "
    "separate GPU VRAM. No explicit CPU\u2192GPU data transfers needed for many "
    "operations. Memory pressure from CPU workloads reduces GPU-available "
    "memory. Parabricks fq2bam peaks at ~40 GB, DeepVariant at ~60 GB."
)

H2("Storage Requirements")
add_table(
    ["Dataset", "Size", "Notes"],
    [
        ["GRCh38 reference", "3.1 GB", "Pre-indexed for BWA-MEM2"],
        ["FASTQ input (30\u00d7 WGS)", "~200 GB", "HG002 paired-end"],
        ["BAM intermediate", "~100 GB", "Temporary, deleted after VCF"],
        ["ClinVar database", "~1.2 GB", "4.1M clinical variants"],
        ["AlphaMissense database", "~4 GB", "71M predictions"],
        ["Milvus index", "~2 GB", "3.5M \u00d7 384-dim vectors"],
        ["BioNeMo model cache", "~10 GB", "MolMIM + DiffDock weights"],
        ["Total minimum", "~320 GB", "Plus OS and Docker layers"],
    ],
)
spacer()

H2("Deployment Progression")
add_table(
    ["Phase", "Hardware", "Price", "Scope"],
    [
        ["1 \u2014 Proof Build", "DGX Spark", "$4,699", "Single patient, Docker Compose"],
        ["2 \u2014 Departmental", "1\u20132\u00d7 DGX B200", "$500K\u2013$1M", "Multiple concurrent patients, Kubernetes"],
        ["3 \u2014 Enterprise", "DGX SuperPOD", "$7M\u2013$60M+", "Thousands concurrent, FLARE federated"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. REPOSITORY LAYOUT
# ══════════════════════════════════════════════════════════════
H1("3. Repository Layout")

code_block(
    "hcls-ai-factory-public/\n"
    "\u251c\u2500\u2500 README.md                           # Project overview\n"
    "\u251c\u2500\u2500 LICENSE                             # Apache 2.0\n"
    "\u251c\u2500\u2500 docker-compose.yml                  # All services\n"
    "\u251c\u2500\u2500 start-services.sh                   # Startup orchestration\n"
    "\u251c\u2500\u2500 .env.example                        # Environment template\n"
    "\u2502\n"
    "\u251c\u2500\u2500 hcls-orchestrator/                   # Nextflow pipeline\n"
    "\u2502   \u251c\u2500\u2500 main.nf                         # DSL2 entry point\n"
    "\u2502   \u251c\u2500\u2500 nextflow.config                 # Profiles and parameters\n"
    "\u2502   \u251c\u2500\u2500 run_pipeline.py                 # Python CLI launcher\n"
    "\u2502   \u2514\u2500\u2500 modules/                        # genomics/rag_chat/drug_discovery/reporting\n"
    "\u2502\n"
    "\u251c\u2500\u2500 core/engines/genomic-foundation/                  # Stage 1: Parabricks\n"
    "\u2502   \u251c\u2500\u2500 src/run_parabricks.py           # fq2bam + DeepVariant\n"
    "\u2502   \u2514\u2500\u2500 src/web_portal.py               # Flask portal (:5000)\n"
    "\u2502\n"
    "\u251c\u2500\u2500 core/engines/precision-intelligence/                  # Stage 2: RAG + Claude\n"
    "\u2502   \u251c\u2500\u2500 src/rag_engine.py               # Core RAG (23 KB)\n"
    "\u2502   \u251c\u2500\u2500 src/milvus_client.py            # Vector DB client (13 KB)\n"
    "\u2502   \u251c\u2500\u2500 src/annotator.py                # ClinVar+AM+VEP (23 KB)\n"
    "\u2502   \u251c\u2500\u2500 src/knowledge.py                # 201 genes (88 KB)\n"
    "\u2502   \u2514\u2500\u2500 src/streamlit_chat.py           # Chat UI (:8501)\n"
    "\u2502\n"
    "\u251c\u2500\u2500 core/engines/therapeutic-discovery/            # Stage 3: BioNeMo + RDKit\n"
    "\u2502   \u251c\u2500\u2500 src/pipeline.py                 # 10-stage orchestration (18 KB)\n"
    "\u2502   \u251c\u2500\u2500 src/nim_clients.py              # MolMIM+DiffDock clients (15 KB)\n"
    "\u2502   \u251c\u2500\u2500 src/molecule_generator.py       # SMILES generation (11 KB)\n"
    "\u2502   \u251c\u2500\u2500 src/cryoem_evidence.py          # Cryo-EM scoring (6 KB)\n"
    "\u2502   \u2514\u2500\u2500 src/models.py                   # Pydantic models (8 KB)\n"
    "\u2502\n"
    "\u251c\u2500\u2500 landing-page/                       # Entry point (:8080)\n"
    "\u251c\u2500\u2500 monitoring/                         # Prometheus + Grafana\n"
    "\u2514\u2500\u2500 docs/                               # Documentation (122 KB+)"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. DOCKER COMPOSE SERVICES
# ══════════════════════════════════════════════════════════════
H1("4. Docker Compose Services")

H2("Port Allocation")
add_table(
    ["Service", "Port", "Protocol", "Stage"],
    [
        ["Landing Page", "8080", "HTTP (Flask)", "Orchestration"],
        ["Genomics Portal", "5000", "HTTP (Flask)", "Stage 1"],
        ["RAG REST API", "5001", "HTTP REST", "Stage 2"],
        ["Milvus Vector DB", "19530", "gRPC", "Stage 2"],
        ["Attu (Milvus UI)", "8000", "HTTP", "Stage 2"],
        ["Streamlit Chat", "8501", "HTTP", "Stage 2"],
        ["MolMIM NIM", "8001", "HTTP REST", "Stage 3"],
        ["DiffDock NIM", "8002", "HTTP REST", "Stage 3"],
        ["Discovery UI", "8505", "HTTP (Streamlit)", "Stage 3"],
        ["Discovery Portal", "8510", "HTTP", "Stage 3"],
        ["Grafana", "3000", "HTTP", "Monitoring"],
        ["Prometheus", "9099", "HTTP", "Monitoring"],
        ["Node Exporter", "9100", "HTTP", "Monitoring"],
        ["DCGM Exporter", "9400", "HTTP", "Monitoring"],
    ],
)
spacer()

H2("Key Container Images")
add_table(
    ["Service", "Image", "Notes"],
    [
        ["Parabricks", "nvcr.io/nvidia/clara/clara-parabricks:4.6.0-1", "GPU-accelerated genomics"],
        ["Milvus", "milvusdb/milvus:v2.4-latest", "Vector database"],
        ["MolMIM", "nvcr.io/nvidia/clara/bionemo-molmim:1.0", "Molecule generation NIM"],
        ["DiffDock", "nvcr.io/nvidia/clara/diffdock:1.0", "Molecular docking NIM"],
        ["Grafana", "grafana/grafana:10.2.2", "Dashboards"],
        ["Prometheus", "prom/prometheus:v2.48.0", "Metrics TSDB"],
    ],
)
spacer()

H2("Service Startup Order")
body("The start-services.sh script orchestrates startup in dependency order:")
bullet("1.", "Infrastructure (Milvus, monitoring)")
bullet("2.", "Stage 1 services (Parabricks, genomics portal)")
bullet("3.", "Stage 2 services (RAG engine, Streamlit chat)")
bullet("4.", "Stage 3 services (BioNeMo NIMs, discovery UI)")
bullet("5.", "Landing page (health monitor for all 10 services)")

H2("Health Monitoring")
body("The landing page at port 8080 monitors 10 services:")
add_table(
    ["Service", "Health Endpoint", "Interval"],
    [
        ["Parabricks", "Port 5000 /health", "30s"],
        ["Milvus", "Port 19530 gRPC ping", "30s"],
        ["RAG API", "Port 5001 /health", "30s"],
        ["Chat UI", "Port 8501 /healthz", "30s"],
        ["MolMIM NIM", "Port 8001 /v1/health/ready", "30s"],
        ["DiffDock NIM", "Port 8002 /v1/health/ready", "30s"],
        ["Discovery UI", "Port 8505 /healthz", "30s"],
        ["Grafana", "Port 3000 /api/health", "30s"],
        ["Prometheus", "Port 9099 /-/healthy", "30s"],
        ["DCGM Exporter", "Port 9400 /metrics", "30s"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. STAGE 1: GENOMICS PIPELINE
# ══════════════════════════════════════════════════════════════
H1("5. Stage 1: Genomics Pipeline")

H2("Overview")
body(
    "Stage 1 takes raw FASTQ files from an Illumina sequencer and produces "
    "a Variant Call Format (VCF) file using NVIDIA Parabricks \u2014 a "
    "GPU-accelerated implementation of industry-standard bioinformatics tools."
)

H2("Input Specifications")
add_table(
    ["Parameter", "Value"],
    [
        ["Sample", "HG002 (GIAB reference standard)"],
        ["Coverage", "30\u00d7 whole-genome sequencing (WGS)"],
        ["Read Length", "2\u00d7250 bp paired-end"],
        ["File Size", "~200 GB (FASTQ pair)"],
        ["Reference Genome", "GRCh38 (3.1 GB, pre-indexed)"],
        ["Format", "FASTQ (gzip-compressed)"],
    ],
)
spacer()

H3("Step 1: BWA-MEM2 Alignment (fq2bam)")
code_block(
    "pbrun fq2bam \\\n"
    "  --ref /reference/GRCh38.fa \\\n"
    "  --in-fq /data/HG002_R1.fastq.gz /data/HG002_R2.fastq.gz \\\n"
    "  --out-bam /output/HG002.bam \\\n"
    "  --num-gpus 1",
    "bash"
)
spacer()
add_table(
    ["Metric", "Value"],
    [
        ["Duration", "20-45 minutes"],
        ["GPU Utilization", "70-90%"],
        ["Peak Memory", "~40 GB"],
        ["Output", "Sorted BAM + BAI index"],
        ["Algorithm", "BWA-MEM2 (GPU-accelerated)"],
    ],
)
spacer()

H3("Step 2: DeepVariant Variant Calling")
code_block(
    "pbrun deepvariant \\\n"
    "  --ref /reference/GRCh38.fa \\\n"
    "  --in-bam /output/HG002.bam \\\n"
    "  --out-variants /output/HG002.vcf.gz \\\n"
    "  --num-gpus 1",
    "bash"
)
spacer()
add_table(
    ["Metric", "Value"],
    [
        ["Duration", "10-35 minutes"],
        ["GPU Utilization", "80-95%"],
        ["Peak Memory", "~60 GB"],
        ["Output", "VCF (gzip-compressed + tabix index)"],
        ["Algorithm", "Google DeepVariant (CNN-based, >99% accuracy)"],
    ],
)
spacer()

H2("VCF Output Statistics")
add_table(
    ["Metric", "Count"],
    [
        ["Total Variants", "~11.7M"],
        ["High-Quality (QUAL>30)", "~3.5M"],
        ["SNPs", "~4.2M"],
        ["Indels", "~1.0M"],
        ["Coding Region Variants", "~35,000"],
        ["Multi-allelic Sites", "~150,000"],
    ],
)
spacer()

H2("Parabricks Container")
body("Image: nvcr.io/nvidia/clara/clara-parabricks:4.6.0-1")
body("GPU: Required (CUDA). Volumes: /reference, /data, /output. Port: 5000 (Flask web portal).")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. STAGE 2: RAG/CHAT PIPELINE
# ══════════════════════════════════════════════════════════════
H1("6. Stage 2: RAG/Chat Pipeline")

H2("Overview")
body(
    "Stage 2 annotates VCF variants with clinical and functional databases, "
    "indexes them in a Milvus vector database, and uses Anthropic Claude with "
    "RAG to identify druggable gene targets supported by evidence."
)

H2("Architecture")
code_block(
    "VCF (11.7M variants)\n"
    "  \u2192 Quality filter (QUAL>30) \u2192 3.5M variants\n"
    "  \u2192 ClinVar annotation \u2192 clinical significance\n"
    "  \u2192 AlphaMissense annotation \u2192 pathogenicity prediction\n"
    "  \u2192 VEP annotation \u2192 functional consequences\n"
    "  \u2192 BGE-small-en-v1.5 embedding \u2192 384-dim vectors\n"
    "  \u2192 Milvus IVF_FLAT indexing \u2192 3.5M searchable embeddings\n"
    "  \u2192 Claude RAG query \u2192 target hypothesis with evidence chain"
)
spacer()

H2("Annotation Funnel")
add_table(
    ["Stage", "Variant Count", "Filter"],
    [
        ["Raw VCF", "~11.7M", "\u2014"],
        ["Quality filter", "~3.5M", "QUAL > 30"],
        ["ClinVar match", "~35,616", "Clinical significance annotated"],
        ["AlphaMissense match", "~6,831", "AI pathogenicity predicted"],
        ["Coding + pathogenic", "~2,400", "Actionable subset"],
    ],
)
spacer()

H2("Embedding Model")
add_table(
    ["Parameter", "Value"],
    [
        ["Model", "BGE-small-en-v1.5"],
        ["Dimensions", "384"],
        ["Index Type", "IVF_FLAT"],
        ["Index Params", "nlist=1024"],
        ["Search Params", "nprobe=16"],
        ["Distance Metric", "COSINE"],
        ["Total Embeddings", "~3.5M"],
    ],
)
spacer()

H2("Query Flow")
bullet("1.", "User asks a natural language question in the Streamlit chat")
bullet("2.", "Query is expanded using 10 therapeutic area keyword maps")
bullet("3.", "BGE-small-en-v1.5 embeds the expanded query")
bullet("4.", "Milvus performs approximate nearest-neighbor search (top_k=20)")
bullet("5.", "Retrieved variant contexts are assembled into a RAG prompt")
bullet("6.", "Claude processes the prompt with knowledge base grounding")
bullet("7.", "Response includes gene target, evidence chain, and confidence")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. MILVUS VECTOR DATABASE SCHEMA
# ══════════════════════════════════════════════════════════════
H1("7. Milvus Vector Database Schema")

H2("Collection: genomic_evidence")
body("17 fields capturing genomic position, annotation, and embedding:")
add_table(
    ["Field", "Type", "Description"],
    [
        ["id", "INT64 (PK, auto)", "Primary key"],
        ["embedding", "FLOAT_VECTOR(384)", "BGE-small-en-v1.5 embedding"],
        ["chrom", "VARCHAR(10)", "Chromosome (chr1-22, chrX, chrY)"],
        ["pos", "INT64", "Genomic position"],
        ["ref", "VARCHAR(1000)", "Reference allele"],
        ["alt", "VARCHAR(1000)", "Alternate allele"],
        ["qual", "FLOAT", "Variant quality score"],
        ["gene", "VARCHAR(100)", "Gene symbol"],
        ["consequence", "VARCHAR(200)", "Functional consequence"],
        ["impact", "VARCHAR(20)", "HIGH, MODERATE, LOW, MODIFIER"],
        ["genotype", "VARCHAR(10)", "Sample genotype (0/1, 1/1)"],
        ["text_summary", "VARCHAR(2000)", "Human-readable description"],
        ["clinical_significance", "VARCHAR(200)", "ClinVar classification"],
        ["rsid", "VARCHAR(20)", "dbSNP identifier"],
        ["disease_associations", "VARCHAR(2000)", "Associated diseases"],
        ["am_pathogenicity", "FLOAT", "AlphaMissense score (0-1)"],
        ["am_class", "VARCHAR(20)", "pathogenic/ambiguous/benign"],
    ],
)
spacer()

H2("Index Configuration")
code_block(
    'index_params = {\n'
    '    "index_type": "IVF_FLAT",\n'
    '    "metric_type": "COSINE",\n'
    '    "params": {"nlist": 1024}\n'
    '}\n'
    '\n'
    'search_params = {\n'
    '    "metric_type": "COSINE",\n'
    '    "params": {"nprobe": 16}\n'
    '}',
    "python"
)
spacer()

H2("Milvus Infrastructure")
add_table(
    ["Component", "Port", "Purpose"],
    [
        ["Milvus standalone", "19530", "gRPC vector operations"],
        ["Attu UI", "8000", "Web-based Milvus management"],
        ["etcd", "2379", "Metadata storage"],
        ["MinIO", "9000", "Object storage for indexes"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. VARIANT ANNOTATION PIPELINE
# ══════════════════════════════════════════════════════════════
H1("8. Variant Annotation Pipeline")

H2("ClinVar Integration")
add_table(
    ["Parameter", "Value"],
    [
        ["Database", "ClinVar (NCBI)"],
        ["Total Variants", "4.1M clinical variants"],
        ["Match Rate", "~35,616 / 3.5M variants (1.0%)"],
        ["Classifications", "Pathogenic, Likely pathogenic, VUS, Likely benign, Benign"],
        ["Update Frequency", "Monthly releases"],
    ],
)
spacer()

H2("AlphaMissense Integration")
add_table(
    ["Parameter", "Value"],
    [
        ["Database", "AlphaMissense (DeepMind)"],
        ["Total Predictions", "71,697,560 missense variant predictions"],
        ["Match Rate", "~6,831 / 35,616 ClinVar variants (19.2%)"],
        ["Model", "AlphaFold-derived protein structure features"],
        ["Output", "Pathogenicity score (0.0-1.0)"],
    ],
)
spacer()

H3("AlphaMissense Thresholds")
add_table(
    ["Class", "Score Range", "Interpretation"],
    [
        ["Pathogenic", "> 0.564", "Likely disease-causing"],
        ["Ambiguous", "0.34 \u2013 0.564", "Uncertain significance"],
        ["Benign", "< 0.34", "Likely neutral"],
    ],
)
spacer()

H2("Ensembl VEP Integration")
add_table(
    ["Parameter", "Value"],
    [
        ["Tool", "Ensembl Variant Effect Predictor (VEP)"],
        ["Purpose", "Functional consequence annotation"],
        ["Impact Levels", "HIGH, MODERATE, LOW, MODIFIER"],
        ["Key Consequences", "missense_variant, stop_gained, frameshift_variant, splice_donor_variant"],
    ],
)
spacer()

H2("Annotation Pipeline Code Pattern")
code_block(
    "def annotate_variants(vcf_path: str) -> List[AnnotatedVariant]:\n"
    '    """VCF \u2192 ClinVar \u2192 AlphaMissense \u2192 VEP \u2192 Annotated variants"""\n'
    "    variants = parse_vcf(vcf_path, min_qual=30)        # ~3.5M pass\n"
    "    variants = annotate_clinvar(variants)                # Clinical significance\n"
    "    variants = annotate_alphamissense(variants)          # AI pathogenicity\n"
    "    variants = annotate_vep(variants)                    # Functional consequences\n"
    "    return variants",
    "python"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. KNOWLEDGE BASE
# ══════════════════════════════════════════════════════════════
H1("9. Knowledge Base \u2014 201 Genes, 13 Therapeutic Areas")

H2("Gene Distribution")
add_table(
    ["Therapeutic Area", "Count", "Example Genes"],
    [
        ["Neurology", "36", "VCP, APP, PSEN1, MAPT, SOD1, FUS, C9orf72"],
        ["Oncology", "27", "EGFR, BRAF, KRAS, TP53, BRCA1, BRCA2, PIK3CA"],
        ["Metabolic", "22", "GCK, PPARG, SLC2A2, ABCA1, PCSK9"],
        ["Infectious Disease", "21", "ACE2, CCR5, IFITM3, TLR4, TMPRSS2"],
        ["Respiratory", "13", "CFTR, SERPINA1, MUC5B, TERT"],
        ["Rare Disease", "12", "VCP, HTT, SMN1, DMD, CFTR"],
        ["Hematology", "12", "HBB, HBA1, F5, JAK2, CALR"],
        ["GI/Hepatology", "12", "HFE, ATP7B, NOD2, SERPINA1"],
        ["Pharmacogenomics", "11", "CYP2D6, CYP2C19, CYP3A4, DPYD, TPMT"],
        ["Ophthalmology", "11", "RHO, RPE65, RS1, ABCA4"],
        ["Cardiovascular", "10", "LDLR, PCSK9, SCN5A, MYBPC3, KCNQ1"],
        ["Immunology", "9", "HLA-B, TNF, IL6, JAK1, CTLA4"],
        ["Dermatology", "9", "FLG, MC1R, TYR, KRT14"],
    ],
)
spacer()
body("Total: 201 genes, 171 druggable targets (85% druggability rate).")

H2("Knowledge Base Entry Structure")
code_block(
    '{\n'
    '    "gene": "VCP",\n'
    '    "uniprot": "P55072",\n'
    '    "therapeutic_area": "Neurology",\n'
    '    "diseases": ["Frontotemporal Dementia", "ALS", "IBMPFD"],\n'
    '    "druggability": "High",\n'
    '    "drug_targets": ["D2 ATPase domain", "N-D1 interface"],\n'
    '    "known_inhibitors": ["CB-5083", "NMS-873"],\n'
    '    "variant_hotspots": ["R155H", "R191Q", "A232E"],\n'
    '    "pathway": "Ubiquitin-proteasome system",\n'
    '    "mechanism": "AAA+ ATPase, protein homeostasis"\n'
    '}',
    "python"
)
spacer()

H2("Query Expansion Maps")
body(
    "10 therapeutic area query expansion maps enrich user queries with "
    "domain-specific terminology for improved Milvus retrieval."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 10. ANTHROPIC CLAUDE LLM INTEGRATION
# ══════════════════════════════════════════════════════════════
H1("10. Anthropic Claude LLM Integration")

H2("Configuration")
add_table(
    ["Parameter", "Value"],
    [
        ["Model", "claude-sonnet-4-20250514"],
        ["Temperature", "0.3"],
        ["Max Tokens", "4096"],
        ["API", "Anthropic Messages API"],
        ["Role", "RAG-grounded clinical reasoning"],
    ],
)
spacer()

H2("RAG Prompt Structure")
code_block(
    'system_prompt = """You are a clinical genomics specialist\n'
    "analyzing patient variant data. Ground all responses in\n"
    "the retrieved variant evidence and knowledge base. Cite\n"
    "specific variants, genes, and clinical classifications.\n"
    "When recommending drug targets, explain the evidence\n"
    'chain from variant to disease mechanism to druggability."""\n'
    "\n"
    'user_prompt = f"""\n'
    "## Retrieved Variant Evidence (top {top_k} matches)\n"
    "{formatted_variants}\n"
    "\n"
    "## Knowledge Base Context\n"
    "{knowledge_context}\n"
    "\n"
    "## User Question\n"
    '{user_question}\n'
    '"""',
    "python"
)
spacer()

H2("Response Format")
body("Claude generates structured target hypotheses including gene, "
     "confidence level, evidence chain, therapeutic area, diseases, "
     "and recommended action for downstream drug discovery.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 11. STAGE 3: DRUG DISCOVERY PIPELINE
# ══════════════════════════════════════════════════════════════
H1("11. Stage 3: Drug Discovery Pipeline")

H2("Overview")
body(
    "Stage 3 takes a target gene hypothesis from Stage 2 and produces 100 "
    "ranked novel drug candidates using BioNeMo generative chemistry, "
    "molecular docking, and drug-likeness scoring."
)

H2("10-Stage Pipeline")
add_table(
    ["Stage", "Process", "Description"],
    [
        ["1", "Initialize", "Load target hypothesis, validate inputs"],
        ["2", "Normalize Target", "Map gene \u2192 UniProt ID \u2192 PDB structures"],
        ["3", "Structure Discovery", "Query RCSB PDB for Cryo-EM/X-ray structures"],
        ["4", "Structure Preparation", "Score and rank structures, select best site"],
        ["5", "Molecule Generation", "MolMIM generates novel SMILES from seed"],
        ["6", "Chemistry QC", "RDKit validates chemical feasibility"],
        ["7", "Conformer Generation", "RDKit 3D conformer embedding (ETKDG)"],
        ["8", "Molecular Docking", "DiffDock predicts binding poses and affinities"],
        ["9", "Composite Ranking", "30% gen + 40% dock + 30% QED weighted scoring"],
        ["10", "Reporting", "PDF report generation (ReportLab)"],
    ],
)
spacer()

H2("Pipeline Configuration")
code_block(
    'PIPELINE_CONFIG = {\n'
    '    "num_candidates": 100,\n'
    '    "molmim_endpoint": "http://localhost:8001/v1/generate",\n'
    '    "diffdock_endpoint": "http://localhost:8002/v1/dock",\n'
    '    "min_qed": 0.3,\n'
    '    "min_dock_score": -6.0,         # kcal/mol\n'
    '    "scoring_weights": {\n'
    '        "generation": 0.30,\n'
    '        "docking": 0.40,\n'
    '        "qed": 0.30\n'
    '    }\n'
    '}',
    "python"
)
spacer()

H2("UniProt Mappings")
add_table(
    ["Gene", "UniProt ID", "Function"],
    [
        ["VCP", "P55072", "AAA+ ATPase, protein homeostasis"],
        ["EGFR", "P00533", "Receptor tyrosine kinase"],
        ["BRAF", "P15056", "Serine/threonine kinase"],
        ["KRAS", "P01116", "GTPase signaling"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 12. BIONEMO NIM SERVICES
# ══════════════════════════════════════════════════════════════
H1("12. BioNeMo NIM Services")

H2("MolMIM (Port 8001) \u2014 Molecule Generation")
add_table(
    ["Parameter", "Value"],
    [
        ["Endpoint", "POST http://localhost:8001/v1/generate"],
        ["Model", "MolMIM (Molecular Masked Inverse Model)"],
        ["Input", "Seed SMILES string"],
        ["Output", "Novel SMILES candidates"],
        ["Container", "nvcr.io/nvidia/clara/bionemo-molmim:1.0"],
    ],
)
spacer()
H3("MolMIM Request/Response")
code_block(
    '# Request\n'
    '{"smiles": "CC(=O)Nc1ccc(O)cc1",\n'
    ' "num_molecules": 100,\n'
    ' "temperature": 0.7, "top_k": 50}\n'
    '\n'
    '# Response\n'
    '{"molecules": [\n'
    '  {"smiles": "CC(=O)Nc1ccc(O)c(F)c1", "score": 0.85},\n'
    '  {"smiles": "CC(=O)Nc1ccc(O)c(Cl)c1", "score": 0.82}\n'
    ']}',
    "json"
)
spacer()

H2("DiffDock (Port 8002) \u2014 Molecular Docking")
add_table(
    ["Parameter", "Value"],
    [
        ["Endpoint", "POST http://localhost:8002/v1/dock"],
        ["Model", "DiffDock (diffusion-based docking)"],
        ["Input", "Ligand SMILES + protein PDB structure"],
        ["Output", "Binding pose + affinity score (kcal/mol)"],
        ["Container", "nvcr.io/nvidia/clara/diffdock:1.0"],
    ],
)
spacer()

H2("Docking Score Interpretation")
add_table(
    ["Score (kcal/mol)", "Interpretation"],
    [
        ["-12 to -8", "Excellent binding affinity"],
        ["-8 to -6", "Good binding affinity"],
        ["-6 to -4", "Moderate binding affinity"],
        ["> -4", "Weak binding affinity"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 13. DRUG-LIKENESS SCORING
# ══════════════════════════════════════════════════════════════
H1("13. Drug-Likeness Scoring")

H2("Lipinski\u2019s Rule of Five")
add_table(
    ["Rule", "Threshold", "Description"],
    [
        ["Molecular Weight", "\u2264 500 Da", "Oral absorption limit"],
        ["LogP", "\u2264 5", "Lipophilicity"],
        ["H-Bond Donors", "\u2264 5", "NH + OH groups"],
        ["H-Bond Acceptors", "\u2264 10", "N + O atoms"],
    ],
)
spacer()

H2("QED (Quantitative Estimate of Drug-likeness)")
add_table(
    ["Range", "Interpretation"],
    [
        ["> 0.67", "Drug-like (favorable properties)"],
        ["0.49 \u2013 0.67", "Moderate drug-likeness"],
        ["< 0.49", "Less drug-like"],
    ],
)
spacer()

H2("TPSA (Topological Polar Surface Area)")
add_table(
    ["Range (\u00c5\u00b2)", "Interpretation"],
    [
        ["< 140", "Good oral bioavailability"],
        ["60\u201390", "Optimal range"],
        ["> 140", "Poor oral absorption"],
    ],
)
spacer()

H2("Composite Scoring Formula")
code_block(
    "def compute_composite_score(gen_score, dock_score, qed_score):\n"
    '    """30% generation + 40% docking + 30% QED"""\n'
    "    dock_normalized = max(0.0, min(1.0, (10.0 + dock_score) / 20.0))\n"
    "    composite = (\n"
    "        0.30 * gen_score +\n"
    "        0.40 * dock_normalized +\n"
    "        0.30 * qed_score\n"
    "    )\n"
    "    return composite",
    "python"
)
spacer()

H2("RDKit Property Calculation")
code_block(
    "from rdkit import Chem\n"
    "from rdkit.Chem import Descriptors, QED\n"
    "\n"
    "def calculate_properties(smiles: str) -> dict:\n"
    "    mol = Chem.MolFromSmiles(smiles)\n"
    "    return {\n"
    '        "molecular_weight": Descriptors.MolWt(mol),\n'
    '        "logp": Descriptors.MolLogP(mol),\n'
    '        "hbd": Descriptors.NumHDonors(mol),\n'
    '        "hba": Descriptors.NumHAcceptors(mol),\n'
    '        "tpsa": Descriptors.TPSA(mol),\n'
    '        "qed": QED.qed(mol),\n'
    '        "lipinski_pass": all([\n'
    "            Descriptors.MolWt(mol) <= 500,\n"
    "            Descriptors.MolLogP(mol) <= 5,\n"
    "            Descriptors.NumHDonors(mol) <= 5,\n"
    "            Descriptors.NumHAcceptors(mol) <= 10,\n"
    "        ])\n"
    "    }",
    "python"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 14. CRYO-EM STRUCTURE EVIDENCE
# ══════════════════════════════════════════════════════════════
H1("14. Cryo-EM Structure Evidence")

H2("Structure Scoring Algorithm")
body("The pipeline automatically retrieves and scores PDB structures:")
code_block(
    "def score_structure(structure: StructureInfo) -> float:\n"
    '    """Score PDB structure for drug discovery suitability.\n'
    "    - Resolution: lower is better (max 5 \u00c5 cutoff)\n"
    "    - Inhibitor-bound: +3 bonus\n"
    "    - Druggable pockets: +0.5 per pocket\n"
    '    - Cryo-EM method: +0.5"""\n'
    "    score += max(0, 5.0 - resolution)\n"
    "    if has_inhibitor_bound: score += 3.0\n"
    "    score += num_druggable_pockets * 0.5\n"
    "    if 'Cryo-EM' in method: score += 0.5\n"
    "    return score",
    "python"
)
spacer()

H2("VCP Structures (Demo)")
add_table(
    ["PDB ID", "Resolution", "Method", "Description"],
    [
        ["8OOI", "2.9 \u00c5", "Cryo-EM", "WT VCP hexamer"],
        ["9DIL", "3.2 \u00c5", "Cryo-EM", "Mutant VCP"],
        ["7K56", "2.5 \u00c5", "Cryo-EM", "VCP complex"],
        ["5FTK", "2.3 \u00c5", "X-ray", "VCP + CB-5083 inhibitor"],
    ],
)
spacer()

H2("VCP Binding Site")
add_table(
    ["Parameter", "Value"],
    [
        ["Domain", "D2 ATPase domain"],
        ["Mechanism", "ATP-competitive inhibition"],
        ["Pocket Volume", "~450 \u00c5\u00b3"],
        ["Druggability Score", "0.92"],
        ["Key Residues", "ALA464, GLY479, ASP320, GLY215"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 15. VCP/FTD DEMO WALKTHROUGH
# ══════════════════════════════════════════════════════════════
H1("15. VCP/FTD Demo Walkthrough")

H2("Demo Target: Valosin-Containing Protein (VCP/p97)")
add_table(
    ["Parameter", "Value"],
    [
        ["Gene", "VCP"],
        ["Protein", "p97 / Valosin-Containing Protein"],
        ["UniProt", "P55072"],
        ["Function", "AAA+ ATPase, ubiquitin-proteasome pathway"],
        ["Diseases", "Frontotemporal Dementia (FTD), ALS, IBMPFD"],
        ["Variant", "rs188935092 (chr9:35065263 G>A)"],
        ["ClinVar", "Pathogenic"],
        ["AlphaMissense", "0.87 (pathogenic, >0.564 threshold)"],
        ["Seed Compound", "CB-5083 (Phase I clinical VCP inhibitor)"],
    ],
)
spacer()

H2("Demo Flow")

H3("Stage 1 \u2014 Genomics (Demo Mode: ~20 min)")
bullet("1.", "Load pre-processed HG002 FASTQ subset")
bullet("2.", "Run Parabricks fq2bam alignment")
bullet("3.", "Run DeepVariant variant calling")
bullet("4.", "Output VCF with ~11.7M variants including rs188935092")

H3("Stage 2 \u2014 RAG/Chat (Interactive)")
bullet("1.", "VCF annotated: ClinVar flags rs188935092 as pathogenic in VCP")
bullet("2.", "AlphaMissense scores the missense variant at 0.87 (pathogenic)")
bullet("3.", "3.5M variants embedded and indexed in Milvus")
bullet("4.", 'User queries: "What are the most promising drug targets?"')
bullet("5.", "Claude identifies VCP with full evidence chain")
bullet("6.", "Target hypothesis: VCP \u2192 FTD \u2192 druggable D2 ATPase domain")

H3("Stage 3 \u2014 Drug Discovery (~10 min)")
bullet("1.", "VCP \u2192 UniProt P55072 \u2192 PDB structure retrieval")
bullet("2.", "Cryo-EM structures scored: 8OOI, 9DIL, 7K56, 5FTK")
bullet("3.", "5FTK selected (inhibitor-bound, highest score)")
bullet("4.", "CB-5083 seed SMILES \u2192 MolMIM generates 100 novel analogs")
bullet("5.", "RDKit validates Lipinski, QED, TPSA")
bullet("6.", "DiffDock docks each candidate against VCP D2 domain")
bullet("7.", "Composite ranking: 30% gen + 40% dock + 30% QED")
bullet("8.", "Top candidates: novel VCP inhibitors with improved drug-likeness")
bullet("9.", "PDF report generated via ReportLab")

H2("Expected Demo Output")
code_block(
    "Pipeline: HCLS AI Factory \u2014 VCP/FTD Demo\n"
    "Target: VCP (P55072) \u2014 Frontotemporal Dementia\n"
    "Seed: CB-5083 (ATP-competitive VCP inhibitor)\n"
    "Structure: 5FTK (2.3 \u00c5, X-ray, inhibitor-bound)\n"
    "\n"
    "Results:\n"
    "- 100 novel VCP inhibitor candidates generated\n"
    "- 87 pass Lipinski's Rule of Five\n"
    "- 72 have QED > 0.67 (drug-like)\n"
    "- Top 10: docking scores -8.2 to -11.4 kcal/mol\n"
    "- Composite scores range 0.68-0.89"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 16. PYDANTIC DATA MODELS
# ══════════════════════════════════════════════════════════════
H1("16. Pydantic Data Models")

H2("Core Models (from models.py)")
body("All data flows use Pydantic models for validation:")

models_code = [
    ("TargetHypothesis",
     "class TargetHypothesis(BaseModel):\n"
     '    """Output from Stage 2 \u2014 RAG-identified drug target"""\n'
     "    gene: str                     # e.g., 'VCP'\n"
     "    uniprot_id: str               # e.g., 'P55072'\n"
     "    confidence: str               # high, medium, low\n"
     "    evidence_chain: List[str]\n"
     "    therapeutic_area: str\n"
     "    diseases: List[str]\n"
     "    druggability_score: float     # 0-1 scale"),
    ("RankedCandidate",
     "class RankedCandidate(BaseModel):\n"
     '    """Final ranked drug candidate"""\n'
     "    rank: int\n"
     "    smiles: str\n"
     "    generation_score: float\n"
     "    dock_score: float             # kcal/mol\n"
     "    qed: float\n"
     "    composite_score: float        # 30% gen + 40% dock + 30% QED\n"
     "    lipinski_pass: bool\n"
     "    molecular_weight: float\n"
     "    logp: float"),
    ("PipelineConfig",
     "class PipelineConfig(BaseModel):\n"
     '    """Pipeline execution configuration"""\n'
     "    mode: str                     # full, target, drug, demo\n"
     "    num_candidates: int = 100\n"
     "    min_qed: float = 0.3\n"
     "    min_dock_score: float = -6.0\n"
     '    molmim_url: str = "http://localhost:8001/v1/generate"\n'
     '    diffdock_url: str = "http://localhost:8002/v1/dock"'),
]
for name, code in models_code:
    H3(name)
    code_block(code, "python")
    spacer()

body("Additional models: StructureInfo, StructureManifest, MoleculeProperties, "
     "GeneratedMolecule, DockingResult, PipelineRun.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 17. NEXTFLOW DSL2 ORCHESTRATION
# ══════════════════════════════════════════════════════════════
H1("17. Nextflow DSL2 Orchestration")

H2("Pipeline Modes")
add_table(
    ["Mode", "Stages", "Description"],
    [
        ["full", "1 \u2192 2 \u2192 3", "Complete end-to-end pipeline"],
        ["target", "2 \u2192 3", "Skip genomics, use existing VCF"],
        ["drug", "3 only", "Skip to drug discovery with known target"],
        ["demo", "1 \u2192 2 \u2192 3", "Pre-configured VCP/FTD demonstration"],
        ["genomics_only", "1 only", "Run only variant calling"],
    ],
)
spacer()

H2("Main Pipeline Entry (main.nf)")
code_block(
    "#!/usr/bin/env nextflow\n"
    "nextflow.enable.dsl=2\n"
    "\n"
    "include { GENOMICS_PIPELINE } from './modules/genomics'\n"
    "include { RAG_CHAT_PIPELINE } from './modules/rag_chat'\n"
    "include { DRUG_DISCOVERY_PIPELINE } from './modules/drug_discovery'\n"
    "include { REPORTING } from './modules/reporting'\n"
    "\n"
    "workflow {\n"
    "    if (params.mode in ['full', 'demo', 'genomics_only']) {\n"
    "        GENOMICS_PIPELINE(params.fastq_r1, params.fastq_r2, params.reference)\n"
    "    }\n"
    "    if (params.mode in ['full', 'demo', 'target']) {\n"
    "        RAG_CHAT_PIPELINE(...)\n"
    "    }\n"
    "    if (params.mode in ['full', 'demo', 'target', 'drug']) {\n"
    "        DRUG_DISCOVERY_PIPELINE(...)\n"
    "    }\n"
    "    REPORTING(DRUG_DISCOVERY_PIPELINE.out.candidates)\n"
    "}",
    "groovy"
)
spacer()

H2("Nextflow Profiles")
add_table(
    ["Profile", "Description"],
    [
        ["standard", "Default local execution"],
        ["docker", "Docker container execution"],
        ["singularity", "Singularity container execution"],
        ["dgx_spark", "DGX Spark optimized (GPU resources)"],
        ["slurm", "HPC cluster submission"],
        ["test", "Minimal test data"],
    ],
)
spacer()

H2("Pipeline Launcher (run_pipeline.py)")
code_block(
    "# Full pipeline\n"
    "python run_pipeline.py --mode full \\\n"
    "  --fastq-r1 /data/HG002_R1.fastq.gz \\\n"
    "  --fastq-r2 /data/HG002_R2.fastq.gz \\\n"
    "  --reference /reference/GRCh38.fa\n"
    "\n"
    "# Demo mode (pre-configured VCP/FTD)\n"
    "python run_pipeline.py --mode demo\n"
    "\n"
    "# Drug discovery only\n"
    'python run_pipeline.py --mode drug --target-gene VCP',
    "bash"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 18. LANDING PAGE & SERVICE HEALTH
# ══════════════════════════════════════════════════════════════
H1("18. Landing Page & Service Health")

H2("Landing Page (Port 8080)")
body(
    "The Flask-based landing page serves as the entry point for the "
    "HCLS AI Factory, providing a 10-service health status dashboard, "
    "pipeline mode selector, quick-start links, real-time status with "
    "green/red indicators, and pipeline execution history."
)

H2("Service Health Check Implementation")
code_block(
    'SERVICES = [\n'
    '    {"name": "Parabricks Portal", "port": 5000},\n'
    '    {"name": "Milvus Vector DB", "port": 19530},\n'
    '    {"name": "RAG API", "port": 5001},\n'
    '    {"name": "Streamlit Chat", "port": 8501},\n'
    '    {"name": "MolMIM NIM", "port": 8001},\n'
    '    {"name": "DiffDock NIM", "port": 8002},\n'
    '    {"name": "Discovery UI", "port": 8505},\n'
    '    {"name": "Grafana", "port": 3000},\n'
    '    {"name": "Prometheus", "port": 9099},\n'
    '    {"name": "DCGM Exporter", "port": 9400},\n'
    ']',
    "python"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 19. MONITORING STACK
# ══════════════════════════════════════════════════════════════
H1("19. Monitoring Stack")

H2("Grafana (Port 3000)")
add_table(
    ["Parameter", "Value"],
    [
        ["Image", "grafana/grafana:10.2.2"],
        ["Default User", "admin / changeme"],
        ["Dashboards", "HCLS AI Factory (GPU, pipeline, services)"],
        ["Data Source", "Prometheus"],
    ],
)
spacer()

H2("Prometheus (Port 9099)")
add_table(
    ["Parameter", "Value"],
    [
        ["Image", "prom/prometheus:v2.48.0"],
        ["Internal Port", "9090 \u2192 External 9099"],
        ["Retention", "30 days"],
        ["Targets", "Node Exporter, DCGM Exporter, service metrics"],
    ],
)
spacer()

H2("DCGM Exporter (Port 9400)")
add_table(
    ["Metric", "Description"],
    [
        ["DCGM_FI_DEV_GPU_UTIL", "GPU utilization percentage"],
        ["DCGM_FI_DEV_FB_USED", "GPU memory used (bytes)"],
        ["DCGM_FI_DEV_FB_FREE", "GPU memory free (bytes)"],
        ["DCGM_FI_DEV_GPU_TEMP", "GPU temperature (\u00b0C)"],
        ["DCGM_FI_DEV_POWER_USAGE", "GPU power draw (watts)"],
        ["DCGM_FI_DEV_SM_CLOCK", "SM clock frequency (MHz)"],
    ],
)
spacer()

H2("Key Dashboard Panels")
bullet("1.", "GPU Utilization Timeline \u2014 fq2bam \u2192 DeepVariant \u2192 MolMIM/DiffDock bursts")
bullet("2.", "Pipeline Stage Progress \u2014 Stage 1/2/3 completion with timing")
bullet("3.", "Memory Pressure \u2014 Unified memory usage across CPU + GPU")
bullet("4.", "Service Health Grid \u2014 Green/red status for all 10 services")
bullet("5.", "Variant Processing Rate \u2014 Variants annotated per second")
bullet("6.", "Drug Discovery Throughput \u2014 Molecules generated/docked per minute")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 20. CROSS-MODAL INTEGRATION
# ══════════════════════════════════════════════════════════════
H1("20. Cross-Modal Integration")

H2("HCLS AI Factory Ecosystem")
body(
    "The genomics-to-drug-discovery pipeline integrates with the broader "
    "HCLS AI Factory, including the Imaging Intelligence Agent:"
)

H3("Cross-Modal Triggers")
add_table(
    ["Trigger", "Source", "Target", "Action"],
    [
        ["Lung-RADS 4B+", "Imaging Agent", "Genomics Pipeline", "Initiate tumor profiling"],
        ["Pathogenic Variant", "Genomics Pipeline", "Drug Discovery", "Generate targeted therapies"],
        ["Drug Candidates", "Drug Discovery", "Imaging Agent", "Combined clinical report"],
    ],
)
spacer()

H2("NVIDIA FLARE \u2014 Federated Learning")
body(
    "For multi-site deployments (Phase 3), NVIDIA FLARE enables federated "
    "model training. Models train locally at each site; only model updates "
    "(not patient data) are shared. Raw genomic data never leaves the institution."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 21. CONFIGURATION REFERENCE
# ══════════════════════════════════════════════════════════════
H1("21. Configuration Reference")

H2("Environment Variables")
add_table(
    ["Variable", "Default", "Description"],
    [
        ["ANTHROPIC_API_KEY", "(required)", "Anthropic API key for Claude"],
        ["NGC_API_KEY", "(required)", "NVIDIA NGC key for BioNeMo NIMs"],
        ["REFERENCE_GENOME", "/reference/GRCh38.fa", "Path to GRCh38 reference"],
        ["MILVUS_HOST", "localhost", "Milvus server hostname"],
        ["MILVUS_PORT", "19530", "Milvus gRPC port"],
        ["MOLMIM_URL", "http://localhost:8001", "MolMIM NIM endpoint"],
        ["DIFFDOCK_URL", "http://localhost:8002", "DiffDock NIM endpoint"],
        ["CLAUDE_MODEL", "claude-sonnet-4-20250514", "Claude model identifier"],
        ["CLAUDE_TEMPERATURE", "0.3", "LLM temperature"],
        ["PIPELINE_MODE", "full", "Pipeline execution mode"],
        ["NUM_CANDIDATES", "100", "Drug candidates to generate"],
        ["MIN_QED", "0.3", "Minimum QED threshold"],
        ["MIN_DOCK_SCORE", "-6.0", "Minimum docking score (kcal/mol)"],
    ],
)
spacer()

H2("AlphaMissense Thresholds")
code_block(
    "AM_PATHOGENIC_THRESHOLD = 0.564\n"
    "AM_AMBIGUOUS_LOWER = 0.34\n"
    "AM_AMBIGUOUS_UPPER = 0.564\n"
    "AM_BENIGN_THRESHOLD = 0.34",
    "python"
)
spacer()

H2("Scoring Weights")
code_block(
    'SCORING_WEIGHTS = {\n'
    '    "generation": 0.30,   # MolMIM generation confidence\n'
    '    "docking": 0.40,      # DiffDock binding affinity\n'
    '    "qed": 0.30           # RDKit drug-likeness\n'
    '}',
    "python"
)
spacer()

H2("Drug-Likeness Thresholds")
code_block(
    'LIPINSKI = {"max_mw": 500, "max_logp": 5, "max_hbd": 5, "max_hba": 10}\n'
    'QED = {"drug_like": 0.67, "moderate": 0.49}\n'
    'DOCKING = {"excellent": -8.0, "good": -6.0, "moderate": -4.0, "minimum": -6.0}',
    "python"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 22. DEPLOYMENT ROADMAP
# ══════════════════════════════════════════════════════════════
H1("22. Deployment Roadmap")

H2("Phase 1: Proof Build")
add_table(
    ["Parameter", "Value"],
    [
        ["Hardware", "NVIDIA DGX Spark ($4,699)"],
        ["Orchestration", "Docker Compose"],
        ["Scale", "Single patient, sequential processing"],
        ["GPU", "1\u00d7 GB10"],
        ["Memory", "128 GB unified"],
    ],
)
spacer()

H2("Phase 2: Departmental")
add_table(
    ["Parameter", "Value"],
    [
        ["Hardware", "1\u20132\u00d7 DGX B200"],
        ["Orchestration", "Kubernetes"],
        ["Scale", "Multiple concurrent patients"],
        ["GPU", "8\u00d7 B200 per node"],
        ["Memory", "1\u20132 TB HBM3e"],
    ],
)
spacer()

H2("Phase 3: Enterprise / Multi-Site")
add_table(
    ["Parameter", "Value"],
    [
        ["Hardware", "DGX SuperPOD"],
        ["Orchestration", "Kubernetes + NVIDIA FLARE"],
        ["Scale", "Thousands of concurrent patients"],
        ["GPU", "Hundreds of B200 GPUs"],
        ["Privacy", "Federated learning (data stays local)"],
    ],
)
spacer()

H2("Scaling Considerations")
add_table(
    ["Bottleneck", "Phase 1 Solution", "Phase 2+ Solution"],
    [
        ["Genomics throughput", "Sequential (1 sample)", "Parallel Parabricks instances"],
        ["Milvus query latency", "Single-node Milvus", "Milvus cluster with sharding"],
        ["BioNeMo inference", "Single NIM per model", "Multiple NIM replicas"],
        ["Storage I/O", "NVMe direct", "GPUDirect Storage + RAID"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 23. TESTING STRATEGY
# ══════════════════════════════════════════════════════════════
H1("23. Testing Strategy")

H2("Unit Tests")
add_table(
    ["Component", "Test Focus"],
    [
        ["VCF Parser", "Variant extraction, quality filtering"],
        ["Annotator", "ClinVar/AlphaMissense/VEP lookup accuracy"],
        ["Milvus Client", "Index creation, search recall"],
        ["MolMIM Client", "SMILES generation, request format"],
        ["DiffDock Client", "Docking request/response parsing"],
        ["RDKit Scoring", "Lipinski, QED, TPSA calculations"],
        ["Composite Scorer", "Weight application, normalization"],
    ],
)
spacer()

H2("Integration Tests")
add_table(
    ["Test", "Validates"],
    [
        ["VCF \u2192 Annotation \u2192 Milvus", "End-to-end Stage 2 pipeline"],
        ["Target \u2192 PDB \u2192 MolMIM \u2192 DiffDock", "End-to-end Stage 3 pipeline"],
        ["Health check endpoints", "All 10 services responding"],
        ["Nextflow modes", "full, target, drug, demo execution"],
    ],
)
spacer()

H2("Demo Mode Validation")
body(
    "The demo pipeline mode uses pre-configured inputs to validate the "
    "complete pipeline. Input: HG002 FASTQ subset. Expected: VCP identified "
    "as target with rs188935092 evidence. Output: 100 ranked novel VCP "
    "inhibitor candidates."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 24. IMPLEMENTATION SEQUENCE
# ══════════════════════════════════════════════════════════════
H1("24. Implementation Sequence")

H2("Recommended Build Order")
steps = [
    ("1.", "Infrastructure: Docker Compose, Milvus, monitoring stack"),
    ("2.", "Stage 1 \u2014 Genomics: Parabricks container, fq2bam, DeepVariant, VCF output"),
    ("3.", "Stage 2 \u2014 Annotation: ClinVar + AlphaMissense + VEP pipeline"),
    ("4.", "Stage 2 \u2014 Vector DB: Milvus schema, BGE embedding, IVF_FLAT index"),
    ("5.", "Stage 2 \u2014 RAG: Claude integration, knowledge base, query expansion"),
    ("6.", "Stage 2 \u2014 Chat UI: Streamlit interface, REST API"),
    ("7.", "Stage 3 \u2014 Structure: RCSB PDB retrieval, Cryo-EM scoring"),
    ("8.", "Stage 3 \u2014 Generation: MolMIM NIM, molecule generation"),
    ("9.", "Stage 3 \u2014 Docking: DiffDock NIM, binding prediction"),
    ("10.", "Stage 3 \u2014 Scoring: RDKit properties, composite ranking"),
    ("11.", "Stage 3 \u2014 Reporting: PDF generation, Discovery UI"),
    ("12.", "Orchestration: Nextflow DSL2, pipeline modes, landing page"),
    ("13.", "Testing: Unit tests, integration tests, demo mode validation"),
    ("14.", "Monitoring: Grafana dashboards, alerting rules"),
]
for num, desc in steps:
    bullet(num, desc)

spacer()

H2("Key Dependencies")
code_block(
    "GRCh38 reference \u2192 BWA-MEM2 index \u2192 fq2bam alignment\n"
    "ClinVar + AlphaMissense databases \u2192 Annotation pipeline\n"
    "Milvus running \u2192 Embedding indexing \u2192 RAG queries\n"
    "BioNeMo NIMs running \u2192 Molecule generation + docking\n"
    "All services healthy \u2192 Landing page green status"
)

spacer(12)
p = P("", after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
R(p, "This Project Bible is the authoritative technical reference for the "
  "HCLS AI Factory. All other documentation assets derive their technical "
  "details from this source.", italic=True, size=10, color=GRAY_META)

# ══════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════
import pathlib, subprocess, shutil

OUT = pathlib.Path(__file__).resolve().parent
docx_path = OUT / "HCLS_AI_Factory_Project_Bible.docx"
pdf_path = OUT / "HCLS_AI_Factory_Project_Bible.pdf"

doc.save(str(docx_path))
print(f"DOCX saved: {docx_path} ({docx_path.stat().st_size:,} bytes)")

# Attempt PDF conversion
if shutil.which("libreoffice"):
    subprocess.run([
        "libreoffice", "--headless", "--convert-to", "pdf",
        "--outdir", str(OUT), str(docx_path)
    ], capture_output=True, timeout=120)
    if pdf_path.exists():
        print(f"PDF  saved: {pdf_path} ({pdf_path.stat().st_size:,} bytes)")
    else:
        print("PDF conversion completed but file not found.")
elif shutil.which("docx2pdf"):
    subprocess.run(["docx2pdf", str(docx_path), str(pdf_path)],
                    capture_output=True, timeout=120)
    if pdf_path.exists():
        print(f"PDF  saved: {pdf_path} ({pdf_path.stat().st_size:,} bytes)")
else:
    print("No PDF converter found (install libreoffice or docx2pdf).")

print("Done.")
